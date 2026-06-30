import os
from fastapi import FastAPI, HTTPException, Query, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware

class CachedStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        response = await super().get_response(path, scope)
        if "assets/" in path or path.startswith("assets/"):
            response.headers["Cache-Control"] = "public, max-age=31536000, immutable"
        else:
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response
from pydantic import BaseModel
import psycopg2
import psycopg2.extras
from datetime import datetime
from dotenv import load_dotenv
from openai import AsyncOpenAI
from zoneinfo import ZoneInfo
from typing import List, Optional
import hashlib
from docx import Document
from io import BytesIO
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse, RedirectResponse
from ipaddress import ip_address, ip_network
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi import Depends, status
import bcrypt
from jose import JWTError, jwt
from datetime import timedelta
from authlib.integrations.starlette_client import OAuth
from starlette.middleware.sessions import SessionMiddleware
from psycopg2.pool import ThreadedConnectionPool
from contextlib import contextmanager

load_dotenv()

if not os.getenv("OPENAI_API_KEY"):
    print("ERROR: OPENAI_API_KEY not found in environment or .env file!")
else:
    print("OK: OPENAI_API_KEY detected.")

from cache_utils import cache_manager

app = FastAPI(title="GeoPITS Dashboard")

@app.on_event("startup")
async def startup_event():
    # Increase AnyIO worker thread limit for sync routes to prevent thread starvation under high concurrent requests
    from anyio import to_thread
    limiter = to_thread.current_default_thread_limiter()
    limiter.total_tokens = 500
    print(f"AnyIO thread limiter total_tokens set to {limiter.total_tokens}")

# Link new features
from new_features.backend.migrations import run_migrations
from new_features.backend.routes import router as new_features_router

run_migrations()
app.include_router(new_features_router)

origins = [
    "https://api.geovexsight.geopits.com",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(GZipMiddleware, minimum_size=1000)
app.add_middleware(SessionMiddleware, secret_key=os.getenv("JWT_SECRET", "super-secret-key-geopits"))

DEFAULT_ALLOWED = '127.0.0.1'
ALLOWED_NETWORKS_CONFIG = os.getenv("ALLOWED_IP_NETWORKS", DEFAULT_ALLOWED).split(",")

@app.middleware("http")
async def network_restriction_middleware(request: Request, call_next):
    client_ip_str = "127.0.0.1"
    if request.client:
        client_ip_str = request.client.host
        
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        client_ip_str = forwarded.split(",")[0].strip()
        
    try:
        # Normalize localhost string to IP address if passed
        if client_ip_str == "localhost":
            client_ip_str = "127.0.0.1"
            
        client_ip = ip_address(client_ip_str)
        
        # Allow loopback and all private local IP networks (like 192.168.x.x, 10.x.x.x, 172.16.x.x) for seamless local dev/testing
        if client_ip.is_loopback or client_ip.is_private:
            return await call_next(request)
            
        is_allowed = False
        for network_str in ALLOWED_NETWORKS_CONFIG:
            net_str = network_str.strip().replace("[", "").replace("]", "")
            if not net_str: continue
            
            # Skip non-IP network configurations like 'localhost' which fail in ip_network()
            if net_str == "localhost":
                continue
                
            try:
                if client_ip in ip_network(net_str):
                    is_allowed = True
                    break
            except ValueError:
                continue
                
        if not is_allowed:
            print(f"SECURITY DENIED: {client_ip_str} tried to access {request.url.path}")
            return JSONResponse(
                status_code=403, 
                content={"detail": f"GeoPITS Security: Your IP ({client_ip_str}) is not authorized. Access is restricted to trusted office networks."}
            )
    except Exception as e:
        print(f"IP Filter Error identifying client: {e}")
        # Fallback to local bypass for safety in local execution environments to prevent blocking
        return await call_next(request)

    return await call_next(request)

try:
    openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception as e:
    print(f"Failed to initialize OpenAI client: {e}")
    openai_client = None

os.makedirs("static", exist_ok=True)

@app.get("/")
async def serve_index():
    response = FileResponse("static/index.html")
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

app.mount("/static", CachedStaticFiles(directory="static"), name="static")

SECRET_KEY = os.getenv("JWT_SECRET", "super-secret-key-geopits")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

ADMIN_EMAILS = [email.strip().lower() for email in os.getenv("ADMIN_EMAILS", "").split(",") if email.strip()]

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/login")

try:
    db_pool = ThreadedConnectionPool(
        minconn=5,
        maxconn=150,
        host=os.getenv("DB_HOST", "localhost"),
        database=os.getenv("DB_NAME", "Incoming-error-data"),
        user=os.getenv("DB_USER", "postgres"),
        password=os.getenv("DB_PASSWORD", "y7UMhWmLcqSJzmhTGDyK"),
        port=os.getenv("DB_PORT", "5432")
    )
    print("Database Connection Pool initialized.")
except Exception as e:
    print(f"CRITICAL ERROR: Could not initialize database pool: {e}")
    db_pool = None

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

@contextmanager
def get_db_connection():
    """Context manager for database connections from the pool."""
    if not db_pool:
        conn = psycopg2.connect(
            host=os.getenv("DB_HOST", "localhost"),
            database=os.getenv("DB_NAME", "Incoming-error-data"),
            user=os.getenv("DB_USER", "postgres"),
            password=os.getenv("DB_PASSWORD", "y7UMhWmLcqSJzmhTGDyK"),
            port=os.getenv("DB_PORT", "5432")
        )
        try:
            yield conn
        finally:
            conn.close()
    else:
        conn = db_pool.getconn()
        try:
            yield conn
        finally:
            db_pool.putconn(conn)

def resolve_user_role(email, username, current_role):
    """
    Dynamically determines user role based on assignments in the system_admins and leads table.
    - If email exists in 'system_admins' and is active, user becomes 'admin'.
    - If assigned 'Global' technology in leads, user becomes 'admin' (backward compatibility).
    - If assigned any other technology as lead, user becomes 'lead'.
    """
    email_strip = (email or "").strip().lower()
    username_strip = (username or "").strip().lower()
    
    if not email_strip and not username_strip:
        return current_role
        
    try:
        if current_role == "admin":
            return "admin"
            
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            if email_strip:
                cur.execute("SELECT status FROM system_admins WHERE LOWER(email) = %s AND status = 'active'", (email_strip,))
                if cur.fetchone():
                    cur.close()
                    print(f"DEBUG_ROLE: Promoted '{email_strip}' to admin via system_admins table")
                    return "admin"
            
            cur.execute("""
                SELECT technology, is_lead FROM leads 
                WHERE (TRIM(LOWER(email)) = %s OR TRIM(LOWER(email)) = %s)
                AND status = 'active'
            """, (email_strip, username_strip))
            rows = cur.fetchall()
            cur.close()
            
            if not rows:
                return current_role
                
            techs = [row[0] for row in rows]
            lead_techs = [row[0] for row in rows if row[1] is True]
            
            if any(t in ['Global', 'Global Admin'] for t in techs):
                print(f"DEBUG_ROLE: Promoted '{username_strip}' to admin via Global assignment in leads")
                return "admin"
            
            if lead_techs:
                return "lead"
            
            return current_role
            
    except Exception as e:
        print(f"DEBUG_ROLE_ERROR: Failed to resolve role for {email}/{username}: {e}")
        return current_role

def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
        
    assigned_techs = []
    allowed_clients = []
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT id, username, role, full_name, profile_pic, email FROM users WHERE username = %s;", (username,))
        user_row = cur.fetchone()
        if user_row:
            email_addr = (user_row[5] or "").lower()
            cur.execute("SELECT DISTINCT technology FROM leads WHERE LOWER(email) = LOWER(%s) AND status = 'active';", (email_addr,))
            assigned_techs = [r[0] for r in cur.fetchall()]
            
            # Fetch explicitly permitted client names for user-based access control
            cur.execute("""
                SELECT DISTINCT c.client_name 
                FROM user_clients uc
                JOIN admin_clients c ON uc.client_id = c.id
                WHERE uc.user_id = %s;
            """, (user_row[0],))
            allowed_clients = [r[0] for r in cur.fetchall() if r[0]]
        cur.close()

    if user_row is None:
        raise credentials_exception
        
    user_data = {
        "id": user_row[0],
        "username": user_row[1],
        "role": user_row[2],
        "full_name": user_row[3],
        "profile_pic": user_row[4],
        "email": user_row[5]
    }
    
    email_addr = (user_row[5] or "").lower()
    user_data["role"] = resolve_user_role(email_addr, user_data["username"], user_data["role"])
    
    user_data["isAdmin"] = (user_data["role"] == "admin") or (email_addr in ADMIN_EMAILS)

    user_data["fullName"] = user_row[3]
    user_data["profilePic"] = user_row[4]
    user_data["assigned_techs"] = assigned_techs
    user_data["allowed_clients"] = allowed_clients

    has_access, access_type = check_user_access(username)
    if not has_access:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="GeoPITS Access Restricted: Your account has not been assigned any monitoring privileges. Please contact an administrator."
        )

    print(f"TELEMETRY: Identified user {user_data['username']} (Admin: {user_data['isAdmin']}, Email: {user_data['email']}, Techs: {user_data['assigned_techs']})")
    return user_data

@app.get("/api/me")
def get_me(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("UPDATE users SET last_active_at = %s WHERE username = %s", (datetime.now(ZoneInfo("Asia/Kolkata")), current_user["username"]))
            conn.commit()
            cur.close()
    except Exception as e:
        print(f"Error updating last_active_at: {e}")
    
    print(f"DEBUG_ME_ROLE: User '{current_user.get('username')}' ({current_user.get('email')}) has role: '{current_user.get('role')}' (isAdmin: {current_user.get('isAdmin')})")
    
    return current_user

class ShareRecordRequest(BaseModel):
    notes: str
    platform: str
    content_type: str
    client_name: Optional[str] = ''
    server_name: Optional[str] = ''
    log_message: Optional[str] = ''
    status: Optional[str] = ''
    owner: Optional[str] = ''
    client_visibility: Optional[str] = ''
    ticket_status: Optional[str] = ''
    next_action: Optional[str] = ''
    db_type: Optional[str] = ''

@app.post("/api/share/record")
def record_share(req: ShareRecordRequest, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute(
                """INSERT INTO share_history (username, notes, platform, content_type, client_name, server_name, log_message, status, owner, client_visibility, ticket_status, next_action, db_type) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                (current_user["username"], req.notes, req.platform, req.content_type,
                 req.client_name, req.server_name, req.log_message, req.status, req.owner, req.client_visibility, req.ticket_status, req.next_action, req.db_type)
            )
            conn.commit()
        return {"status": "success"}
    except Exception as e:
        print(f"Share record error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/share/history")
def get_share_history(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = """
                SELECT id, username, notes, platform, content_type, shared_at AS created_at, 
                       client_name, server_name, log_message, status, owner, 
                       client_visibility, ticket_status, next_action, db_type 
                FROM share_history 
                WHERE 1=1
            """
            params = []
            
            allowed_techs = get_user_allowed_technologies(current_user)
            if allowed_techs is not None:
                query += " AND (TRIM(LOWER(db_type)) = ANY(%s) OR username = %s)"
                tech_params = [t.lower().strip() for t in allowed_techs]
                params.append(tech_params)
                params.append(current_user["username"])
            else:
                pass

            query += " ORDER BY created_at DESC LIMIT 50"
            cur.execute(query, tuple(params))
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
            cur.close()
        return {"history": results}
    except Exception as e:
        print(f"Share history fetch error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def check_user_access(username_or_email: str):
    """
    Checks if a user has access to the application.
    Admins (by email or role) have full access.
    Normal users must have at least one active assignment in the 'leads' table.
    """
    username_or_email = username_or_email.lower()
    
    if username_or_email in ADMIN_EMAILS:
        return True, "admin"

    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            if username_or_email in ADMIN_EMAILS:
                cur.close()
                return True, "admin"
                
            cur.execute("SELECT status FROM system_admins WHERE LOWER(email) = %s AND status = 'active'", (username_or_email,))
            if cur.fetchone():
                cur.close()
                return True, "admin"
            
            cur.execute("SELECT role, email FROM users WHERE username = %s OR email = %s", (username_or_email, username_or_email))
            user_row = cur.fetchone()
            if user_row:
                role, email = user_row
                if role == "admin" or (email and email.lower() in ADMIN_EMAILS):
                    cur.close()
                    return True, "admin"
                    
                cur.execute("SELECT COUNT(*) FROM leads WHERE (LOWER(email) = %s OR LOWER(email) = (SELECT LOWER(email) FROM users WHERE username = %s)) AND status = 'active'", (username_or_email, username_or_email))
                count = cur.fetchone()[0]
                
                cur.execute("""
                    SELECT COUNT(*) FROM user_clients uc
                    JOIN users u ON uc.user_id = u.id
                    WHERE u.username = %s OR LOWER(u.email) = %s;
                """, (username_or_email, username_or_email))
                client_count = cur.fetchone()[0]
                cur.close()
                
                if count > 0 or client_count > 0:
                    return True, "user"
                    
        return False, "restricted"
    except Exception as e:
        print(f"Error checking user access: {e}")
        return False, "error"

@app.post("/api/login")
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT hashed_password, email, role FROM users WHERE username = %s;", (form_data.username,))
        row = cur.fetchone()
        cur.close()
    
    if not row or not bcrypt.checkpw(form_data.password.encode('utf-8'), row[0].encode('utf-8')):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    has_access, access_type = check_user_access(form_data.username)
    if not has_access:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="GeoPITS Access Restricted: Your account has not been assigned any monitoring privileges. Please contact an administrator."
        )
        
    resolved_role = resolve_user_role(row[1], form_data.username, row[2])
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": form_data.username, "role": resolved_role}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}
oauth = OAuth()

TENANT_ID = os.getenv("APP_TENANT")
oauth.register(
    name='microsoft',
    client_id=os.getenv("APP_CLIENT"),
    client_secret=os.getenv("APP_SECRET"),
    server_metadata_url=f"https://login.microsoftonline.com/{TENANT_ID}/v2.0/.well-known/openid-configuration" if TENANT_ID else "https://login.microsoftonline.com/common/v2.0/.well-known/openid-configuration",
    client_kwargs={'scope': 'openid email profile', 'prompt': 'select_account'}
)

@app.get("/api/auth/login/{provider}")
async def auth_login(provider: str, request: Request):
    redirect_uri = os.getenv("APP_REDIRECT_URI")
    if not redirect_uri:
        host = request.headers.get("host", "127.0.0.1:8000")
        if "127.0.0.1" in host or "localhost" in host:
            redirect_uri = f"http://{host}/api/auth/callback/{provider}"
        else:
            redirect_uri = f"https://api.geovexsight.geopits.com/api/auth/callback/{provider}"
        
    if provider == 'microsoft':
        return await oauth.microsoft.authorize_redirect(request, redirect_uri)
    else:
        raise HTTPException(status_code=400, detail="Invalid provider")

@app.get("/api/auth/callback/{provider}")
async def auth_callback(provider: str, request: Request):
    try:
        redirect_uri = os.getenv("APP_REDIRECT_URI")
        if not redirect_uri:
            host = request.headers.get("host", "127.0.0.1:8000")
            if "127.0.0.1" in host or "localhost" in host:
                redirect_uri = f"http://{host}/api/auth/callback/{provider}"
            else:
                redirect_uri = f"https://api.geovexsight.geopits.com/api/auth/callback/{provider}"

        if provider == 'microsoft':
            token = await oauth.microsoft.authorize_access_token(request, redirect_uri=redirect_uri)
            user_info = token.get('userinfo')
        else:
            raise HTTPException(status_code=400, detail="Invalid provider")

        if not user_info:
            raise HTTPException(status_code=401, detail="Could not retrieve user info")

        print(f"OAuth {provider} userinfo: {dict(user_info)}")

        email = user_info.get('email') or user_info.get('preferred_username')
        full_name = user_info.get('name') or user_info.get('given_name', '') + ' ' + user_info.get('family_name', '') or email
        full_name = full_name.strip() if full_name else email
        profile_pic = user_info.get('picture')
        username = email.split('@')[0] if email else user_info.get('name', 'oauth_user')

        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT username, email, role FROM users WHERE username = %s OR username = %s", (username, email))
            row = cur.fetchone()
            
            if not row:
                user_role = "user"
                cur.execute(
                    "INSERT INTO users (username, hashed_password, full_name, profile_pic, email, role) VALUES (%s, %s, %s, %s, %s, %s)", 
                    (username, "OAUTH_LOGIN_NO_PASSWORD", full_name, profile_pic, email, user_role)
                )
                conn.commit()
                print(f"Created new OAuth user: {username} ({email})")
            else:
                username = row[0]
                user_role = row[2]
                cur.execute(
                    "UPDATE users SET full_name = %s, profile_pic = %s, email = %s WHERE username = %s",
                    (full_name, profile_pic, email, username)
                )
                conn.commit()
                
            cur.close()

        has_access, access_type = check_user_access(email)
        if not has_access:
             return RedirectResponse(url="/#/login?error=restricted")

        resolved_role = resolve_user_role(email, username, user_role)

        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": username, "role": resolved_role},
            expires_delta=access_token_expires
        )
        
        return RedirectResponse(url=f"/#/?token={access_token}")

    except HTTPException:
        raise
    except Exception as e:
        print(f"OAuth Callback Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))




def get_user_allowed_technologies(user: dict):
    """Returns a list of allowed technologies for the user. Returns None if unrestricted."""
    if user.get("isAdmin") or user.get("role") == "admin":
        return None
    
    email = user.get("email")
    if not email:
        return None
        
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = LOWER(%s) AND status = 'active';", (email,))
            techs = [row[0] for row in cur.fetchall()]
            cur.close()
        return techs if techs else None 
    except Exception as e:
        print(f"Error fetching user privileges: {e}")
        return None

@app.get("/api/users/by-tech")
def get_users_by_tech(db_type: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("SELECT id, username, email, full_name, role FROM users")
            all_users = cur.fetchall()
            
            cur.execute("SELECT email, technology FROM leads WHERE status = 'active'")
            lead_rows = cur.fetchall()
            
            lead_techs = {}
            for row in lead_rows:
                email = (row['email'] or "").lower()
                if email not in lead_techs:
                    lead_techs[email] = set()
                lead_techs[email].add(row['technology'])
            
            caller_email = (current_user.get("email") or "").lower()
            caller_is_admin = (current_user.get("role") == "admin") or (caller_email in ADMIN_EMAILS)
            
            effective_techs = set()
            if db_type:
                effective_techs.add(db_type)
            elif not caller_is_admin and caller_email in lead_techs:
                effective_techs = lead_techs[caller_email]
                
            valid_users = []
            for u in all_users:
                u_email = (u['email'] or "").lower()
                is_admin = (u['role'] == "admin") or (u_email in ADMIN_EMAILS)
                
                if effective_techs and not is_admin:
                    user_techs = lead_techs.get(u_email, set())
                    if not user_techs.intersection(effective_techs):
                        continue
                        
                valid_users.append({
                    "username": u['username'],
                    "label": u['full_name'] or u['username']
                })
                
            valid_users.sort(key=lambda x: x['label'].lower())
            
            cur.close()
        return {"users": valid_users}
        
    except Exception as e:
        print(f"Error fetching users by tech: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/filters")
def get_filters(current_user: dict = Depends(get_current_user)):
    try:
        allowed_techs = get_user_allowed_technologies(current_user)
        is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
        allowed_clients = current_user.get("allowed_clients", [])

        # Thread-safe TTL cache key construction
        techs_key = ",".join(sorted([str(t) for t in (allowed_techs or [])]))
        clients_key = ",".join(sorted([str(c) for c in (allowed_clients or [])]))
        cache_key = f"filters:{is_admin}:{techs_key}:{clients_key}"

        cached_val = cache_manager.get(cache_key)
        if cached_val is not None:
            res_copy = cached_val.copy()
            res_copy["current_ist"] = datetime.now(ZoneInfo("Asia/Kolkata")).isoformat()
            return res_copy

        # Increase lookback to 30 days to ensure filters are populated even if recent logs are sparse
        filter_lookback = (datetime.now(ZoneInfo("Asia/Kolkata")) - timedelta(days=30)).strftime("%Y-%m-%d %H:%M:%S")

        db_types_set = set()
        clients_set = set()
        db_server_map = {}
        db_client_map = {}
        client_server_map = {}
        server_logtype_map = {}

        if not is_admin and not allowed_clients and allowed_techs is not None and len(allowed_techs) == 0:
            return {
                "db_types": [],
                "db_server_map": {},
                "db_client_map": {},
                "clients": [],
                "client_server_map": {},
                "server_logtype_map": {},
                "current_ist": datetime.now(ZoneInfo("Asia/Kolkata")).isoformat()
            }

        with get_db_connection() as conn:
            cur = conn.cursor()
            
            base_all_logs = """
                (
                     SELECT db_type, client_name, server_name, log_type, log_time_ist FROM db_monitoring_logs
                     UNION ALL
                     SELECT db_type, client_name, server_name, log_type, log_time_ist FROM db_archived_logs
                ) as all_logs
            """
            
            query = f"""
                SELECT DISTINCT db_type, client_name, server_name, log_type 
                FROM {base_all_logs}
                WHERE log_time_ist >= %s
            """
            params = [filter_lookback]
            
            if not is_admin:
                if allowed_clients:
                    query += " AND (client_name = ANY(%s) OR server_name = ANY(%s))"
                    params.append(allowed_clients)
                    params.append(allowed_clients)
                if allowed_techs is not None:
                    query += " AND db_type IN %s"
                    params.append(tuple(allowed_techs))
            elif allowed_techs is not None:
                query += " AND db_type IN %s"
                params.append(tuple(allowed_techs))
                
            cur.execute(query, tuple(params))
            rows = cur.fetchall()
            cur.close()

        for db_type, client_name, server_name, log_type in rows:
            if db_type:
                db_types_set.add(db_type)
            if client_name:
                clients_set.add(client_name)
            
            if db_type and server_name:
                if db_type not in db_server_map:
                    db_server_map[db_type] = set()
                db_server_map[db_type].add(server_name)
                
            if db_type and client_name:
                if db_type not in db_client_map:
                    db_client_map[db_type] = set()
                db_client_map[db_type].add(client_name)
                
            if client_name and server_name:
                if client_name not in client_server_map:
                    client_server_map[client_name] = set()
                client_server_map[client_name].add(server_name)
                
            if server_name and log_type:
                if server_name not in server_logtype_map:
                    server_logtype_map[server_name] = set()
                server_logtype_map[server_name].add(log_type)

        # Convert sets to sorted lists for the response
        db_types = sorted(list(db_types_set))
        clients = sorted(list(clients_set))
        db_server_map = {k: sorted(list(v)) for k, v in db_server_map.items()}
        db_client_map = {k: sorted(list(v)) for k, v in db_client_map.items()}
        client_server_map = {k: sorted(list(v)) for k, v in client_server_map.items()}
        server_logtype_map = {k: sorted(list(v)) for k, v in server_logtype_map.items()}

        ist_now = datetime.now(ZoneInfo("Asia/Kolkata")).isoformat()
        
        result = {
            "db_types": db_types,
            "db_server_map": db_server_map,
            "db_client_map": db_client_map,
            "clients": clients, 
            "client_server_map": client_server_map, 
            "server_logtype_map": server_logtype_map,
            "current_ist": ist_now
        }
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except Exception as e:
        import traceback
        print(f"ERROR in get_filters: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

class LogMetadataUpdate(BaseModel):
    client_name: str
    server_name: str
    log_message: str
    log_hash: str
    status: str
    owner: str
    client_visibility: str
    ticket_status: str
    next_action: str
    severity: Optional[str] = None

@app.patch("/api/logs/metadata")
def update_log_metadata(update: LogMetadataUpdate, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            terminal_statuses = ['Resolved', 'Ignored', 'No action Required']
            is_terminal = update.status in terminal_statuses
            
            search_hash = update.log_hash.strip()
            cur.execute("SELECT 1 FROM db_monitoring_logs WHERE TRIM(log_hash) = %s", (search_hash,))
            in_main = cur.fetchone() is not None
            
            cur.execute("SELECT 1 FROM db_archived_logs WHERE TRIM(log_hash) = %s", (search_hash,))
            in_archive = cur.fetchone() is not None
    
            target_table = "db_monitoring_logs"
            if in_main:
                target_table = "db_monitoring_logs"
            elif in_archive:
                target_table = "db_archived_logs"
    
            now_ist = datetime.now(ZoneInfo("Asia/Kolkata"))
            
            common_cols = "client_name, server_name, db_type, log_type, log_source, log_time, log_time_utc, log_time_ist, log_level, log_message, occurrence_count, raw_log, email_subject, email_received_time, log_hash, created_at, status, owner, client_visibility, ticket_status, next_action, severity, status_updated_at"

            if is_terminal and in_main:
                print(f"Archiving log: {update.client_name} | {update.server_name} | {update.log_hash}")
                cur.execute(f"""
                    INSERT INTO db_archived_logs ({common_cols})
                    SELECT {common_cols} FROM db_monitoring_logs 
                    WHERE TRIM(log_hash) = %s
                """, (search_hash,))
                cur.execute("DELETE FROM db_monitoring_logs WHERE TRIM(log_hash) = %s", (search_hash,))
                target_table = "db_archived_logs"
            elif not is_terminal and in_archive:
                print(f"Un-archiving log: {update.client_name} | {update.server_name} | {search_hash}")
                cur.execute(f"""
                    INSERT INTO db_monitoring_logs ({common_cols})
                    SELECT {common_cols} FROM db_archived_logs 
                    WHERE TRIM(log_hash) = %s
                """, (search_hash,))
                cur.execute("DELETE FROM db_archived_logs WHERE TRIM(log_hash) = %s", (search_hash,))
                target_table = "db_monitoring_logs"
            cur.execute(f"""
                UPDATE {target_table}
                SET status = %s,
                    owner = %s,
                    client_visibility = %s,
                    ticket_status = %s,
                    next_action = %s,
                    severity = %s,
                    status_updated_at = %s
                WHERE TRIM(log_hash) = %s
            """, (
                update.status, update.owner, update.client_visibility, update.ticket_status, update.next_action, 
                update.severity, now_ist, search_hash
            ))
            
            print(f"DEBUG_METADATA_UPDATE: Table: {target_table}, Hash: {search_hash[:10]}..., Status: {update.status}, Affected: {cur.rowcount}")
            
            conn.commit()
            cache_manager.invalidate("logs:")
            cache_manager.invalidate("owner-counts:")
            cache_manager.invalidate("filters:")
            return {"success": True}
    except Exception as e:
            print(f"Error updating metadata: {e}")
            if 'conn' in locals(): conn.rollback()
            raise HTTPException(status_code=500, detail=str(e))
    finally:
            if 'cur' in locals(): cur.close()




@app.get("/api/logs")
def get_logs(
    start_time: Optional[str] = None,
    end_time: Optional[str] = None,
    client_name: Optional[str] = None,
    db_type: Optional[str] = None,
    server_name: Optional[str] = None,
    log_type: List[str] = Query(None),
    severity: List[str] = Query(None),
    log_status: List[str] = Query(None),

    limit: int = Query(100, le=1000),
    offset: int = Query(0),
    owner: Optional[str] = None,
    username: Optional[str] = None,
    log_id: Optional[int] = None,
    current_user: dict = Depends(get_current_user)
):
    print(f"DEBUG: get_logs called. user={current_user.get('email')}, status={log_status}, client={client_name}, tech={db_type}, log_id={log_id}, owner={owner}, username={username}")
    
    # Thread-safe TTL cache check
    log_type_key = ",".join(sorted(log_type or []))
    severity_key = ",".join(sorted(severity or []))
    log_status_key = ",".join(sorted(log_status or []))
    user_email = current_user.get("email", "")

    cache_key = f"logs:{start_time}:{end_time}:{client_name}:{db_type}:{server_name}:{log_type_key}:{severity_key}:{log_status_key}:{limit}:{offset}:{owner}:{username}:{log_id}:{user_email}"

    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    if username and not owner:
        owner = username

    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)



            
            log_table = "db_monitoring_logs"
            if log_status and any(s in ['Resolved', 'Ignored', 'No action Required'] for s in log_status):
                log_table = "db_archived_logs"
            
            print(f"DEBUG: Using log_table={log_table}")
    
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            base_query = f"""
                FROM {log_table} 
                WHERE 1=1
            """
            params = []
            
            if not start_time and log_table == "db_monitoring_logs":
                start_time = (datetime.now(ZoneInfo("Asia/Kolkata")) - timedelta(hours=1)).strftime("%Y-%m-%d %H:%M:%S")
                print(f"DEBUG: No start_time provided. Defaulting to last 1 hour: {start_time}")

            if start_time:
                base_query += " AND log_time_ist >= %s"
                params.append(start_time)
            if end_time:
                adj_end_time = end_time
                if len(adj_end_time) == 16:
                    adj_end_time += ":59"
                elif adj_end_time.endswith(":00"):
                    adj_end_time = adj_end_time[:-3] + ":59"
                    
                base_query += " AND log_time_ist <= %s"
                params.append(adj_end_time)
            if client_name:
                base_query += " AND client_name = %s"
                params.append(client_name)
            if db_type:
                base_query += " AND db_type ILIKE %s"
                params.append(db_type)
            if server_name:
                base_query += " AND server_name = %s"
                params.append(server_name)
            if log_type:
                base_query += " AND log_type IN %s"
                params.append(tuple(log_type) if isinstance(log_type, list) else (log_type,))
            if severity:
                base_query += " AND severity IN %s"
                params.append(tuple(severity) if isinstance(severity, list) else (severity,))
                
                
            if log_status:
                base_query += " AND LOWER(status) IN %s"
                params.append(tuple(s.lower() for s in log_status) if isinstance(log_status, list) else (log_status.lower(),))
            else:
                base_query += " AND (status IS NULL OR TRIM(status) = '' OR LOWER(status) NOT IN ('resolved', 'ignored', 'no action required'))"
                
            if owner:
                if owner.lower() == "unassigned":
                    base_query += " AND (owner IS NULL OR owner = '' OR owner = 'Unassigned' OR owner = 'None')"
                elif owner.lower() == "assigned":
                    base_query += " AND (owner IS NOT NULL AND owner != '' AND owner != 'Unassigned' AND owner != 'None')"
                else:
                    base_query += " AND owner ILIKE %s"
                    params.append(f"%{owner}%")
            
            if log_id:
                base_query += " AND id = %s"
                params.append(log_id)

            
            # Apply strict client and technology constraints for standard users
            is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
            if not is_admin:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                
                # Apply allowed_clients constraint
                if allowed_clients:
                    base_query += " AND (client_name = ANY(%s) OR server_name = ANY(%s))"
                    params.append(allowed_clients)
                    params.append(allowed_clients)
                    
                # Apply allowed_techs constraint
                if allowed_techs is not None:
                    if not allowed_techs:
                        return {"logs": [], "total": 0}
                    if db_type and db_type not in allowed_techs:
                        return {"logs": [], "total": 0}
                    base_query += " AND TRIM(LOWER(db_type)) = ANY(%s)"
                    tech_params = [t.lower().strip() for t in allowed_techs]
                    params.append(tech_params)
                    
                # If a standard user has no privileges assigned, return empty results
                if not allowed_clients and allowed_techs is not None and not allowed_techs:
                    return {"logs": [], "total": 0}
            group_by = " GROUP BY client_name, server_name, db_type, log_type, log_message, status, owner, client_visibility, ticket_status, next_action, severity, log_hash"
            
            count_query = f"SELECT COUNT(*) as total_count FROM (SELECT 1 {base_query} {group_by}) as sub"
            print(f"ULTRA_VERBOSE_SQL_COUNT: {count_query} | Params: {params}")
            cur.execute(count_query, params)
            count_res = cur.fetchone()
            total_records = count_res['total_count'] if count_res else 0
            print(f"ULTRA_VERBOSE_SQL_COUNT: Found {total_records} records")
                
            query = f"""
                SELECT 
                    client_name, 
                    server_name, 
                    db_type, 
                    log_type, 
                    MAX(log_time_ist) as log_time_ist, 
                    log_message, 
                    status,
                    owner,
                    client_visibility,
                    ticket_status,
                    next_action,
                    severity,
                    log_hash,
                    MAX(id) as id,
                    MAX(status_updated_at) as status_updated_at,
                    SUM(occurrence_count) as occurrence_count 
                {base_query} 
                {group_by}
                ORDER BY log_time_ist DESC 
                LIMIT %s OFFSET %s
            """
            
            print(f"DEBUG_LOGS: main_query: {query} | params: {params + [limit, offset]}")
            cur.execute(query, params + [limit, offset])
            logs = cur.fetchall()
            
            results = []
            for r in logs:
                row_dict = dict(r)
                for key, value in row_dict.items():
                    if isinstance(value, (datetime, ZoneInfo)):
                        row_dict[key] = value.isoformat()
                results.append(row_dict)
                
            response_data = {"logs": results, "total": total_records}
            cache_manager.set(cache_key, response_data, ttl_seconds=5)
            return response_data
    except Exception as e:
            import traceback
            print(f"ERROR in get_logs: {e}")
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=str(e))
    finally:
            if 'cur' in locals(): cur.close()
@app.get("/api/owner-counts")
def get_owner_counts(
    start_time: Optional[str] = None,
    end_time: Optional[str] = None,
    client_name: Optional[str] = None,
    db_type: Optional[str] = None,
    server_name: Optional[str] = None,
    log_type: List[str] = Query(None),
    severity: List[str] = Query(None),
    log_status: List[str] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    # Thread-safe TTL cache check
    log_type_key = ",".join(sorted(log_type or []))
    severity_key = ",".join(sorted(severity or []))
    log_status_key = ",".join(sorted(log_status or []))
    user_email = current_user.get("email", "")

    cache_key = f"owner-counts:{start_time}:{end_time}:{client_name}:{db_type}:{server_name}:{log_type_key}:{severity_key}:{log_status_key}:{user_email}"

    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            log_table = "db_monitoring_logs"
            if log_status and any(s in ['Resolved', 'Ignored', 'No action Required'] for s in log_status):
                log_table = "db_archived_logs"
            
            base_query = f"FROM {log_table} WHERE 1=1"
            params = []
            
            if not start_time and log_table == "db_monitoring_logs":
                start_time = (datetime.now(ZoneInfo("Asia/Kolkata")) - timedelta(hours=1)).strftime("%Y-%m-%d %H:%M:%S")

            if start_time:
                base_query += " AND log_time_ist >= %s"
                params.append(start_time)
            if end_time:
                base_query += " AND log_time_ist <= %s"
                params.append(end_time)
            if client_name:
                base_query += " AND client_name = %s"
                params.append(client_name)
            if db_type:
                base_query += " AND db_type ILIKE %s"
                params.append(db_type)
            if server_name:
                base_query += " AND server_name = %s"
                params.append(server_name)
            if log_type:
                base_query += " AND log_type IN %s"
                params.append(tuple(log_type) if isinstance(log_type, list) else (log_type,))
            if severity:
                base_query += " AND severity IN %s"
                params.append(tuple(severity) if isinstance(severity, list) else (severity,))
            if log_status:
                base_query += " AND LOWER(status) IN %s"
                params.append(tuple(s.lower() for s in log_status) if isinstance(log_status, list) else (log_status.lower(),))
            else:
                base_query += " AND (status IS NULL OR TRIM(status) = '' OR LOWER(status) NOT IN ('resolved', 'ignored', 'no action required'))"

            allowed_techs = get_user_allowed_technologies(current_user)
            if allowed_techs is not None:
                base_query += " AND TRIM(LOWER(db_type)) = ANY(%s)"
                params.append([t.lower().strip() for t in allowed_techs])

            # Group by owner to get per-person counts
            # We use a subquery to group similar logs first (like get_logs does) to be consistent with the UI counts
            group_cols = "client_name, server_name, db_type, log_type, log_message, status, owner, client_visibility, ticket_status, next_action, severity, log_hash"
            
            owner_query = f"""
                SELECT owner, COUNT(*) as count 
                FROM (SELECT owner {base_query} GROUP BY {group_cols}) as sub
                GROUP BY owner
            """
            
            cur.execute(owner_query, params)
            rows = cur.fetchall()
            
            owner_counts = {}
            total_assigned = 0
            total_unassigned = 0
            
            for row in rows:
                owner = row['owner']
                count = row['count']
                if owner and owner.strip() and owner.lower() != 'none':
                    owner_counts[owner] = count
                    total_assigned += count
                else:
                    total_unassigned += count
            
            result = {
                "owner_counts": owner_counts,
                "total_assigned": total_assigned,
                "total_unassigned": total_unassigned,
                "total_all": total_assigned + total_unassigned
            }
            cache_manager.set(cache_key, result, ttl_seconds=5)
            return result
    except Exception as e:
        print(f"ERROR in get_owner_counts: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if 'cur' in locals(): cur.close()

class LogSummaryRequest(BaseModel):
    logs: List[str]
    filters: dict = {}

@app.post("/api/summarize")
async def summarize_logs(req: LogSummaryRequest, current_user: dict = Depends(get_current_user)):
    if not openai_client:
        raise HTTPException(status_code=500, detail="OpenAI client is not configured. Is OPENAI_API_KEY in .env?")
    
    if not req.logs:
        raise HTTPException(status_code=400, detail="No logs provided for summary.")
    
    try:
        start_time = datetime.now()
        print(f"AI Summary Start: Requested by {current_user.get('username', 'unknown')} ({len(req.logs)} logs)")
        
        capped_logs = req.logs[:100]
        log_text = "\n".join(capped_logs)
        
        filters = req.filters or {}
        start_time_val = filters.get('start', 'N/A')
        end_time_val = filters.get('end', 'N/A')
        
        if client_name := filters.get('client'):
            system_prompt = (
                "You are an elite, highly experienced Database Administrator and Systems Engineer AI. "
                "Your task is to analyze database capacity growth patterns, tablespace metrics, and "
                "server resource performance telemetry (CPU, Memory, Disk, IO). "
                f"Compile a stunning, professional, and detailed Capacity & Performance Diagnostics Report for client {client_name}."
            )
            
            user_prompt = f"""
Analyze the following telemetry and capacity details for client '{client_name}' over the period {start_time_val} to {end_time_val}.

You must output exactly a comprehensive point-by-point diagnostic report containing exactly between 15 to 20 lines in total (excluding headers/dividers). 
Make the report rich, deeply professional, and highly detailed. Cover all the following points:
1. Current overall database cluster sizes and daily database capacity growth trends (mention positive growth and no growth databases explicitly).
2. Table growth metrics, highlighting the heaviest tables and tablespace allocations.
3. CPU, Memory, Disk, and IO performance telemetry logs (if live/available in the data below. If these resource logs are not available or not open, explicitly mention that server hardware consoles are currently not online or data is pending, but do not make up hardware values).

You MUST format your output exactly as a single list of point-by-point bullets, with a header:
### 📈 Expert Growth & Resource Diagnostics Report

Ensure it is point-by-point, highly technical, actionable, and contains exactly 15 to 20 detailed bulleted points (lines). Do not include any introductory or concluding conversational filler.

--- TELEMETRY DATA ---
{log_text}
--- END TELEMETRY DATA ---
"""
        else:
            system_prompt = (
                "You are an expert Database Administrator AI. Your task is to analyze postgres, MSSQL, Mysql, "
                "Mongodb, RDS, and Windows Event logs securely. Provide a clear, professional diagnostic report "
                "in strict Markdown format. You must strictly decline processing any requests, logs, or queries "
                "that are off-topic or unrelated to database, software, or data engineering."
            )
            
            user_prompt = f"""
Analyze the following batch of logs concisely for the period: {start_time_val} to {end_time_val}.
Do not exceed 200 words and do not include any introductory or concluding conversational filler.

The log entries are prefixed with their occurrence counts (e.g., [Count: 5]). 
Analyze the specific time range, occurrence frequency, and the log messages themselves effectively to provide a detailed summary of the issue patterns.

If the logs contain no discernible actionable data, output exactly: "I was unable to analyze the Logs."
Don't hallucinate.
Format your response exactly with these Markdown headers:
### 📊 1. Overall Summary
(Provide a high-level, easy-to-read overview of what occurred)

### 🔍 2. Root Cause Analysis
(Explain the underlying issue clearly, using `code blocks` for any specific error codes or paths)

### ⚠️ 3. Severity Assessment
(Explicitly count and list Critical, Medium, and Low issues using clear bullet points)

### 💡 4. Action Recommendations
(Provide clear, actionable, numbered steps to resolve the issue or optimize the system)

--- LOG DATA ---
{log_text}
--- END LOG DATA ---
"""
        
        model_name = os.getenv("OPENAI_MODEL_NAME", "gpt-4o-mini")
        
        def get_logs_hash(logs: List[str]) -> str:
            import re
            normalized = []
            for log in logs:
                clean = log.strip()
                clean = re.sub(r'^\[Count:\s*\d+\]\s*', '', clean)
                normalized.append(clean.strip())
            text = "".join(sorted(normalized))
            return hashlib.sha256(text.encode()).hexdigest()

        logs_hash = get_logs_hash(req.logs)
        
        try:
            with get_db_connection() as conn:
                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                cur.execute("""
                    SELECT summary_text FROM ai_summary_history 
                    WHERE logs_hash = %s AND created_at > NOW() - INTERVAL '7 days'
                    ORDER BY created_at DESC LIMIT 1
                """, (logs_hash,))
                cached = cur.fetchone()
                cur.close()
            
            if cached:
                print(f"AI CACHE HIT: Returning existing summary for logs_hash {logs_hash[:10]}...")
                return {"summary": cached['summary_text'], "cached": True}
        except Exception as cache_err:
            print(f"Cache check failed (ignoring): {cache_err}")

        try:
            response = await openai_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_content if 'system_content' in locals() else system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_completion_tokens=2000
            )
        except Exception as api_err:
            if "model_not_found" in str(api_err).lower() or "not found" in str(api_err).lower():
                print(f"⚠️ Preferred model '{model_name}' not found. Falling back to gpt-4o-mini...")
                response = await openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_content if 'system_content' in locals() else system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    max_completion_tokens=2000
                )
            else:
                raise api_err

        summary_text = response.choices[0].message.content
        
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        print(f"AI Summary Complete: Generated in {duration:.2f}s")
        
        try:
            from psycopg2.extras import Json
            import json
            filters_json = req.filters or {}
            filters_json["username"] = current_user["username"]
            
            with get_db_connection() as conn:
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO ai_summary_history (summary_text, filters_json, username, logs_hash)
                    VALUES (%s, %s, %s, %s)
                """, (summary_text, json.dumps(filters_json), current_user["username"], logs_hash))
                conn.commit()
                cur.close()
        except Exception as db_err:
            print(f"Error saving to history: {db_err}")

        return {"summary": summary_text, "cached": False}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI API Error: {str(e)}")

@app.get("/api/history")
def get_summary_history(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = "SELECT id, created_at, username, filters_json FROM ai_summary_history WHERE 1=1"
            params = []
            
            is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
            if not is_admin:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                if allowed_clients:
                    query += " AND (filters_json->>'client_name' = ANY(%s) OR username = %s)"
                    params.append(allowed_clients)
                    params.append(current_user["username"])
                elif allowed_techs is not None:
                    query += " AND (TRIM(LOWER(filters_json->>'technology')) = ANY(%s) OR username = %s)"
                    tech_params = [t.lower().strip() for t in allowed_techs]
                    params.append(tech_params)
                    params.append(current_user["username"])
                else:
                    query += " AND username = %s"
                    params.append(current_user["username"])
            
            query += " ORDER BY created_at DESC LIMIT 50"
            cur.execute(query, tuple(params))
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                
                if row.get('filters_json') is None:
                    row['filters_json'] = {}
                elif isinstance(row['filters_json'], str):
                    try:
                        row['filters_json'] = json.loads(row['filters_json'])
                    except:
                        pass
                results.append(row)
                
            cur.close()
        return {"history": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
@app.get("/api/history/{history_id}")
def get_history_detail(history_id: int, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SELECT summary_text, created_at, filters_json, username FROM ai_summary_history WHERE id = %s", (history_id,))
            result = cur.fetchone()
            cur.close()
        
        if not result:
            raise HTTPException(status_code=404, detail="History entry not found")
        
        row = dict(result)
        
        # Enforce strict multi-tenant boundary checks for standard users
        is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
        if not is_admin:
            filters = row.get("filters_json", {})
            if isinstance(filters, str):
                try:
                    filters = json.loads(filters)
                except:
                    filters = {}
            row_username = row.get("username", "")
            row_client = (filters or {}).get("client_name")
            row_tech = ((filters or {}).get("technology") or "").lower().strip()
            
            is_owner = row_username.lower() == current_user["username"].lower()
            has_permission = is_owner
            
            if not has_permission:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                if allowed_clients:
                    if row_client in allowed_clients:
                        has_permission = True
                elif allowed_techs is not None:
                    if row_tech in [t.lower().strip() for t in allowed_techs]:
                        has_permission = True
            
            if not has_permission:
                raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this history entry.")
        
        if row.get('filters_json') is None:
            row['filters_json'] = {}
        elif isinstance(row['filters_json'], str):
            try:
                row['filters_json'] = json.loads(row['filters_json'])
            except:
                pass
        
        if isinstance(row.get('created_at'), datetime):
             row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")

        return row
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class ExportRequest(BaseModel):
    title: str
    content: str
    client_name: Optional[str] = "N/A"
    server_name: Optional[str] = "N/A"
    db_type: Optional[str] = "N/A"
    severity: Optional[str] = "N/A"
    generated_on: Optional[str] = None

@app.post("/api/export/docx")
async def export_docx(request: ExportRequest, current_user: dict = Depends(get_current_user)):
    try:
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        
        title = doc.add_heading(request.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Client Name", request.client_name),
            ("Server Instance", request.server_name),
            ("Log Technology", request.db_type),
            ("Severity", request.severity),
            ("Generated On", request.generated_on or datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")),
            ("Generated By", current_user.get("full_name") or current_user.get("username", "System"))
        ]
        
        for label, value in metadata:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            row_cells[0].paragraphs[0].runs[0].bold = True
            
        doc.add_paragraph()
        
        import re
        
        lines = request.content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line).strip(), style='List Number')
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"GeoPITS AI Log Analyzer © {datetime.now().year} | Prepared By SANJAY G | Confidential Technical Report"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_client = re.sub(r'[^a-zA-Z0-9]', '_', request.client_name or "Log")
        filename = f"Report_{safe_client}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        print(f"Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def get_admin_user(current_user: dict = Depends(get_current_user)):
    if not current_user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin privileges required")
    return current_user

@app.get("/api/admin/users")
def get_admin_users(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SELECT id, username, email, full_name, profile_pic, role, last_active_at FROM users ORDER BY last_active_at DESC NULLS LAST")
            users = cur.fetchall()
            cur.close()
        
        results = []
        for r in users:
            row = dict(r)
            if row.get('last_active_at') and isinstance(row['last_active_at'], datetime):
                row['last_active_at'] = row['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"users": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class RoleUpdateRequest(BaseModel):
    email: Optional[str] = None
    username: Optional[str] = None
    role: str

@app.patch("/api/admin/users/role")
def update_user_role(req: RoleUpdateRequest, admin_user: dict = Depends(get_admin_user)):
    if req.role not in ['admin', 'user']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    identifier = req.email or req.username
    if not identifier:
        raise HTTPException(status_code=400, detail="Email or username required")

    if (req.username == admin_user.get("username") or req.email == admin_user.get("email")) and req.role == 'user':
        raise HTTPException(status_code=400, detail="You cannot demote yourself from admin status")

    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            target_email = None
            if req.email:
                cur.execute("SELECT email FROM users WHERE LOWER(email) = LOWER(%s)", (req.email,))
            else:
                cur.execute("SELECT email FROM users WHERE username = %s", (req.username,))
            
            email_row = cur.fetchone()
            if email_row:
                target_email = (email_row[0] or "").lower()
            
            if req.email:
                cur.execute("UPDATE users SET role = %s WHERE LOWER(email) = LOWER(%s)", (req.role, req.email))
            else:
                cur.execute("UPDATE users SET role = %s WHERE username = %s", (req.role, req.username))
            
            if cur.rowcount == 0:
                cur.close()
                raise HTTPException(status_code=404, detail="User not found in system")
            
            if target_email:
                if req.role == 'admin':
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (target_email,))
                else:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (target_email,))
                    cur.execute("UPDATE leads SET status = 'removed' WHERE LOWER(email) = %s AND technology IN ('Global', 'Global Admin')", (target_email,))
            
            conn.commit()
            cur.close()
        return {"status": "success", "message": f"User role updated to {req.role}"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.get("/api/admin/summaries")
def get_admin_summaries(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("""
                SELECT ah.id, ah.summary_text, ah.created_at, ah.username, u.full_name, u.profile_pic
                FROM ai_summary_history ah
                LEFT JOIN users u ON ah.username = u.username
                ORDER BY ah.created_at DESC LIMIT 100
            """)
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
                
            cur.close()
        return {"history": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
class LeadCreateRequest(BaseModel):
    email: str
    technology: Optional[str] = None
    technologies: Optional[List[str]] = None
    is_lead: Optional[bool] = False

@app.get("/api/admin/leads")
def get_admin_leads(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("""
                SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, u.role 
                FROM leads l
                LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                ORDER BY l.created_at DESC
            """)
            leads = cur.fetchall()
            cur.close()
        
        results = []
        for r in leads:
            row = dict(r)
            if isinstance(row['created_at'], datetime):
                row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"leads": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/leads")
def create_admin_lead(req: LeadCreateRequest, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            tech_list = []
            if req.technologies:
                tech_list.extend(req.technologies)
            if req.technology:
                tech_list.append(req.technology)
            
            tech_list = list(set(tech_list))
            
            if not tech_list:
                tech_list = ['Global']
            
            inserted_ids = []
            for tech in tech_list:
                cur.execute("""
                    INSERT INTO leads (email, technology, status, is_lead) 
                    VALUES (%s, %s, 'active', %s) 
                    ON CONFLICT (email, technology) DO UPDATE SET is_lead = EXCLUDED.is_lead, status = 'active'
                    RETURNING id
                """, (req.email, tech, req.is_lead))
                res = cur.fetchone()
                if res:
                    inserted_ids.append(res[0])
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (req.email.lower(),))
                    
            conn.commit()
            cur.close()
        return {"status": "success", "ids": inserted_ids}
    except Exception as e:
        print(f"Lead Creation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/api/admin/leads/{lead_id}/status")
def toggle_lead_status(lead_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT email, technology, status FROM leads WHERE id = %s", (lead_id,))
            row = cur.fetchone()
            if not row:
                cur.close()
                raise HTTPException(status_code=404, detail="Lead not found")
            
            email, tech, current_status = row
            new_status = 'removed' if current_status == 'active' else 'active'
            
            cur.execute("UPDATE leads SET status = %s WHERE id = %s", (new_status, lead_id))
            
            if tech in ['Global', 'Global Admin'] and email:
                cur.execute("""
                    INSERT INTO system_admins (email, status)
                    VALUES (%s, %s)
                    ON CONFLICT (email) DO UPDATE SET status = EXCLUDED.status
                """, (email.lower(), new_status))
                
            conn.commit()
            cur.close()
        return {"status": "updated", "new_status": new_status}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/admin/leads/{lead_id}")
def delete_admin_lead(lead_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (lead_id,))
            row = cur.fetchone()
            if row:
                email, tech = row
                cur.execute("DELETE FROM leads WHERE id = %s", (lead_id,))
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (email.lower(),))
            
            conn.commit()
            cur.close()
        return {"status": "deleted"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/lead/technologies")
def get_lead_technologies(current_user: dict = Depends(get_current_user)):
    """Returns technologies for which the current user is a lead."""
    try:
        email = current_user.get('email', '').lower()
        print(f"DEBUG_LEAD: Checking technologies for email: '{email}'")
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            techs = [row[0] for row in cur.fetchall()]
            print(f"DEBUG_LEAD: Found techs for '{email}': {techs}")
            cur.close()
        return {"technologies": techs}
    except Exception as e:
        print(f"DEBUG_LEAD_ERROR: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/lead/users")
def get_lead_users(current_user: dict = Depends(get_current_user)):
    """Returns users assigned to technologies for which the current user is a lead."""
    try:
        email = current_user.get('email', '').lower()
        print(f"DEBUG_LEAD_USERS: Fetching users for lead email: '{email}'")
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row['technology'] for row in cur.fetchall()]
            
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if is_global_admin:
                cur.execute("""
                    SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, 
                           u.full_name, u.profile_pic, u.last_active_at, u.username
                    FROM leads l
                    LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                    WHERE LOWER(l.email) != %s
                    ORDER BY u.last_active_at DESC NULLS LAST
                """, (email,))
            elif my_techs:
                cur.execute("""
                    SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, 
                           u.full_name, u.profile_pic, u.last_active_at, u.username
                    FROM leads l
                    LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                    WHERE l.technology IN %s AND LOWER(l.email) != %s
                    ORDER BY u.last_active_at DESC NULLS LAST
                """, (tuple(my_techs), email))
            else:
                cur.close()
                return {"users": []}
            
            users = cur.fetchall()
            results = []
            for r in users:
                row = dict(r)
                if row.get('created_at') and isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                if row.get('last_active_at') and isinstance(row['last_active_at'], datetime):
                    row['last_active_at'] = row['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
            
            cur.close()
        return {"users": results}
    except Exception as e:
        print(f"Error fetching lead users: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/lead/users")
def create_lead_user(req: LeadCreateRequest, current_user: dict = Depends(get_current_user)):
    """Allows a lead to assign a user to a technology they lead."""
    try:
        email = current_user.get("email")
        if not email:
             raise HTTPException(status_code=403, detail="Email required")
             
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            
            is_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            requested_techs = req.technologies if req.technologies else ([req.technology] if req.technology else [])
            if not is_admin:
                for tech in requested_techs:
                    if tech not in my_techs:
                        cur.close()
                        raise HTTPException(status_code=403, detail=f"You do not have lead privilege for {tech}")
        
        cur.execute("SELECT role, email FROM users WHERE email = %s", (req.email,))
        target_row = cur.fetchone()
        
        is_target_admin = False
        if target_row:
            target_role, target_email = target_row
            is_target_admin = (target_role == 'admin') or ((target_email or "").lower() in ADMIN_EMAILS)
        
        if is_target_admin:
             cur.close()
             conn.close()
             raise HTTPException(status_code=403, detail="you cannot assign admin as user")

        inserted_ids = []
        for tech in requested_techs:
            cur.execute("""
                INSERT INTO leads (email, technology, status, is_lead) 
                VALUES (%s, %s, 'active', FALSE) 
                ON CONFLICT (email, technology) DO UPDATE SET status = 'active'
                RETURNING id
            """, (req.email, tech))
            res = cur.fetchone()
            if res:
                inserted_id = res[0]
                inserted_ids.append(inserted_id)
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (req.email.lower(),))
                
        conn.commit()
        cur.close()
        conn.close()
        return {"status": "success", "ids": inserted_ids}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Lead User Creation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/api/lead/users/{assignment_id}/status")
def toggle_lead_user_status(assignment_id: int, current_user: dict = Depends(get_current_user)):
    """Allows a lead to toggle 'active'/'revoked' status for an assignment they manage."""
    try:
        email = current_user.get("email", "").lower()
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if not is_global_admin:
                cur.execute("""
                    SELECT l_target.technology 
                    FROM leads l_target
                    JOIN leads l_lead ON l_target.technology = l_lead.technology
                    WHERE l_target.id = %s AND LOWER(l_lead.email) = %s AND l_lead.is_lead = TRUE AND l_lead.status = 'active'
                """, (assignment_id, email))
                
                if not cur.fetchone():
                    cur.close()
                    conn.close()
                    raise HTTPException(status_code=403, detail="You do not have lead privileges for this technology or user.")
            
            cur.execute("""
                UPDATE leads 
                SET status = CASE WHEN status = 'active' THEN 'revoked' ELSE 'active' END 
                WHERE id = %s
                RETURNING status
            """, (assignment_id,))
            new_status = cur.fetchone()[0]
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (assignment_id,))
            sync_row = cur.fetchone()
            if sync_row:
                s_email, s_tech = sync_row
                if s_tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, %s)
                        ON CONFLICT (email) DO UPDATE SET status = EXCLUDED.status
                    """, (s_email.lower(), new_status))
            
            conn.commit()
            cur.close()
        return {"status": "success", "new_status": new_status}
    except HTTPException: raise
    except Exception as e:
        print(f"Error toggling lead user status: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/lead/users/{assignment_id}")
def delete_lead_user(assignment_id: int, current_user: dict = Depends(get_current_user)):
    """Allows a lead to delete a user assignment for a technology they lead."""
    try:
        email = current_user.get("email")
        if not email:
            raise HTTPException(status_code=403, detail="Email required")
            
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE id = %s", (assignment_id,))
            row = cur.fetchone()
            if not row:
                cur.close()
                raise HTTPException(status_code=404, detail="Assignment not found")
            tech = row[0]
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if not is_global_admin:
                if tech not in my_techs:
                    cur.close()
                    conn.close()
                    raise HTTPException(status_code=403, detail=f"You do not have lead privilege for {tech}")
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (assignment_id,))
            sync_row = cur.fetchone()
            if sync_row:
                s_email, s_tech = sync_row
                cur.execute("DELETE FROM leads WHERE id = %s", (assignment_id,))
                
                if s_tech in ['Global', 'Global Admin']:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (s_email.lower(),))
            
            conn.commit()
            cur.close()
        return {"status": "deleted"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting lead user: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/lead-activity")
def get_lead_activity(current_user: dict = Depends(get_current_user)):
    """Monitor all user assignments overseen by leads (Admin only)."""
    if not current_user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin access required")
        
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("""
                SELECT l.id, l.email as user_email, l.technology, l.status, l.created_at,
                       u.full_name as user_name, u.profile_pic, u.last_active_at, u.username,
                       (
                           SELECT STRING_AGG(email, ', ') 
                           FROM leads 
                           WHERE technology = l.technology AND is_lead = TRUE AND status = 'active'
                       ) as lead_emails
                FROM leads l
                LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                WHERE l.is_lead = FALSE
                ORDER BY l.created_at DESC
            """)
            
            activity = []
            for row in cur.fetchall():
                res = dict(row)
                if res.get('created_at'):
                    res['created_at'] = res['created_at'].strftime("%Y-%m-%d %H:%M:%S")
                if res.get('last_active_at') and str(res.get('last_active_at')) != 'NaT':
                    try:
                        if hasattr(res['last_active_at'], 'astimezone'):
                            res['last_active_at'] = res['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                        else:
                            res['last_active_at'] = res['last_active_at'].strftime("%Y-%m-%d %H:%M:%S")
                    except:
                        res['last_active_at'] = str(res['last_active_at'])
                else:
                    res['last_active_at'] = 'Never'
                activity.append(res)
                
            cur.close()
        return {"activity": activity}
    except Exception as e:
        print(f"Error fetching lead activity: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ==============================================================================
# AUTOMATED DAILY TELEMETRY BACKGROUND SCHEDULER
# ==============================================================================
import threading
import time
from new_features.backend.telemetry_parser import run_telemetry_sync
from new_features.backend.utilization_sync import sync_utilization_history

def get_setting(key: str, default: str = None) -> str:
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT value FROM system_settings WHERE key = %s", (key,))
            row = cur.fetchone()
            cur.close()
            return row[0] if row else default
    except Exception as e:
        print(f"Error fetching setting {key}: {e}")
        return default

def set_setting(key: str, value: str):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO system_settings (key, value)
                VALUES (%s, %s)
                ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value
            """, (key, str(value)))
            conn.commit()
            cur.close()
    except Exception as e:
        print(f"Error setting {key}: {e}")

# Shared status
telemetry_sync_in_progress = False

def run_manual_sync_in_background():
    global telemetry_sync_in_progress
    if telemetry_sync_in_progress:
        return
    telemetry_sync_in_progress = True
    set_setting("telemetry_scheduler_status", "Syncing...")
    try:
        now_str = datetime.now(ZoneInfo("Asia/Kolkata")).strftime('%Y-%m-%d %H:%M:%S IST')
        print(f"[SCHEDULER] ⏰ Manual trigger start at {now_str}")
        run_telemetry_sync()
        try:
            sync_utilization_history()
        except Exception as ex:
            print(f"[SCHEDULER] Server utilization sync error: {ex}")
        set_setting("telemetry_last_sync_time", now_str)
        set_setting("telemetry_last_sync_status", "Success")
    except Exception as e:
        print(f"[SCHEDULER] Manual sync failed: {e}")
        set_setting("telemetry_last_sync_status", f"Failed: {str(e)}")
    finally:
        telemetry_sync_in_progress = False
        set_setting("telemetry_scheduler_status", "Idle")

def telemetry_scheduler_daemon():
    print("[SCHEDULER] Automated daily database size telemetry scheduler daemon started.")
    set_setting("telemetry_scheduler_status", "Idle")
    while True:
        try:
            # Check scheduler trigger time settings
            trigger_hour = int(get_setting("telemetry_scheduler_hour", "14"))
            trigger_minute = int(get_setting("telemetry_scheduler_minute", "0"))
            
            now = datetime.now(ZoneInfo("Asia/Kolkata"))
            last_run_date_str = get_setting("telemetry_last_run_date", "")
            
            if now.hour == trigger_hour and now.minute == trigger_minute and last_run_date_str != str(now.date()):
                global telemetry_sync_in_progress
                if not telemetry_sync_in_progress:
                    telemetry_sync_in_progress = True
                    set_setting("telemetry_scheduler_status", "Syncing...")
                    try:
                        print(f"[SCHEDULER] ⏰ Triggering automated daily telemetry mail sync at {now.strftime('%Y-%m-%d %H:%M:%S IST')}...")
                        run_telemetry_sync()
                        try:
                            sync_utilization_history()
                        except Exception as ex:
                            print(f"[SCHEDULER] Server utilization sync error: {ex}")
                        
                        set_setting("telemetry_last_run_date", str(now.date()))
                        set_setting("telemetry_last_sync_time", now.strftime('%Y-%m-%d %H:%M:%S IST'))
                        set_setting("telemetry_last_sync_status", "Success")
                    except Exception as e:
                        print(f"[SCHEDULER] Ingestion Daemon run failed: {e}")
                        set_setting("telemetry_last_sync_status", f"Failed: {str(e)}")
                    finally:
                        telemetry_sync_in_progress = False
                        set_setting("telemetry_scheduler_status", "Idle")
        except Exception as e:
            print(f"[SCHEDULER] Daemon loop error: {e}")
        time.sleep(30)

# Start thread as a background daemon
t = threading.Thread(target=telemetry_scheduler_daemon, daemon=True)
t.start()

# API Endpoints for Scheduler Management
@app.get("/api/admin/scheduler/status")
def get_scheduler_status(admin_user: dict = Depends(get_admin_user)):
    try:
        hour = get_setting("telemetry_scheduler_hour", "14")
        minute = get_setting("telemetry_scheduler_minute", "0")
        status = get_setting("telemetry_scheduler_status", "Idle")
        last_sync = get_setting("telemetry_last_sync_time", "Never")
        last_status = get_setting("telemetry_last_sync_status", "N/A")
        
        return {
            "trigger_hour": int(hour),
            "trigger_minute": int(minute),
            "status": "Syncing..." if telemetry_sync_in_progress else status,
            "last_sync_time": last_sync,
            "last_sync_status": last_status,
            "sync_in_progress": telemetry_sync_in_progress
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class SchedulerSettingsRequest(BaseModel):
    trigger_hour: int
    trigger_minute: int

@app.post("/api/admin/scheduler/settings")
def update_scheduler_settings(req: SchedulerSettingsRequest, admin_user: dict = Depends(get_admin_user)):
    if not (0 <= req.trigger_hour <= 23) or not (0 <= req.trigger_minute <= 59):
        raise HTTPException(status_code=400, detail="Invalid hour (0-23) or minute (0-59)")
    try:
        set_setting("telemetry_scheduler_hour", str(req.trigger_hour))
        set_setting("telemetry_scheduler_minute", str(req.trigger_minute))
        return {"status": "success", "message": "Scheduler trigger settings updated successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/scheduler/trigger")
def trigger_scheduler_sync(admin_user: dict = Depends(get_admin_user)):
    global telemetry_sync_in_progress
    if telemetry_sync_in_progress:
        raise HTTPException(status_code=409, detail="A telemetry sync operation is already in progress.")
    try:
        threading.Thread(target=run_manual_sync_in_background, daemon=True).start()
        return {"status": "success", "message": "Telemetry sync initiated successfully in the background."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




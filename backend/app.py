import os
import sys
import threading
import time as _time
# Add backend directory to path for absolute imports resolution
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from fastapi import FastAPI, HTTPException, Query, Request, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware

class CachedStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        response = await super().get_response(path, scope)
        if "assets/" in path or path.startswith("assets/"):
            response.headers["Cache-Control"] = "public, max-age=31536000, immutable"
        else:
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response
from pydantic import BaseModel
import psycopg2
import psycopg2.extras
from datetime import datetime, timezone
from dotenv import load_dotenv
from openai import AsyncOpenAI
from zoneinfo import ZoneInfo
from typing import List, Optional
import hashlib
from docx import Document
from io import BytesIO
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse, RedirectResponse
from ipaddress import ip_address, ip_network
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi import Depends, status
import bcrypt
from jose import JWTError, jwt
from datetime import timedelta
from authlib.integrations.starlette_client import OAuth
from starlette.middleware.sessions import SessionMiddleware
from psycopg2.pool import ThreadedConnectionPool
from contextlib import contextmanager

load_dotenv()

if not os.getenv("OPENAI_API_KEY"):
    print("ERROR: OPENAI_API_KEY not found in environment or .env file!")
else:
    print("OK: OPENAI_API_KEY detected.")

from cache_utils import cache_manager

import json
from functools import lru_cache

FILTER_CACHE = {}
FILTER_CACHE_TTL = 300
_USER_CACHE_TTL = 300

from log_utils import audit_logger

app = FastAPI(title="GeoPITS Dashboard")

def check_client_alert_thresholds():
    import psycopg2
    import psycopg2.extras
    import json
    import os
    from datetime import datetime, timedelta
    from zoneinfo import ZoneInfo
    from routes import send_email_outlook

    print("[ALERT DAEMON] Starting client alert threshold checking sweep...")
    db_host = os.getenv("DB_HOST", "localhost")
    db_name = os.getenv("DB_NAME", "Incoming-error-data")
    db_user = os.getenv("DB_USER", "postgres")
    db_password = os.getenv("DB_PASSWORD", "y7UMhWmLcqSJzmhTGDyK")
    db_port = os.getenv("DB_PORT", "5432")

    try:
        conn = psycopg2.connect(host=db_host, database=db_name, user=db_user, password=db_password, port=db_port)
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # 1. Fetch all alert settings
        cur.execute("SELECT * FROM client_alert_settings;")
        settings = cur.fetchall()

        now = datetime.now()
        one_hour_ago = now - timedelta(hours=1)

        for setting in settings:
            client = setting["client_name"]
            db_type = setting["db_type"]
            last_sent = setting["last_summary_sent"]

            # If last summary sent was within the last hour, skip to avoid double alerting
            if last_sent and (now - last_sent) < timedelta(hours=1):
                continue

            # Check CPU, Memory, Disk, IO spikes
            # Search db_monitoring_logs for this client, db_type, in the last hour
            cur.execute("""
                SELECT log_type, log_message, log_time_ist, server_name
                FROM db_monitoring_logs
                WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                  AND LOWER(TRIM(db_type)) = LOWER(TRIM(%s))
                  AND log_time_ist >= %s;
            """, (client, db_type, one_hour_ago))
            logs = cur.fetchall()

            # Fallback to server_utilization_history for servers
            cur.execute("""
                SELECT DISTINCT server_name FROM admin_clients 
                WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                  AND LOWER(TRIM(%s)) = ANY(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ','));
            """, (client, db_type))
            servers = [r["server_name"] for r in cur.fetchall() if r["server_name"]]

            util_logs = []
            if servers:
                cur.execute("""
                    SELECT server_name, cpu_utilization, memory_utilization, disk_utilization, io_utilization, captured_at
                    FROM server_utilization_history
                    WHERE server_name = ANY(%s) AND captured_at >= %s;
                """, (servers, one_hour_ago))
                util_logs = cur.fetchall()

            # Analyze max values
            max_cpu = 0.0
            max_mem = 0.0
            max_disk = 0.0
            max_io = 0.0

            def parse_val(val_str):
                if not val_str:
                    return 0.0
                try:
                    return float(str(val_str).replace("%", "").strip())
                except:
                    return 0.0

            # Parse from db_monitoring_logs
            for l in logs:
                lt = l["log_type"].lower()
                val = parse_val(l["log_message"])
                if "cpu" in lt:
                    max_cpu = max(max_cpu, val)
                elif "memory" in lt or "mem" in lt:
                    max_mem = max(max_mem, val)
                elif "disk" in lt:
                    max_disk = max(max_disk, val)
                elif "io" in lt:
                    max_io = max(max_io, val)

            # Parse from server_utilization_history
            for u in util_logs:
                if u["cpu_utilization"]:
                    max_cpu = max(max_cpu, float(u["cpu_utilization"]))
                if u["memory_utilization"]:
                    max_mem = max(max_mem, float(u["memory_utilization"]))
                if u["disk_utilization"]:
                    max_disk = max(max_disk, float(u["disk_utilization"]))
                if u["io_utilization"]:
                    max_io = max(max_io, float(u["io_utilization"]))

            # Check threshold breach
            cpu_breach = max_cpu > float(setting["cpu_threshold"])
            mem_breach = max_mem > float(setting["memory_threshold"])
            disk_breach = max_disk > float(setting["disk_threshold"])
            io_breach = max_io > float(setting["io_threshold"])

            spiked_resources = []
            if cpu_breach: spiked_resources.append(f"CPU (Max: {max_cpu}%, Threshold: {setting['cpu_threshold']}%)")
            if mem_breach: spiked_resources.append(f"Memory (Max: {max_mem}%, Threshold: {setting['memory_threshold']}%)")
            if disk_breach: spiked_resources.append(f"Disk (Max: {max_disk}%, Threshold: {setting['disk_threshold']}%)")
            if io_breach: spiked_resources.append(f"IO (Max: {max_io}%, Threshold: {setting['io_threshold']}%)")

            # 2. Server Down Alert
            server_down_triggered = False
            offline_servers = []
            if setting["server_down_alert"]:
                cur.execute("""
                    SELECT DISTINCT server_name, service_name, status, captured_at
                    FROM db_uptime_history
                    WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s))
                      AND LOWER(TRIM(db_type)) = LOWER(TRIM(%s))
                      AND captured_at >= %s
                      AND status IN ('OFFLINE', 'STOPPED');
                """, (client, db_type, one_hour_ago))
                offline_rows = cur.fetchall()
                if offline_rows:
                    server_down_triggered = True
                    for row in offline_rows:
                        offline_servers.append(f"{row['server_name']} - {row['service_name']} is {row['status']} at {row['captured_at']}")

            # 3. Critical Error Logs Alert
            critical_errors_found = []
            if setting["critical_error_alert"]:
                cur.execute("""
                    SELECT log_message, log_time_ist, severity, server_name
                    FROM db_monitoring_logs
                    WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s))
                      AND LOWER(TRIM(db_type)) = LOWER(TRIM(%s))
                      AND log_time_ist >= %s
                      AND LOWER(log_type) = 'error_log'
                      AND LOWER(severity) IN ('critical', 'error', 'fatal', 'high');
                """, (client, db_type, one_hour_ago))
                crit_rows = cur.fetchall()
                for row in crit_rows:
                    critical_errors_found.append(f"[{row['severity']}] {row['server_name']}: {row['log_message']} ({row['log_time_ist']})")

            # Resolve client-specific contact details
            cur.execute("""
                SELECT client_email, phone_number FROM admin_clients 
                WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                  AND LOWER(TRIM(%s)) = ANY(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ','))
                LIMIT 1;
            """, (client, db_type))
            client_row = cur.fetchone()
            client_email_val = client_row["client_email"] if (client_row and client_row["client_email"]) else None
            phone_val = client_row["phone_number"] if (client_row and client_row["phone_number"]) else None

            default_cc = ""
            try:
                cur.execute("SELECT alert_email FROM technology_alerts_config WHERE LOWER(technology) = LOWER(%s) AND alert_email IS NOT NULL AND alert_email <> '';", (db_type,))
                row = cur.fetchone()
                if row and row.get("alert_email"):
                    default_cc = row["alert_email"]
            except Exception as db_err:
                print(f"Error querying technology_alerts_config: {db_err}")

            client_emails_str = setting["client_emails"] or client_email_val or ""
            if not client_emails_str:
                print("[ALERT ROUTING] there is no client mail")
            if not default_cc:
                print("[ALERT ROUTING] there is no technology alert mail")

            # Resolve Emails — send to both client email and technology alert email
            import re
            to_list = ["dccagent@geopits.com"]
            if client_emails_str:
                for email in re.split(r'[;,]', client_emails_str):
                    email = email.strip()
                    if email and email not in to_list:
                        to_list.append(email)
            if default_cc:
                for email in re.split(r'[;,]', default_cc):
                    email = email.strip()
                    if email and email not in to_list:
                        to_list.append(email)

            to_emails = ", ".join(to_list)

            cc_list = []
            if setting["cc_emails"]:
                for email in re.split(r'[;,]', setting["cc_emails"]):
                    email = email.strip()
                    if email and email not in to_list and email not in cc_list:
                        cc_list.append(email)
            cc_emails = ", ".join(cc_list) if cc_list else None

            # Process Spikes Summary
            if spiked_resources:
                # Find peak of utilization in the last 1 hour
                peak_time = now
                max_val = 0.0
                for u in util_logs:
                    try:
                        cu = float(u["cpu_utilization"] or 0)
                        mu = float(u["memory_utilization"] or 0)
                        du = float(u["disk_utilization"] or 0)
                        iu = float(u["io_utilization"] or 0)
                        total_u = cu + mu + du + iu
                        if total_u > max_val:
                            max_val = total_u
                            peak_time = u["captured_at"]
                    except:
                        pass

                if isinstance(peak_time, str):
                    try:
                        peak_time = datetime.fromisoformat(peak_time)
                    except:
                        peak_time = now
                if peak_time.tzinfo is None:
                    peak_time = now.astimezone(pytz.utc).replace(tzinfo=None)

                start_window = peak_time - timedelta(hours=1)
                end_window = peak_time + timedelta(hours=1)

                parsed_long_queries = []
                parsed_error_logs = []
                parsed_blocking = []

                def safe_load_json(val):
                    if not val:
                        return []
                    if isinstance(val, dict):
                        return [val]
                    if isinstance(val, list):
                        return val
                    try:
                        loaded = json.loads(val)
                        if isinstance(loaded, dict):
                            return [loaded]
                        if isinstance(loaded, list):
                            return loaded
                    except:
                        pass
                    return []

                # Query diagnosticdata_long_queries (telemetry table)
                try:
                    cur.execute("""
                        SELECT captured_at, raw_data, server_name 
                        FROM diagnosticdata_long_queries 
                        WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                          AND captured_at BETWEEN %s AND %s
                        ORDER BY captured_at DESC LIMIT 50;
                    """, (client, start_window, end_window))
                    for r in cur.fetchall():
                        for item in safe_load_json(r["raw_data"]):
                            parsed_long_queries.append({
                                "time": r["captured_at"],
                                "server_name": r["server_name"],
                                "duration_ms": item.get("duration") or item.get("duration_ms") or (item.get("duration_us", 0)/1000) or item.get("Duration") or 0,
                                "cpu_ms": item.get("cpu") or item.get("cpu_time") or item.get("CPU") or 0,
                                "reads": item.get("reads") or item.get("logical_reads") or item.get("Reads") or 0,
                                "writes": item.get("writes") or item.get("logical_writes") or item.get("Writes") or 0,
                                "sql": item.get("sql_text") or item.get("query_text") or item.get("statement_text") or item.get("SQLText") or item.get("sql") or "N/A"
                            })
                except Exception as e:
                    print(f"Error querying diagnosticdata_long_queries: {e}")

                # Query diagnosticdata_error_logs
                try:
                    cur.execute("""
                        SELECT captured_at, raw_data, server_name 
                        FROM diagnosticdata_error_logs 
                        WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                          AND captured_at BETWEEN %s AND %s
                        ORDER BY captured_at DESC LIMIT 50;
                    """, (client, start_window, end_window))
                    for r in cur.fetchall():
                        for item in safe_load_json(r["raw_data"]):
                            parsed_error_logs.append({
                                "time": r["captured_at"],
                                "server_name": r["server_name"],
                                "process": item.get("process_info") or item.get("ProcessInfo") or item.get("process") or "System",
                                "message": item.get("text") or item.get("message") or item.get("LogText") or item.get("LogMessage") or item.get("raw_text") or "N/A"
                            })
                except Exception as e:
                    print(f"Error querying diagnosticdata_error_logs: {e}")

                # Query diagnosticdata_blocking
                try:
                    cur.execute("""
                        SELECT captured_at, raw_data, server_name 
                        FROM diagnosticdata_blocking 
                        WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) 
                          AND captured_at BETWEEN %s AND %s
                        ORDER BY captured_at DESC LIMIT 50;
                    """, (client, start_window, end_window))
                    for r in cur.fetchall():
                        for item in safe_load_json(r["raw_data"]):
                            parsed_blocking.append({
                                "time": r["captured_at"],
                                "server_name": r["server_name"],
                                "blocking_spid": item.get("blocking_session_id") or item.get("blocking_spid") or 0,
                                "blocked_spid": item.get("blocked_session_id") or item.get("blocked_spid") or 0,
                                "wait_time_ms": item.get("wait_time") or item.get("wait_time_ms") or item.get("duration") or 0,
                                "sql": item.get("sql_text") or item.get("blocking_sql_text") or item.get("sql") or "N/A"
                            })
                except Exception as e:
                    print(f"Error querying diagnosticdata_blocking: {e}")

                # Add legacy tickets and comments slow queries
                slow_queries = []
                cur.execute("""
                    SELECT ticket_name, description, created_at FROM tickets
                    WHERE LOWER(TRIM(company)) = LOWER(TRIM(%s))
                      AND LOWER(TRIM(business_unit)) = LOWER(TRIM(%s))
                      AND created_at >= %s
                      AND (ticket_name ILIKE '%%Long Running%%' or ticket_name ILIKE '%%Slow Query%%' or ticket_name ILIKE '%%Transaction%%' or description ILIKE '%%Executing SQL%%');
                """, (client, db_type, one_hour_ago))
                ticket_slows = cur.fetchall()
                for ts in ticket_slows:
                    slow_queries.append({
                        "source": "Ticket",
                        "title": ts["ticket_name"],
                        "details": ts["description"],
                        "time": ts["created_at"]
                    })

                cur.execute("""
                    SELECT tc.content, tc.created_at, t.ticket_name FROM ticket_comments tc
                    JOIN tickets t ON tc.ticket_id = t.id
                    WHERE LOWER(TRIM(t.company)) = LOWER(TRIM(%s))
                      AND LOWER(TRIM(t.business_unit)) = LOWER(TRIM(%s))
                      AND tc.created_at >= %s
                      AND tc.comment_type = 'log'
                      AND tc.content LIKE 'MSSQL_LOG_DATA:%%';
                """, (client, db_type, one_hour_ago))
                log_comments = cur.fetchall()
                for lc in log_comments:
                    try:
                        raw_data = lc["content"].replace("MSSQL_LOG_DATA:", "", 1)
                        parsed = json.loads(raw_data)
                        slow_queries.append({
                            "source": "Log Comment",
                            "title": lc["ticket_name"],
                            "details": f"SPID: {parsed.get('spid')}\nSQL: {parsed.get('executing_sql') or parsed.get('sql_text')}",
                            "time": lc["created_at"]
                        })
                    except:
                        pass

                # Perform Root Cause Analysis (RCA) Heuristics
                rca_findings = []
                if cpu_breach:
                    high_cpu_query = None
                    max_c = 0.0
                    for q in parsed_long_queries:
                        try:
                            c_val = float(q.get("cpu_ms") or 0)
                        except:
                            c_val = 0.0
                        if c_val > max_c:
                            max_c = c_val
                            high_cpu_query = q
                    if high_cpu_query:
                        sql_trunc = str(high_cpu_query['sql'])[:300] + "..." if len(str(high_cpu_query['sql'])) > 300 else str(high_cpu_query['sql'])
                        reason = f"High CPU utilization of {max_c} ms detected. This query is likely causing the resource starvation due to high compilation costs, lack of index, or heavy CPU-bound operations (e.g. hash match, sorting)."
                        rca_findings.append(f"🔴 <strong>CPU Spike Suspect:</strong> Query on server '{high_cpu_query['server_name']}' consumed {max_c} ms CPU.<br/><strong>Reason:</strong> {reason}<br/><strong>Query:</strong> <code>{sql_trunc}</code>")
                    else:
                        rca_findings.append("⚠️ <strong>CPU Spike:</strong> CPU threshold was breached, but no specific long-running queries with high CPU time were found in this 2-hour window.")

                if mem_breach:
                    high_read_query = None
                    max_r = 0
                    for q in parsed_long_queries:
                        try:
                            r_val = int(q.get("reads") or 0)
                        except:
                            r_val = 0
                        if r_val > max_r:
                            max_r = r_val
                            high_read_query = q
                    if high_read_query:
                        sql_trunc = str(high_read_query['sql'])[:300] + "..." if len(str(high_read_query['sql'])) > 300 else str(high_read_query['sql'])
                        reason = f"Heavy memory utilization detected. Query generated {max_r} page reads, forcing buffer pool eviction and memory grant pressure."
                        rca_findings.append(f"🔴 <strong>Memory Spike Suspect:</strong> Query on server '{high_read_query['server_name']}' performed {max_r} logical reads.<br/><strong>Reason:</strong> {reason}<br/><strong>Query:</strong> <code>{sql_trunc}</code>")
                    else:
                        mem_err = None
                        for err in parsed_error_logs:
                            msg = str(err["message"]).lower()
                            if "memory" in msg or "buffer pool" in msg or "allocate" in msg or "out of memory" in msg:
                                mem_err = err
                                break
                        if mem_err:
                            rca_findings.append(f"🔴 <strong>Memory Spike Suspect:</strong> Memory pressure warning/error: '{mem_err['message']}' logged on '{mem_err['server_name']}' at {mem_err['time']}")
                        else:
                            rca_findings.append("⚠️ <strong>Memory Spike:</strong> Memory threshold breached, but no heavy read queries or memory allocation errors were isolated in this 2-hour window.")

                if io_breach or disk_breach:
                    high_io_query = None
                    max_io_score = 0
                    for q in parsed_long_queries:
                        try:
                            score = int(q.get("writes") or 0) + int(q.get("reads") or 0)
                        except:
                            score = 0
                        if score > max_io_score:
                            max_io_score = score
                            high_io_query = q
                    if high_io_query:
                        sql_trunc = str(high_io_query['sql'])[:300] + "..." if len(str(high_io_query['sql'])) > 300 else str(high_io_query['sql'])
                        reason = f"Disk I/O / space pressure detected. Query generated heavy page transfer score ({max_io_score}) leading to storage queue congestion."
                        rca_findings.append(f"🔴 <strong>IO/Disk Spike Suspect:</strong> Query on server '{high_io_query['server_name']}' generated heavy read/write score ({max_io_score}).<br/><strong>Reason:</strong> {reason}<br/><strong>Query:</strong> <code>{sql_trunc}</code>")
                    else:
                        latency_err = None
                        for err in parsed_error_logs:
                            msg = str(err["message"]).lower()
                            if "taking longer than" in msg or "io requests" in msg or "latency" in msg:
                                latency_err = err
                                break
                        if latency_err:
                            rca_findings.append(f"🔴 <strong>IO/Disk Spike Suspect:</strong> Storage I/O alert logged: '{latency_err['message']}'")
                        else:
                            rca_findings.append("⚠️ <strong>IO/Disk Spike:</strong> Threshold breached, but no heavy read/write query or disk latency error was isolated in this 2-hour window.")

                # Construct detailed tables HTML
                spikes_html = "".join([f"<li><strong>{r}</strong></li>" for r in spiked_resources])

                rca_html = ""
                if rca_findings:
                    rca_html = "<h3>Automated Root Cause Analysis (RCA):</h3>"
                    rca_html += "<div style='background-color:#fffbeb; border:1px solid #fef3c7; border-radius:8px; padding:16px; margin-bottom:16px;'>"
                    rca_html += "<ul style='margin:0; padding-left:20px;'>" + "".join([f"<li style='margin-bottom:10px;'>{r}</li>" for r in rca_findings]) + "</ul>"
                    rca_html += "</div>"

                long_queries_html = ""
                if parsed_long_queries:
                    long_queries_html += "<h3>Slow / Long Running Queries (within +-1 hr of peak):</h3>"
                    long_queries_html += "<table border='1' cellpadding='8' style='border-collapse:collapse; width:100%; border-color:#e2e8f0; margin-bottom:16px; font-size:12px;'>"
                    long_queries_html += "<tr style='background-color:#f8fafc; text-align:left;'><th>Time</th><th>Server</th><th>Duration (ms)</th><th>CPU (ms)</th><th>Reads/Writes</th><th>Query SQL</th></tr>"
                    for sq in parsed_long_queries[:5]:
                        sql_clean = str(sq["sql"]).replace("\n", "<br/>")
                        long_queries_html += f"<tr><td>{sq['time']}</td><td>{sq['server_name']}</td><td>{sq['duration_ms']}</td><td>{sq['cpu_ms']}</td><td>{sq['reads']} / {sq['writes']}</td><td><code style='font-family:monospace; color:#0f172a;'>{sql_clean}</code></td></tr>"
                    long_queries_html += "</table>"

                err_logs_html = ""
                if parsed_error_logs:
                    err_logs_html += "<h3>Critical Database Error Logs (within +-1 hr of peak):</h3>"
                    err_logs_html += "<table border='1' cellpadding='8' style='border-collapse:collapse; width:100%; border-color:#e2e8f0; margin-bottom:16px; font-size:12px;'>"
                    err_logs_html += "<tr style='background-color:#f8fafc; text-align:left;'><th>Time</th><th>Server</th><th>Process</th><th>Log Message</th></tr>"
                    for err in parsed_error_logs[:5]:
                        msg_clean = str(err["message"]).replace("\n", "<br/>")
                        err_logs_html += f"<tr><td>{err['time']}</td><td>{err['server_name']}</td><td>{err['process']}</td><td>{msg_clean}</td></tr>"
                    err_logs_html += "</table>"

                blocking_html = ""
                if parsed_blocking:
                    blocking_html += "<h3>Active Blocking / Locking Sessions (within +-1 hr of peak):</h3>"
                    blocking_html += "<table border='1' cellpadding='8' style='border-collapse:collapse; width:100%; border-color:#e2e8f0; margin-bottom:16px; font-size:12px;'>"
                    blocking_html += "<tr style='background-color:#f8fafc; text-align:left;'><th>Time</th><th>Server</th><th>Blocking SPID</th><th>Blocked SPID</th><th>Wait Time (ms)</th><th>Query</th></tr>"
                    for bl in parsed_blocking[:5]:
                        sql_clean = str(bl["sql"]).replace("\n", "<br/>")
                        blocking_html += f"<tr><td>{bl['time']}</td><td>{bl['server_name']}</td><td>{bl['blocking_spid']}</td><td>{bl['blocked_spid']}</td><td>{bl['wait_time_ms']}</td><td><code style='font-family:monospace; color:#0f172a;'>{sql_clean}</code></td></tr>"
                    blocking_html += "</table>"

                legacy_slows_html = ""
                if slow_queries:
                    legacy_slows_html += "<h3>Slow Queries / Long Running Queries from Legacy Logs:</h3>"
                    legacy_slows_html += "<table border='1' cellpadding='8' style='border-collapse:collapse; width:100%; border-color:#e2e8f0; margin-bottom:16px; font-size:12px;'>"
                    legacy_slows_html += "<tr style='background-color:#f8fafc; text-align:left;'><th>Time</th><th>Source</th><th>Incident Info</th><th>Details</th></tr>"
                    for sq in slow_queries[:5]:
                        details_clean = str(sq["details"]).replace("\n", "<br/>")
                        legacy_slows_html += f"<tr><td>{sq['time']}</td><td>{sq['source']}</td><td><strong>{sq['title']}</strong></td><td><code style='font-family:monospace;'>{details_clean}</code></td></tr>"
                    legacy_slows_html += "</table>"

                ticket_desc = f"Resource Spike Summary Alert:\n\nSpiked Resources:\n" + "\n".join(spiked_resources)
                if rca_findings:
                    ticket_desc += "\n\nAutomated Root Cause Analysis:\n" + "\n".join([f"- {r.replace('<strong>', '').replace('</strong>', '').replace('<br/>', ' ').replace('<code>', '').replace('</code>', '')}" for r in rca_findings])
                ticket_desc += f"\n\nDetails:\n- Slow Queries in window (+-1hr): {len(parsed_long_queries)}\n- Error Logs in window (+-1hr): {len(parsed_error_logs)}\n- Blocking Sessions in window (+-1hr): {len(parsed_blocking)}\n- Legacy Slow Queries: {len(slow_queries)}\n\nAssigned Agent: SYSTEM"
                
                ticket_name = f"{client} - {db_type} Resource Spike Alert Summary"
                cur.execute("""
                    INSERT INTO tickets (business_unit, company, contact, ticket_name, category, status, priority, agent, description, created_by, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'System', NOW())
                    RETURNING id;
                """, (db_type, client, to_emails, ticket_name, 'Alert', 'OPEN', 'High', 'SYSTEM', ticket_desc))
                new_t_id = cur.fetchone()[0]

                subject = f"[Ticket #{new_t_id}] [Spike Notification] {client} ({db_type}) - Resource Threshold Exceeded"
                email_body = f"""
                <html>
                <body style="font-family:-apple-system, BlinkMacSystemFont, sans-serif; color:#1e293b; line-height:1.6; max-width:900px; margin:0 auto; padding:20px;">
                    <div style="background: linear-gradient(135deg, #ef4444 0%, #f97316 100%); padding:24px; border-radius:12px; color:white;">
                        <h2 style="margin:0; font-size:24px;">⚠️ Resource Spike Alert Summary</h2>
                        <p style="margin:8px 0 0 0; font-size:14px; opacity:0.9;">Ticket #{new_t_id} | Hourly Diagnostic Telemetry Report for {client} - {db_type}</p>
                    </div>
                    <div style="margin-top:24px; background:#ffffff; border:1px solid #e2e8f0; border-radius:12px; padding:24px;">
                        <p>Hello,</p>
                        <p>This is an automated diagnostic summary indicating that one or more system resource thresholds have been exceeded in the last hour:</p>
                        <ul>{spikes_html}</ul>
                        
                        <p><strong>Ticket ID:</strong> #{new_t_id}<br/><strong>Assigned Agent:</strong> SYSTEM</p>
                        
                        {rca_html}
                        {long_queries_html}
                        {err_logs_html}
                        {blocking_html}
                        {legacy_slows_html}
                        
                        <p style="margin-top:24px; font-size:12px; color:#64748b; border-top:1px solid #e2e8f0; padding-top:12px;">
                            This alert was triggered automatically by the GeoMon monitoring service. Please log in to the GeoMon Incident Center for real-time analytics. Always from mailbox: dccagent@geopits.com.
                        </p>
                    </div>
                </body>
                </html>
                """

                if to_emails:
                    send_email_outlook(to_emails, cc_emails, subject, email_body, sender_email="dccagent@geopits.com")
                else:
                    print(f"[ALERT DAEMON] No recipient configured for {client} ({db_type}) — spike alert email skipped.")

                cur.execute("""
                    INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
                    VALUES (%s, 'System', 'log', %s, '');
                """, (new_t_id, f"Resource Spike Alert: " + ", ".join(spiked_resources)))

                # Log email routing outcome
                if to_emails:
                    email_log_msg = f"Alert email sent to: {to_emails}"
                    if cc_emails:
                        email_log_msg += f" (CC: {cc_emails})"
                else:
                    email_log_msg = f"Email alert skipped: No recipient email configured for client {client} or database technology {db_type}."
                cur.execute("""
                    INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
                    VALUES (%s, 'System', 'log', %s, '');
                """, (new_t_id, email_log_msg))

                cur.execute("UPDATE client_alert_settings SET last_summary_sent = NOW() WHERE id = %s;", (setting["id"],))
                conn.commit()
                print(f"[ALERT DAEMON] Created ticket #{new_t_id} & sent summary email for {client} ({db_type}).")

            # Handle Server Down Alert
            if server_down_triggered and offline_servers:
                server_desc = "Offline Servers detected:\n" + "\n".join(offline_servers)
                ticket_name = f"[Server Down Alert] {client} ({db_type}) - Servers/Services Offline"
                
                cur.execute("""
                    SELECT id FROM tickets 
                    WHERE company = %s AND business_unit = %s AND ticket_name = %s AND status != 'CLOSED' AND status != 'RESOLVED';
                """, (client, db_type, ticket_name))
                existing_down_t = cur.fetchone()
                
                if not existing_down_t:
                    cur.execute("""
                        INSERT INTO tickets (business_unit, company, contact, ticket_name, category, status, priority, agent, description, created_by, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'System', NOW())
                        RETURNING id;
                    """, (db_type, client, to_emails, ticket_name, 'Alert', 'OPEN', 'Critical', 'SYSTEM', server_desc + "\nAssigned Agent: SYSTEM"))
                    new_down_id = cur.fetchone()[0]

                    email_down_body = f"""
                    <html>
                    <body style="font-family:-apple-system, BlinkMacSystemFont, sans-serif; color:#1e293b; line-height:1.6; max-width:800px; margin:0 auto; padding:20px;">
                        <div style="background:#ef4444; padding:24px; border-radius:12px; color:white;">
                            <h2 style="margin:0; font-size:24px;">🔴 Critical: Server Down Alert</h2>
                            <p style="margin:8px 0 0 0; font-size:14px; opacity:0.9;">Ticket #{new_down_id} | Service Status Incident Report for {client} - {db_type}</p>
                        </div>
                        <div style="margin-top:24px; background:#ffffff; border:1px solid #e2e8f0; border-radius:12px; padding:24px;">
                            <p>Hello,</p>
                            <p>One or more database instances / services are reported as OFFLINE or STOPPED:</p>
                            <pre style="background:#f1f5f9; padding:16px; border-radius:8px; font-family:monospace; white-space:pre-wrap;">{server_desc}</pre>
                            <p><strong>Ticket ID:</strong> #{new_down_id}<br/><strong>Assigned Agent:</strong> SYSTEM</p>
                            <p>An incident ticket #{new_down_id} has been opened automatically. Please check immediately.</p>
                            <p style="margin-top:24px; font-size:12px; color:#64748b; border-top:1px solid #e2e8f0; padding-top:12px;">
                                This alert was triggered automatically by the GeoMon monitoring service. Always from mailbox: dccagent@geopits.com.
                            </p>
                        </div>
                    </body>
                    </html>
                    """
                    if to_emails:
                        send_email_outlook(to_emails, cc_emails, f"[Ticket #{new_down_id}] {ticket_name}", email_down_body, sender_email="dccagent@geopits.com")
                    else:
                        print(f"[ALERT DAEMON] No recipient configured for {client} ({db_type}) — server-down email skipped.")

                    # Log email routing outcome
                    if to_emails:
                        email_log_msg = f"Alert email sent to: {to_emails}"
                        if cc_emails:
                            email_log_msg += f" (CC: {cc_emails})"
                    else:
                        email_log_msg = f"Email alert skipped: No recipient email configured for client {client} or database technology {db_type}."
                    cur.execute("""
                        INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
                        VALUES (%s, 'System', 'log', %s, '');
                    """, (new_down_id, email_log_msg))
                    conn.commit()
                    print(f"[ALERT DAEMON] Created server down ticket #{new_down_id} & sent email for {client} ({db_type}).")

            # Handle Critical Error Log Alerts
            if critical_errors_found:
                for err_msg in critical_errors_found:
                    ticket_name = f"[Critical Error] {client} ({db_type}) - Error Log Alert"
                    cur.execute("""
                        SELECT id FROM tickets 
                        WHERE company = %s AND business_unit = %s AND ticket_name = %s AND description = %s AND status != 'CLOSED' AND status != 'RESOLVED';
                    """, (client, db_type, ticket_name, err_msg))
                    existing_err_t = cur.fetchone()

                    if not existing_err_t:
                        cur.execute("""
                            INSERT INTO tickets (business_unit, company, contact, ticket_name, category, status, priority, agent, description, created_by, created_at)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'System', NOW())
                            RETURNING id;
                        """, (db_type, client, to_emails, ticket_name, 'Alert', 'OPEN', 'High', 'SYSTEM', err_msg + "\nAssigned Agent: SYSTEM"))
                        new_err_id = cur.fetchone()[0]

                        email_err_body = f"""
                        <html>
                        <body style="font-family:-apple-system, BlinkMacSystemFont, sans-serif; color:#1e293b; line-height:1.6; max-width:800px; margin:0 auto; padding:20px;">
                            <div style="background:#ef4444; padding:24px; border-radius:12px; color:white;">
                                <h2 style="margin:0; font-size:24px;">🚨 Critical Error Log Detected</h2>
                                <p style="margin:8px 0 0 0; font-size:14px; opacity:0.9;">Ticket #{new_err_id} | Database Error Log Event for {client} - {db_type}</p>
                            </div>
                            <div style="margin-top:24px; background:#ffffff; border:1px solid #e2e8f0; border-radius:12px; padding:24px;">
                                <p>Hello,</p>
                                <p>The following critical database error log was captured by the agent monitor:</p>
                                <pre style="background:#f1f5f9; padding:16px; border-radius:8px; font-family:monospace; white-space:pre-wrap;">{err_msg}</pre>
                                <p><strong>Ticket ID:</strong> #{new_err_id}<br/><strong>Assigned Agent:</strong> SYSTEM</p>
                                <p>An incident ticket #{new_err_id} has been opened automatically. Please check immediately.</p>
                                <p style="margin-top:24px; font-size:12px; color:#64748b; border-top:1px solid #e2e8f0; padding-top:12px;">
                                    This alert was triggered automatically by the GeoMon monitoring service. Always from mailbox: dccagent@geopits.com.
                                </p>
                            </div>
                        </body>
                        </html>
                        """
                        if to_emails:
                            send_email_outlook(to_emails, cc_emails, f"[Ticket #{new_err_id}] {ticket_name}", email_err_body, sender_email="dccagent@geopits.com")
                        else:
                            print(f"[ALERT DAEMON] No recipient configured for {client} ({db_type}) — critical error email skipped.")

                        # Log email routing outcome
                        if to_emails:
                            email_log_msg = f"Alert email sent to: {to_emails}"
                            if cc_emails:
                                email_log_msg += f" (CC: {cc_emails})"
                        else:
                            email_log_msg = f"Email alert skipped: No recipient email configured for client {client} or database technology {db_type}."
                        cur.execute("""
                            INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
                            VALUES (%s, 'System', 'log', %s, '');
                        """, (new_err_id, email_log_msg))
                        conn.commit()
                        print(f"[ALERT DAEMON] Created critical error ticket #{new_err_id} & sent email for {client} ({db_type}).")

        cur.close()
        conn.close()
    except Exception as sweep_err:
        print(f"[ALERT DAEMON] Exception in alert checking sweep: {sweep_err}")

def client_alert_settings_loop():
    import time
    print("[ALERT DAEMON] Daemon thread running every 5 minutes...")
    while True:
        try:
            check_client_alert_thresholds()
        except Exception as e:
            print(f"[ALERT DAEMON] Loop error: {e}")
        time.sleep(300)


def _mail_reader_loop():
    """Background daemon: reads DB-uptime (and other telemetry) mails every hour."""
    try:
        from email_extracter import read_mail
        from exchangelib.protocol import Protocol
    except Exception as import_err:
        print(f"[MAIL SCHEDULER] Import error – mail loop disabled: {import_err}")
        return

    print("[MAIL SCHEDULER] DB-uptime mail reader started (interval: 1 hour).")
    while True:
        try:
            read_mail()
            try:
                from email_extracter import account, simulate_local_uptime_update
                if not account:
                    simulate_local_uptime_update()
            except Exception as sim_err:
                print(f"[MAIL SCHEDULER] Simulation update error: {sim_err}")
            print("[MAIL SCHEDULER] Mail sweep complete.")
        except Exception as err:
            err_msg = str(err).lower()
            print(f"[MAIL SCHEDULER] Error during sweep: {err}")
            if any(kw in err_msg for kw in ["refresh_token", "aadsts", "expired"]):
                print("[MAIL SCHEDULER] OAuth token expired – attempting reconnect...")
                try:
                    from email_extracter import get_account
                    import email_extracter as _em
                    Protocol.clear_cache()
                    _em.account = get_account()
                    print("[MAIL SCHEDULER] Reconnected successfully.")
                except Exception as reconnect_err:
                    print(f"[MAIL SCHEDULER] Reconnect failed: {reconnect_err}")
            elif any(kw in err_msg for kw in ["cannot service", "try again later"]):
                print("[MAIL SCHEDULER] Exchange busy – backing off 2 min.")
                _time.sleep(120)
                continue
        _time.sleep(3600)  # wait 1 hour before next sweep


@app.on_event("startup")
async def startup_event():
    # Increase AnyIO worker thread limit for sync routes to prevent thread starvation under high concurrent requests
    from anyio import to_thread
    limiter = to_thread.current_default_thread_limiter()
    limiter.total_tokens = 500
    print(f"AnyIO thread limiter total_tokens set to {limiter.total_tokens}")

    # Launch the hourly DB-uptime mail reader as a background daemon thread
    mail_thread = threading.Thread(target=_mail_reader_loop, name="MailReaderDaemon", daemon=True)
    mail_thread.start()
    print("[MAIL SCHEDULER] Background mail reader thread launched.")

    # Launch the client alert threshold checking daemon thread
    alert_settings_thread = threading.Thread(target=client_alert_settings_loop, name="AlertSettingsDaemon", daemon=True)
    alert_settings_thread.start()
    print("[ALERT DAEMON] Background alert settings checker thread launched.")

    # Run database maintenance cleanup for processed emails older than 30 days
    try:
        from cleanup_processed_emails import run_cleanup
        cleanup_thread = threading.Thread(target=run_cleanup, name="DbMaintenanceCleanup", daemon=True)
        cleanup_thread.start()
        print("[CLEANUP] Database maintenance cleanup thread started.")
    except Exception as cleanup_err:
        print(f"[CLEANUP ERROR] Failed to start cleanup thread: {cleanup_err}")

from migrations import run_migrations
from routes import router as new_features_router

run_migrations()
app.include_router(new_features_router)

#origins = [
#    "https://api.geomon.geopits.com",
#    "https://api.geovexsight.geopits.com",
#]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(GZipMiddleware, minimum_size=1000)
app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("JWT_SECRET", "super-secret-key-geopits"),
    same_site="none",      # Required: Microsoft OAuth redirects cross-site
    https_only=True,       # Required when same_site=none (browser enforces this)
    max_age=300,           # Session lives 5 minutes — enough for OAuth flow
)

DEFAULT_ALLOWED = '127.0.0.1'
ALLOWED_NETWORKS_CONFIG = os.getenv("ALLOWED_IP_NETWORKS", DEFAULT_ALLOWED).split(",")

@app.middleware("http")
async def network_restriction_middleware(request: Request, call_next):
    client_ip_str = "127.0.0.1"
    if request.client:
        client_ip_str = request.client.host
        
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        client_ip_str = forwarded.split(",")[0].strip()
        
    try:
        # Normalize localhost string to IP address if passed
        if client_ip_str == "localhost":
            client_ip_str = "127.0.0.1"
            
        client_ip = ip_address(client_ip_str)
        
        # Allow loopback and all private local IP networks (like 192.168.x.x, 10.x.x.x, 172.16.x.x) for seamless local dev/testing
        if client_ip.is_loopback or client_ip.is_private:
            return await call_next(request)
            
        is_allowed = False
        for network_str in ALLOWED_NETWORKS_CONFIG:
            net_str = network_str.strip().replace("[", "").replace("]", "")
            if not net_str: continue
            
            # Skip non-IP network configurations like 'localhost' which fail in ip_network()
            if net_str == "localhost":
                continue
                
            try:
                if client_ip in ip_network(net_str):
                    is_allowed = True
                    break
            except ValueError:
                continue
                
        if not is_allowed:
            print(f"SECURITY DENIED: {client_ip_str} tried to access {request.url.path}")
            return JSONResponse(
                status_code=403, 
                content={"detail": f"GeoPITS Security: Your IP ({client_ip_str}) is not authorized. Access is restricted to trusted office networks."}
            )
    except Exception as e:
        print(f"IP Filter Error identifying client: {e}")
        # Fallback to local bypass for safety in local execution environments to prevent blocking
        return await call_next(request)

    return await call_next(request)

try:
    openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception as e:
    print(f"Failed to initialize OpenAI client: {e}")
    openai_client = None

os.makedirs("static", exist_ok=True)

@app.get("/")
async def serve_index():
    response = FileResponse("static/index.html")
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

app.mount("/static", CachedStaticFiles(directory="static"), name="static")

SECRET_KEY = os.getenv("JWT_SECRET", "super-secret-key-geopits")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

ADMIN_EMAILS = [email.strip().lower() for email in os.getenv("ADMIN_EMAILS", "").split(",") if email.strip()]

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/login")

try:
    db_pool = ThreadedConnectionPool(
        minconn=5,
        maxconn=150,
        host=os.getenv("DB_HOST", "localhost"),
        database=os.getenv("DB_NAME", "Incoming-error-data"),
        user=os.getenv("DB_USER", "postgres"),
        password=os.getenv("DB_PASSWORD", "y7UMhWmLcqSJzmhTGDyK"),
        port=os.getenv("DB_PORT", "5432")
    )
    print("Database Connection Pool initialized.")
except Exception as e:
    print(f"CRITICAL ERROR: Could not initialize database pool: {e}")
    db_pool = None

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

@contextmanager
def get_db_connection():
    """Context manager for database connections from the pool."""
    if not db_pool:
        conn = psycopg2.connect(
            host=os.getenv("DB_HOST", "localhost"),
            database=os.getenv("DB_NAME", "Incoming-error-data"),
            user=os.getenv("DB_USER", "postgres"),
            password=os.getenv("DB_PASSWORD", "y7UMhWmLcqSJzmhTGDyK"),
            port=os.getenv("DB_PORT", "5432")
        )
        try:
            yield conn
        except Exception:
            try:
                conn.rollback()
            except Exception:
                pass
            raise
        finally:
            try:
                conn.close()
            except Exception:
                pass
    else:
        conn = db_pool.getconn()
        try:
            yield conn
        except Exception:
            try:
                conn.rollback()
            except Exception:
                pass
            raise
        finally:
            try:
                db_pool.putconn(conn)
            except Exception:
                pass

# Global lock to prevent concurrent refreshes of the materialized view
_mv_refresh_lock = False

def refresh_combined_logs_mv():
    """Helper to refresh the materialized view in the background."""
    global _mv_refresh_lock
    if _mv_refresh_lock:
        print("[MV REFRESH] Skip: Refresh already in progress.")
        return

    _mv_refresh_lock = True
    print("[MV REFRESH] Start: Refreshing combined_logs_mv...")
    try:
        # Use a new connection for the refresh
        with get_db_connection() as conn:
            # Refresh must run outside a transaction block for CONCURRENTLY if possible, 
            # or simply commit immediately after.
            conn.set_isolation_level(0) # AUTOCOMMIT
            with conn.cursor() as cur:
                cur.execute("REFRESH MATERIALIZED VIEW CONCURRENTLY combined_logs_mv;")
            print("[MV REFRESH] Success: View updated.")
    except Exception as e:
        print(f"[MV REFRESH] Error: {str(e)}")
    finally:
        _mv_refresh_lock = False


def resolve_user_role(email, username, current_role):
    """
    Dynamically determines user role based on assignments in the system_admins and leads table.
    - If email exists in 'system_admins' and is active, user becomes 'admin'.
    - If assigned 'Global' technology in leads, user becomes 'admin' (backward compatibility).
    - If assigned any other technology as lead, user becomes 'lead'.
    """
    email_strip = (email or "").strip().lower()
    username_strip = (username or "").strip().lower()
    
    if not email_strip and not username_strip:
        return current_role
        
    try:
        if current_role == "admin":
            return "admin"
            
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            if email_strip:
                cur.execute("SELECT status FROM system_admins WHERE LOWER(email) = %s AND status = 'active'", (email_strip,))
                if cur.fetchone():
                    cur.close()
                    print(f"DEBUG_ROLE: Promoted '{email_strip}' to admin via system_admins table")
                    return "admin"
            
            cur.execute("""
                SELECT technology, is_lead FROM leads 
                WHERE (TRIM(LOWER(email)) = %s OR TRIM(LOWER(email)) = %s)
                AND status = 'active'
            """, (email_strip, username_strip))
            rows = cur.fetchall()
            cur.close()
            
            if not rows:
                return current_role
                
            techs = [row[0] for row in rows]
            lead_techs = [row[0] for row in rows if row[1] is True]
            
            if any(t in ['Global', 'Global Admin'] for t in techs):
                print(f"DEBUG_ROLE: Promoted '{username_strip}' to admin via Global assignment in leads")
                return "admin"
            
            if lead_techs:
                return "lead"
            
            return current_role
            
    except Exception as e:
        print(f"DEBUG_ROLE_ERROR: Failed to resolve role for {email}/{username}: {e}")
        return current_role

def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
        
    assigned_techs = []
    allowed_clients = []
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT id, username, role, full_name, profile_pic, email FROM users WHERE username = %s;", (username,))
        user_row = cur.fetchone()
        if user_row:
            email_addr = (user_row[5] or "").lower()
            cur.execute("SELECT DISTINCT technology FROM leads WHERE LOWER(email) = LOWER(%s) AND status = 'active';", (email_addr,))
            assigned_techs = [r[0] for r in cur.fetchall()]
            
            # Fetch explicitly permitted client names for user-based access control
            cur.execute("""
                SELECT DISTINCT c.client_name 
                FROM user_clients uc
                JOIN admin_clients c ON uc.client_id = c.id
                WHERE uc.user_id = %s;
            """, (user_row[0],))
            allowed_clients = [r[0] for r in cur.fetchall() if r[0]]
        cur.close()

    if user_row is None:
        raise credentials_exception
        
    user_data = {
        "id": user_row[0],
        "username": user_row[1],
        "role": user_row[2],
        "full_name": user_row[3],
        "profile_pic": user_row[4],
        "email": user_row[5]
    }
    
    email_addr = (user_row[5] or "").lower()
    user_data["role"] = resolve_user_role(email_addr, user_data["username"], user_data["role"])
    
    user_data["isAdmin"] = (user_data["role"] == "admin") or (email_addr in ADMIN_EMAILS)

    user_data["fullName"] = user_row[3]
    user_data["profilePic"] = user_row[4]
    user_data["assigned_techs"] = assigned_techs
    user_data["allowed_clients"] = allowed_clients

    # Check client access status
    client_disabled = False
    is_client_user = False
    if email_addr:
        try:
            with get_db_connection() as conn:
                cur = conn.cursor()
                cur.execute("SELECT status FROM client_access WHERE LOWER(client_email) = %s", (email_addr,))
                row = cur.fetchone()
                if row:
                    is_client_user = True
                    if row[0] == 'disabled':
                        client_disabled = True
                cur.close()
        except Exception as e:
            print(f"Error checking client access status: {e}")
            
    user_data["clientAccessDisabled"] = client_disabled
    user_data["isClientUser"] = is_client_user

    has_access, access_type = check_user_access(username)
    if not has_access:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="GeoPITS Access Restricted: Your account has not been assigned any monitoring privileges. Please contact an administrator."
        )

    print(f"TELEMETRY: Identified user {user_data['username']} (Admin: {user_data['isAdmin']}, Email: {user_data['email']}, Techs: {user_data['assigned_techs']})")
    return user_data

@app.get("/api/me")
def get_me(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("UPDATE users SET last_active_at = %s WHERE username = %s", (datetime.now(ZoneInfo("Asia/Kolkata")), current_user["username"]))
            conn.commit()
            cur.close()
    except Exception as e:
        print(f"Error updating last_active_at: {e}")
    
    print(f"DEBUG_ME_ROLE: User '{current_user.get('username')}' ({current_user.get('email')}) has role: '{current_user.get('role')}' (isAdmin: {current_user.get('isAdmin')})")
    
    return current_user

class ShareRecordRequest(BaseModel):
    notes: str
    platform: str
    content_type: str
    client_name: Optional[str] = ''
    server_name: Optional[str] = ''
    log_message: Optional[str] = ''
    status: Optional[str] = ''
    owner: Optional[str] = ''
    client_visibility: Optional[str] = ''
    ticket_status: Optional[str] = ''
    next_action: Optional[str] = ''
    db_type: Optional[str] = ''

@app.post("/api/share/record")
def record_share(req: ShareRecordRequest, current_user: dict = Depends(get_current_user)):
    if current_user.get("isClientUser") or current_user.get("clientAccessDisabled"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Client Access Restricted: Your account has client access mode and cannot perform modifications."
        )
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute(
                """INSERT INTO share_history (username, notes, platform, content_type, client_name, server_name, log_message, status, owner, client_visibility, ticket_status, next_action, db_type) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                (current_user["username"], req.notes, req.platform, req.content_type,
                 req.client_name, req.server_name, req.log_message, req.status, req.owner, req.client_visibility, req.ticket_status, req.next_action, req.db_type)
            )
            conn.commit()
        return {"status": "success"}
    except Exception as e:
        print(f"Share record error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/share/history")
def get_share_history(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = """
                SELECT id, username, notes, platform, content_type, shared_at AS created_at, 
                       client_name, server_name, log_message, status, owner, 
                       client_visibility, ticket_status, next_action, db_type 
                FROM share_history 
                WHERE 1=1
            """
            params = []
            
            allowed_techs = get_user_allowed_technologies(current_user)
            if allowed_techs is not None:
                query += " AND (TRIM(LOWER(db_type)) = ANY(%s) OR username = %s)"
                tech_params = [t.lower().strip() for t in allowed_techs]
                params.append(tech_params)
                params.append(current_user["username"])
            else:
                pass

            query += " ORDER BY created_at DESC LIMIT 50"
            cur.execute(query, tuple(params))
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
            cur.close()
        return {"history": results}
    except Exception as e:
        print(f"Share history fetch error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def check_user_access(username_or_email: str):
    """
    Checks if a user has access to the application.
    Admins (by email or role) have full access.
    Normal users must have at least one active assignment in the 'leads' table.
    """
    username_or_email = username_or_email.lower()
    
    if username_or_email in ADMIN_EMAILS:
        return True, "admin"

    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            if username_or_email in ADMIN_EMAILS:
                cur.close()
                return True, "admin"
                
            cur.execute("SELECT status FROM system_admins WHERE LOWER(email) = %s AND status = 'active'", (username_or_email,))
            if cur.fetchone():
                cur.close()
                return True, "admin"
            
            cur.execute("SELECT role, email FROM users WHERE username = %s OR email = %s", (username_or_email, username_or_email))
            user_row = cur.fetchone()
            if user_row:
                role, email = user_row
                if role == "admin" or (email and email.lower() in ADMIN_EMAILS):
                    cur.close()
                    return True, "admin"
                    
                cur.execute("SELECT COUNT(*) FROM leads WHERE (LOWER(email) = %s OR LOWER(email) = (SELECT LOWER(email) FROM users WHERE username = %s)) AND status = 'active'", (username_or_email, username_or_email))
                count = cur.fetchone()[0]
                
                cur.execute("""
                    SELECT COUNT(*) FROM user_clients uc
                    JOIN users u ON uc.user_id = u.id
                    WHERE u.username = %s OR LOWER(u.email) = %s;
                """, (username_or_email, username_or_email))
                client_count = cur.fetchone()[0]
                cur.close()
                
                if count > 0 or client_count > 0:
                    return True, "user"
                    
        return False, "restricted"
    except Exception as e:
        print(f"Error checking user access: {e}")
        return False, "error"

@app.post("/api/login")
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT hashed_password, email, role FROM users WHERE username = %s;", (form_data.username,))
        row = cur.fetchone()
        cur.close()
    
    if not row or not bcrypt.checkpw(form_data.password.encode('utf-8'), row[0].encode('utf-8')):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    has_access, access_type = check_user_access(form_data.username)
    if not has_access:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="GeoPITS Access Restricted: Your account has not been assigned any monitoring privileges. Please contact an administrator."
        )
        
    resolved_role = resolve_user_role(row[1], form_data.username, row[2])
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": form_data.username, "role": resolved_role}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}
oauth = OAuth()

TENANT_ID = os.getenv("APP_TENANT")
oauth.register(
    name='microsoft',
    client_id=os.getenv("APP_CLIENT"),
    client_secret=os.getenv("APP_SECRET"),
    server_metadata_url=f"https://login.microsoftonline.com/{TENANT_ID}/v2.0/.well-known/openid-configuration" if TENANT_ID else "https://login.microsoftonline.com/common/v2.0/.well-known/openid-configuration",
    client_kwargs={'scope': 'openid email profile', 'prompt': 'select_account'}
)

@app.get("/api/auth/login/{provider}")
async def auth_login(provider: str, request: Request):
    redirect_uri = os.getenv("APP_REDIRECT_URI")
    if not redirect_uri:
        redirect_uri = f"https://api.geomon.geopits.com/api/auth/callback/{provider}"
    if provider == 'microsoft':
        return await oauth.microsoft.authorize_redirect(request, redirect_uri)
    else:
        raise HTTPException(status_code=400, detail="Invalid provider")

@app.get("/api/auth/callback/{provider}")
async def auth_callback(provider: str, request: Request):
    try:
        redirect_uri = os.getenv("APP_REDIRECT_URI")
        if not redirect_uri:
            redirect_uri = f"https://api.geomon.geopits.com/api/auth/callback/{provider}"
        if provider == 'microsoft':
            try:
                token = await oauth.microsoft.authorize_access_token(request, redirect_uri=redirect_uri)
            except Exception as state_err:
                err_str = str(state_err).lower()
                if "mismatching_state" in err_str or "csrf" in err_str or "state" in err_str:
                    # Session cookie was lost during the redirect (common with proxies / SameSite issues).
                    # Clear stale state from session and retry without CSRF validation.
                    print(f"[OAuth] State mismatch detected — retrying without session state: {state_err}")
                    request.session.pop("_state_microsoft", None)
                    request.session.pop("_state", None)
                    # Re-request token directly using the code in the query string
                    code = request.query_params.get("code")
                    if not code:
                        return RedirectResponse(url="/#/login?error=oauth_failed")
                    from authlib.integrations.httpx_client import AsyncOAuth2Client
                    import httpx
                    client_id = os.getenv("APP_CLIENT")
                    client_secret = os.getenv("APP_SECRET")
                    tenant_id = os.getenv("APP_TENANT", "common")
                    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
                    async with httpx.AsyncClient() as client:
                        resp = await client.post(token_url, data={
                            "grant_type": "authorization_code",
                            "code": code,
                            "redirect_uri": redirect_uri,
                            "client_id": client_id,
                            "client_secret": client_secret,
                            "scope": "openid email profile",
                        })
                    if resp.status_code != 200:
                        print(f"[OAuth] Token exchange failed: {resp.text}")
                        return RedirectResponse(url="/#/login?error=oauth_failed")
                    token = resp.json()
                    # Decode the id_token to get user info
                    import base64, json as _json
                    id_token = token.get("id_token", "")
                    try:
                        payload_b64 = id_token.split(".")[1]
                        payload_b64 += "=" * (4 - len(payload_b64) % 4)
                        user_info = _json.loads(base64.b64decode(payload_b64).decode())
                    except Exception:
                        user_info = {}
                else:
                    raise
            else:
                user_info = token.get('userinfo')
        else:
            raise HTTPException(status_code=400, detail="Invalid provider")

        if not user_info:
            raise HTTPException(status_code=401, detail="Could not retrieve user info")

        print(f"OAuth {provider} userinfo: {dict(user_info)}")

        email = user_info.get('email') or user_info.get('preferred_username')
        full_name = user_info.get('name') or user_info.get('given_name', '') + ' ' + user_info.get('family_name', '') or email
        full_name = full_name.strip() if full_name else email
        profile_pic = user_info.get('picture')
        username = email.split('@')[0] if email else user_info.get('name', 'oauth_user')

        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT username, email, role FROM users WHERE username = %s OR username = %s", (username, email))
            row = cur.fetchone()
            
            if not row:
                user_role = "user"
                cur.execute(
                    "INSERT INTO users (username, hashed_password, full_name, profile_pic, email, role) VALUES (%s, %s, %s, %s, %s, %s)", 
                    (username, "OAUTH_LOGIN_NO_PASSWORD", full_name, profile_pic, email, user_role)
                )
                conn.commit()
                print(f"Created new OAuth user: {username} ({email})")
            else:
                username = row[0]
                user_role = row[2]
                cur.execute(
                    "UPDATE users SET full_name = %s, profile_pic = %s, email = %s WHERE username = %s",
                    (full_name, profile_pic, email, username)
                )
                conn.commit()
                
            cur.close()

        has_access, access_type = check_user_access(email)
        if not has_access:
             return RedirectResponse(url="/#/login?error=restricted")

        resolved_role = resolve_user_role(email, username, user_role)

        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": username, "role": resolved_role},
            expires_delta=access_token_expires
        )
        
        return RedirectResponse(url=f"/#/?token={access_token}")

    except HTTPException:
        raise
    except Exception as e:
        print(f"OAuth Callback Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))





def get_user_allowed_technologies(user: dict):
    """Returns a list of allowed technologies for the user. Returns None if unrestricted."""
    if user.get("isAdmin") or user.get("role") == "admin":
        return None
    
    email = user.get("email")
    if not email:
        return None
        
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = LOWER(%s) AND status = 'active';", (email,))
            techs = [row[0] for row in cur.fetchall()]
            cur.close()
        return techs if techs else None 
    except Exception as e:
        print(f"Error fetching user privileges: {e}")
        return None

def get_client_allowed_filters(user: dict):
    """
    Returns a list of tuples (technology, client_name, server_name)
    assigned to the client user. Returns None if they are not a client user.
    """
    if not user.get("isClientUser"):
        return None
        
    email = user.get("email")
    if not email:
        return []
        
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT technology, client_name, server_name FROM client_access WHERE LOWER(client_email) = LOWER(%s) AND status = 'enabled';",
                (email,)
            )
            filters = cur.fetchall()
            cur.close()
        return filters
    except Exception as e:
        print(f"Error fetching client allowed filters: {e}")
        return []


@app.get("/api/users/by-tech")
def get_users_by_tech(db_type: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("SELECT id, username, email, full_name, role FROM users")
            all_users = cur.fetchall()
            
            cur.execute("SELECT email, technology FROM leads WHERE status = 'active'")
            lead_rows = cur.fetchall()

            # Build set of client-portal emails — never show as assignable owners
            cur.execute("SELECT DISTINCT LOWER(client_email) AS ce FROM client_access")
            client_emails_set = {row['ce'] for row in cur.fetchall()}
            
            lead_techs = {}
            for row in lead_rows:
                email = (row['email'] or "").lower()
                if email not in lead_techs:
                    lead_techs[email] = set()
                lead_techs[email].add(row['technology'])
            
            caller_email = (current_user.get("email") or "").lower()
            caller_is_admin = (current_user.get("role") == "admin") or (caller_email in ADMIN_EMAILS)
            
            effective_techs = set()
            if db_type:
                effective_techs.add(db_type)
            elif not caller_is_admin and caller_email in lead_techs:
                effective_techs = lead_techs[caller_email]
                
            valid_users = []
            for u in all_users:
                u_email = (u['email'] or "").lower()
                is_admin = (u['role'] == "admin") or (u_email in ADMIN_EMAILS)

                # Exclude client-portal accounts from the owner assignment dropdown
                if u_email in client_emails_set and not is_admin:
                    continue
                
                if effective_techs and not is_admin:
                    user_techs = lead_techs.get(u_email, set())
                    if not user_techs.intersection(effective_techs):
                        continue
                        
                valid_users.append({
                    "username": u['username'],
                    "label": u['full_name'] or u['username']
                })
                
            valid_users.sort(key=lambda x: x['label'].lower())
            
            cur.close()
        return {"users": valid_users}
        
    except Exception as e:
        print(f"Error fetching users by tech: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/filters")

@app.get("/api/logs/filters")
def get_filters(current_user: dict = Depends(get_current_user)):
    cache_key = f"filters:{current_user.get('email')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val:
        return cached_val

    allowed_techs = get_user_allowed_technologies(current_user)
    allowed_clients = current_user.get("allowed_clients", [])
    is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")

    try:
        now = datetime.now(ZoneInfo("Asia/Kolkata"))
        tech_lookback = "2000-01-01 00:00:00"
        filter_lookback = (now - timedelta(days=7)).strftime("%Y-%m-%d %H:%M:%S")

        with get_db_connection() as conn:
            cur = conn.cursor()

            tech_filter_clause = ""
            if allowed_techs is not None:
                tech_filter_clause = "AND db_type = ANY(%s)"
            
            client_filter_clause = ""
            if not is_admin and allowed_clients:
                client_filter_clause = "AND (client_name = ANY(%s) OR server_name = ANY(%s))"

            skip_scan_query = f"""
                WITH RECURSIVE skip_scan AS (
                    (SELECT db_type FROM db_monitoring_logs 
                     WHERE log_time_ist >= %s 
                       AND db_type IS NOT NULL
                       {tech_filter_clause}
                       {client_filter_clause}
                     ORDER BY db_type LIMIT 1)
                    UNION ALL
                    (SELECT (SELECT db_type FROM db_monitoring_logs 
                             WHERE db_type > s.db_type 
                               AND log_time_ist >= %s 
                               AND db_type IS NOT NULL
                               {tech_filter_clause}
                               {client_filter_clause}
                             ORDER BY db_type LIMIT 1)
                     FROM skip_scan s
                     WHERE s.db_type IS NOT NULL)
                )
                SELECT db_type FROM skip_scan WHERE db_type IS NOT NULL;
            """
            
            skip_params = []
            # First SELECT
            skip_params.append(tech_lookback)
            if allowed_techs is not None:
                skip_params.append(allowed_techs)
            if not is_admin and allowed_clients:
                skip_params.extend([allowed_clients, allowed_clients])
            # Second SELECT
            skip_params.append(tech_lookback)
            if allowed_techs is not None:
                skip_params.append(allowed_techs)
            if not is_admin and allowed_clients:
                skip_params.extend([allowed_clients, allowed_clients])

            cur.execute(skip_scan_query, skip_params)
            discovered_techs = [r[0] for r in cur.fetchall()]

            main_filter_clause = ""
            main_params = [filter_lookback]
            if allowed_techs is not None:
                main_filter_clause += " AND db_type = ANY(%s)"
                main_params.append(allowed_techs)
            if not is_admin and allowed_clients:
                main_filter_clause += " AND (client_name = ANY(%s) OR server_name = ANY(%s))"
                main_params.extend([allowed_clients, allowed_clients])

            cur.execute(f"""
                SELECT DISTINCT db_type, client_name, server_name, log_type
                FROM db_monitoring_logs
                WHERE log_time_ist >= %s
                  AND db_type IS NOT NULL
                  AND client_name IS NOT NULL
                  AND server_name IS NOT NULL
                  AND (
                    CASE 
                        WHEN time_bucket LIKE '%%%%_08AM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08AM', ' 08:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                        WHEN time_bucket LIKE '%%%%_08PM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08PM', ' 20:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                        ELSE log_time_ist::TIMESTAMP 
                    END <= (NOW() AT TIME ZONE 'Asia/Kolkata')::TIMESTAMP
                  )
                  {main_filter_clause}
            """, main_params)

            rows_main = cur.fetchall()
            
            archive_filter_clause = ""
            archive_params = []
            if allowed_techs is not None:
                archive_filter_clause += " AND db_type = ANY(%s)"
                archive_params.append(allowed_techs)
            if not is_admin and allowed_clients:
                archive_filter_clause += " AND (client_name = ANY(%s) OR server_name = ANY(%s))"
                archive_params.extend([allowed_clients, allowed_clients])

            cur.execute(f"""
                SELECT DISTINCT db_type, client_name, server_name, log_type
                FROM db_archived_logs
                WHERE db_type IS NOT NULL
                  AND client_name IS NOT NULL
                  AND server_name IS NOT NULL
                  {archive_filter_clause}
            """, archive_params)
            rows_archive = cur.fetchall()
            
            rows = list(set(rows_main + rows_archive))
            
            for tech in discovered_techs:
                if not any(r[0] == tech for r in rows):
                    rows.append((tech, "Unknown", "Unknown", "N/A"))
            
            rows.sort()
            cur.close()

        db_types, clients = set(), set()
        db_server_map, db_client_map = {}, {}
        client_server_map, server_logtype_map = {}, {}

        for db_type, client_name, server_name, log_type in rows:
            db_types.add(db_type)
            clients.add(client_name)
            db_server_map.setdefault(db_type, set()).add(server_name)
            db_client_map.setdefault(db_type, set()).add(client_name)
            client_server_map.setdefault(client_name, set()).add(server_name)
            if log_type:
                server_logtype_map.setdefault(server_name, set()).add(log_type)

        result = {
            "db_types": sorted(db_types),
            "clients": sorted(clients),
            "db_server_map":    {k: sorted(v) for k, v in db_server_map.items()},
            "db_client_map":    {k: sorted(v) for k, v in db_client_map.items()},
            "client_server_map":{k: sorted(v) for k, v in client_server_map.items()},
            "server_logtype_map":{k: sorted(v) for k, v in server_logtype_map.items()},
            "current_ist": now.isoformat()
        }
        
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class LogBulkAction(BaseModel):
    log_hashes: List[str]
    status: Optional[str] = "None"
    owner: Optional[str] = "None"
    client_visibility: Optional[str] = ""
    ticket_status: Optional[str] = ""
    next_action: Optional[str] = ""
    severity: Optional[str] = None


def create_auto_ticket_for_log(cur, client_name: str, db_type: str, severity: str, log_message: str, log_hash: str, server_name: str = None, log_type: str = None):
    # Deduplicate: check if a ticket with this log hash already exists and is not resolved/closed
    try:
        cur.execute("""
            SELECT id FROM tickets 
            WHERE description LIKE %s 
              AND UPPER(status) NOT IN ('RESOLVED', 'CLOSED') 
            LIMIT 1
        """, (f"%Log Hash: {log_hash}%",))
        row_ex_t = cur.fetchone()
        if row_ex_t and row_ex_t[0]:
            print(f"[DEDUPLICATE] Existing ticket found: {row_ex_t[0]} for hash {log_hash}")
            return row_ex_t[0]

        # Also check if there's an active ticket with the same client, db_type, and containing the log message
        cur.execute("""
            SELECT id FROM tickets 
            WHERE LOWER(TRIM(company)) = LOWER(TRIM(%s)) 
              AND LOWER(TRIM(business_unit)) = LOWER(TRIM(%s)) 
              AND UPPER(status) NOT IN ('RESOLVED', 'CLOSED') 
              AND description LIKE %s 
            LIMIT 1
        """, (client_name, db_type, f"%{log_message}%"))
        row_ex_msg = cur.fetchone()
        if row_ex_msg and row_ex_msg[0]:
            print(f"[DEDUPLICATE] Existing active ticket found by message: {row_ex_msg[0]}")
            return row_ex_msg[0]
    except Exception as ex_dedup:
        print(f"Error checking existing ticket: {ex_dedup}")

    from db_manager import get_alert_contacts
    contacts = get_alert_contacts(cur, client_name, db_type)
    contact_emails = contacts["to_emails"]
    
    ticket_name = f"[AUTO] {severity} Alert: {log_type or 'Incident'} on {server_name or 'any'}"
    description = f"Auto-generated ticket for log event.\nMessage: {log_message}\nLog Hash: {log_hash}"
    
    cur.execute("""
        INSERT INTO tickets (
            business_unit, company, contact, ticket_name, category, 
            status, priority, agent, description, created_by, created_at
        )
        VALUES (%s, %s, %s, %s, 'Logs', 'OPEN', %s, 'Unassigned', %s, 'System', NOW())
        RETURNING id;
    """, (db_type, client_name, contact_emails, ticket_name, severity, description))
    t_id = cur.fetchone()[0]

    # Send email notification with Ticket ID in subject (Disabled per user request for consolidated 8 AM / 8 PM emails)
    pass

    return t_id

def ensure_ticket_for_high_critical_log(cur, target_table: str, log_hash: str, new_severity: str = None):
    cur.execute(f"""
        SELECT client_name, server_name, db_type, log_type, log_message, ticket_id, severity 
        FROM {target_table} 
        WHERE TRIM(log_hash) = TRIM(%s)
        LIMIT 1
    """, (log_hash,))
    row = cur.fetchone()
    if not row:
        return
        
    client_name, server_name, db_type, log_type, log_message, ticket_id, severity = row
    final_severity = new_severity if new_severity else severity
    
    if final_severity in ["High", "Critical"] and not ticket_id:
        t_id = create_auto_ticket_for_log(cur, client_name, db_type or "MSSQL", final_severity, log_message, log_hash, server_name, log_type)
        cur.execute(f"""
            UPDATE {target_table}
            SET ticket_id = %s, ticket_status = 'OPEN'
            WHERE TRIM(log_hash) = TRIM(%s)
        """, (t_id, log_hash))

class LogMetadataUpdate(BaseModel):
    client_name: str
    server_name: str
    log_message: str
    log_hash: str
    status: Optional[str] = "None"
    owner: Optional[str] = "None"
    client_visibility: Optional[str] = ""
    ticket_status: Optional[str] = ""
    next_action: Optional[str] = ""
    severity: Optional[str] = None


@app.patch("/api/logs/metadata")

@app.patch("/api/logs/metadata")
def update_log_metadata(
    update: LogMetadataUpdate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    print(f"DEBUG_METADATA_UPDATE: Received request: {update}")
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            terminal_statuses = ['Resolved', 'Ignored', 'No action Required']
            is_terminal = update.status in terminal_statuses
            
            search_hash = update.log_hash.strip()
            cur.execute("SELECT 1 FROM db_monitoring_logs WHERE TRIM(log_hash) = %s", (search_hash,))
            in_main = cur.fetchone() is not None
            
            cur.execute("SELECT 1 FROM db_archived_logs WHERE TRIM(log_hash) = %s", (search_hash,))
            in_archive = cur.fetchone() is not None
    
            target_table = "db_monitoring_logs"
            if in_main:
                target_table = "db_monitoring_logs"
            elif in_archive:
                target_table = "db_archived_logs"
    
            now_ist = datetime.now(ZoneInfo("Asia/Kolkata"))
            
            common_cols = "client_name, server_name, db_type, log_type, log_source, log_time, log_time_utc, log_time_ist, log_level, log_message, occurrence_count, raw_log, email_subject, email_received_time, log_hash, created_at, status, owner, client_visibility, ticket_status, next_action, severity, status_updated_at, is_semantic, semantic_count, semantic_hash, time_bucket, ticket_id"
 
            if is_terminal and in_main:
                print(f"Archiving log: {update.client_name} | {update.server_name} | {update.log_hash}")
                cur.execute(f"""
                    INSERT INTO db_archived_logs ({common_cols})
                    SELECT {common_cols} FROM db_monitoring_logs 
                    WHERE TRIM(log_hash) = %s
                    ON CONFLICT (log_hash) DO NOTHING
                """, (search_hash,))
                cur.execute("DELETE FROM db_monitoring_logs WHERE TRIM(log_hash) = %s", (search_hash,))
                target_table = "db_archived_logs"
            elif not is_terminal and in_archive:
                print(f"Un-archiving log: {update.client_name} | {update.server_name} | {search_hash}")
                cur.execute(f"""
                    INSERT INTO db_monitoring_logs ({common_cols})
                    SELECT {common_cols} FROM db_archived_logs 
                    WHERE TRIM(log_hash) = %s
                    ON CONFLICT (log_hash) DO NOTHING
                """, (search_hash,))
                cur.execute("DELETE FROM db_archived_logs WHERE TRIM(log_hash) = %s", (search_hash,))
                target_table = "db_monitoring_logs"
            cur.execute(f"""
                UPDATE {target_table}
                SET status = %s,
                    owner = %s,
                    client_visibility = %s,
                    ticket_status = %s,
                    next_action = %s,
                    severity = %s,
                    status_updated_at = %s
                WHERE TRIM(log_hash) = %s
            """, (
                update.status, update.owner, update.client_visibility, update.ticket_status, update.next_action, 
                update.severity, now_ist, search_hash
            ))
            
            # Ensure ticket is created for High/Critical severity logs
            ensure_ticket_for_high_critical_log(cur, target_table, search_hash, update.severity)
            
            # Bidirectional sync: update the linked ticket in tickets table if ticket_id exists
            cur.execute(f"SELECT ticket_id FROM {target_table} WHERE TRIM(log_hash) = %s", (search_hash,))
            row = cur.fetchone()
            if row and row[0]:
                t_id = row[0]
                new_t_status = update.ticket_status.upper() if update.ticket_status else 'OPEN'
                if update.status == 'Resolved':
                    new_t_status = 'RESOLVED'
                elif update.status == 'In Progress':
                    new_t_status = 'IN PROGRESS'
                elif update.status == 'Pending':
                    new_t_status = 'PENDING'
                elif update.status == 'Open':
                    new_t_status = 'OPEN'
                
                priority_map = {
                    'Critical': 'Critical',
                    'High': 'High',
                    'Medium': 'Medium',
                    'Low': 'Low'
                }
                new_priority = priority_map.get(update.severity, 'Medium') if update.severity else 'Medium'
                
                cur.execute("SELECT status, resolved_by, resolved_at FROM tickets WHERE id = %s", (t_id,))
                t_row = cur.fetchone()
                if t_row:
                    t_curr_status, t_curr_res_by, t_curr_res_at = t_row
                    res_by = t_curr_res_by
                    res_at = t_curr_res_at
                    if new_t_status == 'RESOLVED' and t_curr_status != 'RESOLVED':
                        res_by = current_user.get('username', 'System')
                        res_at = now_ist
                    elif new_t_status != 'RESOLVED':
                        res_by = None
                        res_at = None

                    new_agent = update.owner if update.owner and update.owner != 'None' and update.owner != 'Unassigned' else 'Unassigned'

                    cur.execute("""
                        UPDATE tickets
                        SET status = %s,
                            priority = %s,
                            agent = %s,
                            resolved_by = %s,
                            resolved_at = %s
                        WHERE id = %s
                    """, (new_t_status, new_priority, new_agent, res_by, res_at, t_id))
            
            print(f"DEBUG_METADATA_UPDATE: Table: {target_table}, Hash: {search_hash[:10]}..., Status: {update.status}, Affected: {cur.rowcount}")
            
            conn.commit()
            background_tasks.add_task(refresh_combined_logs_mv)
            
            # Invalidate caches
            cache_manager.invalidate("logs:")
            cache_manager.invalidate("owner-counts:")
            cache_manager.invalidate("filters:")
            cache_manager.invalidate("tickets:")
            cache_manager.invalidate("ticket-stats:")
            global OBSERVABILITY_CACHE
            OBSERVABILITY_CACHE.clear()
            return {"success": True}
    except Exception as e:
            print(f"Error updating metadata: {e}")
            if 'conn' in locals(): conn.rollback()
            raise HTTPException(status_code=500, detail=str(e))
    finally:
            if 'cur' in locals(): cur.close()




@app.get("/api/logs")

@app.get("/api/logs")
def get_logs(
    start_time: Optional[str] = None,
    end_time: Optional[str] = None,
    client_name: Optional[str] = None,
    db_type: Optional[str] = None,
    server_name: Optional[str] = None,
    log_type: List[str] = Query(None),
    severity: List[str] = Query(None),
    log_status: List[str] = Query(None),

    limit: int = Query(100, le=1000),
    offset: int = Query(0),
    owner: Optional[str] = None,
    username: Optional[str] = None,
    log_id: Optional[int] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    print(f"DEBUG: get_logs called. user={current_user.get('email')}, status={log_status}, client={client_name}, tech={db_type}, log_id={log_id}, owner={owner}, username={username}")

    cache_key = f"logs:{current_user.get('email')}:{start_time}:{end_time}:{client_name}:{db_type}:{server_name}:{log_type}:{severity}:{log_status}:{limit}:{offset}:{owner}:{log_id}:{search}"
    cached_val = cache_manager.get(cache_key)
    if cached_val:
        return cached_val

    if username and not owner:
        owner = username

    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            log_table = "db_monitoring_logs"
            if log_status and any(s in ['Resolved', 'Ignored', 'No action Required'] for s in log_status):
                log_table = "db_archived_logs"
            
            print(f"DEBUG: Using log_table={log_table}")
    
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            base_query = f"""
                FROM {log_table} 
                WHERE (
                    CASE 
                        WHEN time_bucket LIKE '%%%%_08AM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08AM', ' 08:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                        WHEN time_bucket LIKE '%%%%_08PM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08PM', ' 20:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                        ELSE log_time_ist::TIMESTAMP 
                    END <= (NOW() AT TIME ZONE 'Asia/Kolkata')::TIMESTAMP
                )
            """
            params = []

            if not start_time:
                # Default to last 1 hours if no filter is active
                now_ist = datetime.now(ZoneInfo("Asia/Kolkata"))
                start_time = (now_ist - timedelta(hours=1)).strftime("%Y-%m-%d %H:%M:%S")
                print(f"DEBUG: No start_time provided. Defaulting to last 1 hours: {start_time}")

            if start_time:
                base_query += " AND log_time_ist >= %s"
                params.append(start_time)
            if end_time:
                adj_end_time = end_time
                if len(adj_end_time) == 16:
                    adj_end_time += ":59"
                elif adj_end_time.endswith(":00"):
                    adj_end_time = adj_end_time[:-3] + ":59"
                    
                base_query += " AND log_time_ist <= %s"
                params.append(adj_end_time)
            if client_name:
                base_query += " AND client_name = %s"
                params.append(client_name)
            if db_type:
                base_query += " AND db_type ILIKE %s"
                params.append(db_type)
            if server_name:
                base_query += " AND server_name = %s"
                params.append(server_name)
            if log_type:
                base_query += " AND log_type IN %s"
                params.append(tuple(log_type) if isinstance(log_type, list) else (log_type,))
            if severity:
                base_query += " AND severity IN %s"
                params.append(tuple(severity) if isinstance(severity, list) else (severity,))
                
            if log_status:
                base_query += " AND LOWER(status) IN %s"
                params.append(tuple(s.lower() for s in log_status) if isinstance(log_status, list) else (log_status.lower(),))
            else:
                base_query += " AND (status IS NULL OR TRIM(status) = '' OR LOWER(status) NOT IN ('resolved', 'ignored', 'no action required'))"
                
            if owner:
                if owner.lower() == "unassigned":
                    base_query += " AND (owner IS NULL OR owner = '' OR owner = 'Unassigned' OR owner = 'None')"
                elif owner.lower() == "assigned":
                    base_query += " AND (owner IS NOT NULL AND owner != '' AND owner != 'Unassigned' AND owner != 'None')"
                else:
                    base_query += " AND owner ILIKE %s"
                    params.append(f"%{owner}%")
            
            if log_id:
                base_query += " AND id = %s"
                params.append(log_id)

            if search:
                base_query += " AND (log_message ILIKE %s OR CAST(id AS TEXT) ILIKE %s)"
                search_pattern = f"%{search}%"
                params.extend([search_pattern, search_pattern])

            # Apply strict client and technology constraints for standard/client users
            is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
            is_client = current_user.get("isClientUser")
            if is_client:
                allowed_filters = get_client_allowed_filters(current_user)
                if not allowed_filters:
                    print(f"DEBUG_CLIENT: No allowed filters for client {current_user.get('email')}")
                    return {"logs": [], "total": 0}
                
                filter_clauses = []
                filter_params = []
                for f_tech, f_client, f_server in allowed_filters:
                    filter_clauses.append("(TRIM(LOWER(db_type)) = %s AND client_name = %s AND server_name = %s)")
                    filter_params.extend([f_tech.lower().strip(), f_client, f_server])
                
                base_query += " AND (" + " OR ".join(filter_clauses) + ")"
                params.extend(filter_params)
                print(f"DEBUG_CLIENT: Applied client filters for {current_user.get('email')}: {allowed_filters}")
            elif not is_admin:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                
                # Apply allowed_clients constraint
                if allowed_clients:
                    base_query += " AND (client_name = ANY(%s) OR server_name = ANY(%s))"
                    params.append(allowed_clients)
                    params.append(allowed_clients)
                    
                # Apply allowed_techs constraint
                if allowed_techs is not None:
                    if not allowed_techs:
                        return {"logs": [], "total": 0}
                    if db_type and db_type not in allowed_techs:
                        return {"logs": [], "total": 0}
                    base_query += " AND TRIM(LOWER(db_type)) = ANY(%s)"
                    tech_params = [t.lower().strip() for t in allowed_techs]
                    params.append(tech_params)
                    
                # If a standard user has no privileges assigned, return empty results
                if not allowed_clients and allowed_techs is not None and not allowed_techs:
                    return {"logs": [], "total": 0}

            # Optimization: db_archived_logs has a unique index on log_hash.
            # We can skip GROUP BY and use COUNT(*) if log_table is db_archived_logs.
            if log_table == "db_archived_logs":
                group_by = ""
                count_query = f"SELECT COUNT(*) as total_count {base_query}"
                select_clause = """
                    client_name, 
                    server_name, 
                    db_type, 
                    log_type, 
                    log_time_ist, 
                    log_message, 
                    status,
                    owner,
                    client_visibility,
                    ticket_status,
                    next_action,
                    severity,
                    log_hash,
                    id,
                    status_updated_at,
                    occurrence_count,
                    is_semantic,
                    semantic_count,
                    semantic_hash,
                    time_bucket
                """
            else:
                group_by = " GROUP BY client_name, server_name, db_type, log_type, log_message, status, owner, client_visibility, ticket_status, next_action, severity, log_hash"
                count_query = f"SELECT COUNT(DISTINCT log_hash) as total_count {base_query}"
                select_clause = """
                    client_name, 
                    server_name, 
                    db_type, 
                    log_type, 
                    MAX(log_time_ist) as log_time_ist, 
                    log_message, 
                    status,
                    owner,
                    client_visibility,
                    ticket_status,
                    next_action,
                    severity,
                    log_hash,
                    MAX(id) as id,
                    MAX(status_updated_at) as status_updated_at,
                    SUM(occurrence_count) as occurrence_count,
                    bool_or(is_semantic) as is_semantic,
                    sum(semantic_count) as semantic_count,
                    max(semantic_hash) as semantic_hash,
                    max(time_bucket) as time_bucket
                """
            
            print(f"ULTRA_VERBOSE_SQL_COUNT: {count_query} | Params: {params}")
            cur.execute(count_query, params)
            count_res = cur.fetchone()
            total_records = count_res['total_count'] if count_res else 0
            print(f"ULTRA_VERBOSE_SQL_COUNT: Found {total_records} records")
                
            query = f"""
                SELECT 
                    {select_clause}
                {base_query} 
                {group_by}
                ORDER BY log_time_ist DESC 
                LIMIT %s OFFSET %s
            """
            
            print(f"DEBUG_LOGS: main_query: {query} | params: {params + [limit, offset]}")
            cur.execute(query, params + [limit, offset])
            logs = cur.fetchall()
            
            results = []
            for r in logs:
                row_dict = dict(r)
                for key, value in row_dict.items():
                    if isinstance(value, (datetime, ZoneInfo)):
                        row_dict[key] = value.isoformat()
                results.append(row_dict)
                
            response_data = {"logs": results, "total": total_records}
            cache_manager.set(cache_key, response_data, ttl_seconds=5)
            return response_data
    except Exception as e:
            import traceback
            print(f"ERROR in get_logs: {e}")
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=str(e))
    finally:
            if 'cur' in locals(): cur.close()
@app.get("/api/owner-counts")
def get_owner_counts(
    start_time: Optional[str] = None,
    end_time: Optional[str] = None,
    client_name: Optional[str] = None,
    db_type: Optional[str] = None,
    server_name: Optional[str] = None,
    log_type: List[str] = Query(None),
    severity: List[str] = Query(None),
    log_status: List[str] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    # Thread-safe TTL cache check
    log_type_key = ",".join(sorted(log_type or []))
    severity_key = ",".join(sorted(severity or []))
    log_status_key = ",".join(sorted(log_status or []))
    user_email = current_user.get("email", "")

    cache_key = f"owner-counts:{start_time}:{end_time}:{client_name}:{db_type}:{server_name}:{log_type_key}:{severity_key}:{log_status_key}:{user_email}"

    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            log_table = "db_monitoring_logs"
            if log_status and any(s in ['Resolved', 'Ignored', 'No action Required'] for s in log_status):
                log_table = "db_archived_logs"
            
            base_query = f"FROM {log_table} WHERE 1=1"
            params = []
            
            if not start_time and log_table == "db_monitoring_logs":
                start_time = (datetime.now(ZoneInfo("Asia/Kolkata")) - timedelta(hours=1)).strftime("%Y-%m-%d %H:%M:%S")

            if start_time:
                base_query += " AND log_time_ist >= %s"
                params.append(start_time)
            if end_time:
                base_query += " AND log_time_ist <= %s"
                params.append(end_time)
            if client_name:
                base_query += " AND client_name = %s"
                params.append(client_name)
            if db_type:
                base_query += " AND db_type ILIKE %s"
                params.append(db_type)
            if server_name:
                base_query += " AND server_name = %s"
                params.append(server_name)
            if log_type:
                base_query += " AND log_type IN %s"
                params.append(tuple(log_type) if isinstance(log_type, list) else (log_type,))
            if severity:
                base_query += " AND severity IN %s"
                params.append(tuple(severity) if isinstance(severity, list) else (severity,))
            if log_status:
                base_query += " AND LOWER(status) IN %s"
                params.append(tuple(s.lower() for s in log_status) if isinstance(log_status, list) else (log_status.lower(),))
            else:
                base_query += " AND (status IS NULL OR TRIM(status) = '' OR LOWER(status) NOT IN ('resolved', 'ignored', 'no action required'))"

            is_client = current_user.get("isClientUser")
            if is_client:
                allowed_filters = get_client_allowed_filters(current_user)
                if not allowed_filters:
                    return {"owner_counts": {}, "total_assigned": 0, "total_unassigned": 0}
                
                filter_clauses = []
                filter_params = []
                for f_tech, f_client, f_server in allowed_filters:
                    filter_clauses.append("(TRIM(LOWER(db_type)) = %s AND client_name = %s AND server_name = %s)")
                    filter_params.extend([f_tech.lower().strip(), f_client, f_server])
                
                base_query += " AND (" + " OR ".join(filter_clauses) + ")"
                params.extend(filter_params)
            else:
                allowed_techs = get_user_allowed_technologies(current_user)
                if allowed_techs is not None:
                    base_query += " AND TRIM(LOWER(db_type)) = ANY(%s)"
                    params.append([t.lower().strip() for t in allowed_techs])

            # Group by owner to get per-person counts
            # We use a subquery to group similar logs first (like get_logs does) to be consistent with the UI counts
            group_cols = "client_name, server_name, db_type, log_type, log_message, status, owner, client_visibility, ticket_status, next_action, severity, log_hash"
            
            owner_query = f"""
                SELECT owner, COUNT(*) as count 
                FROM (SELECT owner {base_query} GROUP BY {group_cols}) as sub
                GROUP BY owner
            """
            
            cur.execute(owner_query, params)
            rows = cur.fetchall()
            
            owner_counts = {}
            total_assigned = 0
            total_unassigned = 0
            
            for row in rows:
                owner = row['owner']
                count = row['count']
                if owner and owner.strip() and owner.lower() != 'none':
                    owner_counts[owner] = count
                    total_assigned += count
                else:
                    total_unassigned += count
            
            result = {
                "owner_counts": owner_counts,
                "total_assigned": total_assigned,
                "total_unassigned": total_unassigned,
                "total_all": total_assigned + total_unassigned
            }
            cache_manager.set(cache_key, result, ttl_seconds=5)
            return result
    except Exception as e:
        print(f"ERROR in get_owner_counts: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if 'cur' in locals(): cur.close()

class LogSummaryRequest(BaseModel):
    logs: List[str]
    filters: dict = {}

@app.post("/api/summarize")
async def summarize_logs(req: LogSummaryRequest, current_user: dict = Depends(get_current_user)):
    if not openai_client:
        raise HTTPException(status_code=500, detail="OpenAI client is not configured. Is OPENAI_API_KEY in .env?")
    
    if not req.logs:
        raise HTTPException(status_code=400, detail="No logs provided for summary.")
    
    try:
        start_time = datetime.now()
        print(f"AI Summary Start: Requested by {current_user.get('username', 'unknown')} ({len(req.logs)} logs)")
        
        capped_logs = req.logs[:100]
        log_text = "\n".join(capped_logs)
        
        filters = req.filters or {}
        start_time_val = filters.get('start', 'N/A')
        end_time_val = filters.get('end', 'N/A')
        
        if client_name := filters.get('client'):
            system_prompt = (
                "You are an elite, highly experienced Database Administrator and Systems Engineer AI. "
                "Your task is to analyze database capacity growth patterns, tablespace metrics, and "
                "server resource performance telemetry (CPU, Memory, Disk, IO). "
                f"Compile a stunning, professional, and detailed Capacity & Performance Diagnostics Report for client {client_name}."
            )
            
            user_prompt = f"""
Analyze the following telemetry and capacity details for client '{client_name}' over the period {start_time_val} to {end_time_val}.

You must output exactly a comprehensive point-by-point diagnostic report containing exactly between 15 to 20 lines in total (excluding headers/dividers). 
Make the report rich, deeply professional, and highly detailed. Cover all the following points:
1. Current overall database cluster sizes and daily database capacity growth trends (mention positive growth and no growth databases explicitly).
2. Table growth metrics, highlighting the heaviest tables and tablespace allocations.
3. CPU, Memory, Disk, and IO performance telemetry logs (if live/available in the data below. If these resource logs are not available or not open, explicitly mention that server hardware consoles are currently not online or data is pending, but do not make up hardware values).

You MUST format your output exactly as a single list of point-by-point bullets, with a header:
### 📈 Expert Growth & Resource Diagnostics Report

Ensure it is point-by-point, highly technical, actionable, and contains exactly 15 to 20 detailed bulleted points (lines). Do not include any introductory or concluding conversational filler.

--- TELEMETRY DATA ---
{log_text}
--- END TELEMETRY DATA ---
"""
        else:
            system_prompt = (
                "You are an expert Database Administrator AI. Your task is to analyze postgres, MSSQL, Mysql, "
                "Mongodb, RDS, and Windows Event logs securely. Provide a clear, professional diagnostic report "
                "in strict Markdown format. You must strictly decline processing any requests, logs, or queries "
                "that are off-topic or unrelated to database, software, or data engineering."
            )
            
            user_prompt = f"""
Analyze the following batch of logs concisely for the period: {start_time_val} to {end_time_val}.
Do not exceed 200 words and do not include any introductory or concluding conversational filler.

The log entries are prefixed with their occurrence counts (e.g., [Count: 5]). 
Analyze the specific time range, occurrence frequency, and the log messages themselves effectively to provide a detailed summary of the issue patterns.

If the logs contain no discernible actionable data, output exactly: "I was unable to analyze the Logs."
Don't hallucinate.
Format your response exactly with these Markdown headers:
### 📊 1. Overall Summary
(Provide a high-level, easy-to-read overview of what occurred)

### 🔍 2. Root Cause Analysis
(Explain the underlying issue clearly, using `code blocks` for any specific error codes or paths)

### ⚠️ 3. Severity Assessment
(Explicitly count and list Critical, Medium, and Low issues using clear bullet points)

### 💡 4. Action Recommendations
(Provide clear, actionable, numbered steps to resolve the issue or optimize the system)

--- LOG DATA ---
{log_text}
--- END LOG DATA ---
"""
        
        model_name = os.getenv("OPENAI_MODEL_NAME", "gpt-4o-mini")
        
        def get_logs_hash(logs: List[str]) -> str:
            import re
            normalized = []
            for log in logs:
                clean = log.strip()
                clean = re.sub(r'^\[Count:\s*\d+\]\s*', '', clean)
                normalized.append(clean.strip())
            text = "".join(sorted(normalized))
            return hashlib.sha256(text.encode()).hexdigest()

        logs_hash = get_logs_hash(req.logs)
        
        try:
            with get_db_connection() as conn:
                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                cur.execute("""
                    SELECT summary_text FROM ai_summary_history 
                    WHERE logs_hash = %s AND created_at > NOW() - INTERVAL '7 days'
                    ORDER BY created_at DESC LIMIT 1
                """, (logs_hash,))
                cached = cur.fetchone()
                cur.close()
            
            if cached:
                print(f"AI CACHE HIT: Returning existing summary for logs_hash {logs_hash[:10]}...")
                return {"summary": cached['summary_text'], "cached": True}
        except Exception as cache_err:
            print(f"Cache check failed (ignoring): {cache_err}")

        try:
            response = await openai_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_content if 'system_content' in locals() else system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_completion_tokens=2000
            )
        except Exception as api_err:
            if "model_not_found" in str(api_err).lower() or "not found" in str(api_err).lower():
                print(f"⚠️ Preferred model '{model_name}' not found. Falling back to gpt-4o-mini...")
                response = await openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_content if 'system_content' in locals() else system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    max_completion_tokens=2000
                )
            else:
                raise api_err

        summary_text = response.choices[0].message.content
        
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        print(f"AI Summary Complete: Generated in {duration:.2f}s")
        
        try:
            from psycopg2.extras import Json
            import json
            filters_json = req.filters or {}
            filters_json["username"] = current_user["username"]
            
            with get_db_connection() as conn:
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO ai_summary_history (summary_text, filters_json, username, logs_hash)
                    VALUES (%s, %s, %s, %s)
                """, (summary_text, json.dumps(filters_json), current_user["username"], logs_hash))
                conn.commit()
                cur.close()
        except Exception as db_err:
            print(f"Error saving to history: {db_err}")

        return {"summary": summary_text, "cached": False}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI API Error: {str(e)}")

@app.get("/api/history")
def get_summary_history(current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = "SELECT id, created_at, username, filters_json FROM ai_summary_history WHERE 1=1"
            params = []
            
            is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
            if not is_admin:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                if allowed_clients:
                    query += " AND (filters_json->>'client_name' = ANY(%s) OR username = %s)"
                    params.append(allowed_clients)
                    params.append(current_user["username"])
                elif allowed_techs is not None:
                    query += " AND (TRIM(LOWER(filters_json->>'technology')) = ANY(%s) OR username = %s)"
                    tech_params = [t.lower().strip() for t in allowed_techs]
                    params.append(tech_params)
                    params.append(current_user["username"])
                else:
                    query += " AND username = %s"
                    params.append(current_user["username"])
            
            query += " ORDER BY created_at DESC LIMIT 50"
            cur.execute(query, tuple(params))
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                
                if row.get('filters_json') is None:
                    row['filters_json'] = {}
                elif isinstance(row['filters_json'], str):
                    try:
                        row['filters_json'] = json.loads(row['filters_json'])
                    except Exception:
                        pass
                results.append(row)
                
            cur.close()
        return {"history": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
@app.get("/api/history/{history_id}")
def get_history_detail(history_id: int, current_user: dict = Depends(get_current_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SELECT summary_text, created_at, filters_json, username FROM ai_summary_history WHERE id = %s", (history_id,))
            result = cur.fetchone()
            cur.close()
        
        if not result:
            raise HTTPException(status_code=404, detail="History entry not found")
        
        row = dict(result)
        
        # Enforce strict multi-tenant boundary checks for standard users
        is_admin = current_user.get("role") == "admin" or current_user.get("isAdmin")
        if not is_admin:
            filters = row.get("filters_json", {})
            if isinstance(filters, str):
                try:
                    filters = json.loads(filters)
                except Exception:
                    filters = {}
            row_username = row.get("username", "")
            row_client = (filters or {}).get("client_name")
            row_tech = ((filters or {}).get("technology") or "").lower().strip()
            
            is_owner = row_username.lower() == current_user["username"].lower()
            has_permission = is_owner
            
            if not has_permission:
                allowed_clients = current_user.get("allowed_clients", [])
                allowed_techs = get_user_allowed_technologies(current_user)
                if allowed_clients:
                    if row_client in allowed_clients:
                        has_permission = True
                elif allowed_techs is not None:
                    if row_tech in [t.lower().strip() for t in allowed_techs]:
                        has_permission = True
            
            if not has_permission:
                raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this history entry.")
        
        if row.get('filters_json') is None:
            row['filters_json'] = {}
        elif isinstance(row['filters_json'], str):
            try:
                row['filters_json'] = json.loads(row['filters_json'])
            except Exception:
                pass
        
        if isinstance(row.get('created_at'), datetime):
             row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")

        return row
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class ExportRequest(BaseModel):
    title: str
    content: str
    client_name: Optional[str] = "N/A"
    server_name: Optional[str] = "N/A"
    db_type: Optional[str] = "N/A"
    severity: Optional[str] = "N/A"
    generated_on: Optional[str] = None

@app.post("/api/export/docx")
async def export_docx(request: ExportRequest, current_user: dict = Depends(get_current_user)):
    try:
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        
        title = doc.add_heading(request.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Client Name", request.client_name),
            ("Server Instance", request.server_name),
            ("Log Technology", request.db_type),
            ("Severity", request.severity),
            ("Generated On", request.generated_on or datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")),
            ("Generated By", current_user.get("full_name") or current_user.get("username", "System"))
        ]
        
        for label, value in metadata:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            row_cells[0].paragraphs[0].runs[0].bold = True
            
        doc.add_paragraph()
        
        import re
        
        lines = request.content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line).strip(), style='List Number')
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"GeoPITS AI Log Analyzer © {datetime.now().year} | Prepared By SANJAY G | Confidential Technical Report"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_client = re.sub(r'[^a-zA-Z0-9]', '_', request.client_name or "Log")
        filename = f"Report_{safe_client}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        print(f"Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def get_admin_user(current_user: dict = Depends(get_current_user)):
    if not current_user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin privileges required")
    return current_user

@app.get("/api/admin/users")
def get_admin_users(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SELECT id, username, email, full_name, profile_pic, role, last_active_at FROM users ORDER BY last_active_at DESC NULLS LAST")
            users = cur.fetchall()
            cur.close()
        
        results = []
        for r in users:
            row = dict(r)
            if row.get('last_active_at') and isinstance(row['last_active_at'], datetime):
                row['last_active_at'] = row['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"users": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class RoleUpdateRequest(BaseModel):
    email: Optional[str] = None
    username: Optional[str] = None
    role: str

@app.patch("/api/admin/users/role")
def update_user_role(req: RoleUpdateRequest, admin_user: dict = Depends(get_admin_user)):
    if req.role not in ['admin', 'user']:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    identifier = req.email or req.username
    if not identifier:
        raise HTTPException(status_code=400, detail="Email or username required")

    if (req.username == admin_user.get("username") or req.email == admin_user.get("email")) and req.role == 'user':
        raise HTTPException(status_code=400, detail="You cannot demote yourself from admin status")

    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            target_email = None
            if req.email:
                cur.execute("SELECT email FROM users WHERE LOWER(email) = LOWER(%s)", (req.email,))
            else:
                cur.execute("SELECT email FROM users WHERE username = %s", (req.username,))
            
            email_row = cur.fetchone()
            if email_row:
                target_email = (email_row[0] or "").lower()
            
            if req.email:
                cur.execute("UPDATE users SET role = %s WHERE LOWER(email) = LOWER(%s)", (req.role, req.email))
            else:
                cur.execute("UPDATE users SET role = %s WHERE username = %s", (req.role, req.username))
            
            if cur.rowcount == 0:
                cur.close()
                raise HTTPException(status_code=404, detail="User not found in system")
            
            if target_email:
                if req.role == 'admin':
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (target_email,))
                else:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (target_email,))
                    cur.execute("UPDATE leads SET status = 'removed' WHERE LOWER(email) = %s AND technology IN ('Global', 'Global Admin')", (target_email,))
            
            conn.commit()
            cur.close()
        return {"status": "success", "message": f"User role updated to {req.role}"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.get("/api/admin/summaries")
def get_admin_summaries(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("""
                SELECT ah.id, ah.summary_text, ah.created_at, ah.username, u.full_name, u.profile_pic
                FROM ai_summary_history ah
                LEFT JOIN users u ON ah.username = u.username
                ORDER BY ah.created_at DESC LIMIT 100
            """)
            history = cur.fetchall()
            
            results = []
            for r in history:
                row = dict(r)
                if isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
                
            cur.close()
        return {"history": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
class LeadCreateRequest(BaseModel):
    email: str
    technology: Optional[str] = None
    technologies: Optional[List[str]] = None
    is_lead: Optional[bool] = False

class ClientAccessRequest(BaseModel):
    client_email: str
    technology: str
    client_name: str
    server_name: str

@app.get("/api/admin/leads")
def get_admin_leads(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("""
                SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, u.role 
                FROM leads l
                LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                ORDER BY l.created_at DESC
            """)
            leads = cur.fetchall()
            cur.close()
        
        results = []
        for r in leads:
            row = dict(r)
            if isinstance(row['created_at'], datetime):
                row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"leads": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/leads")
def create_admin_lead(req: LeadCreateRequest, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            tech_list = []
            if req.technologies:
                tech_list.extend(req.technologies)
            if req.technology:
                tech_list.append(req.technology)
            
            tech_list = list(set(tech_list))
            
            if not tech_list:
                tech_list = ['Global']
            
            inserted_ids = []
            for tech in tech_list:
                cur.execute("""
                    INSERT INTO leads (email, technology, status, is_lead) 
                    VALUES (%s, %s, 'active', %s) 
                    ON CONFLICT (email, technology) DO UPDATE SET is_lead = EXCLUDED.is_lead, status = 'active'
                    RETURNING id
                """, (req.email, tech, req.is_lead))
                res = cur.fetchone()
                if res:
                    inserted_ids.append(res[0])
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (req.email.lower(),))
                    
            conn.commit()
            cur.close()
        return {"status": "success", "ids": inserted_ids}
    except Exception as e:
        print(f"Lead Creation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/api/admin/leads/{lead_id}/status")
def toggle_lead_status(lead_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT email, technology, status FROM leads WHERE id = %s", (lead_id,))
            row = cur.fetchone()
            if not row:
                cur.close()
                raise HTTPException(status_code=404, detail="Lead not found")
            
            email, tech, current_status = row
            new_status = 'removed' if current_status == 'active' else 'active'
            
            cur.execute("UPDATE leads SET status = %s WHERE id = %s", (new_status, lead_id))
            
            if tech in ['Global', 'Global Admin'] and email:
                cur.execute("""
                    INSERT INTO system_admins (email, status)
                    VALUES (%s, %s)
                    ON CONFLICT (email) DO UPDATE SET status = EXCLUDED.status
                """, (email.lower(), new_status))
                
            conn.commit()
            cur.close()
        return {"status": "updated", "new_status": new_status}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/admin/leads/{lead_id}")
def delete_admin_lead(lead_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (lead_id,))
            row = cur.fetchone()
            if row:
                email, tech = row
                cur.execute("DELETE FROM leads WHERE id = %s", (lead_id,))
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (email.lower(),))
            
            conn.commit()
            cur.close()
        return {"status": "deleted"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# Data Freshness Alerts & Monitored Clients
# ---------------------------------------------------------------------------
MONITORED_CLIENTS = [
    {"client_name": "360tf",     "server_name": None},
    {"client_name": "artfine",   "server_name": None},
    {"client_name": "credopay",  "server_name": "cpprodmysqldbsrv"},
]
FRESHNESS_HOURS = 8  # Alert threshold in hours

@app.get("/api/alerts/data-freshness")
def get_data_freshness_alerts(current_user: dict = Depends(get_current_user)):
    """
    Returns a list of monitored clients that have not received log data
    in the last FRESHNESS_HOURS hours. Results are filtered by the
    current user's allowed technologies.
    """
    allowed_techs = get_user_allowed_technologies(current_user)  # None = unrestricted (admin)
    is_client = current_user.get("isClientUser")

    alerts = []
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            now_ist = datetime.now(ZoneInfo("Asia/Kolkata"))

            for entry in MONITORED_CLIENTS:
                client_name = entry["client_name"]
                server_name = entry["server_name"]

                # Build query to get the most recent log time + technology
                if server_name:
                    cur.execute(
                        """
                        SELECT MAX(log_time_ist) AS last_log, db_type
                        FROM db_monitoring_logs
                        WHERE LOWER(client_name) = LOWER(%s)
                          AND LOWER(server_name) = LOWER(%s)
                        GROUP BY db_type
                        ORDER BY MAX(log_time_ist) DESC NULLS LAST
                        LIMIT 1
                        """,
                        (client_name, server_name)
                    )
                else:
                    cur.execute(
                        """
                        SELECT MAX(log_time_ist) AS last_log, db_type
                        FROM db_monitoring_logs
                        WHERE LOWER(client_name) = LOWER(%s)
                        GROUP BY db_type
                        ORDER BY MAX(log_time_ist) DESC NULLS LAST
                        LIMIT 1
                        """,
                        (client_name,)
                    )

                row = cur.fetchone()

                # Determine last log time and staleness
                if row and row["last_log"]:
                    last_log_dt = row["last_log"]
                    # Ensure timezone-aware
                    if last_log_dt.tzinfo is None:
                        last_log_dt = last_log_dt.replace(tzinfo=ZoneInfo("Asia/Kolkata"))
                    delta = now_ist - last_log_dt
                    hours_since = round(delta.total_seconds() / 3600, 1)
                    db_type = row["db_type"] or "Unknown"
                    is_stale = hours_since >= FRESHNESS_HOURS
                    last_log_str = last_log_dt.strftime("%Y-%m-%d %H:%M:%S")
                else:
                    hours_since = None
                    db_type = "Unknown"
                    is_stale = True
                    last_log_str = None

                if not is_stale:
                    continue

                # Skip if user doesn't have access to this technology
                if is_client:
                    continue  # Client users don't see freshness alerts
                if allowed_techs is not None and db_type not in allowed_techs:
                    continue

                # Determine the display server name (for "any" entries, find actual latest server)
                display_server = server_name
                if not display_server and row:
                    # Get the server name that had the most recent log
                    if server_name is None:
                        cur.execute(
                            """
                            SELECT server_name FROM db_monitoring_logs
                            WHERE LOWER(client_name) = LOWER(%s)
                            ORDER BY log_time_ist DESC NULLS LAST
                            LIMIT 1
                            """,
                            (client_name,)
                        )
                        srv_row = cur.fetchone()
                        display_server = srv_row["server_name"] if srv_row else "N/A"

                if hours_since is not None:
                    msg = f"No data loaded for '{client_name}' in the last {hours_since} hrs"
                else:
                    msg = f"No data has ever been loaded for '{client_name}'"

                alerts.append({
                    "client_name": client_name,
                    "server_name": display_server or entry.get("server_name") or "any",
                    "db_type": db_type,
                    "last_log_time": last_log_str,
                    "hours_since": hours_since,
                    "message": msg
                })

            cur.close()
    except Exception as e:
        import traceback
        print(f"ERROR in get_data_freshness_alerts: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

    return {"alerts": alerts, "count": len(alerts)}


@app.get("/api/notification-alerts")
def get_notification_alerts(current_user: dict = Depends(get_current_user)):
    """Check if Artfine, 360tf, or Credopay logs have not been uploaded for past 8 hours."""
    try:
        alerts = []
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            # Query the latest upload time for target clients
            query = """
                SELECT client_name, MAX(last_upload) as last_upload
                FROM (
                    SELECT client_name, MAX(created_at) as last_upload
                    FROM db_monitoring_logs
                    WHERE client_name IN ('Artfine', '360tf', 'Credopay')
                    GROUP BY client_name
                    UNION ALL
                    SELECT client_name, MAX(created_at) as last_upload
                    FROM db_archived_logs
                    WHERE client_name IN ('Artfine', '360tf', 'Credopay')
                    GROUP BY client_name
                ) as combined
                GROUP BY client_name
            """
            cur.execute(query)
            rows = cur.fetchall()
            
            # Map database results
            last_uploads = {row["client_name"]: row["last_upload"] for row in rows}
            
            for client in ["360tf", "Artfine", "Credopay"]:
                last_time = last_uploads.get(client)
                
                # If no log exists ever, we treat it as never uploaded
                if not last_time:
                    alerts.append({
                        "client": client,
                        "status": "warning",
                        "hours_since": None,
                        "message": f"Client {client} logs have never been uploaded to the database."
                    })
                    continue
                
                if last_time.tzinfo:
                    now = datetime.now(last_time.tzinfo)
                else:
                    now = datetime.now()
                diff_hours = (now - last_time).total_seconds() / 3600.0
                
                if diff_hours >= 8.0:
                    display_time = last_time
                    if hasattr(last_time, 'astimezone'):
                        display_time = last_time.astimezone(ZoneInfo("Asia/Kolkata"))
                    
                    time_str = display_time.strftime("%Y-%m-%d %H:%M:%S")
                    alerts.append({
                        "client": client,
                        "status": "warning",
                        "hours_since": round(diff_hours, 1),
                        "message": f"Client {client} logs have not been uploaded for the past {round(diff_hours, 1)} hours (Last upload: {time_str})."
                    })
            cur.close()
            
        return {"alerts": alerts}
    except Exception as e:
        print(f"Error fetching notification alerts: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/clients")
def get_admin_clients(admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SELECT id, client_email, technology, client_name, server_name, status, created_at FROM client_access ORDER BY created_at DESC")
            clients = cur.fetchall()
            cur.close()
        
        results = []
        for r in clients:
            row = dict(r)
            if isinstance(row['created_at'], datetime):
                row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"clients": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/admin/clients")
def create_admin_client(req: ClientAccessRequest, admin_user: dict = Depends(get_admin_user)):
    email = req.client_email.strip().lower()
    tech = req.technology.strip()
    name = req.client_name.strip()
    server = req.server_name.strip()
    
    if not email or not tech or not name or not server:
        raise HTTPException(status_code=400, detail="All client details (Email, Technology, Name, Server) are required.")
    
    try:
        user_created = False
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO client_access (client_email, technology, client_name, server_name, status) 
                VALUES (%s, %s, %s, %s, 'enabled') 
                ON CONFLICT (client_email, technology, client_name, server_name) DO UPDATE SET status = 'enabled'
                RETURNING id
            """, (email, tech, name, server))
            res = cur.fetchone()
            inserted_id = res[0] if res else None
            
            # Check/Pre-create user record in users table
            cur.execute("SELECT id, role FROM users WHERE LOWER(email) = LOWER(%s);", (email,))
            user_row = cur.fetchone()
            if user_row:
                # If they exist, ensure client role is updated if not admin
                existing_role = user_row[1]
                if existing_role not in ['admin', 'client']:
                    cur.execute("UPDATE users SET role = 'client' WHERE id = %s;", (user_row[0],))
            else:
                # Pre-create the client user securely
                username = email.split("@")[0]
                import bcrypt
                hashed_pwd = bcrypt.hashpw("geopits123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                cur.execute("""
                    INSERT INTO users (username, email, full_name, hashed_password, role)
                    VALUES (%s, %s, %s, %s, 'client')
                    ON CONFLICT (username) DO NOTHING;
                """, (username, email, username.capitalize(), hashed_pwd))
                user_created = True
                
            conn.commit()
            cur.close()
            
        audit_msg = f"Client Access Created: {name} ({email}) for tech {tech} on server {server} (status: enabled) by admin {admin_user.get('username')}"
        print(audit_msg)
        audit_logger.info(audit_msg)
        
        # Send Email
        try:
            from routes import send_email_outlook, build_gorgeous_html_email
            username = email.split("@")[0]
            greeting = f"Hello {username.capitalize()},"
            if user_created:
                subject = "[GeoMon Portal] Invitation"
                lead_text = "You have been invited to join the GeoMon Enterprise Observability Portal. A new account has been pre-created for you. Please log in using the temporary credentials below and update your password upon first entry."
                details = {
                    "Username": username,
                    "Temporary Password": "geopits123",
                    "Assigned Role": "CLIENT",
                    "Mapped Client": name,
                    "Database Technology": tech,
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Account Invitation"
            else:
                subject = f"[GeoMon Portal] Client Environment Assigned: {name}"
                lead_text = "An administrator has registered or updated your client access mapping. Please review your updated access configuration below."
                details = {
                    "Username": username,
                    "Client Name": name,
                    "Database Technology": tech,
                    "Server Assigned": server,
                    "Status": "ENABLED",
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Client Access Configured"
                
            body = build_gorgeous_html_email(
                title=title,
                greeting=greeting,
                lead_text=lead_text,
                details=details,
                action_url="http://localhost:8000/#/login",
                action_text="Access Observability Portal"
            )
            send_email_outlook(to_emails=email, cc_emails=None, subject=subject, body=body, exclude_dccagent=True)
            print(f"[CLIENT NOTIFICATION SENT] Emailed {email} about client mapping for {name}")
        except Exception as mail_err:
            print(f"[CLIENT NOTIFICATION ERROR] Failed to send email to {email}: {mail_err}")
            
        return {"status": "success", "id": inserted_id}
    except Exception as e:
        print(f"Client Access Creation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/api/admin/clients/{client_id}/status")
def toggle_client_status(client_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT client_name, client_email, status FROM client_access WHERE id = %s", (client_id,))
            row = cur.fetchone()
            if not row:
                cur.close()
                raise HTTPException(status_code=404, detail="Client not found")
            
            name, email, current_status = row
            new_status = 'disabled' if current_status == 'enabled' else 'enabled'
            
            cur.execute("UPDATE client_access SET status = %s WHERE id = %s", (new_status, client_id))
            conn.commit()
            cur.close()
            
        audit_msg = f"Client Access Status Updated: {name} ({email}) set to status '{new_status}' by admin {admin_user.get('username')}"
        print(audit_msg)
        audit_logger.info(audit_msg)
        
        return {"status": "updated", "new_status": new_status}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/admin/clients/{client_id}")
def delete_admin_client(client_id: int, admin_user: dict = Depends(get_admin_user)):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT client_name, client_email FROM client_access WHERE id = %s", (client_id,))
            row = cur.fetchone()
            if row:
                name, email = row
                cur.execute("DELETE FROM client_access WHERE id = %s", (client_id,))
                conn.commit()
                
                audit_msg = f"Client Access Deleted: {name} ({email}) by admin {admin_user.get('username')}"
                print(audit_msg)
                audit_logger.info(audit_msg)
                
            cur.close()
        return {"status": "deleted"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/lead/technologies")
def get_lead_technologies(current_user: dict = Depends(get_current_user)):
    """Returns technologies for which the current user is a lead."""
    try:
        email = current_user.get('email', '').lower()
        print(f"DEBUG_LEAD: Checking technologies for email: '{email}'")
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            techs = [row[0] for row in cur.fetchall()]
            print(f"DEBUG_LEAD: Found techs for '{email}': {techs}")
            cur.close()
        return {"technologies": techs}
    except Exception as e:
        print(f"DEBUG_LEAD_ERROR: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/lead/users")
def get_lead_users(current_user: dict = Depends(get_current_user)):
    """Returns users assigned to technologies for which the current user is a lead."""
    try:
        email = current_user.get('email', '').lower()
        print(f"DEBUG_LEAD_USERS: Fetching users for lead email: '{email}'")
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row['technology'] for row in cur.fetchall()]
            
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if is_global_admin:
                cur.execute("""
                    SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, 
                           u.full_name, u.profile_pic, u.last_active_at, u.username
                    FROM leads l
                    LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                    WHERE LOWER(l.email) != %s
                    ORDER BY u.last_active_at DESC NULLS LAST
                """, (email,))
            elif my_techs:
                cur.execute("""
                    SELECT l.id, l.email, l.technology, l.status, l.created_at, l.is_lead, 
                           u.full_name, u.profile_pic, u.last_active_at, u.username
                    FROM leads l
                    LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                    WHERE l.technology IN %s AND LOWER(l.email) != %s
                    ORDER BY u.last_active_at DESC NULLS LAST
                """, (tuple(my_techs), email))
            else:
                cur.close()
                return {"users": []}
            
            users = cur.fetchall()
            results = []
            for r in users:
                row = dict(r)
                if row.get('created_at') and isinstance(row['created_at'], datetime):
                    row['created_at'] = row['created_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                if row.get('last_active_at') and isinstance(row['last_active_at'], datetime):
                    row['last_active_at'] = row['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                results.append(row)
            
            cur.close()
        return {"users": results}
    except Exception as e:
        print(f"Error fetching lead users: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/lead/users")
def create_lead_user(req: LeadCreateRequest, current_user: dict = Depends(get_current_user)):
    """Allows a lead to assign a user to a technology they lead."""
    try:
        email = current_user.get("email")
        if not email:
             raise HTTPException(status_code=403, detail="Email required")
             
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            
            is_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            requested_techs = req.technologies if req.technologies else ([req.technology] if req.technology else [])
            if not is_admin:
                for tech in requested_techs:
                    if tech not in my_techs:
                        cur.close()
                        raise HTTPException(status_code=403, detail=f"You do not have lead privilege for {tech}")
        
        cur.execute("SELECT role, email FROM users WHERE email = %s", (req.email,))
        target_row = cur.fetchone()
        
        is_target_admin = False
        if target_row:
            target_role, target_email = target_row
            is_target_admin = (target_role == 'admin') or ((target_email or "").lower() in ADMIN_EMAILS)
        
        if is_target_admin:
             cur.close()
             conn.close()
             raise HTTPException(status_code=403, detail="you cannot assign admin as user")

        inserted_ids = []
        for tech in requested_techs:
            cur.execute("""
                INSERT INTO leads (email, technology, status, is_lead) 
                VALUES (%s, %s, 'active', FALSE) 
                ON CONFLICT (email, technology) DO UPDATE SET status = 'active'
                RETURNING id
            """, (req.email, tech))
            res = cur.fetchone()
            if res:
                inserted_id = res[0]
                inserted_ids.append(inserted_id)
                
                if tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, 'active')
                        ON CONFLICT (email) DO UPDATE SET status = 'active'
                    """, (req.email.lower(),))
                
        conn.commit()
        cur.close()
        conn.close()
        return {"status": "success", "ids": inserted_ids}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Lead User Creation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/api/lead/users/{assignment_id}/status")
def toggle_lead_user_status(assignment_id: int, current_user: dict = Depends(get_current_user)):
    """Allows a lead to toggle 'active'/'revoked' status for an assignment they manage."""
    try:
        email = current_user.get("email", "").lower()
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if not is_global_admin:
                cur.execute("""
                    SELECT l_target.technology 
                    FROM leads l_target
                    JOIN leads l_lead ON l_target.technology = l_lead.technology
                    WHERE l_target.id = %s AND LOWER(l_lead.email) = %s AND l_lead.is_lead = TRUE AND l_lead.status = 'active'
                """, (assignment_id, email))
                
                if not cur.fetchone():
                    cur.close()
                    conn.close()
                    raise HTTPException(status_code=403, detail="You do not have lead privileges for this technology or user.")
            
            cur.execute("""
                UPDATE leads 
                SET status = CASE WHEN status = 'active' THEN 'revoked' ELSE 'active' END 
                WHERE id = %s
                RETURNING status
            """, (assignment_id,))
            new_status = cur.fetchone()[0]
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (assignment_id,))
            sync_row = cur.fetchone()
            if sync_row:
                s_email, s_tech = sync_row
                if s_tech in ['Global', 'Global Admin']:
                    cur.execute("""
                        INSERT INTO system_admins (email, status)
                        VALUES (%s, %s)
                        ON CONFLICT (email) DO UPDATE SET status = EXCLUDED.status
                    """, (s_email.lower(), new_status))
            
            conn.commit()
            cur.close()
        return {"status": "success", "new_status": new_status}
    except HTTPException: raise
    except Exception as e:
        print(f"Error toggling lead user status: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/lead/users/{assignment_id}")
def delete_lead_user(assignment_id: int, current_user: dict = Depends(get_current_user)):
    """Allows a lead to delete a user assignment for a technology they lead."""
    try:
        email = current_user.get("email")
        if not email:
            raise HTTPException(status_code=403, detail="Email required")
            
        with get_db_connection() as conn:
            cur = conn.cursor()
            
            cur.execute("SELECT technology FROM leads WHERE id = %s", (assignment_id,))
            row = cur.fetchone()
            if not row:
                cur.close()
                raise HTTPException(status_code=404, detail="Assignment not found")
            tech = row[0]
            
            cur.execute("SELECT technology FROM leads WHERE LOWER(email) = %s AND is_lead = TRUE AND status = 'active'", (email,))
            my_techs = [row[0] for row in cur.fetchall()]
            is_global_admin = (current_user.get("role") == 'admin') or ('Global' in my_techs) or ('Global Admin' in my_techs)
            
            if not is_global_admin:
                if tech not in my_techs:
                    cur.close()
                    conn.close()
                    raise HTTPException(status_code=403, detail=f"You do not have lead privilege for {tech}")
            
            cur.execute("SELECT email, technology FROM leads WHERE id = %s", (assignment_id,))
            sync_row = cur.fetchone()
            if sync_row:
                s_email, s_tech = sync_row
                cur.execute("DELETE FROM leads WHERE id = %s", (assignment_id,))
                
                if s_tech in ['Global', 'Global Admin']:
                    cur.execute("UPDATE system_admins SET status = 'removed' WHERE LOWER(email) = %s", (s_email.lower(),))
            
            conn.commit()
            cur.close()
        return {"status": "deleted"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting lead user: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/lead-activity")
def get_lead_activity(current_user: dict = Depends(get_current_user)):
    """Monitor all user assignments overseen by leads (Admin only)."""
    if not current_user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin access required")
        
    try:
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            cur.execute("""
                SELECT l.id, l.email as user_email, l.technology, l.status, l.created_at,
                       u.full_name as user_name, u.profile_pic, u.last_active_at, u.username,
                       (
                           SELECT STRING_AGG(email, ', ') 
                           FROM leads 
                           WHERE technology = l.technology AND is_lead = TRUE AND status = 'active'
                       ) as lead_emails
                FROM leads l
                LEFT JOIN users u ON LOWER(l.email) = LOWER(u.email)
                WHERE l.is_lead = FALSE
                ORDER BY l.created_at DESC
            """)
            
            activity = []
            for row in cur.fetchall():
                res = dict(row)
                if res.get('created_at'):
                    res['created_at'] = res['created_at'].strftime("%Y-%m-%d %H:%M:%S")
                if res.get('last_active_at') and str(res.get('last_active_at')) != 'NaT':
                    try:
                        if hasattr(res['last_active_at'], 'astimezone'):
                            res['last_active_at'] = res['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
                        else:
                            res['last_active_at'] = res['last_active_at'].strftime("%Y-%m-%d %H:%M:%S")
                    except Exception:
                        res['last_active_at'] = str(res['last_active_at'])
                else:
                    res['last_active_at'] = 'Never'
                activity.append(res)
                
            cur.close()
        return {"activity": activity}
    except Exception as e:
        print(f"Error fetching lead activity: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ==============================================================================
# AUTOMATED DAILY TELEMETRY BACKGROUND SCHEDULER
# ==============================================================================
import threading
import time
from telemetry_parser import run_telemetry_sync
from utilization_sync import sync_utilization_history

def get_setting(key: str, default: str = None) -> str:
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT value FROM system_settings WHERE key = %s", (key,))
            row = cur.fetchone()
            cur.close()
            return row[0] if row else default
    except Exception as e:
        print(f"Error fetching setting {key}: {e}")
        return default

def set_setting(key: str, value: str):
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO system_settings (key, value)
                VALUES (%s, %s)
                ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value
            """, (key, str(value)))
            conn.commit()
            cur.close()
    except Exception as e:
        print(f"Error setting {key}: {e}")

# Shared status
telemetry_sync_in_progress = False

def run_manual_sync_in_background():
    global telemetry_sync_in_progress
    if telemetry_sync_in_progress:
        return
    telemetry_sync_in_progress = True
    set_setting("telemetry_scheduler_status", "Syncing...")
    try:
        now_str = datetime.now(ZoneInfo("Asia/Kolkata")).strftime('%Y-%m-%d %H:%M:%S IST')
        print(f"[SCHEDULER] ⏰ Manual trigger start at {now_str}")
        run_telemetry_sync()
        try:
            sync_utilization_history()
        except Exception as ex:
            print(f"[SCHEDULER] Server utilization sync error: {ex}")
        set_setting("telemetry_last_sync_time", now_str)
        set_setting("telemetry_last_sync_status", "Success")
    except Exception as e:
        print(f"[SCHEDULER] Manual sync failed: {e}")
        set_setting("telemetry_last_sync_status", f"Failed: {str(e)}")
    finally:
        telemetry_sync_in_progress = False
        set_setting("telemetry_scheduler_status", "Idle")

def telemetry_scheduler_daemon():
    print("[SCHEDULER] Automated daily database size telemetry scheduler daemon started.")
    set_setting("telemetry_scheduler_status", "Idle")
    while True:
        try:
            # Check scheduler trigger time settings
            trigger_hour = int(get_setting("telemetry_scheduler_hour", "14"))
            trigger_minute = int(get_setting("telemetry_scheduler_minute", "0"))
            
            now = datetime.now(ZoneInfo("Asia/Kolkata"))
            last_run_date_str = get_setting("telemetry_last_run_date", "")
            
            if now.hour == trigger_hour and now.minute == trigger_minute and last_run_date_str != str(now.date()):
                global telemetry_sync_in_progress
                if not telemetry_sync_in_progress:
                    telemetry_sync_in_progress = True
                    set_setting("telemetry_scheduler_status", "Syncing...")
                    try:
                        print(f"[SCHEDULER] ⏰ Triggering automated daily telemetry mail sync at {now.strftime('%Y-%m-%d %H:%M:%S IST')}...")
                        run_telemetry_sync()
                        try:
                            sync_utilization_history()
                        except Exception as ex:
                            print(f"[SCHEDULER] Server utilization sync error: {ex}")
                        
                        set_setting("telemetry_last_run_date", str(now.date()))
                        set_setting("telemetry_last_sync_time", now.strftime('%Y-%m-%d %H:%M:%S IST'))
                        set_setting("telemetry_last_sync_status", "Success")
                    except Exception as e:
                        print(f"[SCHEDULER] Ingestion Daemon run failed: {e}")
                        set_setting("telemetry_last_sync_status", f"Failed: {str(e)}")
                    finally:
                        telemetry_sync_in_progress = False
                        set_setting("telemetry_scheduler_status", "Idle")
        except Exception as e:
            print(f"[SCHEDULER] Daemon loop error: {e}")
        time.sleep(30)

# Start thread as a background daemon
t = threading.Thread(target=telemetry_scheduler_daemon, daemon=True)
t.start()

# ==============================================================================
# AUTOMATED DAILY CONSOLIDATED HIGH/CRITICAL ALERTS SCHEDULER
# ==============================================================================
def create_auto_ticket_if_missing(conn, log_id, client, db, server, l_type, msg, severity, log_time_ist, log_hash):
    # Use standard cursor to avoid dict key access issues with RealDictCursor
    cur = conn.cursor()
    try:
        # Check if a ticket with this log hash already exists and is not resolved/closed
        cur.execute("""
            SELECT id FROM tickets 
            WHERE description LIKE %s 
              AND UPPER(status) NOT IN ('RESOLVED', 'CLOSED') 
            LIMIT 1
        """, (f"%%Log Hash: {log_hash}%%",))
        row_ex_t = cur.fetchone()
        if row_ex_t and row_ex_t[0]:
            t_id = row_ex_t[0]
            # Link it to the log in db_monitoring_logs or db_archived_logs
            cur.execute("UPDATE db_monitoring_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
            cur.execute("UPDATE db_archived_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
            return t_id

        # Check if there's an active ticket with the same client, db, and containing the log message
        cur.execute("""
            SELECT id FROM tickets 
            WHERE LOWER(TRIM(company)) = LOWER(TRIM(%s)) 
              AND LOWER(TRIM(business_unit)) = LOWER(TRIM(%s)) 
              AND UPPER(status) NOT IN ('RESOLVED', 'CLOSED') 
              AND description LIKE %s 
            LIMIT 1
        """, (client, db, f"%%{msg}%%"))
        row_ex_msg = cur.fetchone()
        if row_ex_msg and row_ex_msg[0]:
            t_id = row_ex_msg[0]
            cur.execute("UPDATE db_monitoring_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
            cur.execute("UPDATE db_archived_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
            return t_id

        from db_manager import get_alert_contacts
        contacts = get_alert_contacts(cur, client, db)
        contact_emails = contacts["to_emails"]
        
        cur.execute("""
            INSERT INTO tickets (
                business_unit, company, contact, ticket_name, category, 
                status, priority, agent, description, created_by, created_at
            )
            VALUES (%s, %s, %s, %s, 'Logs', 'OPEN', %s, 'Unassigned', %s, 'System', %s)
            RETURNING id;
        """, (
            db, client, contact_emails, f"[AUTO] {severity} Alert: {l_type} on {server}",
            severity, f"Auto-generated ticket for log event.\nMessage: {msg}\nLog Hash: {log_hash}", log_time_ist
        ))
        t_id = cur.fetchone()[0]
        
        cur.execute("UPDATE db_monitoring_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
        cur.execute("UPDATE db_archived_logs SET ticket_id = %s, ticket_status = 'OPEN' WHERE id = %s", (t_id, log_id))
        
        return t_id
    finally:
        cur.close()

def run_daily_consolidated_alerts(conn, cur):
    import re
    from routes import lookup_email_routing_service, send_email_outlook

    cur.execute("SET timezone TO 'Asia/Kolkata';")
    print("[DAILY ALERTS] Fetching high and critical severity logs from the past 12 hours...")
    # Query from both active and archived logs
    cur.execute("""
        SELECT id, client_name, server_name, db_type, log_type, log_message, log_time_ist, severity, log_hash, ticket_id
        FROM db_monitoring_logs
        WHERE UPPER(TRIM(severity)) IN ('HIGH', 'CRITICAL')
          AND log_time_ist >= NOW() - INTERVAL '12 hours'
        UNION ALL
        SELECT id, client_name, server_name, db_type, log_type, log_message, log_time_ist, severity, log_hash, ticket_id
        FROM db_archived_logs
        WHERE UPPER(TRIM(severity)) IN ('HIGH', 'CRITICAL')
          AND log_time_ist >= NOW() - INTERVAL '12 hours'
    """)
    rows = cur.fetchall()
    print(f"[DAILY ALERTS] Found {len(rows)} total logs.")

    # Deduplicate based on client wise and timestamp
    seen_keys = set()
    deduplicated_rows = []
    for row in rows:
        client = row["client_name"]
        t_ist = row["log_time_ist"]
        # Format key as (client_name lower, timestamp string)
        t_str = t_ist.strftime("%Y-%m-%d %H:%M:%S") if hasattr(t_ist, "strftime") else str(t_ist)
        key = (client.strip().lower() if client else "", t_str)
        if key not in seen_keys:
            seen_keys.add(key)
            deduplicated_rows.append(row)
            
    print(f"[DAILY ALERTS] {len(deduplicated_rows)} logs remaining after deduplication by client and timestamp.")

    # Group by client
    client_groups = {}
    for row in deduplicated_rows:
        client = row["client_name"]
        if not client:
            continue
        client = client.strip()
        if client not in client_groups:
            client_groups[client] = []
        client_groups[client].append(row)

    for client, logs in client_groups.items():
        print(f"[DAILY ALERTS] Processing {len(logs)} logs for client: {client}")
        ticket_details = []
        unique_db_types = set()

        for log in logs:
            log_id = log["id"]
            client_name = log["client_name"]
            server_name = log["server_name"]
            db_type = log["db_type"]
            log_type = log["log_type"]
            log_message = log["log_message"]
            severity = log["severity"]
            log_hash = log["log_hash"]
            ticket_id = log["ticket_id"]

            if db_type:
                unique_db_types.add(db_type.strip())

            # Automatically create ticket if missing
            if not ticket_id:
                try:
                    ticket_id = create_auto_ticket_if_missing(
                        conn, log_id, client_name, db_type, server_name,
                        log_type, log_message, severity, log["log_time_ist"], log_hash
                    )
                    conn.commit()
                except Exception as t_err:
                    print(f"[DAILY ALERTS] Error creating ticket for log {log_id}: {t_err}")
                    conn.rollback()

            ticket_details.append({
                "ticket_id": ticket_id,
                "db_type": db_type,
                "server_name": server_name,
                "severity": severity,
                "timestamp": log["log_time_ist"].strftime("%Y-%m-%d %H:%M:%S") if hasattr(log["log_time_ist"], "strftime") else str(log["log_time_ist"]),
                "message": log_message
            })

        if not ticket_details:
            continue

        # Resolve TO and CC emails
        from db_manager import get_alert_contacts
        to_emails_set = set()

        for db_type in unique_db_types:
            try:
                resolved = get_alert_contacts(cur, client, db_type)
                to_e = resolved["to_emails"]
                if to_e:
                    for email in re.split(r'[;,]', to_e):
                        if email.strip(): to_emails_set.add(email.strip())
            except Exception as e_err:
                print(f"[DAILY ALERTS] Email routing lookup failed for {client} / {db_type}: {e_err}")

        to_emails = ", ".join(to_emails_set)
        cc_emails = ""

        if not to_emails:
            to_emails = "dccagent@geopits.com"

        current_hour = datetime.now(ZoneInfo('Asia/Kolkata')).hour
        slot_name = "8 AM" if 4 <= current_hour < 16 else "8 PM"
        subject = f"[GeoMon] {slot_name} Consolidated Incident Report — {client} | {datetime.now(ZoneInfo('Asia/Kolkata')).strftime('%d %b %Y')}"

        critical_count = sum(1 for t in ticket_details if str(t.get("severity","")).upper() == "CRITICAL")
        high_count = sum(1 for t in ticket_details if str(t.get("severity","")).upper() == "HIGH")

        html_body = f"""
<html>
<body style="font-family: Calibri, Arial, sans-serif; background-color: #f1f5f9; color: #1e293b; margin: 0; padding: 20px;">
  <div style="max-width: 720px; margin: 0 auto; background: #ffffff; border-radius: 10px; overflow: hidden; border: 1px solid #e2e8f0; box-shadow: 0 4px 16px rgba(0,0,0,0.08);">
    <!-- Header -->
    <div style="background: linear-gradient(135deg, #1e3a5f 0%, #0f2d56 100%); padding: 28px 32px;">
      <h2 style="margin: 0; color: #ffffff; font-size: 22px; font-weight: 700;">📊 {slot_name} Consolidated Incident Report</h2>
      <p style="margin: 6px 0 0 0; color: rgba(255,255,255,0.8); font-size: 13px;">Client: <strong style="color:#fff;">{client}</strong> &nbsp;|&nbsp; {datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%d %B %Y, %H:%M IST")}</p>
    </div>
    <!-- Greeting -->
    <div style="padding: 28px 32px 0 32px;">
      <p style="margin: 0 0 14px 0; font-size: 15px; color: #334155;">Hi Team,</p>
      <p style="margin: 0 0 20px 0; font-size: 14px; color: #475569; line-height: 1.75;">
        Please find below the <strong>{slot_name} consolidated incident report</strong> for <strong>{client}</strong>.
        Our monitoring system detected <strong>{len(ticket_details)}</strong> high or critical severity
        database event(s) in the past 12 hours — including <strong style="color:#dc2626;">{critical_count} Critical</strong>
        and <strong style="color:#ea580c;">{high_count} High</strong> priority issue(s).
        Tickets have been automatically raised for each incident. Kindly review and coordinate with the DBA team for resolution.
      </p>
      <!-- Summary badges -->
      <div style="display: flex; gap: 12px; margin-bottom: 24px;">
        <span style="background:#fef2f2; color:#dc2626; border:1px solid #fecaca; padding: 6px 16px; border-radius: 6px; font-size: 13px; font-weight: 700;">🔴 Critical: {critical_count}</span>
        <span style="background:#fff7ed; color:#ea580c; border:1px solid #fed7aa; padding: 6px 16px; border-radius: 6px; font-size: 13px; font-weight: 700;">🟠 High: {high_count}</span>
        <span style="background:#f0f9ff; color:#0369a1; border:1px solid #bae6fd; padding: 6px 16px; border-radius: 6px; font-size: 13px; font-weight: 700;">📋 Total: {len(ticket_details)}</span>
      </div>
    </div>
    <!-- Table -->
    <div style="padding: 0 32px 28px 32px;">
      <table style="border-collapse: collapse; width: 100%; font-size: 13px;">
        <thead>
          <tr style="background: #f8fafc;">
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600; white-space: nowrap;">Ticket #</th>
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600;">Severity</th>
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600;">DB Type</th>
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600;">Server</th>
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600; white-space: nowrap;">Time (IST)</th>
            <th style="border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; color: #64748b; font-weight: 600;">Incident Summary</th>
          </tr>
        </thead>
        <tbody>
"""
        for t in ticket_details:
            sev = str(t.get("severity","")).upper()
            if sev == "CRITICAL":
                badge_style = "background:#fef2f2; color:#dc2626; border:1px solid #fecaca;"
                badge_icon = "🔴"
            else:
                badge_style = "background:#fff7ed; color:#ea580c; border:1px solid #fed7aa;"
                badge_icon = "🟠"
            tid = t.get("ticket_id")
            ticket_ref = f"<strong>#{tid}</strong>" if tid else "<em>Pending</em>"
            html_body += f"""
          <tr>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px; color: #0f172a; white-space: nowrap;">{ticket_ref}</td>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px; white-space: nowrap;"><span style="{badge_style} padding: 3px 8px; border-radius: 4px; font-size: 11px; font-weight: 700;">{badge_icon} {t.get("severity","")}</span></td>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px;">{t.get("db_type","")}</td>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px; font-family: monospace; font-size: 12px;">{t.get("server_name","")}</td>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px; white-space: nowrap; color: #64748b; font-size: 12px;">{t.get("timestamp","")}</td>
            <td style="border: 1px solid #e2e8f0; padding: 10px 12px; color: #334155; font-size: 12px;">{str(t.get("message",""))[:200]}</td>
          </tr>
"""

        html_body += f"""
        </tbody>
      </table>
      <p style="margin: 24px 0 0 0; font-size: 13px; color: #64748b; line-height: 1.75;">
        Our DBA team is actively monitoring and working on the above incidents.
        You will receive individual closure notifications once each ticket is resolved.
        For urgent escalations, please reply to this email or contact the support team directly.
      </p>
      <p style="margin: 20px 0 0 0; font-size: 14px; color: #334155;">Regards,<br><strong>GeoMon DBA Support Team</strong><br><span style="color:#94a3b8; font-size: 12px;">Automated Daily Report | Geopits Technologies</span></p>
    </div>
    <!-- Footer -->
    <div style="background: #f8fafc; border-top: 1px solid #e2e8f0; padding: 14px 32px; text-align: center;">
      <p style="margin: 0; font-size: 11px; color: #94a3b8;">This is an automated daily report generated by the GeoMon Obsv Dashboard. Please do not reply directly to this system address.</p>
    </div>
  </div>
</body>
</html>"""

        try:
            print(f"[DAILY ALERTS] Sending daily alert for client {client} to TO: {to_emails}, CC: {cc_emails}")
            send_email_outlook(to_emails, cc_emails, subject, html_body)
            print(f"[DAILY ALERTS] Sent successfully for {client}")
        except Exception as send_err:
            print(f"[DAILY ALERTS] Send email failed for client {client}: {send_err}")

def daily_alert_scheduler_daemon():
    print("[DAILY ALERT SCHEDULER] Automated daily consolidated high/critical severity alerts scheduler daemon started.")
    while True:
        try:
            now = datetime.now(ZoneInfo("Asia/Kolkata"))
            
            # Check for the two fixed runs: 8 AM and 8 PM
            if now.minute == 0:
                if now.hour in [8, 20]:
                    slot = now.strftime("%Y-%m-%d") + f"_{now.hour:02d}PM" if now.hour == 20 else now.strftime("%Y-%m-%d") + f"_{now.hour:02d}AM"
                    last_run_slot = get_setting("daily_alert_last_run_slot", "")
                    if last_run_slot != slot:
                        print(f"[DAILY ALERT SCHEDULER] ⏰ Triggering automated consolidated high/critical severity alerts slot {slot} at {now.strftime('%Y-%m-%d %H:%M:%S IST')}...")
                        try:
                            with get_db_connection() as conn:
                                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                run_daily_consolidated_alerts(conn, cur)
                                conn.commit()
                                cur.close()
                            set_setting("daily_alert_last_run_slot", slot)
                            set_setting("daily_alert_last_sync_time", now.strftime('%Y-%m-%d %H:%M:%S IST'))
                            set_setting("daily_alert_last_sync_status", "Success")
                        except Exception as e:
                            print(f"[DAILY ALERT SCHEDULER] consolidated alerts slot {slot} job failed: {e}")
                            set_setting("daily_alert_last_sync_status", f"Failed: {str(e)}")
            
            # Keep the configurable scheduling capability as a fallback / extra check
            trigger_hour = int(get_setting("daily_alert_scheduler_hour", "8"))
            trigger_minute = int(get_setting("daily_alert_scheduler_minute", "0"))
            last_run_date_str = get_setting("daily_alert_last_run_date", "")

            if now.hour == trigger_hour and now.minute == trigger_minute and last_run_date_str != str(now.date()):
                # Only run if not already covered by the slot runs to avoid sending duplicate emails
                slot = now.strftime("%Y-%m-%d") + f"_{now.hour:02d}"
                last_run_slot = get_setting("daily_alert_last_run_slot", "")
                if last_run_slot != slot:
                    print(f"[DAILY ALERT SCHEDULER] ⏰ Triggering configured daily consolidated high/critical severity alerts at {now.strftime('%Y-%m-%d %H:%M:%S IST')}...")
                    try:
                        with get_db_connection() as conn:
                            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                            run_daily_consolidated_alerts(conn, cur)
                            conn.commit()
                            cur.close()
                        set_setting("daily_alert_last_run_date", str(now.date()))
                        set_setting("daily_alert_last_run_slot", slot)
                        set_setting("daily_alert_last_sync_time", now.strftime('%Y-%m-%d %H:%M:%S IST'))
                        set_setting("daily_alert_last_sync_status", "Success")
                    except Exception as e:
                        print(f"[DAILY ALERT SCHEDULER] daily consolidated alerts job failed: {e}")
                        set_setting("daily_alert_last_sync_status", f"Failed: {str(e)}")
        except Exception as e:
            print(f"[DAILY ALERT SCHEDULER] Daemon loop error: {e}")
        time.sleep(30)

# Start daily consolidated alert thread as a background daemon
t_alert = threading.Thread(target=daily_alert_scheduler_daemon, daemon=True)
t_alert.start()

# API Endpoints for Scheduler Management
@app.get("/api/admin/scheduler/status")
def get_scheduler_status(admin_user: dict = Depends(get_admin_user)):
    try:
        hour = get_setting("telemetry_scheduler_hour", "14")
        minute = get_setting("telemetry_scheduler_minute", "0")
        status = get_setting("telemetry_scheduler_status", "Idle")
        last_sync = get_setting("telemetry_last_sync_time", "Never")
        last_status = get_setting("telemetry_last_sync_status", "N/A")
        
        return {
            "trigger_hour": int(hour),
            "trigger_minute": int(minute),
            "status": "Syncing..." if telemetry_sync_in_progress else status,
            "last_sync_time": last_sync,
            "last_sync_status": last_status,
            "sync_in_progress": telemetry_sync_in_progress
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class SchedulerSettingsRequest(BaseModel):
    trigger_hour: int
    trigger_minute: int

@app.post("/api/admin/scheduler/settings")
def update_scheduler_settings(req: SchedulerSettingsRequest, admin_user: dict = Depends(get_admin_user)):
    if not (0 <= req.trigger_hour <= 23) or not (0 <= req.trigger_minute <= 59):
        raise HTTPException(status_code=400, detail="Invalid hour (0-23) or minute (0-59)")
    try:
        set_setting("telemetry_scheduler_hour", str(req.trigger_hour))
        set_setting("telemetry_scheduler_minute", str(req.trigger_minute))
        return {"status": "success", "message": "Scheduler trigger settings updated successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/scheduler/trigger")
def trigger_scheduler_sync(admin_user: dict = Depends(get_admin_user)):
    global telemetry_sync_in_progress
    if telemetry_sync_in_progress:
        raise HTTPException(status_code=409, detail="A telemetry sync operation is already in progress.")
    try:
        threading.Thread(target=run_manual_sync_in_background, daemon=True).start()
        return {"status": "success", "message": "Telemetry sync initiated successfully in the background."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/scheduler/daily-alerts/status")
def get_daily_alerts_scheduler_status(admin_user: dict = Depends(get_admin_user)):
    try:
        hour = get_setting("daily_alert_scheduler_hour", "8")
        minute = get_setting("daily_alert_scheduler_minute", "0")
        last_sync = get_setting("daily_alert_last_sync_time", "Never")
        last_status = get_setting("daily_alert_last_sync_status", "N/A")
        
        return {
            "trigger_hour": int(hour),
            "trigger_minute": int(minute),
            "last_sync_time": last_sync,
            "last_sync_status": last_status
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class DailyAlertSettingsRequest(BaseModel):
    trigger_hour: int
    trigger_minute: int

@app.post("/api/admin/scheduler/daily-alerts/settings")
def update_daily_alerts_settings(req: DailyAlertSettingsRequest, admin_user: dict = Depends(get_admin_user)):
    if not (0 <= req.trigger_hour <= 23) or not (0 <= req.trigger_minute <= 59):
        raise HTTPException(status_code=400, detail="Invalid hour (0-23) or minute (0-59)")
    try:
        set_setting("daily_alert_scheduler_hour", str(req.trigger_hour))
        set_setting("daily_alert_scheduler_minute", str(req.trigger_minute))
        return {"status": "success", "message": "Daily alert scheduler trigger settings updated successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/scheduler/daily-alerts/trigger")
def trigger_daily_alerts(admin_user: dict = Depends(get_admin_user)):
    try:
        def manual_run():
            try:
                with get_db_connection() as conn:
                    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    run_daily_consolidated_alerts(conn, cur)
                    conn.commit()
                    cur.close()
                now_str = datetime.now(ZoneInfo("Asia/Kolkata")).strftime('%Y-%m-%d %H:%M:%S IST')
                set_setting("daily_alert_last_sync_time", now_str)
                set_setting("daily_alert_last_sync_status", "Success")
            except Exception as e:
                print(f"[DAILY ALERT SCHEDULER] manual daily alerts execution failed: {e}")
                set_setting("daily_alert_last_sync_status", f"Failed: {str(e)}")
        
        threading.Thread(target=manual_run, daemon=True).start()
        return {"status": "success", "message": "Daily consolidated alerts job triggered successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





# ── OBSERVABILITY ENDPOINTS ──────────────────────────────────────────────────

def get_observability_base_query(
    current_user: dict,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = None,
    status: Optional[List[str]] = None,
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = None,
    include_backlog: bool = False
):
    """
    Builds a combined subquery from BOTH tables.
    - db_monitoring_logs  → active/pending statuses (Open, Under Review, etc.)
    - db_archived_logs    → terminal statuses (Resolved, Ignored, No Action Required)
 
    Uses combined_logs_mv materialized view if available, otherwise falls back
    to a direct UNION ALL.
    """
    # Use live UNION ALL to ensure recent shift logs are always instantly visible
    base_subquery = f"""(
        SELECT * FROM (
            SELECT id, client_name, server_name, db_type, log_type, log_time_ist, log_message,
                   status, owner, severity, occurrence_count, status_updated_at, created_at,
                   log_hash, is_semantic, semantic_count, semantic_hash, time_bucket,
                   false AS is_archived
            FROM db_monitoring_logs
            UNION ALL
            SELECT id, client_name, server_name, db_type, log_type, log_time_ist, log_message,
                   status, owner, severity, occurrence_count, status_updated_at, created_at,
                   log_hash, is_semantic, semantic_count, semantic_hash, time_bucket,
                   true AS is_archived
            FROM db_archived_logs
        ) AS combined_logs_mv
        WHERE (
            CASE 
                WHEN time_bucket LIKE '%%%%_08AM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08AM', ' 08:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                WHEN time_bucket LIKE '%%%%_08PM' THEN TO_TIMESTAMP(REPLACE(time_bucket, '_08PM', ' 20:00:00'), 'YYYY-MM-DD HH24:MI:SS')::TIMESTAMP
                ELSE log_time_ist::TIMESTAMP 
            END <= (NOW() AT TIME ZONE 'Asia/Kolkata')::TIMESTAMP
        )
    ) AS bucket_filtered"""
 
    where_clauses = ["1=1"]
    params = []
 
    if start_date and end_date:
        if include_backlog:
            # User defined a specific time window (e.g. a 12h shift).
            # We show logs matching that window OR logs created BEFORE that are still NOT terminal (resolved/ignored).
            where_clauses.append("""
                (
                    (log_time_ist >= %s::TIMESTAMP AND log_time_ist <= %s::TIMESTAMP)
                    OR 
                    (
                        log_time_ist < %s::TIMESTAMP 
                        AND (
                            status IS NULL 
                            OR TRIM(status) = '' 
                            OR LOWER(TRIM(status)) = 'none' 
                            OR LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required')
                        )
                    )
                )
            """)
            params.extend([start_date, end_date, start_date])
        else:
            # Only show logs matching the specific shift for Trends/Heatmaps etc.
            where_clauses.append("log_time_ist >= %s::TIMESTAMP AND log_time_ist <= %s::TIMESTAMP")
            params.extend([start_date, end_date])
    elif start_date:
        where_clauses.append("log_time_ist >= %s::TIMESTAMP")
        params.append(start_date)
    elif end_date:
        where_clauses.append("log_time_ist <= %s::TIMESTAMP")
        params.append(end_date)
    
    if client and client != "All Clients":
        where_clauses.append("client_name = %s")
        params.append(client)
    if server and server != "All Servers":
        where_clauses.append("server_name = %s")
        params.append(server)
    if technology and technology != "All Technologies":
        where_clauses.append("db_type ILIKE %s")
        params.append(technology)
    if severity and "All Severities" not in severity:
        where_clauses.append("severity = ANY(%s)")
        params.append(severity)
    if status and "All Statuses" not in status:
        where_clauses.append("status = ANY(%s)")
        params.append(status)
    if owner and owner not in ("None", "All Owners"):
        if owner.lower() == "unassigned":
            where_clauses.append(
                "(owner IS NULL OR owner = '' OR LOWER(owner) IN ('unassigned','none'))"
            )
        else:
            where_clauses.append("owner ILIKE %s")
            params.append(f"%{owner}%")
    if log_type and "All Log Types" not in log_type:
        where_clauses.append("log_type = ANY(%s)")
        params.append(log_type)
 
    # ── Permission filters ───────────────────────────────────────────────────
    if current_user.get("isClientUser"):
        allowed_filters = get_client_allowed_filters(current_user)
        if not allowed_filters:
            where_clauses.append("1=0")
        else:
            client_clauses = []
            for f_tech, f_client, f_server in allowed_filters:
                client_clauses.append(
                    "(TRIM(LOWER(db_type)) = %s AND client_name = %s AND server_name = %s)"
                )
                params.extend([f_tech.lower().strip(), f_client, f_server])
            where_clauses.append("(" + " OR ".join(client_clauses) + ")")
    else:
        allowed_techs = get_user_allowed_technologies(current_user)
        if allowed_techs is not None:
            if not allowed_techs:
                where_clauses.append("1=0")
            else:
                where_clauses.append("TRIM(LOWER(db_type)) = ANY(%s)")
                params.append([t.lower().strip() for t in allowed_techs])
 
    where_str = " AND ".join(where_clauses)
    return base_subquery, where_str, params

# Global cache for observability analytics (TTL: 60s)
OBSERVABILITY_CACHE = {}
OBSERVABILITY_CACHE_TTL = 60

def serialize_datetime(dt):
    if isinstance(dt, datetime):
        return dt.astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
    return dt

def observability_cache(func):
    """Decorator to cache observability responses based on user and filter params."""
    from functools import wraps
    @wraps(func)
    def wrapper(*args, **kwargs):
        # Create a unique cache key from function name and all arguments
        cache_parts = [func.__name__]
        for k, v in sorted(kwargs.items()):
            if k == 'current_user':
                user_id = v.get('email') or v.get('username') or 'anon'
                techs = get_user_allowed_technologies(v)
                cache_parts.append(f"user:{user_id}:{tuple(sorted(techs)) if techs else 'all'}")
            elif isinstance(v, list):
                cache_parts.append(f"{k}:{tuple(v)}")
            else:
                cache_parts.append(f"{k}:{v}")
        
        cache_key = tuple(cache_parts)
        now = datetime.now(ZoneInfo("Asia/Kolkata"))
        cached = OBSERVABILITY_CACHE.get(cache_key)
        if cached and (now - cached["last_updated"]).total_seconds() < OBSERVABILITY_CACHE_TTL:
            return cached["data"]
            
        result = func(*args, **kwargs)
        OBSERVABILITY_CACHE[cache_key] = {"data": result, "last_updated": now}
        return result
    return wrapper

@app.get("/api/observability/overview")
@observability_cache
def get_observability_overview(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    include_backlog: bool = False,
    current_user: dict = Depends(get_current_user)
):
    """
    Returns KPI counts. When include_backlog=False (the default), only logs
    within the selected shift time window are counted, giving accurate
    shift-wise metrics. Set include_backlog=True to include historical unresolved
    logs from before the selected time window.
    """
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server,
            technology, severity, status, owner, log_type,
            include_backlog=include_backlog
        )
 
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
 
            # ── Single query — group by status + is_archived ─────────────────
            query = f"""
                SELECT
                    COALESCE(NULLIF(TRIM(status), ''), 'None') AS status_label,
                    is_archived,
                    COUNT(*)                   AS unique_count,
                    SUM(occurrence_count)      AS total_occurrences,
                    AVG(
                        CASE
                            WHEN LOWER(TRIM(status)) IN ('resolved','ignored','no action required')
                             AND status_updated_at IS NOT NULL
                             AND created_at IS NOT NULL
                            THEN EXTRACT(EPOCH FROM (status_updated_at - created_at))
                        END
                    ) AS mttr_secs
                FROM {subquery}
                WHERE {where}
                GROUP BY COALESCE(NULLIF(TRIM(status), ''), 'None'), is_archived
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
 
        # ── Aggregate in Python ───────────────────────────────────────────────
        status_counts   = {}   # total per label (both tables)
        severity_counts = {}   # populated by /severity endpoint — not here
        active_counts   = {}   # monitoring table statuses only
        archived_counts = {}   # archived table statuses only
 
        total_logs         = 0
        total_unique       = 0
        total_active_logs  = 0
        total_archived_logs= 0
        mttr_sum           = 0.0
        mttr_count         = 0
        resolved_count     = 0
 
        for r in rows:
            label   = r["status_label"]
            is_arc  = r["is_archived"]
            u_cnt   = int(r["unique_count"]      or 0)
            t_cnt   = int(r["total_occurrences"] or 0)
            mttr_s  = r["mttr_secs"]
 
            status_counts[label] = status_counts.get(label, 0) + u_cnt
            total_unique         += u_cnt
            total_logs           += t_cnt
 
            if is_arc:
                archived_counts[label]  = archived_counts.get(label, 0) + u_cnt
                total_archived_logs     += t_cnt
            else:
                active_counts[label]    = active_counts.get(label, 0) + u_cnt
                total_active_logs       += t_cnt
 
            if mttr_s is not None:
                mttr_sum   += float(mttr_s) * u_cnt
                mttr_count += u_cnt
 
            if label.lower() in ('resolved', 'ignored', 'no action required'):
                resolved_count += u_cnt
 
        dedup_ratio = round(total_logs / total_unique, 1) if total_unique > 0 else 1.0
        mttr_hours  = round((mttr_sum / mttr_count) / 3600.0, 1) if mttr_count > 0 else 0.0
        res_rate    = round((resolved_count / total_unique) * 100.0, 1) if total_unique > 0 else 0.0
 
        return {
            # ── KPI card values (frontend reads these directly) ──────────────
            "status_counts": status_counts,
            "active_counts": active_counts,
            "archived_counts": archived_counts,
 
            # ── Aggregates ───────────────────────────────────────────────────
            "total_logs":           total_logs,
            "unique_issues":        total_unique,
            "total_active_unique":  sum(active_counts.values()),
            "total_active_logs":    total_active_logs,
            "total_archived_unique":sum(archived_counts.values()),
            "total_archived_logs":  total_archived_logs,
            "dedup_ratio":          dedup_ratio,
            "mttr_hours":           mttr_hours,
            "resolution_rate":      res_rate,
        }
 
    except Exception as e:
        import traceback
        print(f"Error in overview: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/observability/status")
@observability_cache
def get_observability_status(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""SELECT COALESCE(NULLIF(TRIM(status), ''), 'None') as status, COUNT(*) as count FROM {subquery}
                WHERE {where}
                GROUP BY COALESCE(NULLIF(TRIM(status), ''), 'None')"""
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
        return {"status_distribution": {r["status"]: r["count"] for r in rows}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/severity")
@observability_cache
def get_observability_severity(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"SELECT COALESCE(severity, 'Unknown') as severity, COUNT(*) as count FROM {subquery} WHERE {where} GROUP BY COALESCE(severity, 'Unknown')"
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
        return {"severity_distribution": {r["severity"]: r["count"] for r in rows}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/client-health")
@observability_cache
def get_observability_client_health(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""
                SELECT 
                    client_name,
                    COUNT(*) as unique_issues,
                    SUM(occurrence_count) as total_logs,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN 1 ELSE 0 END) as critical_count,
                    SUM(CASE WHEN severity ILIKE 'high' THEN 1 ELSE 0 END) as high_count,
                    SUM(CASE WHEN severity ILIKE 'medium' THEN 1 ELSE 0 END) as medium_count,
                    SUM(CASE WHEN severity ILIKE 'low' THEN 1 ELSE 0 END) as low_count,
                    SUM(CASE WHEN LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required') OR status IS NULL OR TRIM(status) = '' THEN 1 ELSE 0 END) as open_count,
                    SUM(CASE WHEN LOWER(TRIM(status)) IN ('action needed from client', 'action needed from dba') THEN 1 ELSE 0 END) as pending_count,
                    SUM(CASE WHEN (LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required') OR status IS NULL OR TRIM(status) = '') AND log_time_ist < NOW() - INTERVAL '24 hours' THEN 1 ELSE 0 END) as aging_count
                FROM {subquery}
                WHERE {where} AND client_name IS NOT NULL
                GROUP BY client_name
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            crit = r["critical_count"] or 0
            high = r["high_count"] or 0
            med = r["medium_count"] or 0
            low = r["low_count"] or 0
            op = r["open_count"] or 0
            pend = r["pending_count"] or 0
            age = r["aging_count"] or 0
            
            # Risk score calculation
            risk_score = (crit * 10) + (high * 6) + (op * 3) + (pend * 4) + (age * 5)
            
            results.append({
                "client_name": r["client_name"],
                "unique_issues": r["unique_issues"],
                "total_logs": r["total_logs"],
                "critical": crit,
                "high": high,
                "medium": med,
                "low": low,
                "risk_score": risk_score
            })
            
        results.sort(key=lambda x: x["risk_score"], reverse=True)
        return {"client_health": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/server-health")
@observability_cache
def get_observability_server_health(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""
                SELECT 
                    server_name,
                    COUNT(*) as unique_issues,
                    SUM(occurrence_count) as total_logs,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN 1 ELSE 0 END) as critical_count,
                    SUM(CASE WHEN severity ILIKE 'high' THEN 1 ELSE 0 END) as high_count,
                    SUM(CASE WHEN severity ILIKE 'medium' THEN 1 ELSE 0 END) as medium_count,
                    SUM(CASE WHEN severity ILIKE 'low' THEN 1 ELSE 0 END) as low_count,
                    SUM(CASE WHEN LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required') OR status IS NULL OR TRIM(status) = '' THEN 1 ELSE 0 END) as open_count,
                    SUM(CASE WHEN LOWER(TRIM(status)) IN ('action needed from client', 'action needed from dba') THEN 1 ELSE 0 END) as pending_count,
                    SUM(CASE WHEN (LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required') OR status IS NULL OR TRIM(status) = '') AND log_time_ist < NOW() - INTERVAL '24 hours' THEN 1 ELSE 0 END) as aging_count
                FROM {subquery}
                WHERE {where} AND server_name IS NOT NULL
                GROUP BY server_name
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            crit = r["critical_count"] or 0
            high = r["high_count"] or 0
            med = r["medium_count"] or 0
            low = r["low_count"] or 0
            op = r["open_count"] or 0
            pend = r["pending_count"] or 0
            age = r["aging_count"] or 0
            
            risk_score = (crit * 10) + (high * 6) + (op * 3) + (pend * 4) + (age * 5)
            
            results.append({
                "server_name": r["server_name"],
                "unique_issues": r["unique_issues"],
                "total_logs": r["total_logs"],
                "critical": crit,
                "high": high,
                "medium": med,
                "low": low,
                "risk_score": risk_score
            })
            
        results.sort(key=lambda x: x["risk_score"], reverse=True)
        return {"server_health": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/client-heatmap")
@observability_cache
def get_observability_client_heatmap(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""
                SELECT client_name, COALESCE(NULLIF(TRIM(status), ''), 'None') as status, COUNT(*) as count 
                FROM {subquery} 
                WHERE {where} AND client_name IS NOT NULL
                GROUP BY client_name, COALESCE(NULLIF(TRIM(status), ''), 'None')
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
        return {"client_heatmap": [dict(r) for r in rows]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/server-heatmap")
@observability_cache
def get_observability_server_heatmap(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""
                SELECT server_name, COALESCE(NULLIF(TRIM(status), ''), 'None') as status, COUNT(*) as count 
                FROM {subquery} 
                WHERE {where} AND server_name IS NOT NULL
                GROUP BY server_name, COALESCE(NULLIF(TRIM(status), ''), 'None')
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
        return {"server_heatmap": [dict(r) for r in rows]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/workflow")
@observability_cache
def get_observability_workflow(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            query = f"""
                SELECT COALESCE(NULLIF(TRIM(status), ''), 'None') as status, COUNT(*) as count 
                FROM {subquery} 
                WHERE {where} 
                GROUP BY COALESCE(NULLIF(TRIM(status), ''), 'None')
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
        return {"workflow": {r["status"]: r["count"] for r in rows}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/bottlenecks")
@observability_cache
def get_observability_bottlenecks(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            # Fetch pending statuses
            query = f"""
                SELECT 
                    COALESCE(NULLIF(TRIM(status), ''), 'None') as status,
                    COUNT(*) as count,
                    AVG(EXTRACT(EPOCH FROM (NOW() - COALESCE(status_updated_at, created_at)))) / 3600.0 as avg_wait_hours
                FROM {subquery}
                WHERE {where} AND LOWER(TRIM(status)) IN ('action needed from client', 'action needed from dba', 'under review', 'monitoring')
                GROUP BY COALESCE(NULLIF(TRIM(status), ''), 'None')
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            
            # Total active issue count to compute percentages
            total_active_query = f"""
                SELECT COUNT(*) as count 
                FROM {subquery} 
                WHERE {where} AND LOWER(status) NOT IN ('resolved', 'ignored', 'no action required')
            """
            cur.execute(total_active_query, params)
            total_active = cur.fetchone()["count"] or 0
            
            cur.close()
            
        results = []
        for r in rows:
            stat = r["status"]
            cnt = r["count"] or 0
            pct = round((cnt / total_active * 100.0), 1) if total_active > 0 else 0.0
            avg_wait = round(r["avg_wait_hours"] or 0.0, 1)
            
            results.append({
                "status": stat,
                "count": cnt,
                "percentage": pct,
                "avg_wait_hours": avg_wait
            })
            
        return {"bottlenecks": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/aging")
@observability_cache
def get_observability_aging(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            query = f"""
                SELECT
                    COUNT(CASE WHEN (LOWER(TRIM(status)) = 'open' OR status IS NULL OR TRIM(status) = '') AND log_time_ist < NOW() - INTERVAL '24 hours' THEN 1 END) as open_24h,
                    COUNT(CASE WHEN (LOWER(TRIM(status)) = 'open' OR status IS NULL OR TRIM(status) = '') AND log_time_ist < NOW() - INTERVAL '48 hours' THEN 1 END) as open_48h,
                    COUNT(CASE WHEN (LOWER(TRIM(status)) = 'open' OR status IS NULL OR TRIM(status) = '') AND log_time_ist < NOW() - INTERVAL '7 days' THEN 1 END) as open_7d,
                    COUNT(CASE WHEN LOWER(TRIM(status)) = 'under review' AND log_time_ist < NOW() - INTERVAL '3 days' THEN 1 END) as review_3d,
                    COUNT(CASE WHEN LOWER(TRIM(status)) = 'action needed from client' AND log_time_ist < NOW() - INTERVAL '7 days' THEN 1 END) as client_7d,
                    COUNT(CASE WHEN LOWER(TRIM(status)) = 'action needed from dba' AND log_time_ist < NOW() - INTERVAL '7 days' THEN 1 END) as dba_7d,
                    COUNT(CASE WHEN LOWER(TRIM(status)) = 'monitoring' AND log_time_ist < NOW() - INTERVAL '14 days' THEN 1 END) as monitoring_14d
                FROM {subquery}
                WHERE {where}
            """
            cur.execute(query, params)
            res = cur.fetchone()
            cur.close()
        return {"aging": dict(res) if res else {}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/trends")
@observability_cache
def get_observability_trends(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            query = f"""
                SELECT 
                    log_time_ist::date as log_date,
                    SUM(CASE WHEN LOWER(TRIM(status)) NOT IN ('resolved', 'ignored', 'no action required') OR status IS NULL OR TRIM(status) = '' THEN occurrence_count ELSE 0 END) as active_count,
                    SUM(CASE WHEN LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') THEN occurrence_count ELSE 0 END) as resolved_count,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN occurrence_count ELSE 0 END) as critical_count,
                    COUNT(DISTINCT log_hash) as unique_issues,
                    SUM(occurrence_count) as total_logs
                FROM {subquery}
                WHERE {where} AND log_time_ist IS NOT NULL
                GROUP BY log_time_ist::date
                ORDER BY log_time_ist::date ASC
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            results.append({
                "date": str(r["log_date"]),
                "active_logs": int(r["active_count"] or 0),
                "resolved_logs": int(r["resolved_count"] or 0),
                "critical_logs": int(r["critical_count"] or 0),
                "unique_issues": int(r["unique_issues"] or 0),
                "total_logs": int(r["total_logs"] or 0)
            })
        return {"trends": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/owners")
@observability_cache
def get_observability_owners(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            query = f"""
                SELECT 
                    owner,
                    SUM(occurrence_count) as assigned_logs,
                    SUM(CASE WHEN LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') THEN occurrence_count ELSE 0 END) as resolved_logs,
                    AVG(CASE WHEN LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') AND status_updated_at IS NOT NULL AND created_at IS NOT NULL THEN EXTRACT(EPOCH FROM (status_updated_at - created_at)) END) / 3600.0 as avg_resolution_hours,
                    SUM(CASE WHEN severity ILIKE 'critical' AND LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') THEN occurrence_count ELSE 0 END) as critical_resolved
                FROM {subquery}
                WHERE {where} AND owner IS NOT NULL AND owner != '' AND owner != 'None' AND owner != 'Unassigned'
                GROUP BY owner
                ORDER BY resolved_logs DESC
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            ass = int(r["assigned_logs"] or 0)
            res = int(r["resolved_logs"] or 0)
            rate = round((res / ass * 100.0), 1) if ass > 0 else 0.0
            avg_res = round(r["avg_resolution_hours"] or 0.0, 1)
            
            results.append({
                "owner": r["owner"],
                "assigned_logs": ass,
                "resolved_logs": res,
                "resolution_rate": rate,
                "avg_resolution_hours": avg_res,
                "critical_resolved": int(r["critical_resolved"] or 0)
            })
        return {"owners": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/critical")
@observability_cache
def get_observability_critical(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            query = f"""
                SELECT 
                    client_name,
                    server_name,
                    severity,
                    log_type,
                    COALESCE(NULLIF(TRIM(status), ''), 'None') as status,
                    SUM(occurrence_count) as occurrences,
                    owner,
                    MIN(log_time_ist) as first_seen,
                    MAX(log_time_ist) as last_seen
                FROM {subquery}
                WHERE {where} AND (severity ILIKE 'critical' OR severity ILIKE 'high')
                GROUP BY client_name, server_name, severity, log_type, COALESCE(NULLIF(TRIM(status), ''), 'None'), owner
                ORDER BY 
                    CASE WHEN severity ILIKE 'critical' THEN 2 WHEN severity ILIKE 'high' THEN 1 ELSE 0 END DESC,
                    occurrences DESC
                LIMIT 50
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            row = dict(r)
            row["first_seen"] = serialize_datetime(row["first_seen"])
            row["last_seen"] = serialize_datetime(row["last_seen"])
            results.append(row)
        return {"critical_issues": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/observability/owners-list")
@observability_cache
def get_observability_owners_list(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Lightweight endpoint — returns only distinct assigned owner names for the filter dropdown."""
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology
        )
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute(f"""
                SELECT DISTINCT owner
                FROM {subquery}
                WHERE {where}
                  AND owner IS NOT NULL
                  AND TRIM(owner) != ''
                  AND LOWER(owner) NOT IN ('none', 'unassigned')
                ORDER BY owner
            """, params)
            owners = [row[0] for row in cur.fetchall()]
            cur.close()
        return {"owners": owners}
    except Exception as e:
        print(f"Error in owners-list: {e}")
        raise HTTPException(status_code=500, detail=str(e))
        
@app.get("/api/observability/recurring")
@observability_cache
def get_observability_recurring(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            query = f"""
                SELECT 
                    log_message,
                    client_name,
                    server_name,
                    COALESCE(severity, 'Unknown') as severity,
                    SUM(occurrence_count) as occurrences,
                    MIN(log_time_ist) as first_seen,
                    MAX(log_time_ist) as last_seen
                FROM {subquery}
                WHERE {where}
                GROUP BY log_message, client_name, server_name, COALESCE(severity, 'Unknown')
                ORDER BY occurrences DESC
                LIMIT 50
            """
            cur.execute(query, params)
            rows = cur.fetchall()
            cur.close()
            
        results = []
        for r in rows:
            row = dict(r)
            row["first_seen"] = serialize_datetime(row["first_seen"])
            row["last_seen"] = serialize_datetime(row["last_seen"])
            results.append(row)
        return {"recurring_issues": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/mttr")
@observability_cache
def get_observability_mttr(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type
        )
        
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            # Helper to run MTTR breakdown query
            def run_mttr_breakdown(group_col):
                q = f"""
                    SELECT {group_col} as name, AVG(EXTRACT(EPOCH FROM (status_updated_at - created_at))) / 3600.0 as mttr
                    FROM {subquery}
                    WHERE {where} AND LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') 
                      AND status_updated_at IS NOT NULL AND created_at IS NOT NULL AND {group_col} IS NOT NULL
                    GROUP BY {group_col}
                    ORDER BY mttr ASC
                """
                cur.execute(q, params)
                return [dict(r) for r in cur.fetchall()]
                
            # Overall MTTR
            overall_q = f"""
                SELECT AVG(EXTRACT(EPOCH FROM (status_updated_at - created_at))) / 3600.0 as overall_mttr
                FROM {subquery}
                WHERE {where} AND LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required')
                  AND status_updated_at IS NOT NULL AND created_at IS NOT NULL
            """
            cur.execute(overall_q, params)
            overall_res = cur.fetchone()
            overall_mttr = round(overall_res["overall_mttr"] or 0.0, 1) if overall_res else 0.0
            
            client_mttr = run_mttr_breakdown("client_name")
            server_mttr = run_mttr_breakdown("server_name")
            tech_mttr = run_mttr_breakdown("db_type")
            owner_mttr = run_mttr_breakdown("owner")
            
            cur.close()
            
        return {
            "overall_mttr": overall_mttr,
            "client_mttr": [{ "name": c["name"], "mttr": round(c["mttr"] or 0.0, 1) } for c in client_mttr],
            "server_mttr": [{ "name": s["name"], "mttr": round(s["mttr"] or 0.0, 1) } for s in server_mttr],
            "technology_mttr": [{ "name": t["name"], "mttr": round(t["mttr"] or 0.0, 1) } for t in tech_mttr],
            "owner_mttr": [{ "name": o["name"], "mttr": round(o["mttr"] or 0.0, 1) } for o in owner_mttr]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/observability/executive-summary")
@observability_cache
def get_observability_executive_summary(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    client: Optional[str] = None,
    server: Optional[str] = None,
    technology: Optional[str] = None,
    severity: Optional[List[str]] = Query(None),
    status: Optional[List[str]] = Query(None),
    owner: Optional[str] = None,
    log_type: Optional[List[str]] = Query(None),
    include_backlog: bool = False,
    current_user: dict = Depends(get_current_user)
):
    try:
        subquery, where, params = get_observability_base_query(
            current_user, start_date, end_date, client, server, technology, severity, status, owner, log_type,
            include_backlog=include_backlog
        )
        with get_db_connection() as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cur.execute("SET timezone TO 'Asia/Kolkata';")
            
            # Fetch summary stats
            sum_query = f"""
                SELECT 
                    COUNT(*) as total_issues,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN 1 ELSE 0 END) as critical_issues,
                    SUM(CASE WHEN LOWER(TRIM(status)) = 'action needed from client' THEN 1 ELSE 0 END) as pending_client,
                    SUM(CASE WHEN LOWER(TRIM(status)) = 'action needed from dba' THEN 1 ELSE 0 END) as pending_dba,
                    AVG(CASE WHEN LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') AND status_updated_at IS NOT NULL AND created_at IS NOT NULL THEN EXTRACT(EPOCH FROM (status_updated_at - created_at)) END) / 3600.0 as mttr,
                    SUM(CASE WHEN LOWER(TRIM(status)) IN ('resolved', 'ignored', 'no action required') THEN 1 ELSE 0 END) as resolved_issues
                FROM {subquery}
                WHERE {where}
            """
            cur.execute(sum_query, params)
            stats = cur.fetchone()
            
            # Highest risk client (by risk score)
            client_q = f"""
                SELECT 
                    client_name,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN 10 WHEN severity ILIKE 'high' THEN 6 ELSE 2 END) as risk
                FROM {subquery}
                WHERE {where} AND client_name IS NOT NULL
                GROUP BY client_name
                ORDER BY risk DESC
                LIMIT 1
            """
            cur.execute(client_q, params)
            top_client = cur.fetchone()
            
            # Highest risk server
            server_q = f"""
                SELECT 
                    server_name,
                    SUM(CASE WHEN severity ILIKE 'critical' THEN 10 WHEN severity ILIKE 'high' THEN 6 ELSE 2 END) as risk
                FROM {subquery}
                WHERE {where} AND server_name IS NOT NULL
                GROUP BY server_name
                ORDER BY risk DESC
                LIMIT 1
            """
            cur.execute(server_q, params)
            top_server = cur.fetchone()
            
            cur.close()
            
        total = stats["total_issues"] or 0
        resolved = stats["resolved_issues"] or 0
        res_rate = round((resolved / total * 100.0), 1) if total > 0 else 0.0
        
        return {
            "total_issues": total,
            "critical_issues": stats["critical_issues"] or 0,
            "pending_client_actions": stats["pending_client"] or 0,
            "pending_dba_actions": stats["pending_dba"] or 0,
            "mttr_hours": round(stats["mttr"] or 0.0, 1),
            "resolution_rate": res_rate,
            "highest_risk_client": top_client["client_name"] if top_client else "N/A",
            "highest_risk_server": top_server["server_name"] if top_server else "N/A"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/observability/clear-cache")
def clear_observability_cache(current_user: dict = Depends(get_current_user)):
    """Clear the observability analytics cache to force fresh data on next request."""
    OBSERVABILITY_CACHE.clear()
    return {"status": "ok", "message": "Cache cleared"}

@app.post("/api/refresh-analytics")
def trigger_analytics_refresh(current_user: dict = Depends(get_current_user)):
    """Manually triggers a refresh of the materialized view and clears filter cache."""
    global FILTER_CACHE
    try:
        # 1. Refresh Materialized View
        refresh_combined_logs_mv()
        
        # 2. Clear Filter Cache
        FILTER_CACHE.clear()
        
        return {"status": "success", "message": "Analytics metrics and filters refreshed."}
    except Exception as e:
        print(f"Error in manual refresh: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.patch("/api/logs/bulk-archive")
def bulk_archive_logs(
    action: LogBulkAction,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    if current_user.get("isClientUser") or current_user.get("clientAccessDisabled"):
        raise HTTPException(status_code=403, detail="Access Restricted")

    if not action.log_hashes:
        return {"message": "No logs selected", "count": 0}

    terminal_statuses = {'Resolved', 'Ignored', 'No action Required'}
    is_terminal = action.status in terminal_statuses
    now_ist = datetime.now(ZoneInfo("Asia/Kolkata"))

    common_cols = """
        client_name, server_name, db_type, log_type, log_source,
        log_time, log_time_utc, log_time_ist, log_level, log_message,
        occurrence_count, raw_log, email_subject, email_received_time,
        log_hash, created_at, status, owner, client_visibility,
        ticket_status, next_action, severity, status_updated_at,
        is_semantic, semantic_count, semantic_hash, time_bucket, ticket_id
    """

    try:
        with get_db_connection() as conn:
            with conn.cursor() as cur:
                # 1. Handle Move (if terminal)
                if is_terminal:
                    print(f"[BULK ARCHIVE] Moving {len(action.log_hashes)} logs to archive...")
                    cur.execute(f"""
                        WITH moved_rows AS (
                            DELETE FROM db_monitoring_logs
                            WHERE log_hash = ANY(%s)
                            RETURNING {common_cols}
                        )
                        INSERT INTO db_archived_logs ({common_cols})
                        SELECT {common_cols} FROM moved_rows
                        ON CONFLICT (log_hash) DO NOTHING;
                    """, (action.log_hashes,))
                
                # 2. Update Status/Metadata in BOTH tables (to be sure)
                for table in ["db_monitoring_logs", "db_archived_logs"]:
                    cur.execute(f"""
                        UPDATE {table}
                        SET status = %s, owner = %s, client_visibility = %s,
                            ticket_status = %s, next_action = %s, 
                            severity = COALESCE(NULLIF(%s, ''), severity),
                            status_updated_at = %s
                        WHERE log_hash = ANY(%s)
                    """, (
                        action.status, action.owner, action.client_visibility,
                        action.ticket_status, action.next_action, action.severity,
                        now_ist, action.log_hashes
                    ))
                    
                    for log_hash in action.log_hashes:
                        ensure_ticket_for_high_critical_log(cur, table, log_hash, action.severity)
            conn.commit()
            
        # Refresh MV once at the end
        background_tasks.add_task(refresh_combined_logs_mv)
        return {"message": "Bulk action completed", "count": len(action.log_hashes)}

    except Exception as e:
        import traceback
        print(f"[FATAL BULK ERROR] {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

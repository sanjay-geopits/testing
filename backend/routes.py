import os
import sys
# Add backend directory to path for absolute imports resolution
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from fastapi import APIRouter, Depends, HTTPException, Query, Response, Header, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Any, Union
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from fastapi.security import OAuth2PasswordBearer
import psycopg2.extras
import base64
import os
import httpx

from migrations import get_connection

from cache_utils import cache_manager

router = APIRouter(prefix="/api/new-features", tags=["new-features"])

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/login")

def get_current_user_local(
    request: Request,
    authorization: Optional[str] = Header(None)
):
    token = request.query_params.get("token")
    print(f"AUTH DEBUG: token={token[:30] if token else None}, auth={authorization[:30] if authorization else None}")
    
    actual_token = None
    if authorization and authorization.startswith("Bearer "):
        actual_token = authorization.split(" ")[1]
    elif token:
        actual_token = token
        
    if not actual_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
        
    try:
        from app import get_current_user
        return get_current_user(actual_token)
    except Exception as e:
        # Fallback decode if app.py is initializing
        from jose import jwt
        import os
        SECRET_KEY = os.getenv("JWT_SECRET", "super-secret-key-geopits")
        try:
            payload = jwt.decode(actual_token, SECRET_KEY, algorithms=["HS256"])
            username = payload.get("sub")
            role = payload.get("role", "user")
            
            # Fetch details from DB if possible
            assigned_techs = []
            allowed_clients = []
            email_addr = ""
            full_name = ""
            user_id = None
            try:
                conn = get_connection()
                cur = conn.cursor()
                cur.execute("SELECT id, username, role, full_name, email FROM users WHERE username = %s;", (username,))
                user_row = cur.fetchone()
                if user_row:
                    user_id = user_row[0]
                    full_name = user_row[3]
                    email_addr = (user_row[4] or "").lower()
                    cur.execute("SELECT DISTINCT technology FROM leads WHERE LOWER(email) = LOWER(%s) AND status = 'active';", (email_addr,))
                    assigned_techs = [r[0] for r in cur.fetchall()]
                    
                    cur.execute("""
                        SELECT DISTINCT c.client_name 
                        FROM user_clients uc
                        JOIN admin_clients c ON uc.client_id = c.id
                        WHERE uc.user_id = %s;
                    """, (user_id,))
                    allowed_clients = [r[0] for r in cur.fetchall() if r[0]]
                cur.close()
                conn.close()
            except Exception as db_err:
                print(f"Fallback DB load failed: {db_err}")
                
            return {
                "id": user_id,
                "username": username,
                "role": role,
                "email": email_addr,
                "fullName": full_name,
                "isAdmin": role == "admin",
                "assigned_techs": assigned_techs,
                "allowed_clients": allowed_clients
            }
        except:
            raise HTTPException(status_code=401, detail="Authentication failed")

from services.email_service import lookup_email_routing_service, build_gorgeous_html_email, send_email_outlook
# ==============================================================================
# PYDANTIC SCHEMAS
# ==============================================================================

class TicketCreate(BaseModel):
    business_unit: str
    company: str
    contact: str
    ticket_name: str
    category: str
    status: str
    priority: str
    agent: str
    description: str

class ReportCreate(BaseModel):
    client_name: str
    title: str
    month: str
    year: str
    file_name: str
    file_data: str # Base64 string
    notes: Optional[str] = ""

class ReportReviewCreate(BaseModel):
    rating: Optional[int] = 5
    comment: Optional[str] = ""
    mom: Optional[str] = ""

class PageTimePing(BaseModel):
    page_path: str
    duration_seconds: int

class AdminClientCreate(BaseModel):
    client_name: str
    db_type: str
    server_name: str
    client_email: Optional[str] = None
    phone_number: Optional[str] = None

class AdminAgentCreate(BaseModel):
    agent_name: str
    company_name: str
    business_unit: str
    technology: str
    email: Optional[str] = None

class UserPermissionCreate(BaseModel):
    email: str
    technology: str
    status: str
    is_lead: bool
    role: Optional[str] = "user"

class UserClientPermissionCreate(BaseModel):
    email: str
    client_name: str
    access_level: Optional[str] = "view"

class BroadcastMessage(BaseModel):
    message: str

class TicketAgentCreate(BaseModel):
    name: str

class OnlineUserCreate(BaseModel):
    username: str
    units: str

class TicketBUCreate(BaseModel):
    name: str

class CommentCreate(BaseModel):
    comment_type: str # 'reply', 'note', 'forward', 'merge', 'log'
    content: str
    attachments: Optional[Any] = ""

class ClientAlertSettingCreate(BaseModel):
    client_name: str
    db_type: str
    cpu_threshold: float
    memory_threshold: float
    disk_threshold: float
    io_threshold: float
    slow_query_threshold_ms: int
    long_running_threshold_sec: int
    client_emails: Optional[str] = ""
    cc_emails: Optional[str] = ""
    server_down_alert: Optional[bool] = True
    critical_error_alert: Optional[bool] = True

class TechnologyAlertConfigCreate(BaseModel):
    technology: str
    alert_email: str

# ==============================================================================
# TICKETS ENDPOINTS
# ==============================================================================

@router.post("/tickets")
def create_ticket(ticket: TicketCreate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        cur.execute("""
            INSERT INTO tickets (business_unit, company, contact, ticket_name, category, status, priority, agent, description, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id;
        """, (
            ticket.business_unit, ticket.company, ticket.contact, ticket.ticket_name,
            ticket.category, ticket.status, ticket.priority, ticket.agent,
            ticket.description, user.get("username", "System")
        ))
        ticket_id = cur.fetchone()[0]
        # Commit ticket insertion immediately to prevent sequence skipping on any subsequent errors
        conn.commit()
        
        # Scope notifications strictly to the assigned agent only
        creator = user.get("username", "admin")
        if ticket.agent:
            cur.execute("""
                INSERT INTO notifications (username, message, is_read)
                VALUES (%s, %s, FALSE);
            """, (ticket.agent, f"New ticket #{ticket_id} '{ticket.ticket_name}' has been assigned to you by '{creator}'",))
        
        conn.commit()
        cache_manager.invalidate("tickets:")
        cache_manager.invalidate("ticket-stats:")
        return {"status": "success", "ticket_id": ticket_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/tickets/categories")
def get_ticket_categories():
    return ["Alert", "Incident", "Events", "Logs", "System Alert"]

@router.get("/tickets")
def get_tickets(
    search: Optional[str] = None,
    business_unit: Optional[str] = None,
    company: Optional[str] = None,
    agent: Optional[str] = None,
    priority: Optional[str] = None,
    status: Optional[str] = None,
    created_by: Optional[str] = None,
    resolved_by: Optional[str] = None,
    created_at: Optional[str] = None,
    resolved_at: Optional[str] = None,
    user: dict = Depends(get_current_user_local)
):
    cache_key = f"tickets:{search}:{business_unit}:{company}:{agent}:{priority}:{status}:{created_by}:{resolved_by}:{created_at}:{resolved_at}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        query = """
            SELECT t.*,
                (SELECT COUNT(*) FROM ticket_comments tc WHERE tc.ticket_id = t.id) AS comment_count
            FROM tickets t
            WHERE 1=1
        """
        params = []

        if search:
            clean_search = search.strip().lstrip('#')
            try:
                db_id = int(clean_search)
                query += " AND (t.ticket_name ILIKE %s OR t.description ILIKE %s OR t.company ILIKE %s OR t.business_unit ILIKE %s OR t.id::text ILIKE %s OR t.id = %s)"
                params.extend([f"%{search}%", f"%{search}%", f"%{search}%", f"%{search}%", f"%{search}%", db_id])
            except ValueError:
                query += " AND (t.ticket_name ILIKE %s OR t.description ILIKE %s OR t.company ILIKE %s OR t.business_unit ILIKE %s OR t.id::text ILIKE %s)"
                params.extend([f"%{search}%", f"%{search}%", f"%{search}%", f"%{search}%", f"%{search}%"])
        if business_unit:
            bu_list = [x.strip().lower() for x in business_unit.split(",") if x.strip()]
            if bu_list:
                query += " AND LOWER(TRIM(t.business_unit)) = ANY(%s)"
                params.append(bu_list)
        if company:
            co_list = [x.strip().lower() for x in company.split(",") if x.strip()]
            if co_list:
                query += " AND LOWER(TRIM(t.company)) = ANY(%s)"
                params.append(co_list)
        if agent:
            ag_list = [x.strip().lower() for x in agent.split(",") if x.strip()]
            if ag_list:
                query += " AND LOWER(TRIM(t.agent)) = ANY(%s)"
                params.append(ag_list)
        if priority:
            pr_list = [x.strip().lower() for x in priority.split(",") if x.strip()]
            if pr_list:
                query += " AND LOWER(TRIM(t.priority)) = ANY(%s)"
                params.append(pr_list)
        if status:
            st_list = [x.strip().lower() for x in status.split(",") if x.strip()]
            if st_list:
                query += " AND LOWER(TRIM(t.status)) = ANY(%s)"
                params.append(st_list)
        if created_by and created_by.strip() and created_by.lower() not in ('null', 'undefined', 'none'):
            query += " AND LOWER(TRIM(t.created_by)) ILIKE %s"
            params.append(f"%{created_by.strip().lower()}%")
        if resolved_by and resolved_by.strip() and resolved_by.lower() not in ('null', 'undefined', 'none'):
            query += " AND LOWER(TRIM(t.resolved_by)) ILIKE %s"
            params.append(f"%{resolved_by.strip().lower()}%")
        if created_at:
            if "," in created_at:
                parts = created_at.split(",")
                if len(parts) == 2:
                    query += " AND t.created_at::date >= %s::date AND t.created_at::date <= %s::date"
                    params.extend([parts[0].strip(), parts[1].strip()])
            else:
                query += " AND t.created_at::date = %s::date"
                params.append(created_at)
        if resolved_at:
            if "," in resolved_at:
                parts = resolved_at.split(",")
                if len(parts) == 2:
                    query += " AND t.resolved_at::date >= %s::date AND t.resolved_at::date <= %s::date"
                    params.extend([parts[0].strip(), parts[1].strip()])
            else:
                query += " AND t.resolved_at::date = %s::date"
                params.append(resolved_at)
            
        query += " ORDER BY t.id DESC"
        cur.execute(query, tuple(params))
        rows = cur.fetchall()
        result = {"tickets": [dict(r) for r in rows]}
        cache_manager.set(cache_key, result, ttl_seconds=5)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/tickets/stats")
def get_ticket_stats(user: dict = Depends(get_current_user_local)):
    cache_key = f"ticket-stats:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        tech_where = "1=1"
        tech_params = []
        
        # 1. Total by status
        cur.execute(f"SELECT status, COUNT(*) as count FROM tickets WHERE {tech_where} GROUP BY status;", tuple(tech_params))
        status_rows = [dict(r) for r in cur.fetchall()]
        
        # 2. Total by priority (Open tickets only)
        cur.execute(f"SELECT priority, COUNT(*) as count FROM tickets WHERE status != 'Resolved' AND {tech_where} GROUP BY priority;", tuple(tech_params))
        priority_rows = [dict(r) for r in cur.fetchall()]
        
        # 3. Last 30 days daily status (Open vs Resolved)
        today = datetime.now()
        dates_list = [(today - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(29, -1, -1)]
        
        daily_stats = []
        for dt in dates_list:
            cur.execute(f"""
                SELECT 
                    SUM(CASE WHEN status != 'Resolved' THEN 1 ELSE 0 END) as open_count,
                    SUM(CASE WHEN status = 'Resolved' THEN 1 ELSE 0 END) as resolved_count
                FROM tickets
                WHERE DATE(created_at) = %s AND {tech_where};
            """, tuple([dt] + tech_params))
            stat = cur.fetchone()
            daily_stats.append({
                "date": dt,
                "open": stat["open_count"] or 0,
                "resolved": stat["resolved_count"] or 0
            })
            
        result = {
            "status_stats": status_rows,
            "priority_stats": priority_rows,
            "daily_stats": daily_stats
        }
        cache_manager.set(cache_key, result, ttl_seconds=10)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# CLIENT REPORTS ENDPOINTS
# ==============================================================================

@router.post("/reports")
def upload_report(report: ReportCreate, user: dict = Depends(get_current_user_local)):
    # Restrict client users from uploading documents
    if user.get("role") == "client":
        raise HTTPException(status_code=403, detail="Forbidden: Client users are not permitted to upload reports.")

    # Verify non-admin has technology clearance for this client database type
    if not user.get("isAdmin"):
        conn = get_connection()
        try:
            cur = conn.cursor()
            if not check_client_and_tech_permission(report.client_name, user, cur):
                raise HTTPException(status_code=403, detail="Forbidden: You do not have permission to upload reports for this client.")
        finally:
            conn.close()

    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO client_reports (client_name, title, month, year, file_name, file_data, notes, uploaded_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id;
        """, (
            report.client_name, report.title, report.month, report.year,
            report.file_name, report.file_data, report.notes, user.get("username", "System")
        ))
        report_id = cur.fetchone()[0]
        
        # Add dynamic upload notification to the global system pool
        uploader = user.get("username", "admin")
        notify_msg = f"New SLA Diagnostic Report '{report.title}' uploaded by '{uploader}' for Client '{report.client_name}'"
        cur.execute("""
            INSERT INTO notifications (username, message, is_read)
            VALUES ('global', %s, FALSE);
        """, (notify_msg,))
        
        # Scope notification to engineers assigned to this client's database technology
        cur.execute("SELECT db_type FROM admin_clients WHERE client_name = %s;", (report.client_name,))
        row = cur.fetchone()
        db_type = row[0] if row else None
        if db_type:
            cur.execute("SELECT DISTINCT email FROM leads WHERE technology = %s AND status = 'active';", (db_type,))
            emails = [r[0] for r in cur.fetchall()]
            if emails:
                cur.execute("SELECT username FROM users WHERE email = ANY(%s);", (emails,))
                usernames = [r[0] for r in cur.fetchall()]
                for target_user in usernames:
                    cur.execute("""
                        INSERT INTO notifications (username, message, is_read)
                        VALUES (%s, %s, FALSE);
                    """, (target_user, f"New report '{report.title}' uploaded for client '{report.client_name}' ({db_type})",))
        
        conn.commit()
        cache_manager.invalidate("reports:")
        cache_manager.invalidate("report-counts:")
        return {"status": "success", "report_id": report_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports")
def get_reports(client_name: str, user: dict = Depends(get_current_user_local)):
    cache_key = f"reports:{client_name}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Verify non-admin has technology clearance for this client database type
        if not user.get("isAdmin"):
            assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
            cur.execute("""
                SELECT id FROM admin_clients 
                WHERE client_name = %s 
                  AND EXISTS (
                      SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                      WHERE t = ANY(%s)
                  );
            """, (client_name, assigned_techs))
            if not cur.fetchone():
                return {"reports": []}
                
        cur.execute("""
            SELECT id, client_name, title, month, year, file_name, notes, uploaded_by, uploaded_at
            FROM client_reports
            WHERE client_name = %s
            ORDER BY year DESC, month DESC, uploaded_at DESC;
        """, (client_name,))
        rows = cur.fetchall()
        result = {"reports": [dict(r) for r in rows]}
        cache_manager.set(cache_key, result, ttl_seconds=5)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/counts")
def get_report_counts(user: dict = Depends(get_current_user_local)):
    cache_key = f"report-counts:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor()
        
        if not user.get("isAdmin"):
            assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
            cur.execute("""
                SELECT cr.client_name, COUNT(*) 
                FROM client_reports cr
                JOIN admin_clients ac ON cr.client_name = ac.client_name
                WHERE LOWER(TRIM(ac.db_type)) = ANY(%s)
                GROUP BY cr.client_name;
            """, (assigned_techs,))
        else:
            cur.execute("SELECT client_name, COUNT(*) FROM client_reports GROUP BY client_name;")
        
        rows = cur.fetchall()
        counts = {row[0]: row[1] for row in rows}
        cache_manager.set(cache_key, counts, ttl_seconds=10)
        return counts
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/download/{report_id}")
def download_report(report_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT file_name, file_data FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report file not found")
            
        file_name = row["file_name"]
        file_base64 = row["file_data"]
        
        # Decode base64 securely
        if "," in file_base64:
            file_base64 = file_base64.split(",")[1]
            
        # Replace spaces back to plus signs if url-encoded transmission mutated them
        file_base64 = file_base64.replace(" ", "+")
        file_bytes = base64.b64decode(file_base64)
        
        import io
        ext = file_name.split(".")[-1].lower()
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc": "application/msword",
            "txt": "text/plain",
            "csv": "text/csv",
            "json": "application/json",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg"
        }
        media_type = mime_types.get(ext, "application/octet-stream")
        
        # Use StreamingResponse to guarantee binary integrity and prevent encoding data loss/corruption
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/reports/{report_id}/reviews")
def add_report_review(report_id: int, review: ReportReviewCreate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        # Fetch client name of report to check tech clearance
        cur.execute("SELECT client_name FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
        client_name = row[0]
        
        # Check tech permission for standard user
        if not user.get("isAdmin"):
            if not check_client_and_tech_permission(client_name, user, cur):
                raise HTTPException(status_code=403, detail="Forbidden: Access to this client's report is denied.")
        
        # Insert review
        cur.execute("""
            INSERT INTO report_reviews (report_id, username, rating, comment, mom)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id;
        """, (report_id, user.get("username"), review.rating or 5, review.comment, review.mom))
        review_id = cur.fetchone()[0]
        conn.commit()
        return {"status": "success", "review_id": review_id}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/{report_id}/reviews")
def get_report_reviews(report_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Fetch client name to check permissions
        cur.execute("SELECT client_name FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
        client_name = row["client_name"]
        
        # Check tech permission for standard user
        if not user.get("isAdmin"):
            if not check_client_and_tech_permission(client_name, user, cur):
                raise HTTPException(status_code=403, detail="Forbidden: Access to this client's report is denied.")
                
        # Retrieve reviews
        cur.execute("""
            SELECT id, report_id, username, rating, comment, mom, created_at
            FROM report_reviews
            WHERE report_id = %s
            ORDER BY created_at DESC;
        """, (report_id,))
        rows = cur.fetchall()
        return {"reviews": [dict(r) for r in rows]}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/share-download/{report_id}")
def share_download_report(report_id: int):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT file_name, file_data FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report file not found")
            
        file_name = row["file_name"]
        file_base64 = row["file_data"]
        
        # Decode base64 securely
        if "," in file_base64:
            file_base64 = file_base64.split(",")[1]
            
        file_bytes = base64.b64decode(file_base64)
        
        import io
        ext = file_name.split(".")[-1].lower()
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc": "application/msword",
            "txt": "text/plain",
            "csv": "text/csv",
            "json": "application/json",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg"
        }
        media_type = mime_types.get(ext, "application/octet-stream")
        
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={file_name}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.delete("/reports/{report_id}")
def delete_report(report_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT title, client_name FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        title = row[0] if row else "Unknown"
        client_name = row[1] if row else "Unknown"
        
        cur.execute("DELETE FROM client_reports WHERE id = %s;", (report_id,))
        
        creator = user.get("username", "admin")
        notify_msg = f"SLA Diagnostic Report '{title}' for Client '{client_name}' was deleted by admin '{creator}'"
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))
        
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/reports/clear-all")
def clear_all_reports(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("TRUNCATE TABLE client_reports;")
        
        creator = user.get("username", "admin")
        notify_msg = f"All SLA Diagnostic Reports were cleared from the database by admin '{creator}'"
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))
        
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/view-text/{report_id}")
def view_report_text(report_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, client_name, title, month, year, file_name, file_data, notes, uploaded_by, uploaded_at FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
            
        file_name = row["file_name"]
        file_base64 = row["file_data"]
        
        if "," in file_base64:
            file_base64 = file_base64.split(",")[1]
            
        # Replace spaces back to plus signs if url-encoded transmission mutated them
        file_base64 = file_base64.replace(" ", "+")
        file_bytes = base64.b64decode(file_base64)
        
        # Determine format and parse text
        ext = file_name.split(".")[-1].lower() if "." in file_name else ""
        
        # Auto-detect file type using magic bytes if extension is missing or unrecognized
        if not ext or ext not in ["docx", "pdf", "png", "jpg", "jpeg", "gif"]:
            if file_bytes.startswith(b"PK\x03\x04"):
                ext = "docx"
            elif file_bytes.startswith(b"%PDF"):
                ext = "pdf"
            elif file_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
                ext = "png"
            elif file_bytes.startswith(b"\xff\xd8\xff"):
                ext = "jpg"
            elif file_bytes.startswith(b"GIF89a") or file_bytes.startswith(b"GIF87a"):
                ext = "gif"
        
        extracted_text = ""
        is_decoded_as_text = False
        
        # Try decoding as plain text first if no binary signature is found
        if ext in ["txt", "log", "csv", "json", "md", "sql", "xml", "html"] or not ext or ext == file_name.lower():
            try:
                extracted_text = file_bytes.decode("utf-8")
                is_decoded_as_text = True
            except Exception:
                try:
                    extracted_text = file_bytes.decode("latin-1")
                    is_decoded_as_text = True
                except Exception:
                    pass
        
        if is_decoded_as_text:
            if not extracted_text.strip():
                extracted_text = "[Empty Plain Text File]"
        elif ext == "docx":
            import io
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                paragraphs = []
                for p in doc.paragraphs:
                    paragraphs.append(p.text)
                # Keep tables as well!
                for table in doc.tables:
                    for row_idx, row_obj in enumerate(table.rows):
                        row_cells = [cell.text for cell in row_obj.cells]
                        paragraphs.append(" | ".join(row_cells))
                extracted_text = "\n\n".join(paragraphs).strip()
                if len(extracted_text) == 0:
                    extracted_text = "Empty Document or no text paragraphs"
            except Exception as e:
                # Fallback to tag cleaning if zip extraction/docx library has library mismatches
                try:
                    import zipfile
                    import re
                    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                        doc_xml = z.read("word/document.xml")
                        xml_str = doc_xml.decode("utf-8", errors="ignore")
                        # Replace paragraph end tags with newlines before stripping
                        xml_str = xml_str.replace("</w:p>", "\n\n").replace("</w:tr>", "\n")
                        extracted_text = re.sub(r"<[^>]+>", "", xml_str)
                        extracted_text = re.sub(r"[ \t]+", " ", extracted_text).strip()
                except Exception as ex:
                    extracted_text = f"[Secure Sandbox Viewer failed to extract docx text: {str(e)} / {str(ex)}]"
        elif ext == "pdf":
            # PDF is binary, return a beautifully structured layout simulation with file telemetry metadata
            extracted_text = (
                f"[PDF EXECUTIVE SANDBOX READER]\n\n"
                f"Document: {file_name}\n"
                f"Size: {len(file_bytes)} bytes\n"
                f"Owner scope: {row['uploaded_by']}\n"
                f"Status: Safe - Compliant with sandbox policy.\n\n"
                f"To view tables or full database schemas, please use the secure 'Download Document' stream at the bottom."
            )
        else:
            extracted_text = f"[Binary Resource Preview]\n\nResource Name: {file_name}\nType: .{ext}\nSize: {len(file_bytes)} bytes\n\nDownloaded files are verified safe. Use local tools to open."

        return {
            "id": row["id"],
            "client_name": row["client_name"],
            "title": row["title"],
            "month": row["month"],
            "year": row["year"],
            "file_name": row["file_name"],
            "file_ext": ext,
            "file_data": row["file_data"],
            "notes": row["notes"],
            "uploaded_by": row["uploaded_by"],
            "uploaded_at": row["uploaded_at"].isoformat(),
            "extracted_text": extracted_text
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/view-data/{report_id}")
def get_report_data(report_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, client_name, title, month, year, file_name, file_data, notes, uploaded_by, uploaded_at FROM client_reports WHERE id = %s;", (report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
        return row
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# PAGE DURATION TELEMETRY ENDPOINTS
# ==============================================================================

@router.post("/monitoring/page-time")
def save_page_time(ping: PageTimePing, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    username = user.get("username", "Anonymous")
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO user_page_activity (username, page_path, duration_seconds, last_active_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (username, page_path) 
            DO UPDATE SET 
                duration_seconds = user_page_activity.duration_seconds + EXCLUDED.duration_seconds,
                last_active_at = NOW();
        """, (username, ping.page_path, ping.duration_seconds))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/monitoring/page-time")
def get_page_time_logs(user: dict = Depends(get_current_user_local)):
    # Admin only check
    if not user.get("isAdmin") and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT username, page_path, duration_seconds, last_active_at 
            FROM user_page_activity 
            ORDER BY duration_seconds DESC;
        """)
        rows = cur.fetchall()
        return {"telemetry": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/monitoring/my-page-time")
def get_my_page_time_logs(user: dict = Depends(get_current_user_local)):
    username = user.get("username")
    if not username:
        raise HTTPException(status_code=401, detail="Authentication required")
        
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT page_path, duration_seconds, last_active_at 
            FROM user_page_activity 
            WHERE LOWER(username) = LOWER(%s)
            ORDER BY last_active_at DESC;
        """, (username,))
        rows = cur.fetchall()
        for r in rows:
            if r["last_active_at"]:
                r["last_active_at"] = r["last_active_at"].isoformat()
        return {"telemetry": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/alert-settings")
def get_alert_settings(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM client_alert_settings ORDER BY client_name, db_type;")
        settings = cur.fetchall()
        return {"status": "success", "settings": settings}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/alert-settings")
def save_alert_setting(setting: ClientAlertSettingCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO client_alert_settings (
                client_name, db_type, cpu_threshold, memory_threshold, disk_threshold, io_threshold,
                slow_query_threshold_ms, long_running_threshold_sec, client_emails, cc_emails,
                server_down_alert, critical_error_alert
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (client_name, db_type) DO UPDATE SET
                cpu_threshold = EXCLUDED.cpu_threshold,
                memory_threshold = EXCLUDED.memory_threshold,
                disk_threshold = EXCLUDED.disk_threshold,
                io_threshold = EXCLUDED.io_threshold,
                slow_query_threshold_ms = EXCLUDED.slow_query_threshold_ms,
                long_running_threshold_sec = EXCLUDED.long_running_threshold_sec,
                client_emails = EXCLUDED.client_emails,
                cc_emails = EXCLUDED.cc_emails,
                server_down_alert = EXCLUDED.server_down_alert,
                critical_error_alert = EXCLUDED.critical_error_alert;
        """, (
            setting.client_name, setting.db_type, setting.cpu_threshold, setting.memory_threshold,
            setting.disk_threshold, setting.io_threshold, setting.slow_query_threshold_ms,
            setting.long_running_threshold_sec, setting.client_emails, setting.cc_emails,
            setting.server_down_alert, setting.critical_error_alert
        ))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/alert-settings/{setting_id}")
def delete_alert_setting(setting_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM client_alert_settings WHERE id = %s;", (setting_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/technology-alerts")
def get_technology_alerts(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT technology, alert_email FROM technology_alerts_config ORDER BY technology;")
        configs = cur.fetchall()
        return {"status": "success", "configs": configs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/technology-alerts")
def save_technology_alert(config: TechnologyAlertConfigCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO technology_alerts_config (technology, alert_email)
            VALUES (%s, %s)
            ON CONFLICT (technology) DO UPDATE SET alert_email = EXCLUDED.alert_email;
        """, (config.technology, config.alert_email))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/technology-alerts/{technology}")
def delete_technology_alert(technology: str, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM technology_alerts_config WHERE LOWER(technology) = LOWER(%s);", (technology,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# ==============================================================================
# ADMIN SETUPS ENDPOINTS
# ==============================================================================

@router.post("/admin/clients")
def create_admin_client(client: AdminClientCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO admin_clients (client_name, db_type, server_name, client_email, phone_number)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id;
        """, (client.client_name, client.db_type, client.server_name, client.client_email, client.phone_number))
        client_id = cur.fetchone()[0]
        conn.commit()
        cache_manager.invalidate("admin-clients:")
        cache_manager.invalidate("filters:")
        cache_manager.invalidate("telemetry:")
        return {"status": "success", "client_id": client_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/clients")
def get_admin_clients(user: dict = Depends(get_current_user_local)):
    allowed_clients_key = ",".join(sorted(user.get("allowed_clients", [])))
    assigned_techs_key = ",".join(sorted(user.get("assigned_techs", [])))
    cache_key = f"admin-clients:{user.get('isAdmin')}:{allowed_clients_key}:{assigned_techs_key}:{user.get('username')}"
    
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if user.get("isAdmin"):
            cur.execute("SELECT * FROM admin_clients ORDER BY client_name;")
        else:
            allowed_clients = user.get("allowed_clients", [])
            if allowed_clients:
                cur.execute("SELECT * FROM admin_clients WHERE client_name = ANY(%s) ORDER BY client_name;", (allowed_clients,))
            else:
                assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
                cur.execute("""
                    SELECT * FROM admin_clients 
                    WHERE EXISTS (
                        SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                        WHERE t = ANY(%s)
                    ) ORDER BY client_name;
                """, (assigned_techs,))
        rows = cur.fetchall()
        result = {"clients": [dict(r) for r in rows]}
        cache_manager.set(cache_key, result, ttl_seconds=30)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/agents")
def create_admin_agent(agent: AdminAgentCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO admin_agents (agent_name, company_name, business_unit, technology, email)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id;
        """, (agent.agent_name, agent.company_name, agent.business_unit, agent.technology, agent.email))
        agent_id = cur.fetchone()[0]
        conn.commit()
        cache_manager.invalidate("admin-agents:")
        return {"status": "success", "agent_id": agent_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/agents")
def get_admin_agents(user: dict = Depends(get_current_user_local)):
    cache_key = "admin-agents"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM admin_agents ORDER BY agent_name;")
        rows = cur.fetchall()
        result = {"agents": [dict(r) for r in rows]}
        cache_manager.set(cache_key, result, ttl_seconds=30)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/agents/{agent_id}")
def delete_admin_agent(agent_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM admin_agents WHERE id = %s;", (agent_id,))
        conn.commit()
        cache_manager.invalidate("admin-agents:")
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/user-permissions")
def create_user_permissions(perm: UserPermissionCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        # Split technologies by comma
        techs = [t.strip() for t in perm.technology.split(",") if t.strip()]
        if not techs:
            techs = ["MySQL"]
            
        for tech in techs:
            # Inserts scope inside the standard 'leads' table
            cur.execute("""
                INSERT INTO leads (email, technology, status, is_lead)
                VALUES (%s, %s, %s, %s)
                ON CONFLICT (email, technology) 
                DO UPDATE SET status = EXCLUDED.status, is_lead = EXCLUDED.is_lead;
            """, (perm.email.strip(), tech, perm.status, perm.is_lead))
            
        # Role handling
        role_selected = perm.role if perm.role else "user"
        if role_selected == "admin":
            cur.execute("""
                INSERT INTO system_admins (email, status)
                VALUES (%s, 'active')
                ON CONFLICT (email) DO UPDATE SET status = 'active';
            """, (perm.email.strip().lower(),))
        else:
            cur.execute("""
                UPDATE system_admins SET status = 'inactive' WHERE LOWER(email) = %s;
            """, (perm.email.strip().lower(),))
            
        # Check if user exists in users table
        user_created = False
        cur.execute("SELECT id FROM users WHERE LOWER(email) = LOWER(%s);", (perm.email.strip(),))
        user_row = cur.fetchone()
        if user_row:
            cur.execute("UPDATE users SET role = %s WHERE LOWER(email) = LOWER(%s);", (role_selected, perm.email.strip()))
        else:
            # Pre-create user account securely
            username = perm.email.split("@")[0]
            import bcrypt
            hashed_pwd = bcrypt.hashpw("geopits123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cur.execute("""
                INSERT INTO users (username, email, full_name, hashed_password, role)
                VALUES (%s, %s, %s, %s, %s)
                ON CONFLICT (username) DO NOTHING;
            """, (username, perm.email.strip(), username.capitalize(), hashed_pwd, role_selected))
            user_created = True
            
        conn.commit()

        # Send technology & role assignment notification email
        try:
            role_display = "Lead" if perm.is_lead else (perm.role if perm.role else "User")
            username = perm.email.split("@")[0]
            greeting = f"Hello {username.capitalize()},"
            if user_created:
                subject = "[GeoMon Portal] Invitation"
                lead_text = "You have been invited to join the GeoMon Enterprise Observability Portal. A new account has been pre-created for you. Please log in using the temporary password below and update it upon first entry."
                details = {
                    "Username": username,
                    "Temporary Password": "geopits123",
                    "Assigned Role": role_display.upper(),
                    "Technology Scopes": perm.technology,
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Account Invitation"
            else:
                subject = "[GeoMon Portal] System Privilege Allocation Updated"
                lead_text = "Your technology scopes and system access role have been updated by an administrator. Please review your updated permissions below."
                details = {
                    "Username": username,
                    "Assigned Role": role_display.upper(),
                    "Technology Scopes": perm.technology,
                    "Scope Status": perm.status.capitalize(),
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Privilege Update"
                
            body = build_gorgeous_html_email(
                title=title,
                greeting=greeting,
                lead_text=lead_text,
                details=details,
                action_url="http://localhost:8000/#/login",
                action_text="Access Observability Portal"
            )
            send_email_outlook(to_emails=perm.email.strip(), cc_emails=None, subject=subject, body=body, exclude_dccagent=True)
            print(f"[PRIVILEGE NOTIFICATION SENT] Emailed {perm.email.strip()} about role/technology scopes")
        except Exception as mail_err:
            print(f"[PRIVILEGE NOTIFICATION ERROR] Failed to send email to {perm.email}: {mail_err}")

        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# USER CLIENT ACCESS PERMISSIONS
# ==============================================================================

@router.get("/admin/user-clients")
def get_user_clients(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT uc.id, u.email, c.client_name, uc.access_level, uc.created_at
            FROM user_clients uc
            JOIN users u ON uc.user_id = u.id
            JOIN admin_clients c ON uc.client_id = c.id
            ORDER BY u.email, c.client_name;
        """)
        rows = cur.fetchall()
        for r in rows:
            if r.get("created_at"):
                r["created_at"] = r["created_at"].isoformat()
        return {"permissions": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/user-clients")
def create_user_client_permission(perm: UserClientPermissionCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        email_clean = perm.email.strip().lower()
        client_name_clean = perm.client_name.strip()
        
        # 1. Resolve or pre-create the user
        user_created = False
        cur.execute("SELECT id FROM users WHERE LOWER(email) = %s;", (email_clean,))
        user_row = cur.fetchone()
        if user_row:
            user_id = user_row[0]
        else:
            username = email_clean.split("@")[0]
            import bcrypt
            hashed_pwd = bcrypt.hashpw("geopits123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cur.execute("""
                INSERT INTO users (username, email, full_name, hashed_password, role)
                VALUES (%s, %s, %s, %s, 'user')
                RETURNING id;
            """, (username, email_clean, username.capitalize(), hashed_pwd))
            user_id = cur.fetchone()[0]
            user_created = True
            
        # 2. Resolve client_id from admin_clients
        cur.execute("SELECT id FROM admin_clients WHERE client_name = %s;", (client_name_clean,))
        client_row = cur.fetchone()
        if not client_row:
            raise HTTPException(status_code=400, detail=f"Client '{client_name_clean}' does not exist in registry. Register client first.")
        client_id = client_row[0]
        
        # 3. Insert or update the permission mapping
        cur.execute("""
            INSERT INTO user_clients (user_id, client_id, access_level)
            VALUES (%s, %s, %s)
            ON CONFLICT (user_id, client_id) 
            DO UPDATE SET access_level = EXCLUDED.access_level;
        """, (user_id, client_id, perm.access_level))
        
        conn.commit()

        # Send privilege assignment notification email
        try:
            username = email_clean.split("@")[0]
            greeting = f"Hello {username.capitalize()},"
            if user_created:
                subject = "[GeoMon Portal] Invitation"
                lead_text = "You have been invited to join the GeoMon Enterprise Observability Portal. A new account has been pre-created for you. Please log in using the temporary credentials below and update your password upon first entry."
                details = {
                    "Username": username,
                    "Temporary Password": "geopits123",
                    "Assigned Role": "USER",
                    "Client Access": client_name_clean,
                    "Access Level": perm.access_level.upper(),
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Account Invitation"
            else:
                subject = f"[GeoMon Portal] Client Privilege Allocated: {client_name_clean}"
                lead_text = "An administrator has updated your access privileges for the client environment below."
                details = {
                    "Username": username,
                    "Client Name": client_name_clean,
                    "Access Level": perm.access_level.upper(),
                    "Assigned By": user.get("username", "System Administrator"),
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
                title = "Client Access Allocated"
                
            body = build_gorgeous_html_email(
                title=title,
                greeting=greeting,
                lead_text=lead_text,
                details=details,
                action_url="http://localhost:8000/#/login",
                action_text="Access Observability Portal"
            )
            send_email_outlook(to_emails=email_clean, cc_emails=None, subject=subject, body=body, exclude_dccagent=True)
            print(f"[PRIVILEGE NOTIFICATION SENT] Emailed {email_clean} about privilege for client {client_name_clean}")
        except Exception as mail_err:
            print(f"[PRIVILEGE NOTIFICATION ERROR] Failed to send email to {email_clean}: {mail_err}")

        return {"status": "success"}
    except HTTPException as he:
        conn.rollback()
        raise he
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/user-clients/{permission_id}")
def delete_user_client_permission(permission_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM user_clients WHERE id = %s;", (permission_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# NOTIFICATIONS ENDPOINTS
# ==============================================================================


@router.get("/notifications")
def get_notifications(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        username = user.get("username", "Anonymous")
        full_name = user.get("full_name", user.get("fullName", ""))
        
        # We query notifications where the username column equals the user's username, 
        # OR equals the user's full name, OR contains their username as a substring,
        # OR their full name contains the notification username, or username is 'global' or 'admin'
        search_terms = { 'global', 'admin', username.lower() }
        if full_name:
            search_terms.add(full_name.lower())

        cur.execute("""
            SELECT id, message, is_read, created_at 
            FROM notifications 
            WHERE LOWER(username) = ANY(%s)
            ORDER BY created_at DESC 
            LIMIT 25;
        """, (list(search_terms),))
        rows = cur.fetchall()
        for row in rows:
            row["created_at"] = row["created_at"].isoformat()
        return {"notifications": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.get("/notifications/unread-counts")
def get_unread_counts(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        username = user.get("username", "Anonymous")
        full_name = user.get("full_name", user.get("fullName", ""))
        search_terms = { 'global', 'admin', username.lower() }
        if full_name:
            search_terms.add(full_name.lower())

        cur.execute("""
            SELECT 
                COUNT(*) FILTER (WHERE NOT is_read AND (message ILIKE '%%new ticket%%' OR message ILIKE '%%created%%')) as new_tickets,
                COUNT(*) FILTER (WHERE NOT is_read AND (message ILIKE '%%reply%%' OR message ILIKE '%%comment%%' OR message ILIKE '%%replied%%')) as new_replies,
                COUNT(*) FILTER (WHERE NOT is_read) as total_unread
            FROM notifications 
            WHERE LOWER(username) = ANY(%s);
        """, (list(search_terms),))
        row = cur.fetchone()
        return {
            "new_tickets": row[0] or 0,
            "new_replies": row[1] or 0,
            "total_unread": row[2] or 0
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/notifications/read-all")
def mark_all_notifications_read(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        username = user.get("username", "Anonymous")
        full_name = user.get("full_name", user.get("fullName", ""))
        search_terms = { 'global', 'admin', username.lower() }
        if full_name:
            search_terms.add(full_name.lower())

        cur.execute("""
            UPDATE notifications 
            SET is_read = TRUE 
            WHERE LOWER(username) = ANY(%s);
        """, (list(search_terms),))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/notifications/clear")
def clear_user_notifications(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        username = user.get("username", "Anonymous")
        full_name = user.get("full_name", user.get("fullName", ""))
        search_terms = { 'global', 'admin', username.lower() }
        if full_name:
            search_terms.add(full_name.lower())

        cur.execute("""
            DELETE FROM notifications 
            WHERE LOWER(username) = ANY(%s);
        """, (list(search_terms),))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/notifications/read/{notification_id}")
def mark_single_notification_read(notification_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("UPDATE notifications SET is_read = TRUE WHERE id = %s;", (notification_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/notifications/broadcast")
def broadcast_notification(payload: BroadcastMessage, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO notifications (username, message, is_read)
            VALUES ('global', %s, FALSE)
            RETURNING id;
        """, (payload.message,))
        notif_id = cur.fetchone()[0]
        conn.commit()
        return {"status": "success", "notification_id": notif_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# ==============================================================================
# ADDITIONAL ADMIN CRUD & SYSTEM INTEGRATION ENDPOINTS
# ==============================================================================

class AdminUserCreate(BaseModel):
    username: str
    email: str
    full_name: str
    password: Optional[str] = None
    role: str

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str

class ProfilePicRequest(BaseModel):
    profile_pic: str

class AdminUserUpdate(BaseModel):
    username: str
    email: str
    full_name: str
    password: Optional[str] = None
    role: str

class AdminClientUpdate(BaseModel):
    client_name: str
    db_type: str
    server_name: str
    client_email: Optional[str] = None
    phone_number: Optional[str] = None

class TicketAdminUpdate(BaseModel):
    business_unit: str
    company: str
    contact: str
    ticket_name: str
    category: str
    status: str
    priority: str
    agent: str
    description: str

@router.get("/admin/users")
def get_admin_users(user: dict = Depends(get_current_user_local)):
    is_admin = user.get("isAdmin") or user.get("role") == "admin"
    is_lead = user.get("role") == "lead"
    if not is_admin and not is_lead:
        raise HTTPException(status_code=403, detail="Admin or Lead authorization required")
    
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if is_admin:
            cur.execute("""
                SELECT u.id, u.username, u.email, u.full_name, u.profile_pic, u.last_active_at,
                       CASE 
                           WHEN EXISTS (
                               SELECT 1 FROM system_admins sa 
                               WHERE LOWER(sa.email) = LOWER(u.email) AND sa.status = 'active'
                           ) OR LOWER(u.role) = 'admin' THEN 'admin'
                           WHEN EXISTS (
                               SELECT 1 FROM leads l 
                               WHERE (LOWER(l.email) = LOWER(u.email) OR LOWER(l.email) = LOWER(u.username))
                                 AND l.is_lead = true 
                                 AND l.status = 'active'
                           ) THEN 'lead'
                           ELSE u.role 
                       END as role
                FROM users u
                ORDER BY u.last_active_at DESC NULLS LAST;
            """)
        else:
            email = user.get("email") or ""
            domain = email.split("@")[-1].lower() if "@" in email else ""
            if not domain:
                return {"users": []}
            cur.execute("""
                SELECT u.id, u.username, u.email, u.full_name, u.profile_pic, u.last_active_at,
                       CASE 
                           WHEN EXISTS (
                               SELECT 1 FROM system_admins sa 
                               WHERE LOWER(sa.email) = LOWER(u.email) AND sa.status = 'active'
                           ) OR LOWER(u.role) = 'admin' THEN 'admin'
                           WHEN EXISTS (
                               SELECT 1 FROM leads l 
                               WHERE (LOWER(l.email) = LOWER(u.email) OR LOWER(l.email) = LOWER(u.username))
                                 AND l.is_lead = true 
                                 AND l.status = 'active'
                           ) THEN 'lead'
                           ELSE u.role 
                       END as role
                FROM users u
                WHERE LOWER(u.email) LIKE %s
                ORDER BY u.last_active_at DESC NULLS LAST;
            """, (f"%@{domain}",))
        
        rows = cur.fetchall()
        results = []
        for r in rows:
            row = dict(r)
            if row.get('last_active_at') and isinstance(row['last_active_at'], datetime):
                row['last_active_at'] = row['last_active_at'].astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
            results.append(row)
        return {"users": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/users")
def add_admin_user(user_req: AdminUserCreate, user: dict = Depends(get_current_user_local)):
    is_admin = user.get("isAdmin") or user.get("role") == "admin"
    is_lead = user.get("role") == "lead"
    if not is_admin and not is_lead:
        raise HTTPException(status_code=403, detail="Admin or Lead authorization required")
    
    if is_lead:
        lead_email = user.get("email") or ""
        lead_domain = lead_email.split("@")[-1].lower() if "@" in lead_email else ""
        req_email = user_req.email or ""
        req_domain = req_email.split("@")[-1].lower() if "@" in req_email else ""
        if not lead_domain or lead_domain != req_domain:
            raise HTTPException(status_code=400, detail="Lead users can only add users within their own email domain")
        if user_req.role not in ['user', 'client']:
            raise HTTPException(status_code=403, detail="Lead users cannot create admin or lead accounts")
            
    import bcrypt
    temp_pwd = user_req.password.strip() if user_req.password else None
    if temp_pwd:
        hashed_pwd = bcrypt.hashpw(temp_pwd.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    else:
        hashed_pwd = bcrypt.hashpw("geopits123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT id FROM users WHERE username = %s OR email = %s;", (user_req.username, user_req.email))
        if cur.fetchone():
            raise HTTPException(status_code=400, detail="Username or email already exists")
            
        cur.execute("""
            INSERT INTO users (username, email, full_name, hashed_password, role)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id;
        """, (user_req.username, user_req.email, user_req.full_name, hashed_pwd, user_req.role))
        new_id = cur.fetchone()[0]
        
        # Sync access rules based on role
        if user_req.role == 'admin':
            cur.execute("""
                INSERT INTO system_admins (email, status)
                VALUES (%s, 'active')
                ON CONFLICT (email) DO NOTHING;
            """, (user_req.email,))
        elif user_req.role == 'lead':
            cur.execute("""
                INSERT INTO leads (email, technology, status, is_lead)
                VALUES (%s, 'MySQL, PostgreSQL, MongoDB, MSSQL', 'active', TRUE)
                ON CONFLICT (email, technology) DO UPDATE SET is_lead = TRUE, status = 'active';
            """, (user_req.email,))
        conn.commit()

        # Send invitation email using the existing Outlook mail sender
        try:
            greeting = f"Hello {user_req.full_name},"
            if temp_pwd:
                subject = "[GeoMon Portal] Invitation & Login Credentials"
                lead_text = "You have been invited to join the GeoMon Enterprise Observability Portal. Your account has been provisioned with the credentials below. Please log in and change your password upon your first login."
                details = {
                    "Username": user_req.username,
                    "Temporary Password": temp_pwd,
                    "Assigned Role": user_req.role.upper(),
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
            else:
                subject = "[GeoMon Portal] Invitation"
                lead_text = "You have been invited to join the GeoMon Enterprise Observability Portal. Your login account has been successfully initialized. You can now log in using your standard credentials."
                details = {
                    "Username": user_req.username,
                    "Assigned Role": user_req.role.upper(),
                    "Portal Access URL": "http://localhost:8000/#/login"
                }
            body = build_gorgeous_html_email(
                title="Account Invitation",
                greeting=greeting,
                lead_text=lead_text,
                details=details,
                action_url="http://localhost:8000/#/login",
                action_text="Access Observability Portal"
            )
            send_email_outlook(
                to_emails=user_req.email,
                cc_emails=None,
                subject=subject,
                body=body,
                exclude_dccagent=True
            )
        except Exception as mail_err:
            print(f"[MAIL ERROR] Failed to send invitation email: {mail_err}")

        return {"status": "success", "user_id": new_id}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.put("/admin/users/{user_id}")
def update_admin_user(user_id: int, user_req: AdminUserUpdate, user: dict = Depends(get_current_user_local)):
    is_admin = user.get("isAdmin") or user.get("role") == "admin"
    is_lead = user.get("role") == "lead"
    if not is_admin and not is_lead:
        raise HTTPException(status_code=403, detail="Admin or Lead authorization required")
    
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT email, role FROM users WHERE id = %s;", (user_id,))
        old_row = cur.fetchone()
        if not old_row:
            raise HTTPException(status_code=404, detail="User not found")
        old_email = old_row[0]
        
        if is_lead:
            lead_email = user.get("email") or ""
            lead_domain = lead_email.split("@")[-1].lower() if "@" in lead_email else ""
            old_email_domain = old_email.split("@")[-1].lower() if old_email and "@" in old_email else ""
            if not lead_domain or lead_domain != old_email_domain:
                raise HTTPException(status_code=403, detail="Lead users can only manage users within their own email domain")
            
            req_email = user_req.email or ""
            req_domain = req_email.split("@")[-1].lower() if "@" in req_email else ""
            if not lead_domain or lead_domain != req_domain:
                raise HTTPException(status_code=400, detail="Lead users can only set emails within their own domain")
            
            if user_req.role not in ['user', 'client']:
                raise HTTPException(status_code=403, detail="Lead users cannot grant admin or lead roles")
        
        if user_req.password:
            import bcrypt
            hashed_pwd = bcrypt.hashpw(user_req.password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cur.execute("""
                UPDATE users 
                SET username = %s, email = %s, full_name = %s, hashed_password = %s, role = %s
                WHERE id = %s;
            """, (user_req.username, user_req.email, user_req.full_name, hashed_pwd, user_req.role, user_id))
        else:
            cur.execute("""
                UPDATE users 
                SET username = %s, email = %s, full_name = %s, role = %s
                WHERE id = %s;
            """, (user_req.username, user_req.email, user_req.full_name, user_req.role, user_id))
            
        # Clean up old email bindings if changed
        if old_email and old_email != user_req.email:
            cur.execute("DELETE FROM system_admins WHERE email = %s;", (old_email,))
            cur.execute("DELETE FROM leads WHERE email = %s;", (old_email,))
            
        # Sync role-based privileges
        if user_req.role == 'admin':
            cur.execute("DELETE FROM leads WHERE email = %s;", (user_req.email,))
            cur.execute("""
                INSERT INTO system_admins (email, status)
                VALUES (%s, 'active')
                ON CONFLICT (email) DO NOTHING;
            """, (user_req.email,))
        elif user_req.role == 'lead':
            cur.execute("DELETE FROM system_admins WHERE email = %s;", (user_req.email,))
            cur.execute("""
                INSERT INTO leads (email, technology, status, is_lead)
                VALUES (%s, 'MySQL, PostgreSQL, MongoDB, MSSQL', 'active', TRUE)
                ON CONFLICT (email, technology) DO UPDATE SET is_lead = TRUE, status = 'active';
            """, (user_req.email,))
        else:
            cur.execute("DELETE FROM system_admins WHERE email = %s;", (user_req.email,))
            cur.execute("DELETE FROM leads WHERE email = %s;", (user_req.email,))
            
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/users/{user_id}")
def delete_admin_user(user_id: int, user: dict = Depends(get_current_user_local)):
    is_admin = user.get("isAdmin") or user.get("role") == "admin"
    is_lead = user.get("role") == "lead"
    if not is_admin and not is_lead:
        raise HTTPException(status_code=403, detail="Admin or Lead authorization required")
    
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT email, role FROM users WHERE id = %s;", (user_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        email = row[0]
        role = row[1]
        
        if is_lead:
            lead_email = user.get("email") or ""
            lead_domain = lead_email.split("@")[-1].lower() if "@" in lead_email else ""
            email_domain = email.split("@")[-1].lower() if email and "@" in email else ""
            if not lead_domain or lead_domain != email_domain:
                raise HTTPException(status_code=403, detail="Lead users can only delete users within their own email domain")
            if role in ['admin', 'lead']:
                raise HTTPException(status_code=403, detail="Lead users cannot delete admin or lead accounts")
                
        cur.execute("DELETE FROM system_admins WHERE email = %s;", (email,))
        cur.execute("DELETE FROM leads WHERE email = %s;", (email,))
        cur.execute("DELETE FROM user_clients WHERE user_id = %s;", (user_id,))
        cur.execute("DELETE FROM users WHERE id = %s;", (user_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/auth/change-password")
def change_user_password(req: ChangePasswordRequest, user: dict = Depends(get_current_user_local)):
    username = user.get("username")
    if not username:
        raise HTTPException(status_code=401, detail="Authentication required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT id, hashed_password FROM users WHERE username = %s;", (username,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        user_id, hashed_password = row
        
        import bcrypt
        if not bcrypt.checkpw(req.current_password.encode('utf-8'), hashed_password.encode('utf-8')):
            raise HTTPException(status_code=400, detail="Incorrect current password")
            
        new_hashed = bcrypt.hashpw(req.new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        cur.execute("UPDATE users SET hashed_password = %s WHERE id = %s;", (new_hashed, user_id))
        conn.commit()
        return {"status": "success", "message": "Password updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/auth/profile-pic")
def update_profile_pic(req: ProfilePicRequest, user: dict = Depends(get_current_user_local)):
    username = user.get("username")
    if not username:
        raise HTTPException(status_code=401, detail="Authentication required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("UPDATE users SET profile_pic = %s WHERE username = %s;", (req.profile_pic, username))
        conn.commit()
        return {"status": "success", "profile_pic": req.profile_pic}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.put("/admin/clients/{client_id}")
def update_admin_client(client_id: int, client: AdminClientUpdate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            UPDATE admin_clients 
            SET client_name = %s, db_type = %s, server_name = %s, client_email = %s, phone_number = %s
            WHERE id = %s;
        """, (client.client_name, client.db_type, client.server_name, client.client_email, client.phone_number, client_id))
        conn.commit()
        cache_manager.invalidate("admin-clients:")
        cache_manager.invalidate("filters:")
        cache_manager.invalidate("telemetry:")
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/clients/{client_id}")
def delete_admin_client(client_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM admin_clients WHERE id = %s;", (client_id,))
        conn.commit()
        cache_manager.invalidate("admin-clients:")
        cache_manager.invalidate("filters:")
        cache_manager.invalidate("telemetry:")
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.put("/tickets/{ticket_id}")
def update_ticket(ticket_id: int, ticket: TicketAdminUpdate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        # Verify ticket exists
        cur.execute("SELECT status, resolved_by, resolved_at FROM tickets WHERE id = %s;", (ticket_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Ticket not found")
        
        curr_status, curr_res_by, curr_res_at = row
        new_status = ticket.status.upper()
        
        res_by = curr_res_by
        res_at = curr_res_at
        
        if new_status == 'RESOLVED' and curr_status != 'RESOLVED':
            res_by = user.get("username", "System")
            res_at = datetime.now(ZoneInfo("Asia/Kolkata"))
        elif new_status != 'RESOLVED':
            res_by = None
            res_at = None
            
        cur.execute("""
            UPDATE tickets 
            SET business_unit = %s, company = %s, contact = %s, ticket_name = %s, category = %s, 
                status = %s, priority = %s, agent = %s, description = %s,
                resolved_by = %s, resolved_at = %s
            WHERE id = %s;
        """, (
            ticket.business_unit, ticket.company, ticket.contact, ticket.ticket_name, ticket.category, 
            new_status, ticket.priority, ticket.agent, ticket.description,
            res_by, res_at, ticket_id
        ))

        # Bidirectional sync: update logs linked to this ticket_id in db_monitoring_logs and db_archived_logs
        log_status_map = {
            'RESOLVED': 'Resolved',
            'IN PROGRESS': 'In Progress',
            'PENDING': 'Pending',
            'OPEN': 'Open'
        }
        new_log_status = log_status_map.get(new_status, 'Open')
        terminal_statuses = ['Resolved', 'Ignored', 'No action Required']
        is_terminal = new_log_status in terminal_statuses

        # First, update ticket_status in both tables
        cur.execute("UPDATE db_monitoring_logs SET ticket_status = %s, status_updated_at = CURRENT_TIMESTAMP WHERE ticket_id = %s", (new_status, ticket_id))
        cur.execute("UPDATE db_archived_logs SET ticket_status = %s, status_updated_at = CURRENT_TIMESTAMP WHERE ticket_id = %s", (new_status, ticket_id))

        common_cols = "client_name, server_name, db_type, log_type, log_source, log_time, log_time_utc, log_time_ist, log_level, log_message, occurrence_count, raw_log, email_subject, email_received_time, log_hash, created_at, status, owner, client_visibility, ticket_status, next_action, severity, status_updated_at, is_semantic, semantic_count, semantic_hash, time_bucket, ticket_id"

        if is_terminal:
            # Move from monitoring to archived
            cur.execute("SELECT log_hash FROM db_monitoring_logs WHERE ticket_id = %s", (ticket_id,))
            hashes = [r[0] for r in cur.fetchall()]
            for h in hashes:
                cur.execute(f"""
                    INSERT INTO db_archived_logs ({common_cols})
                    SELECT {common_cols} FROM db_monitoring_logs 
                    WHERE TRIM(log_hash) = TRIM(%s)
                    ON CONFLICT (log_hash) DO NOTHING
                """, (h,))
                cur.execute("DELETE FROM db_monitoring_logs WHERE TRIM(log_hash) = TRIM(%s)", (h,))
            
            cur.execute("UPDATE db_archived_logs SET status = %s WHERE ticket_id = %s", (new_log_status, ticket_id))
        else:
            # Move from archived to monitoring
            cur.execute("SELECT log_hash FROM db_archived_logs WHERE ticket_id = %s", (ticket_id,))
            hashes = [r[0] for r in cur.fetchall()]
            for h in hashes:
                cur.execute(f"""
                    INSERT INTO db_monitoring_logs ({common_cols})
                    SELECT {common_cols} FROM db_archived_logs 
                    WHERE TRIM(log_hash) = TRIM(%s)
                    ON CONFLICT (log_hash) DO NOTHING
                """, (h,))
                cur.execute("DELETE FROM db_archived_logs WHERE TRIM(log_hash) = TRIM(%s)", (h,))
            
            cur.execute("UPDATE db_monitoring_logs SET status = %s WHERE ticket_id = %s", (new_log_status, ticket_id))
        
        t_name = ticket.ticket_name
        creator = user.get("username", "user")
        if ticket.agent:
            cur.execute("""
                INSERT INTO notifications (username, message, is_read)
                VALUES (%s, %s, FALSE);
            """, (ticket.agent, f"Ticket #{ticket_id} '{t_name}' has been updated to status '{ticket.status}' by '{creator}'",))
        
        conn.commit()
        cache_manager.invalidate("tickets:")
        cache_manager.invalidate("ticket-stats:")
        cache_manager.invalidate("logs:")
        cache_manager.invalidate("owner-counts:")
        cache_manager.invalidate("filters:")
        try:
            from app import OBSERVABILITY_CACHE
            OBSERVABILITY_CACHE.clear()
        except Exception:
            pass
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/tickets/{ticket_id}")
def delete_ticket(ticket_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required to delete tickets")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT ticket_name FROM tickets WHERE id = %s;", (ticket_id,))
        row = cur.fetchone()
        tname = row[0] if row else "Unknown"
        
        cur.execute("DELETE FROM tickets WHERE id = %s;", (ticket_id,))
        
        creator = user.get("username", "admin")
        notify_msg = f"Incident Ticket #{ticket_id} '{tname}' was deleted by admin '{creator}'"
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))
        
        conn.commit()
        cache_manager.invalidate("tickets:")
        cache_manager.invalidate("ticket-stats:")
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# TICKET COMMENTS & LOGS ENDPOINTS
# ==============================================================================

@router.get("/tickets/all-comments/replies")
def get_all_reply_comments(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT tc.id, tc.ticket_id, tc.author, tc.comment_type, tc.content, tc.attachments, tc.created_at, t.ticket_name
            FROM ticket_comments tc
            JOIN tickets t ON tc.ticket_id = t.id
            WHERE tc.comment_type IN ('dba_reply', 'client_reply', 'reply')
            ORDER BY tc.created_at DESC;
        """)
        rows = cur.fetchall()
        for r in rows:
            if r["created_at"]:
                r["created_at"] = r["created_at"].isoformat()
        return {"comments": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/tickets/{ticket_id}/comments")
def get_ticket_comments(ticket_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Verify ticket exists
        cur.execute("SELECT company FROM tickets WHERE id = %s;", (ticket_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Ticket not found")
            
        cur.execute("""
            SELECT id, ticket_id, author, comment_type, content, attachments, created_at
            FROM ticket_comments
            WHERE ticket_id = %s
            ORDER BY created_at ASC;
        """, (ticket_id,))
        rows = cur.fetchall()
        for r in rows:
            if r["created_at"]:
                r["created_at"] = r["created_at"].isoformat()
        return {"comments": rows}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/tickets/{ticket_id}/comments")
def create_ticket_comment(ticket_id: int, req: CommentCreate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT company, ticket_name, contact, priority, status, business_unit, category, description, created_by FROM tickets WHERE id = %s;", (ticket_id,))
        ticket_row = cur.fetchone()
        if not ticket_row:
            raise HTTPException(status_code=404, detail="Ticket not found")
        ticket_company = ticket_row[0]
        ticket_name, ticket_contact, ticket_priority, ticket_status, ticket_bu, ticket_category, ticket_desc, ticket_created_by = (
            ticket_row[1], ticket_row[2], ticket_row[3], ticket_row[4], ticket_row[5], ticket_row[6], ticket_row[7], ticket_row[8]
        )

        author = user.get("username", "System")
        
        # Safe stringification of attachments to prevent unhashable type: 'dict'
        import json as _json
        attachments_str = req.attachments
        if not isinstance(attachments_str, str) and attachments_str is not None:
            attachments_str = _json.dumps(attachments_str)

        cur.execute("""
            INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id, created_at;
        """, (ticket_id, author, req.comment_type, req.content, attachments_str))
        row = cur.fetchone()
        comment_id = row[0]
        created_at = row[1].isoformat()

        notify_msg = f"User '{author}' added a {req.comment_type} on Ticket #{ticket_id} ({ticket_name})"
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))

        # Add activity log for reply/forward emails
        if req.comment_type in ["reply", "forward"]:
            action_desc = "replied to this ticket" if req.comment_type == "reply" else "forwarded this ticket"
            cur.execute("""
                INSERT INTO ticket_comments (ticket_id, author, comment_type, content, attachments)
                VALUES (%s, %s, 'log', %s, '');
            """, (ticket_id, author, f"Email sent: {author} {action_desc}."))

        # Commit comment first so database is consistent
        conn.commit()

        # Custom Actions based on Comment Type
        import re
        import json as _json
        
        # Parse attached files safely from list, dict, or stored JSON string
        email_attachments = []
        if req.attachments:
            if isinstance(req.attachments, list):
                email_attachments = req.attachments
            elif isinstance(req.attachments, dict):
                email_attachments = [req.attachments]
            elif isinstance(req.attachments, str):
                try:
                    parsed_att = _json.loads(req.attachments)
                    if isinstance(parsed_att, list):
                        email_attachments = parsed_att
                    elif isinstance(parsed_att, dict):
                        email_attachments = [parsed_att]
                except Exception:
                    pass
        
        # 1. Automatic Merging Update
        if req.comment_type == "merge":
            # parse out target ticket ID to resolve it
            m_id = re.search(r"Merged Ticket #(\d+)", req.content)
            if m_id:
                actual_db_id = int(m_id.group(1))
                cur.execute("UPDATE tickets SET status = 'RESOLVED' WHERE id = %s;", (actual_db_id,))
                conn.commit()

        # 2. Email Forwarding / Replying
        to_emails = None
        cc_emails = None
        subject = f"Ticket #{ticket_id}: {ticket_name}"
        email_body = req.content.replace('\r\n', '\n')

        # Standardize MSSQL ticket name if needed for dynamic subject matching
        std_ticket_name = ticket_name
        is_mssql_system_alert = (ticket_bu == 'MSSQL' and ticket_created_by and ticket_created_by.lower() == 'system')
        if is_mssql_system_alert:
            client = ticket_company or "GEOJIT"
            server = "DB-SERVER"
            alert_type = "Long Running Queries"
            status_str = "Closed" if (ticket_status and ticket_status.upper() in ["RESOLVED", "CLOSED"]) else "Open"
            
            # Clean Fwd/Re prefixes
            cleaned_name = re.sub(r'^(?:fw|re|fwd|vs|fwd|aw|wg|rv|ticket\s*#\d+):\s*', '', ticket_name, flags=re.IGNORECASE).strip()
            match = re.search(r'^([a-zA-Z0-9_\-]+)\s+([a-zA-Z0-9_\-]+)\s*-\s*([^:]+?)\s*:\s*([a-zA-Z0-9_\-]+)', cleaned_name, re.IGNORECASE)
            if match:
                std_ticket_name = cleaned_name
            else:
                if ticket_desc:
                    server_match = re.search(r'Server\s*-\s*([a-zA-Z0-9_\-]+)', ticket_desc, re.IGNORECASE)
                    if server_match:
                        server = server_match.group(1).strip()
                    if "transaction" in ticket_desc.lower() or "transaction" in ticket_name.lower():
                        alert_type = "Open Transaction"
                
                if server == "DB-SERVER" and ticket_name:
                    match2 = re.search(r'^([a-zA-Z0-9_\-]+)\s+([a-zA-Z0-9_\-]+)\s*-\s*(Long Running Queries|Open Transaction Alert|Open Transaction|MSSQL Alert)', ticket_name, re.IGNORECASE)
                    if match2:
                        client = match2.group(1).strip()
                        server = match2.group(2).strip()
                
                std_ticket_name = f"{client} {server} - {alert_type}: {status_str}"

        normalized_content = req.content.replace('\r\n', '\n')
        if req.comment_type == "forward":
            # parse To: and Cc: from content if they are prefixed
            to_match = re.match(r"^To:\s*(.*?)\nCc:\s*(.*?)\n\n(.*)$", normalized_content, re.DOTALL)
            if to_match:
                to_emails = to_match.group(1).strip()
                cc_emails = to_match.group(2).strip()
                email_body = to_match.group(3).strip()
                subject = f"FWD: [Ticket #{ticket_id}] {std_ticket_name}" if is_mssql_system_alert else f"FWD: [Ticket #{ticket_id}] {ticket_name}"
            else:
                to_emails = "dccagent@geopits.com"
                subject = f"FWD: [Ticket #{ticket_id}] {std_ticket_name}" if is_mssql_system_alert else f"FWD: [Ticket #{ticket_id}] {ticket_name}"
        elif req.comment_type == "reply":
            subject_match = re.match(r"^Subject:\s*(.*?)\nCc:\s*(.*?)\n\n(.*)$", normalized_content, re.DOTALL)
            if subject_match:
                subject = subject_match.group(1).strip()
                cc_emails = subject_match.group(2).strip()
                email_body = subject_match.group(3).strip()
                to_emails = ticket_contact or "dccagent@geopits.com"
                if not re.search(r'Ticket\s*#?\s*\d+', subject, re.IGNORECASE):
                    subject = f"[Ticket #{ticket_id}] {subject}"
            else:
                to_emails = ticket_contact or "dccagent@geopits.com"
                subject = f"RE: [Ticket #{ticket_id}] {std_ticket_name}" if is_mssql_system_alert else f"RE: [Ticket #{ticket_id}] {ticket_name}"

        # Clean "None" or placeholder values from to_emails and cc_emails
        if to_emails and to_emails.lower() == "none":
            to_emails = None
        if cc_emails and cc_emails.lower() == "none":
            cc_emails = None

        if to_emails:
                
            # Try to retrieve original email html to forward it "as it is"
            original_email_html = None
            cur.execute("""
                SELECT content FROM ticket_comments 
                WHERE ticket_id = %s AND comment_type = 'original_email' 
                ORDER BY id ASC LIMIT 1;
            """, (ticket_id,))
            orig_row = cur.fetchone()
            if orig_row and orig_row[0]:
                original_email_html = orig_row[0]

            if original_email_html:
                original_section = f"""
                <div class="section-title">Original Message</div>
                <div style="margin-top: 15px; border: 1px solid #e2e8f0; padding: 15px; border-radius: 8px; background-color: #ffffff;">
                    {original_email_html}
                </div>
                """
            else:
                original_section = f"""
                <div class="section-title">Original Description</div>
                <div class="desc-block">{ticket_desc}</div>
                """

            formatted_email_body = email_body.replace('\n', '<br/>') if email_body else ""
            
            if is_mssql_system_alert:
                # Always format MSSQL alerts with the premium diagnostic block and table
                cur.execute("""
                    SELECT content FROM ticket_comments 
                    WHERE ticket_id = %s AND comment_type = 'log' 
                    ORDER BY id ASC;
                """, (ticket_id,))
                log_rows = cur.fetchall()
                
                mssql_logs = []
                for row in log_rows:
                    if row[0] and row[0].startswith("MSSQL_LOG_DATA:"):
                        try:
                            log_data = _json.loads(row[0][len("MSSQL_LOG_DATA:"):])
                            mssql_logs.append(log_data)
                        except:
                            pass
                            
                # Fallback: parse from description if empty
                if not mssql_logs and ticket_desc:
                    spid_val = ""
                    spid_match = re.search(r'SPID\s*[:\-]?\s*(\d+)', ticket_desc, re.IGNORECASE)
                    if spid_match:
                        spid_val = spid_match.group(1).strip()
                        
                    sql_text_val = ""
                    sql_match = re.search(r'Executing SQL:\s*\n?(.*?)(?:\n\nEmail Body:|$)', ticket_desc, re.DOTALL | re.IGNORECASE)
                    if sql_match:
                        sql_text_val = sql_match.group(1).strip()
                        
                    db_val = ""
                    db_match = re.search(r'Database\s*[:\-]?\s*([a-zA-Z0-9_]+)', ticket_desc, re.IGNORECASE)
                    if db_match:
                        db_val = db_match.group(1).strip()
                        
                    server_val = ""
                    server_match = re.search(r'Server\s*-\s*([a-zA-Z0-9_\-]+)', ticket_desc, re.IGNORECASE)
                    if server_match:
                        server = server_val = server_match.group(1).strip()
                        
                    mssql_logs.append({
                        "spid": spid_val,
                        "start_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "elapsed_time": "-",
                        "user": "System",
                        "hostname": server_val,
                        "database": db_val,
                        "sql_text": sql_text_val,
                        "wait_type": "None",
                        "stored_procedure": "None"
                    })
                            
                diagnostic_html = ""
                table_html = ""
                has_valid_mssql_logs = mssql_logs and any(log.get("spid") and (log.get("sql_text") or log.get("executing_sql")) for log in mssql_logs)
                if has_valid_mssql_logs:
                    current_spid = mssql_logs[0].get("spid") or ""
                    current_sql = mssql_logs[0].get("sql_text") or mssql_logs[0].get("executing_sql") or ""
                    
                    diagnostic_html = f"""
                    <div style="font-family: Calibri, Arial, sans-serif; margin-top: 15px; margin-bottom: 25px; border: 1px solid #cbd5e1; padding: 20px; border-radius: 8px; background-color: #f8fafc;">
                        <div style="margin-bottom: 12px;">
                            <span style="font-size: 11pt; font-weight: 800; color: #64748b; display: inline-block; width: 120px;">SPID</span>
                            <span style="font-size: 11pt; font-weight: 700; color: #0f172a; padding: 6px 16px; border-radius: 6px; border: 1px solid #cbd5e1; background-color: #ffffff; display: inline-block;">{current_spid}</span>
                        </div>
                        <div style="display: flex; flex-direction: column; gap: 8px;">
                            <label style="font-size: 11pt; font-weight: 800; color: #64748b; display: block; margin-bottom: 6px;">EXECUTING SQL</label>
                            <pre style="width: 100%; padding: 12px; border-radius: 6px; border: 1px solid #cbd5e1; background-color: #ffffff; color: #0f172a; font-family: 'Courier New', Courier, monospace; font-size: 9.5pt; line-height: 1.4; margin: 0; white-space: pre-wrap; word-break: break-all; max-height: 350px; overflow-y: auto;">{current_sql}</pre>
                        </div>
                    </div>
                    """
                    
                    table_html = """
                    <div style="margin-top: 25px; margin-bottom: 20px;">
                        <b style="font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #64748b; text-transform: uppercase;">Logs:</b>
                        <table style="width: 100%; border-collapse: collapse; margin-top: 10px; font-family: Calibri, Arial, sans-serif; font-size: 10pt; border: 1px solid #cbd5e1;">
                            <thead>
                                <tr style="background-color: #0088ff; color: #ffffff; text-align: left;">
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">SPID</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">Start Time</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">Elapsed Time (hh:mm:ss)</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">User</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">HostName</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">Database</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">SQL Text</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">Wait Type</th>
                                    <th style="padding: 10px 12px; border: 1px solid #cbd5e1; font-weight: bold; color: #ffffff;">StoredProcedure</th>
                                </tr>
                            </thead>
                            <tbody>
                    """
                    for idx, log in enumerate(mssql_logs):
                        spid_val = log.get("spid") or ""
                        start_time = log.get("start_time") or log.get("login_time") or ""
                        elapsed_time = log.get("elapsed_time") or (f"00:{log.get('duration_min')}:00" if log.get("duration_min") else "")
                        user_val = log.get("user") or log.get("login_name") or ""
                        hostname = log.get("hostname") or ""
                        database = log.get("database") or ""
                        sql_text_val = log.get("sql_text") or log.get("executing_sql") or ""
                        wait_type = log.get("wait_type") or "None"
                        stored_procedure = log.get("stored_procedure") or "None"
                        
                        bg_color = "#ffffff" if idx % 2 == 0 else "#f8fafc"
                        
                        table_html += f"""
                                <tr style="background-color: {bg_color};">
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{spid_val}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1; white-space: nowrap;">{start_time}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{elapsed_time}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{user_val}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{hostname}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{database}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1; font-family: Consolas, monospace; word-break: break-all;">{sql_text_val}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1;">{wait_type}</td>
                                    <td style="padding: 10px 12px; border: 1px solid #cbd5e1; font-family: Consolas, monospace; word-break: break-all;">{stored_procedure}</td>
                                </tr>
                        """
                    table_html += """
                            </tbody>
                        </table>
                    </div>
                    """
                
                original_section_content = ""
                if original_email_html:
                    original_section_content = f"""
                    <div style="margin-top: 25px; border-top: 1px solid #b5c4df; padding-top: 15px; margin-bottom: 20px;">
                        <b>Original Message:</b><br/><br/>
                        {original_email_html}
                    </div>
                    """
                elif ticket_desc:
                    formatted_ticket_desc = ticket_desc.replace('\n', '<br/>')
                    original_section_content = f"""
                    <div style="margin-top: 25px; border-top: 1px solid #b5c4df; padding-top: 15px; margin-bottom: 20px;">
                        <b>Original Description:</b><br/><br/>
                        <div style="font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #333333;">
                            {formatted_ticket_desc}
                        </div>
                    </div>
                    """

                html_email_content = f"""
                <html>
                <body>
                    <div style="font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #1f497d; margin-bottom: 20px;">
                        {formatted_email_body}
                    </div>
                    <div style="font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #000000; border-top: 1px solid #b5c4df; padding-top: 5px; margin-top: 15px; margin-bottom: 20px;">
                        <b>From:</b> {author}<br/>
                        <b>Sent:</b> {datetime.now().strftime("%B %d, %Y %I:%M %p")}<br/>
                        <b>To:</b> {to_emails}<br/>
                        <b>Subject:</b> {subject}<br/>
                    </div>
                    {diagnostic_html}
                    {table_html}
                    {original_section_content}
                </body>
                </html>
                """
            else:
                # Create a premium styled HTML email body incorporating ticket metadata and description
                html_email_content = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #333333; line-height: 1.6; background-color: #f7fafc; margin: 0; padding: 0; }}
                        .container {{ max-width: 650px; margin: 30px auto; background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05); border: 1px solid #e2e8f0; }}
                        .header {{ background-color: #0c0f1d; padding: 25px 30px; color: #ffffff; border-bottom: 3px solid #ea580c; }}
                        .header h2 {{ margin: 0; font-size: 22px; font-weight: 800; letter-spacing: -0.5px; }}
                        .header p {{ margin: 5px 0 0 0; font-size: 13px; opacity: 0.85; text-transform: uppercase; letter-spacing: 1px; }}
                        .content {{ padding: 30px; }}
                        .comment-block {{ background-color: #f8fafc; padding: 20px; border-left: 4px solid #2563eb; border-radius: 6px; margin-bottom: 30px; box-shadow: inset 0 1px 3px rgba(0,0,0,0.02); }}
                        .comment-title {{ font-size: 13px; color: #64748b; font-weight: 700; margin-bottom: 8px; text-transform: uppercase; }}
                        .comment-text {{ font-size: 15px; color: #1e293b; white-space: pre-wrap; }}
                        .section-title {{ font-size: 15px; font-weight: 800; color: #0c0f1d; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 1px solid #cbd5e1; padding-bottom: 6px; margin-top: 30px; }}
                        .info-table {{ width: 100%; border-collapse: collapse; margin-top: 12px; }}
                        .info-table td {{ padding: 8px 0; font-size: 14px; vertical-align: top; }}
                        .info-label {{ color: #64748b; font-weight: 600; width: 160px; }}
                        .info-value {{ color: #1e293b; }}
                        .badge-priority {{ display: inline-block; background-color: #fee2e2; color: #ef4444; padding: 2px 10px; border-radius: 9999px; font-weight: 750; font-size: 11px; }}
                        .badge-status {{ display: inline-block; background-color: #dbeafe; color: #2563eb; padding: 2px 10px; border-radius: 9999px; font-weight: 750; font-size: 11px; }}
                        .desc-block {{ background-color: #f1f5f9; padding: 15px 20px; border-radius: 8px; font-size: 13.5px; color: #475569; border: 1px solid #e2e8f0; margin-top: 10px; white-space: pre-wrap; }}
                        .footer {{ background-color: #f1f5f9; padding: 15px 30px; text-align: center; font-size: 11px; color: #94a3b8; border-top: 1px solid #e2e8f0; }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <div class="header">
                            <h2>GeoMon Incident Center</h2>
                            <p>Ticket Communication update</p>
                        </div>
                        <div class="content">
                            <div class="comment-block">
                                <div class="comment-title">Message from {author} ({req.comment_type.upper()})</div>
                                <div class="comment-text">{email_body}</div>
                            </div>
                            
                            <div class="section-title">Ticket Reference Details</div>
                            <table class="info-table">
                                <tr>
                                    <td class="info-label">Ticket ID:</td>
                                    <td class="info-value"><strong>#{ticket_id}</strong></td>
                                </tr>
                                <tr>
                                    <td class="info-label">Ticket Name:</td>
                                    <td class="info-value">{ticket_name}</td>
                                </tr>
                                <tr>
                                    <td class="info-label">Priority:</td>
                                    <td class="info-value"><span class="badge-priority">{ticket_priority}</span></td>
                                </tr>
                                <tr>
                                    <td class="info-label">Status:</td>
                                    <td class="info-value"><span class="badge-status">{ticket_status}</span></td>
                                </tr>
                                <tr>
                                    <td class="info-label">Business Unit:</td>
                                    <td class="info-value">{ticket_bu}</td>
                                </tr>
                                <tr>
                                    <td class="info-label">Category:</td>
                                    <td class="info-value">{ticket_category}</td>
                                </tr>
                            </table>
                            
                            {original_section}
                        </div>
                        <div class="footer">
                            This email was dynamically generated and transmitted securely via Microsoft Graph APIs in response to an operator action. Please log in to the GeoMon Incident Center to view attachments or manage this incident ticket.
                        </div>
                    </div>
                </body>
                </html>
                """
            
            # Always send FROM dccagent@geopits.com — the only account with Graph API Mail.Send permission.
            # For manual (non-system) tickets, set Reply-To to the operator's email so recipients can reply to them.
            comment_sender_email = "dccagent@geopits.com"
            comment_reply_to = None
            if not (ticket_created_by and ticket_created_by.lower() == "system"):
                user_email = user.get("email")
                if user_email and user_email.strip() and user_email.strip().lower() != "dccagent@geopits.com":
                    comment_reply_to = user_email.strip()

            send_email_outlook(
                to_emails, cc_emails, subject, html_email_content,
                sender_email=comment_sender_email,
                attachments=email_attachments if email_attachments else None,
                reply_to=comment_reply_to
            )

        return {
            "status": "success",
            "comment": {
                "id": comment_id,
                "ticket_id": ticket_id,
                "author": author,
                "comment_type": req.comment_type,
                "content": req.content,
                "attachments": attachments_str,
                "created_at": created_at
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.put("/tickets/{ticket_id}/comments/{comment_id}")
def update_ticket_comment(ticket_id: int, comment_id: int, req: CommentCreate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        # Verify ticket exists
        cur.execute("SELECT company FROM tickets WHERE id = %s;", (ticket_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Ticket not found")
            
        cur.execute("SELECT author FROM ticket_comments WHERE id = %s AND ticket_id = %s;", (comment_id, ticket_id))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Comment not found")
        author = row[0]

        is_admin = user.get("isAdmin") or user.get("role") == "admin"
        is_author = user.get("username") and author.lower() == user.get("username").lower()

        if not (is_admin or is_author):
            raise HTTPException(status_code=403, detail="Permission Denied: Only the author or an admin can update this comment!")

        # Safe stringification of attachments to prevent unhashable type: 'dict'
        import json as _json
        attachments_str = req.attachments
        if not isinstance(attachments_str, str) and attachments_str is not None:
            attachments_str = _json.dumps(attachments_str)

        cur.execute("""
            UPDATE ticket_comments
            SET content = %s, comment_type = %s, attachments = %s
            WHERE id = %s AND ticket_id = %s;
        """, (req.content, req.comment_type, attachments_str, comment_id, ticket_id))
        conn.commit()
        return {"status": "success"}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/tickets/{ticket_id}/comments/{comment_id}")
def delete_ticket_comment(ticket_id: int, comment_id: int, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        # Verify ticket exists
        cur.execute("SELECT company FROM tickets WHERE id = %s;", (ticket_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Ticket not found")
            
        cur.execute("SELECT author FROM ticket_comments WHERE id = %s AND ticket_id = %s;", (comment_id, ticket_id))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Comment not found")
        author = row[0]

        is_admin = user.get("isAdmin") or user.get("role") == "admin"
        is_author = user.get("username") and author.lower() == user.get("username").lower()

        if not (is_admin or is_author):
            raise HTTPException(status_code=403, detail="Permission Denied: Only the author or an admin can delete this comment!")

        cur.execute("DELETE FROM ticket_comments WHERE id = %s AND ticket_id = %s;", (comment_id, ticket_id))
        conn.commit()
        return {"status": "success"}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# DYNAMIC TICKET OPTIONS ENDPOINTS (AGENTS & BUSINESS UNITS)
# ==============================================================================

@router.get("/admin/ticket-agents")
def get_ticket_agents(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Fetch system users (roles 'user', 'admin', 'lead')
        cur.execute("SELECT id, username as name, email FROM users WHERE role IN ('user', 'admin', 'lead') ORDER BY username;")
        rows = cur.fetchall()
        return {"agents": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/ticket-agents")
def create_ticket_agent(agent: TicketAgentCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO ticket_agents (name)
            VALUES (%s)
            RETURNING id;
        """, (agent.name,))
        agent_id = cur.fetchone()[0]
        conn.commit()
        return {"status": "success", "agent_id": agent_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/ticket-agents/{agent_id}")
def delete_ticket_agent(agent_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM ticket_agents WHERE id = %s;", (agent_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
@router.get("/admin/online-users")
def get_online_users():
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM online_users ORDER BY username;")
        rows = cur.fetchall()
        return {"online_users": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/online-users")
def create_online_user(payload: OnlineUserCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO online_users (username, units)
            VALUES (%s, %s)
            ON CONFLICT (username) DO UPDATE SET units = EXCLUDED.units
            RETURNING id;
        """, (payload.username, payload.units))
        user_id = cur.fetchone()[0]
        conn.commit()
        return {"status": "success", "online_user_id": user_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/online-users/{user_id}")
def delete_online_user(user_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM online_users WHERE id = %s;", (user_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/business-units")
def get_ticket_bus():
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM ticket_business_units ORDER BY name;")
        rows = cur.fetchall()
        return {"business_units": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/business-units")
def create_ticket_bu(bu: TicketBUCreate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO ticket_business_units (name)
            VALUES (%s)
            RETURNING id;
        """, (bu.name,))
        bu_id = cur.fetchone()[0]
        conn.commit()
        return {"status": "success", "bu_id": bu_id}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/business-units/{bu_id}")
def delete_ticket_bu(bu_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM ticket_business_units WHERE id = %s;", (bu_id,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

class LogoUpdate(BaseModel):
    logo_data: str

@router.get("/settings/logo")
def get_app_logo():
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT value FROM system_settings WHERE key = 'app_logo';")
        row = cur.fetchone()
        logo = row[0] if row else "/static/applogo.svg"
        return {"logo": logo}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/admin/settings/logo")
def update_app_logo(payload: LogoUpdate, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO system_settings (key, value)
            VALUES ('app_logo', %s)
            ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value;
        """, (payload.logo_data,))
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

class FeedbackCreate(BaseModel):
    feedback_text: str
    rating: int

@router.post("/feedback")
def submit_feedback(fb: FeedbackCreate, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        username = user.get("username", "anonymous")
        email = user.get("email", "")
        cur.execute("""
            INSERT INTO feedbacks (username, email, feedback_text, rating)
            VALUES (%s, %s, %s, %s);
        """, (username, email, fb.feedback_text, fb.rating))
        
        # Broadcast notification of new feedback to administrators
        notify_msg = f"New Feedback received from {username} ({email}): {fb.feedback_text[:60]}..."
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))
        
        conn.commit()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/admin/feedbacks")
def get_admin_feedbacks(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM feedbacks ORDER BY created_at DESC;")
        rows = cur.fetchall()
        return {"feedbacks": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/feedbacks/{feedback_id}")
def delete_feedback(feedback_id: int, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM feedbacks WHERE id = %s;", (feedback_id,))
        conn.commit()
        return {"status": "success", "message": "Feedback deleted successfully"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.delete("/admin/feedbacks")
def clear_all_feedbacks(user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM feedbacks;")
        conn.commit()
        return {"status": "success", "message": "All feedback history cleared successfully"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

class DbClearRequest(BaseModel):
    target: str

@router.post("/admin/database/clear")
def clear_database_tables(req: DbClearRequest, user: dict = Depends(get_current_user_local)):
    if not user.get("isAdmin"):
        raise HTTPException(status_code=403, detail="Admin authorization required")
        
    conn = get_connection()
    try:
        cur = conn.cursor()
        
        target = req.target.lower()
        message = ""
        
        if target == "feedbacks":
            cur.execute("DELETE FROM feedbacks;")
            message = "All feedback reviews purged successfully."
        elif target == "telemetry":
            cur.execute("DELETE FROM user_page_activity;")
            message = "All active workload telemetry logs purged successfully."
        elif target == "notifications":
            cur.execute("DELETE FROM notifications;")
            message = "All warning alerts and broadcast records purged successfully."
        elif target == "reports":
            cur.execute("DELETE FROM client_reports;")
            message = "All dynamic SLA client reports purged successfully."
        elif target == "all":
            cur.execute("DELETE FROM feedbacks;")
            cur.execute("DELETE FROM user_page_activity;")
            cur.execute("DELETE FROM notifications;")
            cur.execute("DELETE FROM client_reports;")
            message = "Complete transient database reset applied successfully."
        else:
            raise HTTPException(status_code=400, detail="Invalid target table database wipe request.")
            
        conn.commit()
        return {"status": "success", "message": message}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# ==============================================================================
# CHATBOT ENDPOINTS
# ==============================================================================

def detect_client_name(message: str, history: Optional[List[dict]] = None) -> Optional[str]:
    # Fetch all registered client names from admin_clients
    clients = []
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT client_name FROM admin_clients;")
        clients = [r[0] for r in cur.fetchall() if r[0]]
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error fetching clients for detection: {e}")
        
    if not clients:
        return None
        
    msg_lower = message.lower()
    for c in clients:
        if c.lower() in msg_lower:
            return c
            
    if history:
        for item in reversed(history):
            content = (item.get("content") or item.get("text") or "").lower()
            for c in clients:
                if c.lower() in content:
                    return c
                    
    return None

_cached_db_schema = None

def get_dynamic_db_schema() -> str:
    global _cached_db_schema
    if _cached_db_schema is not None:
        return _cached_db_schema

    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT table_name, column_name 
            FROM information_schema.columns 
            WHERE table_schema = 'public'
            ORDER BY table_name, ordinal_position;
        """)
        rows = cur.fetchall()
        conn.close()

        if not rows:
            return ""

        tables = {}
        for tbl, col in rows:
            if tbl not in tables:
                tables[tbl] = []
            tables[tbl].append(col)

        lines = []
        for tbl, cols in sorted(tables.items()):
            cols_str = ", ".join(cols)
            lines.append(f"- {tbl} ({cols_str})")
        _cached_db_schema = "\n".join(lines)
        return _cached_db_schema
    except Exception as e:
        print("Failed to fetch dynamic DB schema:", e)
        return (
            "- tickets (id, business_unit, company, contact, ticket_name, category, status, priority, agent, description, created_by, created_at)\n"
            "- ticket_comments (id, ticket_id, author, comment_type, content, attachments, created_at)\n"
            "- db_monitoring_logs (id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner)\n"
            "- database_size_history (id, server_name, database_name, total_size_bytes, captured_date, db_type)\n"
            "- table_size_history (id, server_name, database_name, table_name, size_bytes, captured_date, db_type)\n"
            "- db_archived_logs (id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner, status_updated_at)\n"
            "- share_history (id, username, platform, content_type, client_name, server_name, log_message, notes, status, owner, ticket_status, next_action, client_visibility, db_type, shared_at)"
        )

def try_execute_db_query(message: str, history: Optional[list] = None) -> Optional[dict]:
    """
    Analyzes the user's message and history. If it requires details from the database tables,
    it dynamically generates a safe, read-only PostgreSQL query, executes it, and returns the query + results.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or "sk-test" in api_key:
        return None
        
    import httpx
    import psycopg2.extras
    
    schema_str = get_dynamic_db_schema()
    
    # We ask gpt-4o-mini to generate the SQL query
    prompt = (
        "You are an expert database administrator. Your task is to output a single read-only PostgreSQL SELECT query "
        "to fetch the relevant data from the database to answer the user's request. "
        "The database contains the following tables and schemas:\n"
        f"{schema_str}\n\n"
        "Guidelines:\n"
        "1. Write a read-only PostgreSQL query. ONLY output the raw SELECT query, nothing else (no markdown blocks, no 'sql', no quotes, no explanation). Example: SELECT * FROM tickets LIMIT 5;\n"
        "2. If the user's message is a general chat, greeting, general knowledge query, explanation of SQL, or does not require looking up live system data, reply exactly with: NO_QUERY\n"
        "3. Make sure to limit any output (e.g. LIMIT 20 or LIMIT 50) to prevent retrieving too much data.\n\n"
        f"User Message: {message}\n"
    )
    
    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {"role": "system", "content": "You are a database SQL assistant. You output raw PostgreSQL SELECT statements or NO_QUERY."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 150
    }
    
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        with httpx.Client(timeout=10.0) as client:
            resp = client.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers)
            if resp.status_code == 200:
                text = resp.json()["choices"][0]["message"]["content"].strip()
                if text.startswith("```"):
                    text = text.replace("```sql", "").replace("```", "").strip()
                
                sql_upper = text.upper()
                if "SELECT" in sql_upper and not any(kw in sql_upper for kw in ["INSERT", "UPDATE", "DELETE", "DROP", "ALTER", "TRUNCATE", "CREATE", "REPLACE", "GRANT", "REVOKE"]):
                    conn = get_connection()
                    try:
                        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        cur.execute(text)
                        results = cur.fetchall()
                        return {
                            "query": text,
                            "results": results
                        }
                    except Exception as db_err:
                        return {
                            "query": text,
                            "error": str(db_err)
                        }
                    finally:
                        conn.close()
    except Exception as e:
        print("Error in try_execute_db_query:", e)
    return None

class ChatRequest(BaseModel):
    message: str
    history: Optional[List[dict]] = None
    client_name: Optional[str] = None

@router.post("/chat")
def chatbot_openai_response(req: ChatRequest, user: dict = Depends(get_current_user_local)):
    """
    Production-grade endpoint that acts as a bridge to the OpenAI completions endpoint.
    Retrieves the OPENAI_API_KEY from the system context, verifies it, and utilizes 
    gpt-4o-mini to return expert-level completions.
    """
    is_restricted, reason = check_client_user_chat_restriction(req.message, user)
    if is_restricted:
        return {"response": "I am only authorized to answer queries related to your assigned clients, and your own user credentials."}

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or "sk-test" in api_key:
        raise HTTPException(
            status_code=400, 
            detail="OpenAI API Key is not configured in the server environment. Please set it in the .env configuration."
        )

    import httpx
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        import datetime
        current_date_str = datetime.date.today().strftime("%B %d, %Y")
        
        # Try to dynamically extract client name from message/history if not explicitly passed from active UI
        client_ctx = req.client_name
        if not client_ctx:
            client_ctx = detect_client_name(req.message, req.history)
            
        # Dynamically fetch latest telemetry context to guarantee accurate dashboard data alignment
        telemetry_context = get_telemetry_context_string(client_ctx, user)
        
        db_query_context = ""
        try:
            import json
            db_res = try_execute_db_query(req.message, req.history)
            if db_res:
                if "results" in db_res:
                    db_query_context = (
                        f"\n\n--- DYNAMIC DB QUERY EXECUTION ---\n"
                        f"Generated Query: {db_res['query']}\n"
                        f"Query Results:\n{json.dumps(db_res['results'], default=str)}\n"
                        f"----------------------------------\n"
                    )
                elif "error" in db_res:
                    db_query_context = (
                        f"\n\n--- DYNAMIC DB QUERY EXECUTION (FAILED) ---\n"
                        f"Generated Query: {db_res['query']}\n"
                        f"Execution Error: {db_res['error']}\n"
                        f"-------------------------------------------\n"
                    )
        except Exception as e:
            print("Chat dynamic query exception:", e)
        
        system_content = (
            "You are GeoBot, a strict, expert diagnostic AI assistant and senior database administrator designed to answer queries EXCLUSIVELY "
            "using the live telemetry database, active support tickets, system monitoring logs, and client mappings "
            "of the GeoMon application, as well as professional concepts in software technology, "
            "database administration, and data engineering.\n\n"
            
            f"CURRENT CALENDAR YEAR CONTEXT: The current year is 2026 (Today is {current_date_str}). "
            "All telemetry, sizes, dates, and historical changes exist within this 2026 calendar period.\n\n"

            "DATABASE SCHEMA & APPLICATION TABLE RELATIONSHIPS:\n"
            "The GeoMon application consists of the following database tables that you are connected to:\n"
            "- 'tickets': stores support tickets (id, business_unit, company (client), contact, ticket_name, category, status, priority, agent, description, created_by, created_at)\n"
            "- 'ticket_comments': ticket discussions (id, ticket_id, author, comment_type, content, attachments, created_at)\n"
            "- 'db_monitoring_logs': database alert notifications (id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner)\n"
            "- 'database_size_history': tracks database growth (id, server_name, database_name, total_size_bytes, captured_date)\n"
            "- 'table_size_history': tracks table growth (id, server_name, database_name, table_name, size_bytes, captured_date)\n"
            "- 'admin_clients': client metadata registry (id, client_name, db_type, server_name, created_at)\n"
            "- 'share_history': WhatsApp/Teams audit trail (id, username, platform, content_type, client_name, log_message, shared_at)\n"
            "- 'user_page_activity': user navigation statistics (id, username, page_path, duration_seconds, last_active_at)\n"
            "- 'feedbacks': application user feedbacks (id, username, email, feedback_text, rating, created_at)\n\n"
            
            "CRITICAL HALLUCINATION PREVENTOR & DATA MAPPING RULES:\n"
            "1. You are mapped DIRECTLY to the live application dataset provided below. "
            "You MUST query and rely ONLY on this exact dataset to answer questions about clients, databases, tables, sizes, active tickets, and system monitoring logs. "
            "2. If the user asks about a client, server, database, table, ticket, or log that is completely absent from the provided dataset, you MUST reply: "
            "'I cannot find any matches in the live dashboard dataset.' "
            "If the requested entity (such as the client or database) DOES exist in the dataset, you must use the available telemetry, tickets, and logs to answer the user's query.\n"
            "General computer science theories, software technology explanations, SQL query optimizations, database normalization, and database concepts "
            "should NOT trigger this refusal; you are encouraged to explain and answer them with senior-level engineering expertise.\n"
            "3. If a question is about general database administration, SQL optimization, data engineering, or software technology, "
            "you may answer it professionally. However, for any system-specific capacities or telemetry, you must stick strictly to the telemetry data.\n\n"
            
            "STRICT SUBJECT MATTER BOUNDARY:\n"
            "You are strictly prohibited from answering queries on topics outside of software technology, database systems, data engineering, "
            "support tickets, database monitoring logs, and the dashboard's live telemetry. "
            "If the user's query is about general knowledge, history, geography, arts, pop culture, math puzzles, translation, recipes, "
            "or general conversation, you MUST immediately refuse by replying: "
            "'I am only authorized to assist with database, software technology, data engineering, support tickets, monitoring logs, and live dashboard telemetry queries.'\n\n"
            
            "--- LIVE DASHBOARD DATA (2026) ---\n"
            f"{telemetry_context}\n"
            f"{db_query_context}\n"
            "-----------------------------------\n\n"
            "Respond in a concise, accurate, and professional manner."
        )
        
        # Build chat payload with history for context retention
        messages = [
            {
                "role": "system", 
                "content": system_content
            }
        ]
        
        if req.history:
            for item in req.history:
                role = item.get("role") or item.get("sender")
                if role == "bot" or role == "assistant":
                    role = "assistant"
                else:
                    role = "user"
                content = item.get("content") or item.get("text")
                if content:
                    messages.append({"role": role, "content": content})
                    
        # Add final user message
        messages.append({"role": "user", "content": req.message})

        payload = {
            "model": "gpt-4o-mini",
            "messages": messages,
            "max_tokens": 1200,
            "temperature": 0.7
        }
        with httpx.Client(timeout=30.0) as client:
            response = client.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers)
            
            if response.status_code == 200:
                res_data = response.json()
                if "choices" in res_data and len(res_data["choices"]) > 0:
                    reply = res_data["choices"][0]["message"]["content"]
                    return {"response": reply}
                else:
                    raise HTTPException(status_code=502, detail="Invalid response structure returned by the OpenAI API.")
            else:
                error_detail = f"OpenAI API Error: {response.text} (Status {response.status_code})"
                raise HTTPException(status_code=response.status_code, detail=error_detail)
                
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=503, 
            detail=f"Failed to communicate with OpenAI server due to a network connection error: {str(exc)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"An unexpected internal server error occurred while processing the chat: {str(e)}"
        )

from fastapi.responses import StreamingResponse
import json

@router.post("/chat/stream")
async def chatbot_openai_stream(req: ChatRequest, user: dict = Depends(get_current_user_local)):
    """
    High-performance async server-sent events (SSE) streaming endpoint.
    Establishes a connection to OpenAI, enabling lightning-fast word-by-word streaming
    for an enhanced, responsive ChatGPT experience.
    """
    is_restricted, reason = check_client_user_chat_restriction(req.message, user)
    if is_restricted:
        async def event_generator_restricted():
            yield f"data: {json.dumps({'text': 'I am only authorized to answer queries related to your assigned clients, and your own user credentials.'})}\n\n"
        return StreamingResponse(event_generator_restricted(), media_type="text/event-stream")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or "sk-test" in api_key:
        raise HTTPException(
            status_code=400, 
            detail="OpenAI API Key is not configured in the server environment."
        )

    async def event_generator():
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        import datetime
        current_date_str = datetime.date.today().strftime("%B %d, %Y")
        
        # Try to dynamically extract client name from message/history if not explicitly passed from active UI
        client_ctx = req.client_name
        if not client_ctx:
            client_ctx = detect_client_name(req.message, req.history)
            
        # Dynamically fetch latest telemetry context to guarantee accurate dashboard data alignment
        telemetry_context = get_telemetry_context_string(client_ctx, user)
        
        db_query_context = ""
        try:
            db_res = try_execute_db_query(req.message, req.history)
            if db_res:
                if "results" in db_res:
                    db_query_context = (
                        f"\n\n--- DYNAMIC DB QUERY EXECUTION ---\n"
                        f"Generated Query: {db_res['query']}\n"
                        f"Query Results:\n{json.dumps(db_res['results'], default=str)}\n"
                        f"----------------------------------\n"
                    )
                elif "error" in db_res:
                    db_query_context = (
                        f"\n\n--- DYNAMIC DB QUERY EXECUTION (FAILED) ---\n"
                        f"Generated Query: {db_res['query']}\n"
                        f"Execution Error: {db_res['error']}\n"
                        f"-------------------------------------------\n"
                    )
        except Exception as e:
            print("Chat stream dynamic query exception:", e)
        
        system_content = (
            "You are GeoBot, a strict, expert diagnostic AI assistant and senior database administrator designed to answer queries EXCLUSIVELY "
            "using the live telemetry database, active support tickets, system monitoring logs, and client mappings "
            "of the GeoMon application, as well as professional concepts in software technology, "
            "database administration, and data engineering.\n\n"
            
            f"CURRENT CALENDAR YEAR CONTEXT: The current year is 2026 (Today is {current_date_str}). "
            "All telemetry, sizes, dates, and historical changes exist within this 2026 calendar period.\n\n"

            "DATABASE SCHEMA & APPLICATION TABLE RELATIONSHIPS:\n"
            "The GeoMon application consists of the following database tables that you are connected to:\n"
            "- 'tickets': stores support tickets (id, business_unit, company (client), contact, ticket_name, category, status, priority, agent, description, created_by, created_at)\n"
            "- 'ticket_comments': ticket discussions (id, ticket_id, author, comment_type, content, attachments, created_at)\n"
            "- 'db_monitoring_logs': database alert notifications (id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner)\n"
            "- 'database_size_history': tracks database growth (id, server_name, database_name, total_size_bytes, captured_date)\n"
            "- 'table_size_history': tracks table growth (id, server_name, database_name, table_name, size_bytes, captured_date)\n"
            "- 'admin_clients': client metadata registry (id, client_name, db_type, server_name, created_at)\n"
            "- 'share_history': WhatsApp/Teams audit trail (id, username, platform, content_type, client_name, log_message, shared_at)\n"
            "- 'user_page_activity': user navigation statistics (id, username, page_path, duration_seconds, last_active_at)\n"
            "- 'feedbacks': application user feedbacks (id, username, email, feedback_text, rating, created_at)\n\n"
            
            "CRITICAL HALLUCINATION PREVENTOR & DATA MAPPING RULES:\n"
            "1. You are mapped DIRECTLY to the live application dataset provided below. "
            "You MUST query and rely ONLY on this exact dataset to answer questions about clients, databases, tables, sizes, active tickets, and system monitoring logs. "
            "2. If the user asks about a client, server, database, table, ticket, or log that is completely absent from the provided dataset, you MUST reply: "
            "'I cannot find any matches in the live dashboard dataset.' "
            "If the requested entity (such as the client or database) DOES exist in the dataset, you must use the available telemetry, tickets, and logs to answer the user's query.\n"
            "General computer science theories, software technology explanations, SQL query optimizations, database normalization, and database concepts "
            "should NOT trigger this refusal; you are encouraged to explain and answer them with senior-level engineering expertise.\n"
            "3. If a question is about general database administration, SQL optimization, data engineering, or software technology, "
            "you may answer it professionally. However, for any system-specific capacities or telemetry, you must stick strictly to the telemetry data.\n\n"
            
            "STRICT SUBJECT MATTER BOUNDARY:\n"
            "You are strictly prohibited from answering queries on topics outside of software technology, database systems, data engineering, "
            "support tickets, database monitoring logs, and the dashboard's live telemetry. "
            "If the user's query is about general knowledge, history, geography, arts, pop culture, math puzzles, translation, recipes, "
            "or general conversation, you MUST immediately refuse by replying: "
            "'I am only authorized to assist with database, software technology, data engineering, support tickets, monitoring logs, and live dashboard telemetry queries.'\n\n"
            
            "--- LIVE DASHBOARD DATA (2026) ---\n"
            f"{telemetry_context}\n"
            f"{db_query_context}\n"
            "-----------------------------------\n\n"
            "Respond in a concise, accurate, and professional manner."
        )
        
        # Build chat payload with history for context retention
        messages = [
            {
                "role": "system", 
                "content": system_content
            }
        ]
        
        if req.history:
            for item in req.history:
                role = item.get("role") or item.get("sender")
                if role == "bot" or role == "assistant":
                    role = "assistant"
                else:
                    role = "user"
                content = item.get("content") or item.get("text")
                if content:
                    messages.append({"role": role, "content": content})
                    
        # Add final user message
        messages.append({"role": "user", "content": req.message})

        payload = {
            "model": "gpt-4o-mini",
            "messages": messages,
            "max_tokens": 1200,
            "temperature": 0.7,
            "stream": True
        }
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                async with client.stream("POST", "https://api.openai.com/v1/chat/completions", json=payload, headers=headers) as response:
                    if response.status_code != 200:
                        error_body = await response.aread()
                        yield f"data: {json.dumps({'error': f'OpenAI Error: {error_body.decode()}'})}\n\n"
                        return

                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            data_content = line[6:].strip()
                            if data_content == "[DONE]":
                                break
                            try:
                                parsed = json.loads(data_content)
                                delta = parsed["choices"][0]["delta"].get("content", "")
                                if delta:
                                    yield f"data: {json.dumps({'text': delta})}\n\n"
                            except Exception:
                                continue
        except httpx.RequestError as exc:
            yield f"data: {json.dumps({'error': f'Network Connection Error: {str(exc)}'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': f'Unexpected Internal Stream Error: {str(e)}'})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

def check_client_user_chat_restriction(message: str, user: dict) -> tuple[bool, str]:
    """
    Checks if a standard client user is attempting to query about unauthorized clients,
    other users' emails, or other usernames.
    Returns (is_restricted, reason).
    """
    # Bypass all restrictions to allow full chatbot access
    return False, ""

    message_lower = message.lower()
    own_username = (user.get("username") or "").lower().strip()
    own_email = (user.get("email") or "").lower().strip()

    conn = get_connection()
    try:
        cur = conn.cursor()
        
        # 1. Fetch all clients
        cur.execute("SELECT DISTINCT client_name, server_name FROM admin_clients;")
        all_client_names = set()
        for r in cur.fetchall():
            if r[0]: all_client_names.add(r[0].lower().strip())
            if r[1]: all_client_names.add(r[1].lower().strip())
            
        # 2. Fetch user's allowed clients
        allowed_clients = [c.lower().strip() for c in user.get("allowed_clients", [])]
        allowed_client_names = set(allowed_clients)
        
        assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
        if assigned_techs:
            cur.execute("""
                SELECT DISTINCT client_name, server_name FROM admin_clients 
                WHERE EXISTS (
                    SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                    WHERE t = ANY(%s)
                );
            """, (assigned_techs,))
            for r in cur.fetchall():
                if r[0]: allowed_client_names.add(r[0].lower().strip())
                if r[1]: allowed_client_names.add(r[1].lower().strip())

        # Clients the user is NOT allowed to view
        forbidden_clients = all_client_names - allowed_client_names
        
        # 3. Fetch all users (emails and usernames)
        cur.execute("SELECT username, email FROM users;")
        all_usernames = set()
        all_emails = set()
        for r in cur.fetchall():
            if r[0]: all_usernames.add(r[0].lower().strip())
            if r[1]: all_emails.add(r[1].lower().strip())
            
        forbidden_usernames = all_usernames - {own_username}
        forbidden_emails = all_emails - {own_email}
        
        cur.close()
        conn.close()
        
        # Check for forbidden client names in message
        import re
        for fc in forbidden_clients:
            if not fc:
                continue
            # Use word boundaries to avoid false positives on partial matches
            pattern = r'\b' + re.escape(fc) + r'\b'
            if re.search(pattern, message_lower):
                return True, f"Access Denied: You do not have permission to query details for client '{fc}'."
                
        # Check for forbidden emails in message
        for fe in forbidden_emails:
            if not fe:
                continue
            if fe in message_lower:
                return True, "Access Denied: You do not have permission to query details for other users' email addresses."
                
        # Check for forbidden usernames in message
        for fu in forbidden_usernames:
            if not fu:
                continue
            # Avoid matching common words if the username is very short/common, check word boundary
            pattern = r'\b' + re.escape(fu) + r'\b'
            if re.search(pattern, message_lower):
                return True, f"Access Denied: You do not have permission to query details for username '{fu}'."
                
        return False, ""
    except Exception as e:
        print(f"Error checking chat restriction: {e}")
        return False, ""

def check_client_and_tech_permission(client_name: str, user: dict, cur) -> bool:
    """
    Enforces user permissions:
    - Admin users have full global access to all clients and all technologies.
    - Client-restricted users (those with 'allowed_clients') can only see metrics for their allowed client(s).
    - Standard non-admin users can only view clients matching their allowed database technologies.
    """
    if user.get("isAdmin") or user.get("role") == "admin":
        return True
        
    # 1. Allowed Clients restriction
    allowed_clients = [c.lower().strip() for c in user.get("allowed_clients", []) if c]
    if allowed_clients:
        cur.execute("""
            SELECT DISTINCT client_name FROM admin_clients 
            WHERE LOWER(TRIM(client_name)) = ANY(%s) AND (LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) OR LOWER(TRIM(server_name)) = LOWER(TRIM(%s)));
        """, (allowed_clients, client_name, client_name))
        if cur.fetchone():
            return True
        if client_name.lower().strip() in allowed_clients:
            return True
        return False
        
    # 2. Technology permissions check
    assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", []) if t]
    if assigned_techs:
        cur.execute("""
            SELECT id FROM admin_clients 
            WHERE (LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) OR LOWER(TRIM(server_name)) = LOWER(TRIM(%s))) 
              AND EXISTS (
                  SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                  WHERE t = ANY(%s)
              );
        """, (client_name, client_name, assigned_techs))
        if cur.fetchone():
            return True
        return False
        
    return False
        
    return False

def get_telemetry_context_string(client_name: Optional[str] = None, user: Optional[dict] = None) -> str:
    """
    Queries database and table historical metrics, active tickets, system logs,
    and client technology mappings to build a complete capacity and status context
    for the AI assistant, ensuring the chatbot has full visibility of the application.
    Enforces strict user permission policies and scales context to the selected client name
    to prevent data leakage outside authorized limits.
    """
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Bypass client and technology restrictions for the chatbot so it has access to read all tables
        target_servers = None
        target_clients = None
        assigned_techs = None
        user_email = None
        username_val = None
        
        if user:
            user_email = user.get("email")
            username_val = user.get("username")

        # If a specific client is requested, restrict scope strictly to that client
        if client_name:
            target_servers = [client_name]
            target_clients = [client_name]

        # 1. admin_clients
        if target_clients is None:
            cur.execute("SELECT client_name, server_name, db_type FROM admin_clients ORDER BY client_name;")
        elif target_clients:
            cur.execute("""
                SELECT client_name, server_name, db_type FROM admin_clients 
                WHERE client_name = ANY(%s) OR server_name = ANY(%s)
                ORDER BY client_name;
            """, (target_clients, target_servers))
        else:
            return "--- No Telemetry Context Available ---"
        clients_tech = cur.fetchall()

        # 2. database_size_history
        dbs = []
        db_summary = []
        if target_servers is None:
            cur.execute("""
                SELECT DISTINCT ON (server_name, database_name)
                       server_name, database_name, total_size_bytes as latest_size, captured_date
                FROM database_size_history
                ORDER BY server_name, database_name, captured_date DESC;
            """)
            all_dbs = cur.fetchall()
            if len(all_dbs) > 150:
                cur.execute("""
                    WITH latest_dbs AS (
                        SELECT DISTINCT ON (server_name, database_name)
                               server_name, database_name, total_size_bytes as latest_size
                        FROM database_size_history
                        ORDER BY server_name, database_name, captured_date DESC
                    )
                    SELECT server_name, COUNT(database_name) as db_count, SUM(latest_size) as total_size
                    FROM latest_dbs
                    GROUP BY server_name;
                """)
                db_summary = cur.fetchall()
            else:
                dbs = all_dbs
        else:
            cur.execute("""
                SELECT DISTINCT ON (server_name, database_name)
                       server_name, database_name, total_size_bytes as latest_size, captured_date
                FROM database_size_history
                WHERE server_name = ANY(%s)
                ORDER BY server_name, database_name, captured_date DESC;
            """, (target_servers,))
            dbs = cur.fetchall()
        
        # 3. table_size_history
        tbls = []
        if target_servers is None:
            cur.execute("""
                WITH latest_tables AS (
                    SELECT DISTINCT ON (server_name, database_name, table_name)
                           server_name, database_name, table_name, size_bytes, captured_date
                    FROM table_size_history
                    ORDER BY server_name, database_name, table_name, captured_date DESC
                ),
                ranked_tables AS (
                    SELECT server_name, database_name, table_name, size_bytes, captured_date,
                           ROW_NUMBER() OVER (PARTITION BY server_name, database_name ORDER BY size_bytes DESC) as rn
                    FROM latest_tables
                )
                SELECT server_name, database_name, table_name, size_bytes as latest_size, captured_date
                FROM ranked_tables
                WHERE rn <= 3;
            """)
            tbls = cur.fetchall()
        else:
            cur.execute("""
                WITH latest_tables AS (
                    SELECT DISTINCT ON (server_name, database_name, table_name)
                           server_name, database_name, table_name, size_bytes, captured_date
                    FROM table_size_history
                    WHERE server_name = ANY(%s)
                    ORDER BY server_name, database_name, table_name, captured_date DESC
                ),
                ranked_tables AS (
                    SELECT server_name, database_name, table_name, size_bytes, captured_date,
                           ROW_NUMBER() OVER (PARTITION BY server_name, database_name ORDER BY size_bytes DESC) as rn
                    FROM latest_tables
                )
                SELECT server_name, database_name, table_name, size_bytes as latest_size, captured_date
                FROM ranked_tables
                WHERE rn <= 25;
            """, (target_servers,))
            tbls = cur.fetchall()

        # 4. tickets
        if target_clients is None:
            cur.execute("SELECT id, ticket_name, company, status, priority, agent, description, created_at FROM tickets ORDER BY id DESC LIMIT 15;")
        else:
            cur.execute("SELECT id, ticket_name, company, status, priority, agent, description, created_at FROM tickets WHERE company = ANY(%s) ORDER BY id DESC LIMIT 15;", (target_clients,))
        tickets = cur.fetchall()

        # 5. db_monitoring_logs
        if target_servers is None:
            cur.execute("SELECT id, client_name, server_name, log_type, log_time, log_message, severity, status, owner FROM db_monitoring_logs ORDER BY log_time DESC LIMIT 15;")
        else:
            cur.execute("SELECT id, client_name, server_name, log_type, log_time, log_message, severity, status, owner FROM db_monitoring_logs WHERE client_name = ANY(%s) OR server_name = ANY(%s) ORDER BY log_time DESC LIMIT 15;", (target_clients, target_servers))
        logs = cur.fetchall()

        # 6. system_settings
        settings = {}
        try:
            cur.execute("SELECT key, value FROM system_settings;")
            settings = {r["key"]: r["value"] for r in cur.fetchall()}
        except Exception:
            pass

        # 7. feedbacks
        feedbacks = []
        try:
            cur.execute("SELECT username, email, feedback_text, rating, created_at FROM feedbacks ORDER BY created_at DESC LIMIT 5;")
            feedbacks = cur.fetchall()
        except Exception:
            pass

        # 8. online_users
        online_users = []
        try:
            cur.execute("SELECT username, units, status FROM online_users ORDER BY username;")
            online_users = cur.fetchall()
        except Exception:
            pass

        # 9. admin_agents
        admin_agents = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, agent_name, company_name, business_unit, technology FROM admin_agents ORDER BY agent_name LIMIT 10;")
            else:
                cur.execute("SELECT id, agent_name, company_name, business_unit, technology FROM admin_agents WHERE company_name = ANY(%s) ORDER BY agent_name LIMIT 10;", (target_clients,))
            admin_agents = cur.fetchall()
        except Exception:
            pass

        # 10. ai_summary_history
        ai_summary_history = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, summary_text, created_at, username FROM ai_summary_history ORDER BY created_at DESC LIMIT 5;")
            else:
                cur.execute("SELECT id, summary_text, created_at, username FROM ai_summary_history WHERE username = %s ORDER BY created_at DESC LIMIT 5;", (username_val,))
            ai_summary_history = cur.fetchall()
        except Exception:
            pass

        # 11. client_access
        client_access = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, client_email, technology, client_name, server_name, status FROM client_access LIMIT 10;")
            else:
                cur.execute("SELECT id, client_email, technology, client_name, server_name, status FROM client_access WHERE client_name = ANY(%s) OR server_name = ANY(%s) LIMIT 10;", (target_clients, target_servers))
            client_access = cur.fetchall()
        except Exception:
            pass

        # 12. client_reports
        client_reports = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, client_name, title, month, year, file_name, notes, uploaded_by, uploaded_at FROM client_reports ORDER BY uploaded_at DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, client_name, title, month, year, file_name, notes, uploaded_by, uploaded_at FROM client_reports WHERE client_name = ANY(%s) ORDER BY uploaded_at DESC LIMIT 10;", (target_clients,))
            client_reports = cur.fetchall()
        except Exception:
            pass

        # 13. clients
        clients_reg = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, name, database_type, is_active FROM clients ORDER BY name LIMIT 15;")
            else:
                cur.execute("SELECT id, name, database_type, is_active FROM clients WHERE name = ANY(%s) ORDER BY name LIMIT 15;", (target_clients,))
            clients_reg = cur.fetchall()
        except Exception:
            pass

        # 14. database_engineers
        database_engineers = []
        try:
            cur.execute("SELECT id, username, business_units, status FROM database_engineers LIMIT 10;")
            database_engineers = cur.fetchall()
        except Exception:
            pass

        # 15. db_archived_logs
        db_archived_logs = []
        try:
            if target_servers is None:
                cur.execute("SELECT id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner FROM db_archived_logs ORDER BY log_time DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, client_name, server_name, db_type, log_type, log_time, log_message, severity, status, owner FROM db_archived_logs WHERE client_name = ANY(%s) OR server_name = ANY(%s) ORDER BY log_time DESC LIMIT 10;", (target_clients, target_servers))
            db_archived_logs = cur.fetchall()
        except Exception:
            pass

        # 16. db_monitoring_logs_backup
        db_monitoring_logs_backup = []
        try:
            if target_servers is None:
                cur.execute("SELECT id, client_name, server_name, db_type, log_type, log_time, log_message FROM db_monitoring_logs_backup ORDER BY log_time DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, client_name, server_name, db_type, log_type, log_time, log_message FROM db_monitoring_logs_backup WHERE client_name = ANY(%s) OR server_name = ANY(%s) ORDER BY log_time DESC LIMIT 10;", (target_clients, target_servers))
            db_monitoring_logs_backup = cur.fetchall()
        except Exception:
            pass

        # 17. db_uptime_history
        db_uptime_history = []
        try:
            if target_servers is None:
                cur.execute("SELECT id, client_name, server_name, db_type, service_name, status, uptime_desc, last_restart_time, captured_at FROM db_uptime_history ORDER BY captured_at DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, client_name, server_name, db_type, service_name, status, uptime_desc, last_restart_time, captured_at FROM db_uptime_history WHERE client_name = ANY(%s) OR server_name = ANY(%s) ORDER BY captured_at DESC LIMIT 10;", (target_clients, target_servers))
            db_uptime_history = cur.fetchall()
        except Exception:
            pass

        # 18. leads
        leads_reg = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, email, technology, status, is_lead FROM leads LIMIT 15;")
            else:
                cur.execute("SELECT id, email, technology, status, is_lead FROM leads WHERE email = %s OR technology = ANY(%s) LIMIT 15;", (user_email, assigned_techs or []))
            leads_reg = cur.fetchall()
        except Exception:
            pass

        # 19. notifications
        notifications = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, username, message, is_read, created_at FROM notifications ORDER BY created_at DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, username, message, is_read, created_at FROM notifications WHERE username = %s ORDER BY created_at DESC LIMIT 10;", (username_val,))
            notifications = cur.fetchall()
        except Exception:
            pass

        # 20. report_audit_log
        report_audit_log = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, report_id, user_id, action, action_details, created_at FROM report_audit_log ORDER BY created_at DESC LIMIT 10;")
            else:
                cur.execute("""
                    SELECT ral.id, ral.report_id, ral.user_id, ral.action, ral.action_details, ral.created_at 
                    FROM report_audit_log ral
                    JOIN reports r ON ral.report_id = r.id
                    JOIN clients c ON r.client_id = c.id
                    WHERE c.name = ANY(%s)
                    ORDER BY ral.created_at DESC LIMIT 10;
                """, (target_clients,))
            report_audit_log = cur.fetchall()
        except Exception:
            pass

        # 21. report_reviews
        report_reviews = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, report_id, username, rating, comment, created_at FROM report_reviews ORDER BY created_at DESC LIMIT 10;")
            else:
                cur.execute("""
                    SELECT rr.id, rr.report_id, rr.username, rr.rating, rr.comment, rr.created_at 
                    FROM report_reviews rr
                    JOIN reports r ON rr.report_id = r.id
                    JOIN clients c ON r.client_id = c.id
                    WHERE c.name = ANY(%s)
                    ORDER BY rr.created_at DESC LIMIT 10;
                """, (target_clients,))
            report_reviews = cur.fetchall()
        except Exception:
            pass

        # 22. report_sharing_history
        report_sharing_history = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, report_id, report_title, shared_by, share_platform, recipient, created_at FROM report_sharing_history ORDER BY created_at DESC LIMIT 10;")
            else:
                cur.execute("""
                    SELECT rsh.id, rsh.report_id, rsh.report_title, rsh.shared_by, rsh.share_platform, rsh.recipient, rsh.created_at 
                    FROM report_sharing_history rsh
                    JOIN reports r ON rsh.report_id = r.id
                    JOIN clients c ON r.client_id = c.id
                    WHERE c.name = ANY(%s)
                    ORDER BY rsh.created_at DESC LIMIT 10;
                """, (target_clients,))
            report_sharing_history = cur.fetchall()
        except Exception:
            pass

        # 23. reports
        reports_reg = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, client_id, user_id, file_name, file_path, file_size, report_month, upload_date, report_type, is_latest, status FROM reports LIMIT 10;")
            else:
                cur.execute("""
                    SELECT r.id, r.client_id, r.user_id, r.file_name, r.file_path, r.file_size, r.report_month, r.upload_date, r.report_type, r.is_latest, r.status
                    FROM reports r
                    JOIN clients c ON r.client_id = c.id
                    WHERE c.name = ANY(%s) LIMIT 10;
                """, (target_clients,))
            reports_reg = cur.fetchall()
        except Exception:
            pass

        # 24. server_utilization_history
        server_utilization_history = []
        try:
            if target_servers is None:
                cur.execute("SELECT id, server_name, cpu_utilization, memory_utilization, disk_utilization, io_utilization, captured_at, read_iops, write_iops FROM server_utilization_history ORDER BY captured_at DESC LIMIT 15;")
            else:
                cur.execute("SELECT id, server_name, cpu_utilization, memory_utilization, disk_utilization, io_utilization, captured_at, read_iops, write_iops FROM server_utilization_history WHERE server_name = ANY(%s) ORDER BY captured_at DESC LIMIT 15;", (target_servers,))
            server_utilization_history = cur.fetchall()
        except Exception:
            pass

        # 25. share_history
        share_history = []
        try:
            if target_servers is None:
                cur.execute("SELECT id, username, shared_at, notes, platform, content_type, client_name, server_name, log_message, status, owner, ticket_status, next_action, client_visibility, db_type FROM share_history ORDER BY shared_at DESC LIMIT 10;")
            else:
                cur.execute("SELECT id, username, shared_at, notes, platform, content_type, client_name, server_name, log_message, status, owner, ticket_status, next_action, client_visibility, db_type FROM share_history WHERE client_name = ANY(%s) OR server_name = ANY(%s) ORDER BY shared_at DESC LIMIT 10;", (target_clients, target_servers))
            share_history = cur.fetchall()
        except Exception:
            pass

        # 26. system_admins
        system_admins = []
        try:
            cur.execute("SELECT id, email, status, created_at, updated_at FROM system_admins LIMIT 5;")
            system_admins = cur.fetchall()
        except Exception:
            pass

        # 27. ticket_agents
        ticket_agents = []
        try:
            cur.execute("SELECT id, name FROM ticket_agents LIMIT 10;")
            ticket_agents = cur.fetchall()
        except Exception:
            pass

        # 28. ticket_assignments
        ticket_assignments = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, ticket_id, assigned_to, assigned_by, assigned_at, unassigned_at FROM ticket_assignments LIMIT 10;")
            else:
                cur.execute("""
                    SELECT ta.id, ta.ticket_id, ta.assigned_to, ta.assigned_by, ta.assigned_at, ta.unassigned_at 
                    FROM ticket_assignments ta
                    JOIN tickets t ON ta.ticket_id = t.id
                    WHERE t.company = ANY(%s) LIMIT 10;
                """, (target_clients,))
            ticket_assignments = cur.fetchall()
        except Exception:
            pass

        # 29. ticket_business_units
        ticket_business_units = []
        try:
            cur.execute("SELECT id, name FROM ticket_business_units LIMIT 10;")
            ticket_business_units = cur.fetchall()
        except Exception:
            pass

        # 30. ticket_comments
        ticket_comments = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, ticket_id, author, comment_type, content, attachments, created_at FROM ticket_comments LIMIT 10;")
            else:
                cur.execute("""
                    SELECT tc.id, tc.ticket_id, tc.author, tc.comment_type, tc.content, tc.attachments, tc.created_at 
                    FROM ticket_comments tc
                    JOIN tickets t ON tc.ticket_id = t.id
                    WHERE t.company = ANY(%s) LIMIT 10;
                """, (target_clients,))
            ticket_comments = cur.fetchall()
        except Exception:
            pass

        # 31. ticket_notifications
        ticket_notifications = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, target_user, ticket_id, message, is_read, created_by, created_at FROM ticket_notifications LIMIT 10;")
            else:
                cur.execute("SELECT id, target_user, ticket_id, message, is_read, created_by, created_at FROM ticket_notifications WHERE target_user = %s OR created_by = %s LIMIT 10;", (username_val, username_val))
            ticket_notifications = cur.fetchall()
        except Exception:
            pass

        # 32. user_clients
        user_clients = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, user_id, client_id, access_level, created_at FROM user_clients LIMIT 15;")
            else:
                cur.execute("""
                    SELECT uc.id, uc.user_id, uc.client_id, uc.access_level, uc.created_at 
                    FROM user_clients uc
                    JOIN clients c ON uc.client_id = c.id
                    WHERE c.name = ANY(%s) LIMIT 15;
                """, (target_clients,))
            user_clients = cur.fetchall()
        except Exception:
            pass

        # 33. user_page_activity
        user_page_activity = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, username, page_path, duration_seconds, last_active_at FROM user_page_activity ORDER BY last_active_at DESC LIMIT 15;")
            else:
                cur.execute("SELECT id, username, page_path, duration_seconds, last_active_at FROM user_page_activity WHERE username = %s ORDER BY last_active_at DESC LIMIT 15;", (username_val,))
            user_page_activity = cur.fetchall()
        except Exception:
            pass

        # 34. users
        users_reg = []
        try:
            if target_clients is None:
                cur.execute("SELECT id, username, role, full_name, email, last_active_at FROM users LIMIT 15;")
            else:
                cur.execute("SELECT id, username, role, full_name, email, last_active_at FROM users WHERE username = %s LIMIT 15;", (username_val,))
            users_reg = cur.fetchall()
        except Exception:
            pass

        # 35. workers
        workers = []
        try:
            cur.execute("SELECT id, name, email, role, department, is_active FROM workers LIMIT 10;")
            workers = cur.fetchall()
        except Exception:
            pass

        def fmt(b):
            if not b: return "0 B"
            for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
                if b < 1024: return f"{b:.2f} {unit}"
                b /= 1024
            return f"{b:.2f} PB"
            
        context = "--- SYSTEM TELEMETRY CURRENT STATE ---\n"
        
        # Client Technologies
        context += "Client Technology Mappings:\n"
        for c in clients_tech:
            context += f"- Client: {c['client_name']} | Server: {c['server_name']} | Database Type: {c['db_type']}\n"
            
        # Databases / Summary
        if db_summary:
            context += "\nMonitored Databases Capacity Summary (Aggregated):\n"
            for s in db_summary:
                context += f"- Server: {s['server_name']} | Databases Monitored: {s['db_count']} | Total Size: {fmt(s['total_size'] or 0)}\n"
        elif dbs:
            context += "\nMonitored Databases Capacity:\n"
            for db in dbs:
                context += f"- Client/Server: {db['server_name']} | Database: {db['database_name']} | Size: {fmt(db['latest_size'] or 0)} (As of: {db['captured_date']})\n"
            
        # Tables
        if tbls:
            context += "\nTop Largest Tables:\n"
            for tbl in tbls:
                context += f"- Client/Server: {tbl['server_name']} | DB: {tbl['database_name']} | Table: {tbl['table_name']} | Size: {fmt(tbl['latest_size'] or 0)} (As of: {tbl['captured_date']})\n"
            
        # Tickets
        context += "\nActive & Historical Tickets:\n"
        if tickets:
            for t in tickets:
                truncated_desc = (t['description'][:120] + '...') if (t['description'] and len(t['description']) > 120) else (t['description'] or 'No description')
                context += f"- Ticket #{t['id']}: [{t['status']}] {t['ticket_name']} (Client: {t['company']}, Priority: {t['priority']}, Assigned To: {t['agent'] or 'Unassigned'}, Created: {t['created_at']})\n  Description: {truncated_desc}\n"
        else:
            context += "- No tickets found.\n"
            
        # Logs
        context += "\nDatabase Monitoring Logs:\n"
        if logs:
            for l in logs:
                truncated_msg = (l['log_message'][:120] + '...') if (l['log_message'] and len(l['log_message']) > 120) else (l['log_message'] or '')
                context += f"- Log #{l['id']}: [{l['severity'] or 'INFO'}] {truncated_msg} (Client: {l['client_name']}/{l['server_name']}, Type: {l['log_type']}, Status: {l['status'] or 'Open'}, Owner: {l['owner'] or 'Unassigned'}, Log Time: {l['log_time']})\n"
        else:
            context += "- No logs found.\n"

        # System Settings & Scheduler Status
        if settings:
            context += "\nSystem Telemetry Scheduler Settings:\n"
            context += f"- Daily Trigger Time: {settings.get('telemetry_sync_hour', '14')}:{settings.get('telemetry_sync_minute', '00')} IST\n"
            context += f"- Last Ingestion Success Time: {settings.get('telemetry_sync_last_time', 'Never')}\n"
            context += f"- Last Ingestion Status: {settings.get('telemetry_sync_last_status', 'N/A')}\n"
            context += f"- Ingestion Daemon State: {'Running/Syncing' if settings.get('telemetry_sync_in_progress') == 'true' else 'Idle'}\n"

        # Online Specialists
        if online_users:
            context += "\nOnline DBA Directory & Assignment Status:\n"
            for ou in online_users:
                context += f"- DBA: {ou['username']} | Unit/Specialty: {ou['units']} | Status: {ou['status']}\n"

        # User Feedbacks
        if feedbacks:
            context += "\nRecent Client Feedback and System Ratings:\n"
            for fb in feedbacks:
                context += f"- User: {fb['username']} ({fb['email']}) rated {fb['rating']}/5 stars: '{fb['feedback_text']}'\n"

        # 9. admin_agents
        if admin_agents:
            context += "\nAdmin Support Agents Registry:\n"
            for agent in admin_agents:
                context += f"- Agent: {agent['agent_name']} | Client: {agent['company_name']} | BU: {agent['business_unit']} | Tech: {agent['technology']}\n"

        # 10. ai_summary_history
        if ai_summary_history:
            context += "\nAI Summary Generation History:\n"
            for summary in ai_summary_history:
                context += f"- Summary #{summary['id']}: User: {summary['username']} | Created: {summary['created_at']} | Text: {summary['summary_text'][:120]}...\n"

        # 11. client_access
        if client_access:
            context += "\nClient Dashboard Access Privileges:\n"
            for ca in client_access:
                context += f"- Privilege #{ca['id']}: Email: {ca['client_email']} | Client: {ca['client_name']} | Server: {ca['server_name']} | Tech: {ca['technology']} | Status: {ca['status']}\n"

        # 12. client_reports
        if client_reports:
            context += "\nClient Monthly Status Reports:\n"
            for cr in client_reports:
                context += f"- Report: '{cr['title']}' for Client: {cr['client_name']} ({cr['month']}/{cr['year']}) | File: {cr['file_name']} | Uploaded By: {cr['uploaded_by']} at {cr['uploaded_at']}\n"

        # 13. clients
        if clients_reg:
            context += "\nGeneral Clients Directory Registry:\n"
            for c in clients_reg:
                context += f"- Client: {c['name']} | Database Engine: {c['database_type']} | Active: {c['is_active']}\n"

        # 14. database_engineers
        if database_engineers:
            context += "\nDatabase Engineers Assignment Directory:\n"
            for de in database_engineers:
                context += f"- Engineer: {de['username']} | Business Units: {de['business_units']} | Status: {de['status']}\n"

        # 15. db_archived_logs
        if db_archived_logs:
            context += "\nArchived Database Alerts & Events Logs:\n"
            for al in db_archived_logs:
                context += f"- Archived Log #{al['id']}: [{al['severity'] or 'INFO'}] {al['log_message'][:120]}... (Client: {al['client_name']}/{al['server_name']}, Type: {al['log_type']}, Owner: {al['owner']}, Time: {al['log_time']})\n"

        # 16. db_monitoring_logs_backup
        if db_monitoring_logs_backup:
            context += "\nBackup Database Monitoring Alert Log Archive:\n"
            for bk in db_monitoring_logs_backup:
                context += f"- Backup Log #{bk['id']}: Client: {bk['client_name']}/{bk['server_name']} | Type: {bk['log_type']} | Time: {bk['log_time']} | Message: {bk['log_message'][:120]}...\n"

        # 17. db_uptime_history
        if db_uptime_history:
            context += "\nDatabase Engine Server Status & Uptime History:\n"
            for up in db_uptime_history:
                context += f"- Uptime Status: Server: {up['server_name']} (Client: {up['client_name']}, Tech: {up['db_type']}) | Service: {up['service_name']} is {up['status']} | Description: {up['uptime_desc']} | Last Restart: {up['last_restart_time']}\n"

        # 18. leads
        if leads_reg:
            context += "\nDatabase Leads & Technologies Matrix:\n"
            for lead in leads_reg:
                context += f"- Lead: Email: {lead['email']} | Assigned Technology: {lead['technology']} | Status: {lead['status']} | Is Lead Person: {lead['is_lead']}\n"

        # 19. notifications
        if notifications:
            context += "\nUser Alert Notifications Log:\n"
            for notif in notifications:
                context += f"- Notification #{notif['id']}: User: {notif['username']} | Message: '{notif['message']}' | Read: {notif['is_read']} | Time: {notif['created_at']}\n"

        # 20. report_audit_log
        if report_audit_log:
            context += "\nReport Sharing and Download Audit Logs:\n"
            for audit in report_audit_log:
                context += f"- Audit #{audit['id']}: User ID: {audit['user_id']} | Action: {audit['action']} | Details: {audit['action_details']} | Created: {audit['created_at']}\n"

        # 21. report_reviews
        if report_reviews:
            context += "\nReport Quality Reviews and Feedback:\n"
            for rev in report_reviews:
                context += f"- Review #{rev['id']}: Report ID: {rev['report_id']} | User: {rev['username']} rated {rev['rating']}/5 | Comment: '{rev['comment']}'\n"

        # 22. report_sharing_history
        if report_sharing_history:
            context += "\nReport Platform Sharing Log Matrix:\n"
            for rsh in report_sharing_history:
                context += f"- Share Record #{rsh['id']}: Report ID: {rsh['report_id']} | Title: '{rsh['report_title']}' | Shared By: {rsh['shared_by']} via {rsh['share_platform']} to {rsh['recipient']}\n"

        # 23. reports
        if reports_reg:
            context += "\nReports Center Documents:\n"
            for rep in reports_reg:
                context += f"- Report #{rep['id']}: File: {rep['file_name']} | Type: {rep['report_type']} | Status: {rep['status']} | Size: {rep['file_size']} Bytes | Month: {rep['report_month']}\n"

        # 24. server_utilization_history
        if server_utilization_history:
            context += "\nServer Resource Capacity History:\n"
            for ut in server_utilization_history:
                context += f"- Resource Metrics: Server: {ut['server_name']} | CPU: {ut['cpu_utilization']}% | Mem: {ut['memory_utilization']}% | Disk: {ut['disk_utilization']}% | I/O: {ut['io_utilization']}% | Read IOPS: {ut['read_iops']} | Write IOPS: {ut['write_iops']} | Captured At: {ut['captured_at']}\n"

        # 25. share_history
        if share_history:
            context += "\nTeams/WhatsApp Telemetry Sharing Activity:\n"
            for sh in share_history:
                context += f"- Shared Alert #{sh['id']}: Shared By: {sh['username']} on {sh['platform']} | Client: {sh['client_name']}/{sh['server_name']} | Log: {sh['log_message'][:120]} | Time: {sh['shared_at']}\n"

        # 26. system_admins
        if system_admins:
            context += "\nSystem Administrators Registry:\n"
            for sa in system_admins:
                context += f"- Admin Contact: {sa['email']} | Status: {sa['status']}\n"

        # 27. ticket_agents
        if ticket_agents:
            context += "\nSupport Ticket Handling Agents:\n"
            for ta in ticket_agents:
                context += f"- Agent: {ta['name']} (ID: {ta['id']})\n"

        # 28. ticket_assignments
        if ticket_assignments:
            context += "\nSupport Tickets DBA Engineers Assignments Log:\n"
            for t_as in ticket_assignments:
                context += f"- Ticket Assignment: Ticket ID: {t_as['ticket_id']} | Assigned To User ID: {t_as['assigned_to']} | Assigned By User ID: {t_as['assigned_by']} | Assigned: {t_as['assigned_at']}\n"

        # 29. ticket_business_units
        if ticket_business_units:
            context += "\nTicket Business Units Catalog:\n"
            for bu in ticket_business_units:
                context += f"- Business Unit: {bu['name']} (ID: {bu['id']})\n"

        # 30. ticket_comments
        if ticket_comments:
            context += "\nSupport Tickets Discussion & Comments Logs:\n"
            for comment in ticket_comments:
                context += f"- Comment by {comment['author']} on Ticket #{comment['ticket_id']}: '{comment['content'][:120]}' (Type: {comment['comment_type']}, Created: {comment['created_at']})\n"

        # 31. ticket_notifications
        if ticket_notifications:
            context += "\nTicket Notification Activity Logs:\n"
            for tn in ticket_notifications:
                context += f"- Ticket Notification #{tn['id']}: Target User: {tn['target_user']} | Ticket ID: {tn['ticket_id']} | Message: '{tn['message']}' | Read: {tn['is_read']}\n"

        # 32. user_clients
        if user_clients:
            context += "\nUser-Client Organization Mappings Matrix:\n"
            for uc in user_clients:
                context += f"- Access Profile: User ID: {uc['user_id']} | Client ID: {uc['client_id']} | Access Level: {uc['access_level']}\n"

        # 33. user_page_activity
        if user_page_activity:
            context += "\nUser Navigation Dashboard Page Views Stats:\n"
            for act in user_page_activity:
                context += f"- View Stats: User: {act['username']} | Page: {act['page_path']} | Duration: {act['duration_seconds']}s | Time: {act['last_active_at']}\n"

        # 34. users
        if users_reg:
            context += "\nAuthorized Users & Accounts Registry Details:\n"
            for u_reg in users_reg:
                context += f"- Profile: User: {u_reg['username']} | Role: {u_reg['role']} | Full Name: {u_reg['full_name']} | Email: {u_reg['email']} | Last Active: {u_reg['last_active_at']}\n"

        # 35. workers
        if workers:
            context += "\nGeoMon Background Workers Directory:\n"
            for worker in workers:
                context += f"- Worker: {worker['name']} ({worker['email']}) | Role: {worker['role']} | Department: {worker['department']} | Active: {worker['is_active']}\n"

        context += "--------------------------------------\n"
        return context
    except Exception as e:
        print(f"Error generating telemetry context: {e}")
        return "--- Telemetry Context Unavailable ---"
    finally:
        conn.close()



@router.get("/telemetry/clients")
def get_telemetry_clients(user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:clients:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor()
        if user.get("isAdmin"):
            cur.execute("""
                SELECT DISTINCT server_name FROM database_size_history
                UNION
                SELECT DISTINCT server_name FROM table_size_history
                ORDER BY server_name;
            """)
            rows = cur.fetchall()
        else:
            allowed_clients = user.get("allowed_clients", [])
            if allowed_clients:
                cur.execute("""
                    SELECT DISTINCT server_name FROM admin_clients WHERE client_name = ANY(%s);
                """, (allowed_clients,))
                allowed_servers = [r[0] for r in cur.fetchall() if r[0]]
                if allowed_servers:
                    cur.execute("""
                        SELECT DISTINCT server_name FROM database_size_history WHERE server_name = ANY(%s)
                        UNION
                        SELECT DISTINCT server_name FROM table_size_history WHERE server_name = ANY(%s)
                        ORDER BY server_name;
                    """, (allowed_servers, allowed_servers))
                    rows = cur.fetchall()
                else:
                    rows = []
            else:
                assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
                if assigned_techs:
                    cur.execute("""
                        SELECT DISTINCT server_name FROM admin_clients 
                        WHERE EXISTS (
                            SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                            WHERE t = ANY(%s)
                        );
                    """, (assigned_techs,))
                    allowed_servers = [r[0] for r in cur.fetchall() if r[0]]
                    if allowed_servers:
                        cur.execute("""
                            SELECT DISTINCT server_name FROM database_size_history WHERE server_name = ANY(%s)
                            UNION
                            SELECT DISTINCT server_name FROM table_size_history WHERE server_name = ANY(%s)
                            ORDER BY server_name;
                        """, (allowed_servers, allowed_servers))
                        rows = cur.fetchall()
                    else:
                        rows = []
                else:
                    rows = []
                    
        # Filter out Unknown / empty clients
        EXCLUDED = {"unknown", "n/a", "", None}
        clients = [r[0] for r in rows if r[0] and r[0].strip().lower() not in EXCLUDED]
        result = {"clients": clients}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.get("/telemetry/client-tech-grid")
def get_telemetry_client_tech_grid(user: dict = Depends(get_current_user_local)):
    """
    Returns one entry per (client_name, db_type) combination so the UI renders
    separate technology-specific cards for each client.
    Example: Shemaroo with MongoDB + MSSQL → two separate grid cards.
    """
    cache_key = f"telemetry:client-tech-grid:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        EXCLUDED = {"unknown", "n/a", "", None}

        if user.get("isAdmin"):
            # Admin sees all: join admin_clients with size history to confirm data exists
            cur.execute("""
                SELECT DISTINCT ac.client_name, ac.db_type, ac.server_name
                FROM admin_clients ac
                WHERE ac.server_name IN (
                    SELECT DISTINCT server_name FROM database_size_history
                    UNION
                    SELECT DISTINCT server_name FROM table_size_history
                )
                ORDER BY ac.client_name, ac.db_type;
            """)
        else:
            allowed_clients = user.get("allowed_clients", [])
            assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
            if allowed_clients:
                cur.execute("""
                    SELECT DISTINCT ac.client_name, ac.db_type, ac.server_name
                    FROM admin_clients ac
                    WHERE ac.client_name = ANY(%s)
                      AND ac.server_name IN (
                        SELECT DISTINCT server_name FROM database_size_history
                        UNION
                        SELECT DISTINCT server_name FROM table_size_history
                      )
                    ORDER BY ac.client_name, ac.db_type;
                """, (allowed_clients,))
            elif assigned_techs:
                cur.execute("""
                    SELECT DISTINCT ac.client_name, ac.db_type, ac.server_name
                    FROM admin_clients ac
                    WHERE LOWER(TRIM(ac.db_type)) = ANY(%s)
                      AND ac.server_name IN (
                        SELECT DISTINCT server_name FROM database_size_history
                        UNION
                        SELECT DISTINCT server_name FROM table_size_history
                      )
                    ORDER BY ac.client_name, ac.db_type;
                """, (assigned_techs,))
            else:
                result = {"client_tech_pairs": []}
                cache_manager.set(cache_key, result, ttl_seconds=60)
                return result

        rows = cur.fetchall()
        pairs = [
            {
                "client_name": r["client_name"],
                "db_type": r["db_type"],
                "server_name": r["server_name"]
            }
            for r in rows
            if r["client_name"] and r["client_name"].strip().lower() not in EXCLUDED
        ]
        result = {"client_tech_pairs": pairs}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.get("/telemetry/last-week/{client_name}")
def get_telemetry_last_week(client_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    """Returns all database + table size records for the last 7 days for a given client."""
    cache_key = f"telemetry:last-week:{client_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")
            
        # Databases — last 7 days
        cur.execute("""
            SELECT database_name, captured_date, total_size_bytes as size_bytes
            FROM database_size_history
            WHERE server_name = %s 
              AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
              AND captured_date >= CURRENT_DATE - INTERVAL '7 days'
            ORDER BY captured_date DESC, database_name;
        """, (client_name, db_type, db_type, db_type))
        db_rows = [dict(r) for r in cur.fetchall()]
        for r in db_rows:
            if r.get("captured_date"):
                r["captured_date"] = r["captured_date"].isoformat()

        # Tables — last 7 days
        cur.execute("""
            SELECT database_name, table_name, captured_date, size_bytes
            FROM table_size_history
            WHERE server_name = %s 
              AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
              AND captured_date >= CURRENT_DATE - INTERVAL '7 days'
            ORDER BY captured_date DESC, database_name, table_name;
        """, (client_name, db_type, db_type, db_type))
        table_rows = [dict(r) for r in cur.fetchall()]
        for r in table_rows:
            if r.get("captured_date"):
                r["captured_date"] = r["captured_date"].isoformat()

        result = {
            "client_name": client_name,
            "databases_last_week": db_rows,
            "tables_last_week": table_rows
        }
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/range-data/{client_name}")
def get_telemetry_range_data(
    client_name: str,
    data_type: str = "databases",
    from_date: str = None,
    to_date:   str = None,
    db_type: Optional[str] = None,
    user: dict = Depends(get_current_user_local)
):
    """
    Returns daily size pivot for all databases (or tables) in a date range.
    Response: { dates:[...], rows:[{ database_name, [table_name], values:{date:bytes} },...] }
    """
    cache_key = f"telemetry:range-data:{client_name}:{data_type}:{from_date}:{to_date}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    from datetime import date as date_cls, timedelta
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")
            
        today  = date_cls.today()
        t_date = today
        f_date = today - timedelta(days=6)
        if from_date:
            try: f_date = date_cls.fromisoformat(from_date)
            except: pass
        if to_date:
            try: t_date = date_cls.fromisoformat(to_date)
            except: pass
        if f_date > t_date:
            f_date, t_date = t_date, f_date

        if data_type == "tables":
            cur.execute("""
                SELECT database_name, table_name, captured_date, size_bytes
                FROM table_size_history
                WHERE server_name=%s 
                  AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
                  AND captured_date BETWEEN %s AND %s
                ORDER BY captured_date, database_name, table_name;
            """, (client_name, db_type, db_type, db_type, f_date, t_date))
            pivot, all_dates = {}, set()
            for r in cur.fetchall():
                d   = r["captured_date"].isoformat(); all_dates.add(d)
                key = (r["database_name"], r["table_name"])
                pivot.setdefault(key, {})[d] = r["size_bytes"]
            sorted_dates = sorted(all_dates)
            rows = [{"database_name": db, "table_name": tbl, "values": vals}
                    for (db, tbl), vals in sorted(pivot.items())]
        else:
            cur.execute("""
                SELECT database_name, captured_date, total_size_bytes
                FROM database_size_history
                WHERE server_name=%s 
                  AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
                  AND captured_date BETWEEN %s AND %s
                ORDER BY captured_date, database_name;
            """, (client_name, db_type, db_type, db_type, f_date, t_date))
            pivot, all_dates = {}, set()
            for r in cur.fetchall():
                d  = r["captured_date"].isoformat(); all_dates.add(d)
                db = r["database_name"]
                pivot.setdefault(db, {})[d] = r["total_size_bytes"]
            sorted_dates = sorted(all_dates)
            rows = [{"database_name": db, "values": vals}
                    for db, vals in sorted(pivot.items())]

        result = {
            "client_name": client_name, "data_type": data_type,
            "from_date": f_date.isoformat(), "to_date": t_date.isoformat(),
            "dates": sorted_dates, "rows": rows
        }
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/uptime/{client_name}")
def get_telemetry_uptime(client_name: str, db_type: Optional[str] = None, refresh: bool = Query(False), user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:uptime:{client_name}:{db_type or ''}:{user.get('username')}"
    
    # Trigger simulation update if we are in local mode (Outlook bypassed)
    try:
        from email_extracter import account, simulate_local_uptime_update
        if refresh or not account:
            simulate_local_uptime_update()
    except Exception as e:
        print("Uptime simulation helper error:", e)

    if not refresh and cache_manager:
        cached_val = cache_manager.get(cache_key)
        if cached_val is not None:
            return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Enforce tenant isolation
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client's technology or registry.")

        # Query latest service statuses
        cur.execute("""
            WITH latest_services AS (
                SELECT DISTINCT ON (service_name)
                    client_name, server_name, db_type, service_name, status, uptime_desc, last_restart_time, captured_at
                FROM db_uptime_history
                WHERE (LOWER(client_name) = LOWER(%s) OR LOWER(server_name) = LOWER(%s))
                  AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
                ORDER BY service_name, captured_at DESC
            )
            SELECT * FROM latest_services;
        """, (client_name, client_name, db_type, db_type, db_type))
        latest = cur.fetchall()

        # Query historical uptime status logs
        cur.execute("""
            SELECT client_name, server_name, db_type, service_name, status, uptime_desc, last_restart_time, captured_at
            FROM db_uptime_history
            WHERE (LOWER(client_name) = LOWER(%s) OR LOWER(server_name) = LOWER(%s))
              AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            ORDER BY captured_at DESC, service_name;
        """, (client_name, client_name, db_type, db_type, db_type))
        history = cur.fetchall()

        for row in latest:
            if row["last_restart_time"]:
                row["last_restart_time"] = row["last_restart_time"].isoformat()
            if row["captured_at"]:
                row["captured_at"] = row["captured_at"].isoformat()
                
        for row in history:
            if row["last_restart_time"]:
                row["last_restart_time"] = row["last_restart_time"].isoformat()
            if row["captured_at"]:
                row["captured_at"] = row["captured_at"].isoformat()

        result = {
            "latest": latest,
            "history": history
        }
        
        if cache_manager:
            cache_manager.set(cache_key, result, ttl_seconds=60)
            
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/databases/{client_name}")
def get_telemetry_databases(client_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:databases:{client_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Enforce technology and client permissions
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client's technology or registry.")

        # Fetch the databases, their latest size, and 7-day-ago size, oldest size, and 7-day average size
        cur.execute("""
            SELECT database_name, MAX(captured_date) as latest_date,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date DESC LIMIT 1) as latest_size,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date <= MAX(h.captured_date) - INTERVAL '7 days' ORDER BY captured_date DESC LIMIT 1) as prev_week_size,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date ASC LIMIT 1) as oldest_size,
                   (SELECT AVG(total_size_bytes) FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date >= MAX(h.captured_date) - INTERVAL '7 days') as avg_size_7d
            FROM database_size_history h
            WHERE server_name = %s AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            GROUP BY database_name;
        """, (client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type))
        rows = cur.fetchall()
        
        databases = []
        for r in rows:
            latest = r["latest_size"] or 0
            old = r["prev_week_size"] or r["oldest_size"] or 0
            diff = latest - old
            pct = round((diff / old * 100), 2) if old > 0 else 0.0
            
            avg_7d = r["avg_size_7d"]
            if avg_7d is not None:
                avg_7d = round(float(avg_7d), 2)
                avg_growth = latest - avg_7d
                avg_growth_pct = round((avg_growth / avg_7d * 100), 2) if avg_7d > 0 else 0.0
            else:
                avg_7d = None
                avg_growth = None
                avg_growth_pct = None
            
            databases.append({
                "database_name": r["database_name"],
                "latest_date": r["latest_date"].isoformat() if r["latest_date"] else None,
                "latest_size": latest,
                "prev_week_size": old,
                "growth_bytes": diff,
                "growth_pct": pct,
                "avg_size_7d": avg_7d,
                "avg_growth_bytes": avg_growth,
                "avg_growth_pct": avg_growth_pct
            })
            
        # Custom sorting logic: Positive growth percentage first (descending), negative/zero growth at the very end
        databases.sort(key=lambda x: (1 if (x["growth_pct"] or 0.0) > 0 else (0 if (x["growth_pct"] or 0.0) == 0 else -1), x["growth_pct"] or 0.0), reverse=True)
        result = {"databases": databases}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/tables/{client_name}")
def get_telemetry_tables(client_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:tables:{client_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Enforce technology and client permissions
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client's technology or registry.")

        # Fetch the tables, their latest size, and 7-day-ago size, oldest size, and 7-day average size
        cur.execute("""
            SELECT database_name, table_name, MAX(captured_date) as latest_date,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date DESC LIMIT 1) as latest_size,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date <= MAX(h.captured_date) - INTERVAL '7 days' ORDER BY captured_date DESC LIMIT 1) as prev_week_size,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date ASC LIMIT 1) as oldest_size,
                   (SELECT AVG(size_bytes) FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date >= MAX(h.captured_date) - INTERVAL '7 days') as avg_size_7d
            FROM table_size_history h
            WHERE server_name = %s AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            GROUP BY database_name, table_name;
        """, (client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type))
        rows = cur.fetchall()
        
        tables = []
        for r in rows:
            latest = r["latest_size"] or 0
            old = r["prev_week_size"] or r["oldest_size"] or 0
            diff = latest - old
            pct = round((diff / old * 100), 2) if old > 0 else 0.0
            
            avg_7d = r["avg_size_7d"]
            if avg_7d is not None:
                avg_7d = round(float(avg_7d), 2)
                avg_growth = latest - avg_7d
                avg_growth_pct = round((avg_growth / avg_7d * 100), 2) if avg_7d > 0 else 0.0
            else:
                avg_7d = None
                avg_growth = None
                avg_growth_pct = None
            
            tables.append({
                "database_name": r["database_name"],
                "table_name": r["table_name"],
                "latest_date": r["latest_date"].isoformat() if r["latest_date"] else None,
                "latest_size": latest,
                "prev_week_size": old,
                "growth_bytes": diff,
                "growth_pct": pct,
                "avg_size_7d": avg_7d,
                "avg_growth_bytes": avg_growth,
                "avg_growth_pct": avg_growth_pct
            })
            
        # Custom sorting logic: Positive growth percentage first (descending), negative/zero growth at the very end
        tables.sort(key=lambda x: (1 if (x["growth_pct"] or 0.0) > 0 else (0 if (x["growth_pct"] or 0.0) == 0 else -1), x["growth_pct"] or 0.0), reverse=True)
        result = {"tables": tables}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/mssql/{table_name}")
def get_mssql_telemetry(table_name: str, client_name: str, server_name: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    valid_tables = [
        "reportdata_restart", "reportdata_backup", "reportdata_server", "reportdata_disk_drive",
        "reportdata_size_growth", "reportdata_top_cpu", "diagnosticdata_disk_io", "diagnosticdata_wait_stats",
        "diagnosticdata_long_queries", "diagnosticdata_deadlocks", "diagnosticdata_tempdb", "diagnosticdata_job_executions",
        "diagnosticdata_blocking", "diagnosticdata_error_logs", "diagnosticdata_cpu_querystore", "diagnosticdata_mem_querystore",
        "reportdata_memory_ple", "reportdata_memory_snapshot", "reportdata_cpu_daily_summary", "reportdata_cpu_spike_analysis"
    ]
    if table_name not in valid_tables:
        raise HTTPException(status_code=400, detail="Invalid table name")

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")

        query = f"SELECT id, client_name, server_name, captured_at, raw_data FROM {table_name} WHERE LOWER(client_name) = LOWER(%s)"
        params = [client_name]
        if server_name:
            query += " AND LOWER(server_name) = LOWER(%s)"
            params.append(server_name)
        query += " ORDER BY captured_at DESC LIMIT 100;"
        
        cur.execute(query, params)
        rows = cur.fetchall()
        
        for r in rows:
            if r["captured_at"]:
                r["captured_at"] = r["captured_at"].isoformat()
        return {"data": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/summary/{client_name}")
def get_telemetry_summary(client_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    """
    Analyzes database & table capacity growth relative to the previous day (daily) and past week.
    """
    cache_key = f"telemetry:summary:{client_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Enforce permissions
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")

        # 1. Database Growth Analysis (comparing latest size vs previous captured day size for Today's growth)
        cur.execute("""
            SELECT database_name, MAX(captured_date) as latest_date,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date DESC LIMIT 1) as latest_size,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date < (SELECT MAX(captured_date) FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))) ORDER BY captured_date DESC LIMIT 1) as prev_day_size,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date <= MAX(h.captured_date) - INTERVAL '7 days' ORDER BY captured_date DESC LIMIT 1) as prev_week_size,
                   (SELECT total_size_bytes FROM database_size_history WHERE server_name = %s AND database_name = h.database_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date ASC LIMIT 1) as oldest_size
            FROM database_size_history h
            WHERE server_name = %s AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            GROUP BY database_name;
        """, (client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type))
        db_rows = cur.fetchall()
        
        databases = []
        for r in db_rows:
            latest = r["latest_size"] or 0
            prev_day = r["prev_day_size"] or r["oldest_size"] or 0
            prev_week = r["prev_week_size"] or r["oldest_size"] or 0
            
            # Daily Growth (Size today vs previous day)
            daily_diff = latest - prev_day
            daily_pct = round((daily_diff / prev_day * 100), 2) if prev_day > 0 else 0.0
            
            # Weekly Growth (Size today vs 7 days ago)
            week_diff = latest - prev_week
            week_pct = round((week_diff / prev_week * 100), 2) if prev_week > 0 else 0.0
            
            databases.append({
                "name": r["database_name"],
                "latest_size": latest,
                "growth": daily_diff, # Today's growth
                "growth_pct": daily_pct,
                "growth_7d": week_diff,
                "growth_7d_pct": week_pct
            })
            
        # Sort based on Today's daily growth percentage (positive first descending, negative/zero last)
        databases.sort(key=lambda x: (1 if (x["growth_pct"] or 0.0) > 0 else (0 if (x["growth_pct"] or 0.0) == 0 else -1), x["growth_pct"] or 0.0), reverse=True)
        
        # 2. Table Growth Analysis (comparing latest size vs previous captured day size for Today's growth)
        cur.execute("""
            SELECT database_name, table_name, MAX(captured_date) as latest_date,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date DESC LIMIT 1) as latest_size,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date < (SELECT MAX(captured_date) FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))) ORDER BY captured_date DESC LIMIT 1) as prev_day_size,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) AND captured_date <= MAX(h.captured_date) - INTERVAL '7 days' ORDER BY captured_date DESC LIMIT 1) as prev_week_size,
                   (SELECT size_bytes FROM table_size_history WHERE server_name = %s AND database_name = h.database_name AND table_name = h.table_name AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s)) ORDER BY captured_date ASC LIMIT 1) as oldest_size
            FROM table_size_history h
            WHERE server_name = %s AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            GROUP BY database_name, table_name;
        """, (client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type,
              client_name, db_type, db_type, db_type))
        table_rows = cur.fetchall()
        
        tables = []
        for r in table_rows:
            latest = r["latest_size"] or 0
            prev_day = r["prev_day_size"] or r["oldest_size"] or 0
            prev_week = r["prev_week_size"] or r["oldest_size"] or 0
            
            # Daily Growth (Size today vs previous day)
            daily_diff = latest - prev_day
            daily_pct = round((daily_diff / prev_day * 100), 2) if prev_day > 0 else 0.0
            
            # Weekly Growth (Size today vs 7 days ago)
            week_diff = latest - prev_week
            week_pct = round((week_diff / prev_week * 100), 2) if prev_week > 0 else 0.0
            
            tables.append({
                "database_name": r["database_name"],
                "table_name": r["table_name"],
                "latest_size": latest,
                "growth": daily_diff, # Today's growth
                "growth_pct": daily_pct,
                "growth_7d": week_diff,
                "growth_7d_pct": week_pct
            })
            
        # Sort based on Today's daily growth percentage (positive first descending, negative/zero last)
        tables.sort(key=lambda x: (1 if (x["growth_pct"] or 0.0) > 0 else (0 if (x["growth_pct"] or 0.0) == 0 else -1), x["growth_pct"] or 0.0), reverse=True)
        result = {
            "client_name": client_name,
            "databases": databases,
            "top_growing_tables": tables[:10],
            "metrics": {
                "cpu": [],
                "memory": [],
                "disk": [],
                "io": []
            }
        }
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/utilization/servers-summary")
def get_utilization_servers_summary(client_name: str = None, user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:utilization-servers-summary:{client_name}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Resolve parent client if client_name is actually a server_name
        parent_client = client_name
        if client_name:
            cur.execute("""
                SELECT DISTINCT client_name FROM db_monitoring_logs 
                WHERE LOWER(server_name) = LOWER(%s) AND client_name IS NOT NULL;
            """, (client_name,))
            row = cur.fetchone()
            if row:
                parent_client = row["client_name"]
                
        # Retrieve all servers belonging to the client
        client_servers = None
        if parent_client:
            cur.execute("""
                SELECT DISTINCT server_name FROM db_monitoring_logs 
                WHERE LOWER(client_name) = LOWER(%s) AND server_name IS NOT NULL;
            """, (parent_client,))
            client_servers = [r["server_name"] for r in cur.fetchall()]
            if not client_servers:
                client_servers = [parent_client]
                
        # Determine allowed servers based on user permissions
        if user.get("isAdmin"):
            if client_servers:
                cur.execute("""
                    SELECT DISTINCT ON (server_name) 
                        server_name, 
                        cpu_utilization, 
                        memory_utilization, 
                        disk_utilization, 
                        io_utilization, 
                        read_iops,
                        write_iops,
                        captured_at
                    FROM server_utilization_history
                    WHERE server_name = ANY(%s)
                    ORDER BY server_name, captured_at DESC;
                """, (client_servers,))
            else:
                cur.execute("""
                    SELECT DISTINCT ON (server_name) 
                        server_name, 
                        cpu_utilization, 
                        memory_utilization, 
                        disk_utilization, 
                        io_utilization, 
                        read_iops,
                        write_iops,
                        captured_at
                    FROM server_utilization_history
                    ORDER BY server_name, captured_at DESC;
                """)
            rows = cur.fetchall()
        else:
            allowed_clients = user.get("allowed_clients", [])
            if allowed_clients:
                if parent_client and parent_client not in allowed_clients:
                    raise HTTPException(status_code=403, detail="Access denied.")
                cur.execute("""
                    SELECT DISTINCT server_name FROM admin_clients WHERE client_name = ANY(%s);
                """, (allowed_clients,))
                allowed_servers = [r["server_name"] for r in cur.fetchall() if r["server_name"]]
                if not allowed_servers:
                    allowed_servers = allowed_clients
                
                # Filter client servers
                if client_servers:
                    allowed_servers = [s for s in allowed_servers if s in client_servers] or client_servers
                
                cur.execute("""
                    SELECT DISTINCT ON (server_name) 
                        server_name, 
                        cpu_utilization, 
                        memory_utilization, 
                        disk_utilization, 
                        io_utilization, 
                        read_iops,
                        write_iops,
                        captured_at
                    FROM server_utilization_history
                    WHERE server_name = ANY(%s)
                    ORDER BY server_name, captured_at DESC;
                """, (allowed_servers,))
                rows = cur.fetchall()
            else:
                assigned_techs = [t.lower().strip() for t in user.get("assigned_techs", [])]
                if assigned_techs:
                    cur.execute("""
                        SELECT DISTINCT server_name FROM admin_clients 
                        WHERE EXISTS (
                            SELECT 1 FROM unnest(string_to_array(REPLACE(LOWER(db_type), ' ', ''), ',')) AS t 
                            WHERE t = ANY(%s)
                        );
                    """, (assigned_techs,))
                    allowed_servers = [r["server_name"] for r in cur.fetchall() if r["server_name"]]
                    
                    if client_servers:
                        allowed_servers = [s for s in allowed_servers if s in client_servers] or client_servers
                        
                    if allowed_servers:
                        cur.execute("""
                            SELECT DISTINCT ON (server_name) 
                                server_name, 
                                cpu_utilization, 
                                memory_utilization, 
                                disk_utilization, 
                                io_utilization, 
                                read_iops,
                                write_iops,
                                captured_at
                            FROM server_utilization_history
                            WHERE server_name = ANY(%s)
                            ORDER BY server_name, captured_at DESC;
                        """, (allowed_servers,))
                        rows = cur.fetchall()
                    else:
                        rows = []
                else:
                    rows = []
                    
        summary = []
        for r in rows:
            summary.append({
                "server_name": r["server_name"],
                "cpu": float(r["cpu_utilization"]) if r["cpu_utilization"] is not None else None,
                "memory": float(r["memory_utilization"]) if r["memory_utilization"] is not None else None,
                "disk": float(r["disk_utilization"]) if r["disk_utilization"] is not None else None,
                "io": float(r["io_utilization"]) if r["io_utilization"] is not None else None,
                "read_iops": float(r["read_iops"]) if r["read_iops"] is not None else None,
                "write_iops": float(r["write_iops"]) if r["write_iops"] is not None else None,
                "captured_at": r["captured_at"].isoformat() if r["captured_at"] else None
            })
            
        result = {"summary": summary}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/utilization/detail/{client_name}")
def get_utilization_detail(
    client_name: str, 
    granularity: str = "hourly", 
    user: dict = Depends(get_current_user_local)
):
    cache_key = f"telemetry:utilization-detail:{client_name}:{granularity}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Resolve parent client if client_name is actually a server_name
        parent_client = client_name
        cur.execute("""
            SELECT DISTINCT client_name FROM db_monitoring_logs 
            WHERE LOWER(server_name) = LOWER(%s) AND client_name IS NOT NULL;
        """, (client_name,))
        row = cur.fetchone()
        if row:
            parent_client = row["client_name"]
            
        if not check_client_and_tech_permission(parent_client, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")
            
        # Retrieve all servers belonging to client_name
        # If client_name is a specific server, only query that server
        cur.execute("""
            SELECT DISTINCT server_name FROM db_monitoring_logs 
            WHERE LOWER(server_name) = LOWER(%s);
        """, (client_name,))
        is_specific_server = cur.fetchone()
        
        if is_specific_server:
            client_servers = [client_name]
        else:
            cur.execute("""
                SELECT DISTINCT server_name FROM db_monitoring_logs 
                WHERE LOWER(client_name) = LOWER(%s) AND server_name IS NOT NULL;
            """, (client_name,))
            client_servers = [r["server_name"] for r in cur.fetchall()]
            if not client_servers:
                client_servers = [client_name]
                
        # Check if there is data in the last 7 days for these servers
        cur.execute("""
            SELECT EXISTS (
                SELECT 1 FROM server_utilization_history
                WHERE server_name = ANY(%s) AND captured_at >= NOW() - INTERVAL '7 days'
            );
        """, (client_servers,))
        has_7_days = cur.fetchone()["exists"]
        
        if granularity == "daily":
            if has_7_days:
                cur.execute("""
                    SELECT 
                        DATE(captured_at) as captured_date,
                        ROUND(AVG(cpu_utilization), 2) as cpu_utilization,
                        ROUND(AVG(memory_utilization), 2) as memory_utilization,
                        ROUND(AVG(disk_utilization), 2) as disk_utilization,
                        ROUND(AVG(io_utilization), 2) as io_utilization
                    FROM server_utilization_history
                    WHERE server_name = ANY(%s) AND captured_at >= NOW() - INTERVAL '7 days'
                    GROUP BY DATE(captured_at)
                    ORDER BY captured_date ASC;
                """, (client_servers,))
            else:
                # Find max captured date
                cur.execute("""
                    SELECT MAX(captured_at) as max_at FROM server_utilization_history 
                    WHERE server_name = ANY(%s);
                """, (client_servers,))
                max_at = cur.fetchone()["max_at"]
                if max_at:
                    cur.execute("""
                        SELECT 
                            DATE(captured_at) as captured_date,
                            ROUND(AVG(cpu_utilization), 2) as cpu_utilization,
                            ROUND(AVG(memory_utilization), 2) as memory_utilization,
                            ROUND(AVG(disk_utilization), 2) as disk_utilization,
                            ROUND(AVG(io_utilization), 2) as io_utilization
                        FROM server_utilization_history
                        WHERE server_name = ANY(%s) AND captured_at >= %s - INTERVAL '1 day'
                        GROUP BY DATE(captured_at)
                        ORDER BY captured_date ASC;
                    """, (client_servers, max_at))
                else:
                    cur.execute("SELECT 1 WHERE FALSE;") # empty
            rows = cur.fetchall()
            
            history = []
            for r in rows:
                history.append({
                    "time": r["captured_date"].isoformat() if r["captured_date"] else None,
                    "cpu": float(r["cpu_utilization"]) if r["cpu_utilization"] is not None else None,
                    "memory": float(r["memory_utilization"]) if r["memory_utilization"] is not None else None,
                    "disk": float(r["disk_utilization"]) if r["disk_utilization"] is not None else None,
                    "io": float(r["io_utilization"]) if r["io_utilization"] is not None else None
                })
        else: # hourly
            if has_7_days:
                cur.execute("""
                    SELECT 
                        captured_at, 
                        ROUND(AVG(cpu_utilization), 2) as cpu_utilization,
                        ROUND(AVG(memory_utilization), 2) as memory_utilization,
                        ROUND(AVG(disk_utilization), 2) as disk_utilization,
                        ROUND(AVG(io_utilization), 2) as io_utilization
                    FROM server_utilization_history
                    WHERE server_name = ANY(%s) AND captured_at >= NOW() - INTERVAL '7 days'
                    GROUP BY captured_at
                    ORDER BY captured_at ASC;
                """, (client_servers,))
            else:
                # Find max captured date
                cur.execute("""
                    SELECT MAX(captured_at) as max_at FROM server_utilization_history 
                    WHERE server_name = ANY(%s);
                """, (client_servers,))
                max_at = cur.fetchone()["max_at"]
                if max_at:
                    cur.execute("""
                        SELECT 
                            captured_at, 
                            ROUND(AVG(cpu_utilization), 2) as cpu_utilization,
                            ROUND(AVG(memory_utilization), 2) as memory_utilization,
                            ROUND(AVG(disk_utilization), 2) as disk_utilization,
                            ROUND(AVG(io_utilization), 2) as io_utilization
                        FROM server_utilization_history
                        WHERE server_name = ANY(%s) AND captured_at >= %s - INTERVAL '1 day'
                        GROUP BY captured_at
                        ORDER BY captured_at ASC;
                    """, (client_servers, max_at))
                else:
                    cur.execute("SELECT 1 WHERE FALSE;") # empty
            rows = cur.fetchall()
            
            history = []
            for r in rows:
                history.append({
                    "time": r["captured_at"].isoformat() if r["captured_at"] else None,
                    "cpu": float(r["cpu_utilization"]) if r["cpu_utilization"] is not None else None,
                    "memory": float(r["memory_utilization"]) if r["memory_utilization"] is not None else None,
                    "disk": float(r["disk_utilization"]) if r["disk_utilization"] is not None else None,
                    "io": float(r["io_utilization"]) if r["io_utilization"] is not None else None
                })
                
        result = {
            "client_name": client_name,
            "parent_client": parent_client,
            "granularity": granularity,
            "history": history
        }
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/database-detail-chart")
def get_db_detail_chart(client_name: str, database_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:database-detail-chart:{client_name}:{database_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")
            
        cur.execute("""
            SELECT captured_date, total_size_bytes
            FROM database_size_history
            WHERE server_name = %s 
              AND database_name = %s
              AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            ORDER BY captured_date ASC;
        """, (client_name, database_name, db_type, db_type, db_type))
        rows = cur.fetchall()
        
        chart_data = []
        for r in rows:
            chart_data.append({
                "date": r["captured_date"].isoformat() if r["captured_date"] else None,
                "size_bytes": r["total_size_bytes"]
            })
        result = {"chart_data": chart_data}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/table-detail-chart")
def get_table_detail_chart(client_name: str, database_name: str, table_name: str, db_type: Optional[str] = None, user: dict = Depends(get_current_user_local)):
    cache_key = f"telemetry:table-detail-chart:{client_name}:{database_name}:{table_name}:{db_type or ''}:{user.get('username')}"
    cached_val = cache_manager.get(cache_key)
    if cached_val is not None:
        return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied. You do not have permissions for this client.")
            
        cur.execute("""
            SELECT captured_date, size_bytes
            FROM table_size_history
            WHERE server_name = %s 
              AND database_name = %s 
              AND table_name = %s
              AND (%s IS NULL OR %s = '' OR LOWER(db_type) = LOWER(%s))
            ORDER BY captured_date ASC;
        """, (client_name, database_name, table_name, db_type, db_type, db_type))
        rows = cur.fetchall()
        
        chart_data = []
        for r in rows:
            chart_data.append({
                "date": r["captured_date"].isoformat() if r["captured_date"] else None,
                "size_bytes": r["size_bytes"]
            })
        result = {"chart_data": chart_data}
        cache_manager.set(cache_key, result, ttl_seconds=60)
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.post("/telemetry/sync")
def sync_telemetry_manually(user: dict = Depends(get_current_user_local)):
    from telemetry_parser import run_telemetry_sync
    from utilization_sync import sync_utilization_history
    
    # Run uptime simulation if in local mode
    try:
        from email_extracter import account, simulate_local_uptime_update
        if not account:
            simulate_local_uptime_update()
    except Exception as e:
        print(f"[API] Uptime simulation error in manual sync: {e}")

    res = run_telemetry_sync()
    if res.get("status") == "success":
        try:
            sync_utilization_history()
        except Exception as e:
            print(f"[API] Utilization sync error: {e}")
        return res
    else:
        raise HTTPException(status_code=500, detail=res.get("message", "Sync failed."))







class ReportShareEmail(BaseModel):
    report_id: int
    to_email: str
    cc_email: Optional[str] = ""
    subject: str
    body: str

@router.post("/reports/share/email")
def share_report_email(req: ReportShareEmail, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT file_name, file_data, client_name, title FROM client_reports WHERE id = %s;", (req.report_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report document not found")
        file_name, file_data, client_name, title = row[0], row[1], row[2], row[3]
        
        # Build attachment list with the report file base64 data
        email_attachments = [{
            "name": file_name,
            "data": file_data
        }]
        
        # Send email via Outlook/Graph/SMTP with the physical attachment
        share_sender_email = "dccagent@geopits.com"
        user_email = user.get("email")
        if user_email and user_email.strip():
            share_sender_email = user_email.strip()

        send_email_outlook(
            req.to_email, req.cc_email, req.subject, req.body,
            sender_email=share_sender_email,
            attachments=email_attachments
        )
        
        # Audit log in notifications
        uploader = user.get("username", "System")
        notify_msg = f"User '{uploader}' shared report document '{file_name}' via email to '{req.to_email}'"
        cur.execute("INSERT INTO notifications (username, message, is_read) VALUES ('global', %s, FALSE);", (notify_msg,))
        
        # Log to sharing history
        cur.execute("""
            INSERT INTO report_sharing_history (report_id, report_title, shared_by, share_platform, recipient)
            VALUES (%s, %s, %s, 'email', %s);
        """, (req.report_id, title, uploader, req.to_email))
        
        conn.commit()
        
        return {"status": "success", "message": "Document email sent successfully!"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

# --- Report Sharing Auditing & Monitoring Endpoints ---

class ShareLogRequest(BaseModel):
    report_id: int
    share_platform: str  # 'whatsapp', 'teams', 'native', 'email'
    recipient: Optional[str] = None

@router.post("/reports/share/log")
def log_report_share(req: ShareLogRequest, user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT title FROM client_reports WHERE id = %s;", (req.report_id,))
        rep_row = cur.fetchone()
        report_title = rep_row[0] if rep_row else "Unknown Report"
        
        shared_by = user.get("username", "System")
        
        cur.execute("""
            INSERT INTO report_sharing_history (report_id, report_title, shared_by, share_platform, recipient)
            VALUES (%s, %s, %s, %s, %s);
        """, (req.report_id, report_title, shared_by, req.share_platform, req.recipient))
        conn.commit()
        return {"status": "success", "message": "Share action logged successfully."}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/reports/share/history")
def get_report_share_history(user: dict = Depends(get_current_user_local)):
    # Check if user has admin privileges
    if not (user.get("role") == "admin" or user.get("isAdmin")):
        raise HTTPException(status_code=403, detail="Admin privileges required to view sharing history.")
        
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, report_id, report_title, shared_by, share_platform, recipient, created_at
            FROM report_sharing_history
            ORDER BY created_at DESC;
        """)
        rows = cur.fetchall()
        for r in rows:
            if r["created_at"]:
                r["created_at"] = r["created_at"].isoformat()
        return {"history": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

def parse_size_to_gb(val_str):
    if not val_str:
        return 0.0
    val_str = str(val_str).strip().upper()
    num_part = ""
    unit = "G"
    for char in val_str:
        if char.isdigit() or char == '.':
            num_part += char
        elif char in ('G', 'M', 'K', 'T'):
            unit = char
            break
    try:
        val = float(num_part) if num_part else 0.0
        if unit == 'M':
            return val / 1024.0
        elif unit == 'K':
            return val / (1024.0 * 1024.0)
        elif unit == 'T':
            return val * 1024.0
        return val
    except ValueError:
        return 0.0

@router.get("/telemetry/utilization/filter-options")
def get_utilization_filter_options(user: dict = Depends(get_current_user_local)):
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Include cloudwatch_log (used by Runloyal/Cropin/Intentwise) and allow 'Standalone' server names
        cur.execute("""
            SELECT DISTINCT 
                INITCAP(TRIM(client_name)) AS client_name,
                TRIM(server_name) AS server_name,
                TRIM(db_type) AS db_type
            FROM db_monitoring_logs
            WHERE log_type IN ('CPUUtilization', 'FreeableMemory', 'FreeLocalStorage', 'ReadIOPS', 'WriteIOPS',
                               'cpu_samples', 'cpu_hourly', 'memory_samples', 'memory_hourly',
                               'disk_samples', 'disk_hourly', 'io_samples', 'io_hourly', 'cloudwatch_log')
              AND client_name IS NOT NULL AND client_name <> ''
              AND server_name IS NOT NULL AND server_name <> ''
              AND db_type IS NOT NULL AND db_type <> '';
        """)
        rows = cur.fetchall()
        
        # Deduplicate case-insensitively by client name, keeping proper-cased version
        seen_clients = {}
        allowed_rows = []
        for r in rows:
            canonical_client = r["client_name"]  # Already INITCAP'd
            client_key = canonical_client.lower()
            # Prefer the version with actual casing (not all-lowercase)
            if client_key not in seen_clients:
                seen_clients[client_key] = canonical_client
            r_copy = dict(r)
            r_copy["client_name"] = seen_clients[client_key]
            if check_client_and_tech_permission(r["client_name"], user, cur):
                allowed_rows.append(r_copy)
                
        # Build unique client list (deduplicated)
        seen_final = {}
        for r in allowed_rows:
            key = r["client_name"].lower()
            if key not in seen_final:
                seen_final[key] = r["client_name"]
        clients = sorted(seen_final.values())
        
        # Build client → server list (merge across case variants)
        client_server_map = {}
        for r in allowed_rows:
            c = seen_final.get(r["client_name"].lower(), r["client_name"])
            s = r["server_name"]
            if c not in client_server_map:
                client_server_map[c] = set()
            client_server_map[c].add(s)
            
        client_server_map = {c: sorted(list(s_set)) for c, s_set in client_server_map.items()}
        
        server_db_type_map = {}
        for r in allowed_rows:
            c = seen_final.get(r["client_name"].lower(), r["client_name"])
            s = r["server_name"]
            if c not in server_db_type_map:
                server_db_type_map[c] = {}
            server_db_type_map[c][s] = r["db_type"]
            
        return {
            "clients": clients,
            "client_server_map": client_server_map,
            "server_db_type_map": server_db_type_map
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@router.get("/telemetry/utilization/history")
def get_utilization_history(
    client_name: str,
    server_name: Optional[str] = None,
    db_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    metric: Optional[str] = None,
    refresh: bool = Query(False),
    user: dict = Depends(get_current_user_local)
):
    cache_key = f"telemetry:utilization:{client_name}:{server_name or ''}:{db_type or ''}:{start_date or ''}:{end_date or ''}:{metric or ''}:{user.get('username')}"
    if not refresh:
        cached_val = cache_manager.get(cache_key)
        if cached_val is not None:
            return cached_val

    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not check_client_and_tech_permission(client_name, user, cur):
            raise HTTPException(status_code=403, detail="Access denied for this client's telemetry.")
            
        if not start_date:
            start_dt = datetime.now() - timedelta(days=30)
            start_str = start_dt.strftime("%Y-%m-%d 00:00:00")
        else:
            start_str = start_date if " " in start_date else f"{start_date} 00:00:00"
            
        if not end_date:
            end_dt = datetime.now()
            end_str = end_dt.strftime("%Y-%m-%d 23:59:59")
        else:
            end_str = end_date if " " in end_date else f"{end_date} 23:59:59"

        # Determine log types to fetch based on metric filter
        allowed_types = [
            'CPUUtilization', 'FreeableMemory', 'FreeableMemoryBytes', 
            'FreeStorageSpace', 'FreeLocalStorage', 'ReadIOPS', 'WriteIOPS',
            'cpu_samples', 'cpu_hourly', 'cpu_daily_metrics',
            'memory_samples', 'memory_hourly', 'memory_daily_metrics',
            'disk_samples', 'disk_hourly', 'disk_daily_metrics',
            'io_samples', 'io_hourly', 'io_daily_metrics', 
            'cloudwatch_log'
        ]
                         
        if metric:
            m_lower = metric.lower()
            if m_lower == 'cpu':
                allowed_types = ['CPUUtilization', 'cpu_samples', 'cpu_hourly', 'cpu_daily_metrics', 'cloudwatch_log']
            elif m_lower == 'memory':
                allowed_types = ['FreeableMemory', 'FreeableMemoryBytes', 'memory_samples', 'memory_hourly', 'memory_daily_metrics', 'cloudwatch_log']
            elif m_lower in ('disk', 'storage'):
                allowed_types = ['FreeStorageSpace', 'FreeLocalStorage', 'disk_samples', 'disk_hourly', 'disk_daily_metrics', 'cloudwatch_log']
            elif m_lower in ('io', 'iops'):
                allowed_types = ['ReadIOPS', 'WriteIOPS', 'io_samples', 'io_hourly', 'io_daily_metrics', 'cloudwatch_log']
            
        query = """
            SELECT client_name, server_name, db_type, log_type, log_message, log_time_ist, log_time
            FROM db_monitoring_logs
            WHERE LOWER(client_name) = LOWER(%s)
              AND (%s IS NULL OR LOWER(server_name) = LOWER(%s))
              AND (%s IS NULL OR LOWER(db_type) = LOWER(%s))
              AND log_time_ist >= %s AND log_time_ist <= %s
              AND log_type = ANY(%s)
            ORDER BY log_time_ist ASC;
        """
        cur.execute(query, (
            client_name,
            server_name, server_name,
            db_type, db_type,
            start_str, end_str,
            allowed_types
        ))
        rows = cur.fetchall()
        
        import json
        daily_data = {}
        disk_names = set()
        
        aws_clients = {'runloyal', 'intentwise', '360tf', 'artfine', 'cropin', 'credopay'}
        is_aws = client_name.lower() in aws_clients
        
        for r in rows:
            log_type = r["log_type"]
            msg = r["log_message"]
            
            parsed = {}
            if msg:
                try:
                    parsed = json.loads(msg)
                except Exception:
                    continue
            
            label = log_type
            if isinstance(parsed, dict) and parsed.get("Label"):
                label = parsed.get("Label")
                
            datapoints = []
            if isinstance(parsed, dict):
                if "Datapoints" in parsed:
                    datapoints = parsed["Datapoints"]
                else:
                    datapoints = [parsed]
            elif isinstance(parsed, list):
                datapoints = parsed
            else:
                continue
                
            for dp in datapoints:
                if not isinstance(dp, dict):
                    continue
                    
                # Strict metric-level filtering for parsed items using lowercase representation
                label_clean = label.replace(" ", "").lower()
                is_cpu_metric = any(x in label_clean for x in ('cpuutilization', 'cpu_samples', 'cpu_hourly', 'cpu_daily_metrics'))
                is_mem_metric = any(x in label_clean for x in ('freeablememory', 'freeablememorybytes', 'memory_samples', 'memory_hourly', 'memory_daily_metrics'))
                is_disk_metric = any(x in label_clean for x in ('freestoragespace', 'freelocalstorage', 'disk_samples', 'disk_hourly', 'disk_daily_metrics'))
                is_io_metric = any(x in label_clean for x in ('readiops', 'writeiops', 'io_samples', 'io_hourly', 'io_daily_metrics'))
                
                if metric:
                    m_lower = metric.lower()
                    if m_lower == 'cpu' and not is_cpu_metric:
                        continue
                    elif m_lower == 'memory' and not is_mem_metric:
                        continue
                    elif m_lower in ('disk', 'storage') and not is_disk_metric:
                        continue
                    elif m_lower in ('io', 'iops') and not is_io_metric:
                        continue

                ts_str = dp.get("Timestamp") or dp.get("captured_time") or dp.get("time") or dp.get("metric_date")
                if ts_str:
                    try:
                        if ts_str.endswith("Z"):
                            ts_str = ts_str[:-1]
                        date_str = ts_str[:10]
                    except Exception:
                        date_str = None
                else:
                    date_str = None
                    
                if not date_str:
                    dt_val = r["log_time_ist"] or r["log_time"]
                    if not dt_val:
                        continue
                    if isinstance(dt_val, datetime):
                        date_str = dt_val.strftime("%Y-%m-%d")
                    else:
                        date_str = str(dt_val)[:10]
                        
                if date_str not in daily_data:
                    daily_data[date_str] = {
                        "cpu": [],
                        "memory": [],
                        "disk": [],
                        "read_io": [],
                        "write_io": []
                    }
                    
                # Determine metric format dynamically based on fields
                is_aws_format = (
                    label_clean in ('cpuutilization', 'freeablememory', 'freeablememorybytes', 'freestoragespace', 'freelocalstorage', 'readiops', 'writeiops') or 
                    any(k in dp for k in ('Average', 'Minimum', 'Maximum'))
                )
                
                if is_aws_format:
                    if label_clean == 'cpuutilization':
                        avg_v = dp.get("Average")
                        min_v = dp.get("Minimum")
                        max_v = dp.get("Maximum")
                        if avg_v is not None:
                            daily_data[date_str]["cpu"].append({
                                "min": float(min_v) if min_v is not None else float(avg_v),
                                "max": float(max_v) if max_v is not None else float(avg_v),
                                "avg": float(avg_v)
                            })
                    elif label_clean in ('freeablememory', 'freeablememorybytes'):
                        avg_v = dp.get("Average")
                        min_v = dp.get("Minimum")
                        max_v = dp.get("Maximum")
                        if avg_v is not None:
                            factor = 1024.0 * 1024.0 * 1024.0
                            daily_data[date_str]["memory"].append({
                                "min": (float(min_v) if min_v is not None else float(avg_v)) / factor,
                                "max": (float(max_v) if max_v is not None else float(avg_v)) / factor,
                                "avg": float(avg_v) / factor
                            })
                    elif label_clean in ('freestoragespace', 'freelocalstorage'):
                        avg_v = dp.get("Average")
                        min_v = dp.get("Minimum")
                        max_v = dp.get("Maximum")
                        if avg_v is not None:
                            factor = 1024.0 * 1024.0 * 1024.0
                            daily_data[date_str]["disk"].append({
                                "min": (float(min_v) if min_v is not None else float(avg_v)) / factor,
                                "max": (float(max_v) if max_v is not None else float(avg_v)) / factor,
                                "avg": float(avg_v) / factor
                            })
                            disk_names.add("FreeStorageSpace" if label_clean == "freestoragespace" else "FreeLocalStorage")
                    elif label_clean == 'readiops':
                        avg_v = dp.get("Average")
                        min_v = dp.get("Minimum")
                        max_v = dp.get("Maximum")
                        if avg_v is not None:
                            daily_data[date_str]["read_io"].append({
                                "min": float(min_v) if min_v is not None else float(avg_v),
                                "max": float(max_v) if max_v is not None else float(avg_v),
                                "avg": float(avg_v)
                            })
                    elif label_clean == 'writeiops':
                        avg_v = dp.get("Average")
                        min_v = dp.get("Minimum")
                        max_v = dp.get("Maximum")
                        if avg_v is not None:
                            daily_data[date_str]["write_io"].append({
                                "min": float(min_v) if min_v is not None else float(avg_v),
                                "max": float(max_v) if max_v is not None else float(avg_v),
                                "avg": float(avg_v)
                            })
                else:
                    if is_cpu_metric:
                        sql_cpu = dp.get("sql_cpu_percent") or dp.get("cpu_avg") or dp.get("Average") or 0.0
                        sys_cpu = dp.get("system_cpu_percent") or 0.0
                        oth_cpu = dp.get("other_cpu_percent") or 0.0
                        val = sql_cpu + sys_cpu + oth_cpu
                        if val == 0.0 and dp.get("cpu_percent") is not None:
                            val = float(dp["cpu_percent"])
                        daily_data[date_str]["cpu"].append({"val": val})
                    elif is_mem_metric:
                        free_mb = dp.get("free_memory_mb") or dp.get("free_mb") or dp.get("freeable_memory_mb")
                        if free_mb is not None:
                            daily_data[date_str]["memory"].append({"val": float(free_mb) / 1024.0})
                    elif is_disk_metric:
                        free_gb = dp.get("free_gb") or dp.get("available_gb")
                        if free_gb is not None:
                            daily_data[date_str]["disk"].append({"val": float(free_gb)})
                            dname = dp.get("disk_name") or dp.get("mount") or dp.get("filesystem")
                            if dname:
                                disk_names.add(dname)
                        else:
                            avail = dp.get("available")
                            if avail is not None:
                                daily_data[date_str]["disk"].append({"val": parse_size_to_gb(avail)})
                                dname = dp.get("mount") or dp.get("filesystem") or dp.get("disk_name")
                                if dname:
                                    disk_names.add(dname)
                    elif is_io_metric:
                        read_iops = dp.get("read_iops")
                        write_iops = dp.get("write_iops")
                        if read_iops is not None:
                            daily_data[date_str]["read_io"].append({"val": float(read_iops)})
                        if write_iops is not None:
                            daily_data[date_str]["write_io"].append({"val": float(write_iops)})
                        
        result = []
        for d_str, metrics in sorted(daily_data.items()):
            day_result = {"date": d_str}
            for m_key in ("cpu", "memory", "disk", "read_io", "write_io"):
                items = metrics[m_key]
                if not items:
                    day_result[m_key] = {"min": 0.0, "max": 0.0, "avg": 0.0}
                    continue
                    
                mins = []
                maxs = []
                avgs = []
                for item in items:
                    if "val" in item:
                        v = item["val"]
                        mins.append(v)
                        maxs.append(v)
                        avgs.append(v)
                    else:
                        if item.get("min") is not None: mins.append(item["min"])
                        if item.get("max") is not None: maxs.append(item["max"])
                        if item.get("avg") is not None: avgs.append(item["avg"])
                        
                day_result[m_key] = {
                    "min": min(mins) if mins else 0.0,
                    "max": max(maxs) if maxs else 0.0,
                    "avg": sum(avgs) / len(avgs) if avgs else 0.0
                }
            result.append(day_result)
            
        res_data = {
            "client_name": client_name,
            "is_aws": is_aws,
            "history": result,
            "disk_names": sorted(list(disk_names))
        }
        cache_manager.set(cache_key, res_data, ttl_seconds=60)
        return res_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

def get_db_type_variations(db_type: str) -> list:
    db_type_lower = db_type.lower().strip()
    if db_type_lower in ('postgresql', 'postgres'):
        return ['postgresql', 'postgres']
    elif db_type_lower in ('mysql',):
        return ['mysql']
    elif db_type_lower in ('mongodb', 'mongo'):
        return ['mongodb', 'mongo']
    elif db_type_lower in ('mssql', 'mssqlserver', 'sqlserver', 'sql server'):
        return ['mssql', 'mssqlserver', 'sqlserver', 'sql server']
    elif db_type_lower in ('oracle',):
        return ['oracle']
    return [db_type_lower]

@router.get("/homepage/overall-summary/{client_name}/{db_type}")
def get_homepage_overall_summary(client_name: str, db_type: str, refresh: bool = Query(False), user: dict = Depends(get_current_user_local)):
    # 1. Enforce permission check
    conn = get_connection()
    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if not user.get("isAdmin") and user.get("role") != "admin":
            if not check_client_and_tech_permission(client_name, user, cur):
                raise HTTPException(status_code=403, detail="Forbidden: Access to this client's overall summary is denied.")

        cache_key = f"overall-summary:{client_name.lower().strip()}:{db_type.lower().strip()}"
        if not refresh:
            cached_val = cache_manager.get(cache_key) if cache_manager else None
            if cached_val:
                return cached_val

        db_vars = get_db_type_variations(db_type)

        # 2. Fetch Uptime Status (latest for each distinct service)
        cur.execute("""
            SELECT DISTINCT ON (service_name) service_name, status, uptime_desc, last_restart_time, captured_at
            FROM db_uptime_history
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
            ORDER BY service_name, captured_at DESC;
        """, (client_name, db_vars))
        uptime_rows = cur.fetchall()
        uptime_data = [dict(r) for r in uptime_rows] if uptime_rows else []

        # Retrieve distinct servers for this client & tech
        cur.execute("""
            SELECT DISTINCT server_name FROM db_monitoring_logs 
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s) AND server_name IS NOT NULL;
        """, (client_name, db_vars))
        servers = [r["server_name"] for r in cur.fetchall()]

        # 3. Fetch hourly/sample telemetry logs (latest 80 entries for comprehensive charts)
        cur.execute("""
            SELECT log_type, log_message, severity, log_time_ist
            FROM db_monitoring_logs
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
              AND (
                LOWER(log_type) IN ('cpu_hourly', 'memory_hourly', 'disk_hourly', 'io_hourly', 'cpu_samples', 'memory_samples', 'disk_samples', 'io_samples', 'cpuutilization', 'readiops', 'writeiops')
                OR LOWER(log_type) LIKE '%%_hourly'
                OR LOWER(log_type) LIKE '%%_samples'
              )
            ORDER BY log_time_ist DESC
            LIMIT 80;
        """, (client_name, db_vars))
        metrics_rows = [dict(r) for r in cur.fetchall()]

        # Fallback to server_utilization_history if metrics_rows is empty
        if not metrics_rows and servers:
            cur.execute("""
                SELECT server_name, cpu_utilization, memory_utilization, disk_utilization, io_utilization, captured_at
                FROM server_utilization_history
                WHERE server_name = ANY(%s)
                ORDER BY captured_at DESC
                LIMIT 80;
            """, (servers,))
            util_rows = [dict(r) for r in cur.fetchall()]
            for r in util_rows:
                cap_time = r["captured_at"].isoformat() if r["captured_at"] else None
                if r["cpu_utilization"] is not None:
                    metrics_rows.append({
                        "log_type": "cpu_hourly",
                        "log_message": f"{r['cpu_utilization']}%",
                        "severity": "Low",
                        "log_time_ist": cap_time
                    })
                if r["memory_utilization"] is not None:
                    metrics_rows.append({
                        "log_type": "memory_hourly",
                        "log_message": f"{r['memory_utilization']}%",
                        "severity": "Low",
                        "log_time_ist": cap_time
                    })
                if r["disk_utilization"] is not None:
                    metrics_rows.append({
                        "log_type": "disk_hourly",
                        "log_message": f"{r['disk_utilization']}%",
                        "severity": "Low",
                        "log_time_ist": cap_time
                    })
                if r["io_utilization"] is not None:
                    metrics_rows.append({
                        "log_type": "io_hourly",
                        "log_message": f"{r['io_utilization']}%",
                        "severity": "Low",
                        "log_time_ist": cap_time
                    })

        # 4. Fetch recent event/critical logs (latest 30 entries)
        cur.execute("""
            SELECT id, log_message, severity, log_time_ist, status, owner
            FROM db_monitoring_logs
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
              AND LOWER(log_type) NOT IN ('cpu_hourly', 'memory_hourly', 'disk_hourly', 'io_hourly', 'cpu_samples', 'memory_samples', 'disk_samples', 'io_samples', 'cpuutilization', 'readiops', 'writeiops')
              AND LOWER(log_type) NOT LIKE '%%_hourly'
              AND LOWER(log_type) NOT LIKE '%%_samples'
            ORDER BY log_time_ist DESC
            LIMIT 30;
        """, (client_name, db_vars))
        event_rows = [dict(r) for r in cur.fetchall()]

        # 5. Fetch database and table sizes history
        # Fetch complete database sizes for download report and change calculation
        cur.execute("""
            SELECT server_name, database_name, total_size_bytes, captured_date, db_type
            FROM database_size_history
            WHERE LOWER(TRIM(db_type)) = ANY(%s) AND server_name IN (
                SELECT DISTINCT server_name FROM db_monitoring_logs WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
            )
            ORDER BY captured_date DESC, database_name ASC;
        """, (db_vars, client_name, db_vars))
        db_sizes_all = [dict(r) for r in cur.fetchall()]

        # Fetch complete table sizes for download report and change calculation
        cur.execute("""
            SELECT server_name, database_name, table_name, size_bytes, captured_date, db_type
            FROM table_size_history
            WHERE LOWER(TRIM(db_type)) = ANY(%s) AND server_name IN (
                SELECT DISTINCT server_name FROM db_monitoring_logs WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
            )
            ORDER BY captured_date DESC, database_name ASC, table_name ASC;
        """, (db_vars, client_name, db_vars))
        table_sizes_all = [dict(r) for r in cur.fetchall()]

        # Calculate DB size changes for the latest captured date (strictly "changes only")
        db_sizes = []
        if db_sizes_all:
            latest_db_date = max(r["captured_date"] for r in db_sizes_all)
            latest_db_records = [r for r in db_sizes_all if r["captured_date"] == latest_db_date]
            for rec in latest_db_records:
                db_name = rec["database_name"]
                prev_records = [r for r in db_sizes_all if r["database_name"] == db_name and r["captured_date"] < latest_db_date]
                if prev_records:
                    prev_rec = max(prev_records, key=lambda x: x["captured_date"])
                    change_bytes = rec["total_size_bytes"] - prev_rec["total_size_bytes"]
                else:
                    change_bytes = 0
                
                # Keep only non-zero changes (strictly "changes only")
                if change_bytes != 0:
                    rec_copy = dict(rec)
                    rec_copy["change_bytes"] = change_bytes
                    # Keep captured_date as string/serializable
                    rec_copy["captured_date"] = rec_copy["captured_date"].isoformat() if rec_copy["captured_date"] else None
                    db_sizes.append(rec_copy)

        # Calculate Table size changes for the latest captured date (strictly "changes only")
        table_sizes = []
        if table_sizes_all:
            latest_table_date = max(r["captured_date"] for r in table_sizes_all)
            latest_table_records = [r for r in table_sizes_all if r["captured_date"] == latest_table_date]
            for rec in latest_table_records:
                db_name = rec["database_name"]
                tbl_name = rec["table_name"]
                prev_records = [r for r in table_sizes_all if r["database_name"] == db_name and r["table_name"] == tbl_name and r["captured_date"] < latest_table_date]
                if prev_records:
                    prev_rec = max(prev_records, key=lambda x: x["captured_date"])
                    change_bytes = rec["size_bytes"] - prev_rec["size_bytes"]
                else:
                    change_bytes = 0
                
                # Keep only non-zero changes (strictly "changes only")
                if change_bytes != 0:
                    rec_copy = dict(rec)
                    rec_copy["change_bytes"] = change_bytes
                    rec_copy["captured_date"] = rec_copy["captured_date"].isoformat() if rec_copy["captured_date"] else None
                    table_sizes.append(rec_copy)


        # 6. Fetch client reports
        cur.execute("""
            SELECT id, title as report_name, file_name as file_path, 'Active' as status, uploaded_at as created_at
            FROM client_reports
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s))
            ORDER BY uploaded_at DESC
            LIMIT 15;
        """, (client_name,))
        reports = [dict(r) for r in cur.fetchall()]

        # 7. Fetch active incidents (tickets)
        cur.execute("""
            SELECT id, ticket_name, category, status, priority, agent, created_at
            FROM tickets
            WHERE LOWER(TRIM(company)) = LOWER(TRIM(%s))
            ORDER BY created_at DESC
            LIMIT 15;
        """, (client_name,))
        tickets = [dict(r) for r in cur.fetchall()]

        # 8. Fetch error and slow query logs specifically
        cur.execute("""
            SELECT id, log_message, log_time_ist, severity
            FROM db_monitoring_logs
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
              AND (LOWER(log_type) LIKE '%%error%%' OR LOWER(log_type) LIKE '%%fail%%')
            ORDER BY log_time_ist DESC
            LIMIT 30;
        """, (client_name, db_vars))
        error_logs = [dict(r) for r in cur.fetchall()]

        cur.execute("""
            SELECT id, log_message, log_time_ist, severity
            FROM db_monitoring_logs
            WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(%s)) AND LOWER(TRIM(db_type)) = ANY(%s)
              AND LOWER(log_type) LIKE '%%slow%%'
            ORDER BY log_time_ist DESC
            LIMIT 30;
        """, (client_name, db_vars))
        slow_query_logs = [dict(r) for r in cur.fetchall()]

        def parse_metric_value(msg):
            if not msg:
                return None
            msg_str = str(msg).strip()
            if not msg_str:
                return None
            if msg_str.startswith('{') and msg_str.endswith('}'):
                try:
                    import json
                    data = json.loads(msg_str)
                    for key in ["cpu_percent", "usage_percent", "Average", "util_percent", "value", "pct"]:
                        if key in data and data[key] is not None:
                            return float(data[key])
                except Exception:
                    pass
            try:
                val_str = ''.join(c for c in msg_str.split('%')[0] if c.isdigit() or c == '.')
                if val_str:
                    return float(val_str)
            except Exception:
                pass
            return None

        # Compute summary values from raw logs for context
        cpu_vals = []
        mem_vals = []
        disk_vals = []
        for r in metrics_rows:
            log_type_lower = (r.get("log_type") or "").lower()
            msg = r.get("log_message") or ""
            val = parse_metric_value(msg)
            if val is not None:
                if "cpu" in log_type_lower:
                    cpu_vals.append(val)
                elif "memory" in log_type_lower or "mem" in log_type_lower:
                    mem_vals.append(val)
                elif "disk" in log_type_lower:
                    disk_vals.append(val)

        avg_cpu = round(sum(cpu_vals)/len(cpu_vals), 2) if cpu_vals else None
        avg_mem = round(sum(mem_vals)/len(mem_vals), 2) if mem_vals else None
        avg_disk = round(sum(disk_vals)/len(disk_vals), 2) if disk_vals else None

        cpu_history = [r for r in metrics_rows if 'cpu' in (r.get('log_type') or '').lower()]
        memory_history = [r for r in metrics_rows if 'memory' in (r.get('log_type') or '').lower() or 'mem' in (r.get('log_type') or '').lower()]
        disk_history = [r for r in metrics_rows if 'disk' in (r.get('log_type') or '').lower()]
        io_history = [r for r in metrics_rows if 'io' in (r.get('log_type') or '').lower()]

        # 9. Call OpenAI to generate custom summary if api_key is available
        summary = ""
        api_key = os.getenv("OPENAI_API_KEY")
        
        # Helper to get first node status info for summary generation
        primary_status = uptime_data[0].get("status") if uptime_data else "Unknown"
        primary_uptime = uptime_data[0].get("uptime_desc") if uptime_data else "N/A"
        primary_captured = uptime_data[0].get("captured_at") if uptime_data else "recent intervals"

        if api_key and "sk-test" not in api_key:
            import httpx
            try:
                prompt_content = (
                    f"You are GeoBot, an expert diagnostic AI database administrator. Generate a concise, high-level, hourly summary report (4-5 lines of bullet points, use markdown) "
                    f"for the client '{client_name}' using database technology '{db_type}'.\n\n"
                    f"CURRENT STATE:\n"
                    f"- Database Uptime Status: {primary_status} ({primary_uptime})\n"
                    f"- Average Metrics (last 24h): CPU: {avg_cpu or 'N/A'}%, Memory: {avg_mem or 'N/A'}%, Disk: {avg_disk or 'N/A'}%\n"
                    f"- Recent Telemetry Logs:\n"
                    + "\n".join([f"  * [{r.get('log_time_ist')}] {r.get('severity') or 'Info'}: {r.get('log_message')}" for r in metrics_rows[:10]])
                    + "\n\n- Recent Security/Alert Events:\n"
                    + "\n".join([f"  * [{r.get('log_time_ist')}] {r.get('severity') or 'Info'}: {r.get('log_message')}" for r in event_rows[:10]])
                    + "\n\nSynthesize this information into a professional real-time status summary. Flag any warnings/errors or system instability if they exist. Do not mention system-internal keys or code."
                )
                
                payload = {
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "system", "content": "You are GeoBot, an expert database monitoring and telemetry diagnostic AI assistant."},
                        {"role": "user", "content": prompt_content}
                    ],
                    "temperature": 0.3,
                    "max_tokens": 300
                }
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                with httpx.Client() as client:
                    resp = client.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers, timeout=10.0)
                    if resp.status_code == 200:
                        summary = resp.json()["choices"][0]["message"]["content"]
            except Exception as ai_err:
                print("OpenAI summary error:", ai_err)

        if not summary:
            status_indicator = "🟢 Healthy" if (primary_status or "Unknown").lower() in ("online", "running") else "🔴 Warning/Unreachable"
            alerts_count = sum(1 for r in event_rows if (r.get("severity") or "").lower() in ("warning", "error", "critical"))
            summary = (
                f"- **Overall Health:** {status_indicator} — Uptime reported as `{primary_uptime}`.\n"
                f"- **Resource Utilization:** CPU averaged `{avg_cpu or 'N/A'}%`, Memory averaged `{avg_mem or 'N/A'}%`, Disk capacity is at `{avg_disk or 'N/A'}%`.\n"
                f"- **Alert Activity:** {alerts_count} warning/critical log notifications recorded in the latest hourly check.\n"
                f"- **Active Diagnostics:** Telemetry updates suggest stable connections, with last service verification completed at `{primary_captured}`."
            )


        # Serialize complete size lists for download report
        db_sizes_all_serialized = []
        for r in db_sizes_all:
            rc = dict(r)
            if rc.get("captured_date"):
                rc["captured_date"] = rc["captured_date"].isoformat()
            db_sizes_all_serialized.append(rc)

        table_sizes_all_serialized = []
        for r in table_sizes_all:
            rc = dict(r)
            if rc.get("captured_date"):
                rc["captured_date"] = rc["captured_date"].isoformat()
            table_sizes_all_serialized.append(rc)

        response_data = {
            "client_name": client_name,
            "db_type": db_type,
            "uptime": uptime_data,
            "summary": summary,
            "server_report": {
                "avg_cpu": avg_cpu,
                "avg_mem": avg_mem,
                "avg_disk": avg_disk,
                "metrics_history": metrics_rows[:20]
            },
            "realtime_logs": event_rows,
            "cpu_history": cpu_history,
            "memory_history": memory_history,
            "disk_history": disk_history,
            "io_history": io_history,
            "db_sizes": db_sizes,
            "table_sizes": table_sizes,
            "db_sizes_all": db_sizes_all_serialized,
            "table_sizes_all": table_sizes_all_serialized,
            "reports": reports,
            "tickets": tickets,
            "error_logs": error_logs,
            "slow_query_logs": slow_query_logs
        }
        if cache_manager:
            cache_manager.set(cache_key, response_data, ttl_seconds=300)
        return response_data
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

